# vysti_api.py

from dotenv import load_dotenv
load_dotenv()

import os
import io
import json
import base64
import hashlib
import time
import random
import pathlib
from io import BytesIO
from scoring import compute_scores as _compute_scores
from pdf_extract import extract_text_from_pdf, PDFExtractionError
from ocr_transcribe import transcribe_scanned_pdf
import urllib.parse

import httpx
import stripe
from collections import Counter
from fastapi import (
    FastAPI,
    File,
    UploadFile,
    Form,
    Depends,
    HTTPException,
    status,
)
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree
import re
import math
import difflib
import datetime
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from starlette.requests import Request
from starlette.middleware.base import BaseHTTPMiddleware


class _SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Add browser-level security headers to every response."""

    _CSP = "; ".join([
        "default-src 'self'",
        "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://challenges.cloudflare.com",
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com",
        "font-src 'self' https://fonts.gstatic.com",
        "img-src 'self' data: blob: https://*.googleusercontent.com https://*.supabase.co",
        "connect-src 'self' https://*.supabase.co https://checkout.stripe.com",
        "frame-src https://challenges.cloudflare.com",
        "object-src 'none'",
    ])

    async def dispatch(self, request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Permissions-Policy"] = (
            "camera=(), microphone=(), geolocation=()"
        )
        response.headers["Content-Security-Policy"] = self._CSP
        if request.url.scheme == "https":
            response.headers["Strict-Transport-Security"] = (
                "max-age=31536000; includeSubDomains"
            )
        return response


def _sanitize_filename(name: str) -> str:
    """Strip path traversal sequences and unsafe characters from a filename.

    Removes directory components (../, /, \\) and control characters so the
    name is safe for use in storage paths and Content-Disposition headers.
    """
    import re as _re
    # Take only the basename (strip any directory path)
    name = name.replace("\\", "/")
    name = name.rsplit("/", 1)[-1]
    # Remove null bytes, control characters, and double-quotes (for headers)
    name = _re.sub(r'[\x00-\x1f\x7f"]+', "", name)
    # Collapse leading/trailing dots and whitespace
    name = name.strip(". ")
    return name or "document.docx"


def _sanitize_for_json(obj):
    """Replace NaN/Infinity floats with None so json.dumps won't choke."""
    if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
        return None
    if isinstance(obj, dict):
        return {k: _sanitize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_sanitize_for_json(v) for v in obj]
    return obj


def _strip_ip_from_issues(issues: list) -> list:
    """Remove proprietary rule explanations from issues before sending to client.

    Replaces 'explanation' with the generalized 'shared_explanation' and
    removes raw detection context fields that reveal engine internals.
    """
    stripped = []
    for issue in issues:
        if not isinstance(issue, dict):
            stripped.append(issue)
            continue
        clean = dict(issue)
        # Replace full explanation with shared (generalized) version
        clean["explanation"] = clean.pop("shared_explanation", "")
        # Keep label, short_explanation, student_guidance, shared_issue, count
        stripped.append(clean)
    return stripped


def _strip_ip_from_examples(examples: list) -> list:
    """Remove engine-internal detection context from examples before sending to client.

    Strips fields that reveal how the marking engine detected issues (found_value,
    topics, thesis, confidence) while keeping what the UI needs (label, sentence,
    paragraph_index, shared_issue, original_phrase).
    """
    _INTERNAL_FIELDS = {"topics", "thesis", "confidence"}
    stripped = []
    for ex in examples:
        if not isinstance(ex, dict):
            stripped.append(ex)
            continue
        clean = {k: v for k, v in ex.items() if k not in _INTERNAL_FIELDS}
        stripped.append(clean)
    return stripped


app = FastAPI(
    title="Vysti Marker API",
    docs_url=None,
    redoc_url=None,
    openapi_url=None,
)

# ===== Rate limiting (per-user via JWT, fallback to IP) =====
def _get_user_rate_limit_key(request: Request) -> str:
    """Extract user ID from JWT or API key prefix for per-caller rate limiting."""
    # API key clients: rate limit by key prefix
    api_key = request.headers.get("X-API-Key", "")
    if api_key:
        return f"apikey:{hashlib.sha256(api_key.encode()).hexdigest()[:16]}"
    # Supabase JWT users: rate limit by user ID
    try:
        auth = request.headers.get("authorization", "")
        if auth.startswith("Bearer "):
            payload_b64 = auth[7:].split(".")[1]
            payload_b64 += "=" * (4 - len(payload_b64) % 4)
            uid = json.loads(base64.urlsafe_b64decode(payload_b64)).get("sub")
            if uid:
                return f"user:{uid}"
    except Exception:
        pass
    return get_remote_address(request)

limiter = Limiter(key_func=_get_user_rate_limit_key)
app.state.limiter = limiter

@app.exception_handler(RateLimitExceeded)
async def _rate_limit_handler(request: Request, exc: RateLimitExceeded):
    return JSONResponse(
        status_code=429,
        content={"error": "Rate limit exceeded. Please slow down."},
    )

_ALLOWED_ORIGINS = [
    "https://app.vysti.org",
    "https://vysti-rules.onrender.com",
    "http://localhost:8000",
    "http://localhost:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-API-Key"],
)
app.add_middleware(_SecurityHeadersMiddleware)

# Mount static files
app.mount("/assets", StaticFiles(directory="assets"), name="assets")
app.mount("/shared", StaticFiles(directory="shared"), name="shared")

# Mount OCR router for mobile handwriting transcription
try:
    from handwriting_ocr_router import ocr_router as _ocr_router
    app.include_router(_ocr_router, prefix="/api/ocr")
except ImportError:
    pass  # OCR dependencies not installed — skip silently

# ===== Debug flag (set VYSTI_DEBUG=1 in env to enable verbose logging) =====
_DEBUG = os.getenv("VYSTI_DEBUG", "").strip() in ("1", "true", "yes")

# ===== Application caps =====
_HARD_WORD_LIMIT = 10_000       # Reject essays exceeding this word count
_SOFT_WORD_LIMIT = 5_000        # Warn (in metadata) for essays above this
_MAX_MARK_EVENTS_PER_USER = 200 # Rolling retention cap per user
_active_marks: set[str] = set() # Concurrency guard: one mark at a time per user

# ===== Lazy engine loader =====
_ENGINE = None

def get_engine():
    """
    Lazy-load the marker engine so the FastAPI app can boot even if spaCy/model
    import is slow or fails. Import errors become a 503 at request-time instead
    of killing the whole ASGI startup.
    """
    global _ENGINE
    if _ENGINE is not None:
        return _ENGINE
    try:
        from marker import mark_docx_bytes, extract_summary_metadata
        _ENGINE = (mark_docx_bytes, extract_summary_metadata)
        return _ENGINE
    except Exception as e:
        print("Failed to import marker engine:", repr(e))
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Marker engine is temporarily unavailable. Please try again later.",
        )

# ===== Supabase config (from environment variables) =====
# ===== Supabase config (from environment variables) =====
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

# ===== Stripe config =====
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")
STRIPE_PRICE_MARK = os.getenv("STRIPE_PRICE_MARK")
STRIPE_PRICE_REVISE = os.getenv("STRIPE_PRICE_REVISE")
STRIPE_PRICE_BOTH = os.getenv("STRIPE_PRICE_BOTH")

if STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY

auth_scheme = HTTPBearer(auto_error=False)


# ===== Pydantic models =====
class RevisionCheckRequest(BaseModel):
    label: str
    label_trimmed: str | None = None
    rewrite: str
    mode: str | None = None
    context_text: str | None = None
    original_sentence: str | None = None
    paragraph_index: int | None = None
    titles: list["TitleInfo"] | None = None


class TitleInfo(BaseModel):
    author: str
    title: str
    is_minor: bool = True


class MarkTextRequest(BaseModel):
    file_name: str
    text: str
    mode: str = "student"
    titles: list[TitleInfo] | None = None
    student_mode: bool = True
    include_summary_table: bool | None = False
    return_metadata: bool = False
    source: str | None = None  # "mobile" when sent from mobile app
    # Teacher rule overrides (optional — sent by teacher recheck)
    forbid_personal_pronouns: bool | None = None
    forbid_audience_reference: bool | None = None
    enforce_closed_thesis: bool | None = None
    require_body_evidence: bool | None = None
    allow_intro_summary_quotes: bool | None = None
    enforce_intro_quote_rule: bool | None = None
    enforce_long_quote_rule: bool | None = None
    enforce_contractions_rule: bool | None = None
    enforce_which_rule: bool | None = None
    enforce_weak_verbs_rule: bool | None = None
    enforce_fact_proof_rule: bool | None = None
    enforce_human_people_rule: bool | None = None
    enforce_vague_terms_rule: bool | None = None
    highlight_thesis_devices: bool | None = None


class ExportDocxRequest(BaseModel):
    file_name: str
    text: str


class ExportTeacherDocxRequest(BaseModel):
    file_name: str = ""
    text: str = ""
    comment: str = ""


class DeleteMarkEventsRequest(BaseModel):
    file_names: list[str]


class SourceWork(BaseModel):
    author: str = ""
    title: str = ""


class UpdateMarkEventRequest(BaseModel):
    file_name: str
    mark_event_id: str | None = None  # UUID — if set, target this specific record instead of file_name
    assignment_name: str | None = None
    essay_title: str | None = None
    source_works: list[SourceWork] | None = None
    notes: str | None = None
    student_name: str | None = None
    class_id: str | None = None
    teacher_comment: str | None = None
    review_status: str | None = None
    score: int | None = None
    created_at: str | None = None


class ErrorReportRequest(BaseModel):
    message: str
    debug_info: dict | None = None
    page_url: str | None = None


class AutoErrorLogRequest(BaseModel):
    error_type: str
    message: str
    details: dict | None = None
    page_url: str | None = None
    build_id: str | None = None
    assignment_name: str | None = None
    essay_title: str | None = None
    source_works: list[SourceWork] | None = None
    notes: str | None = None
    student_name: str | None = None
    class_id: str | None = None
    teacher_comment: str | None = None
    review_status: str | None = None  # 'pending' | 'in_progress' | 'completed' | 'archived'
    score: int | None = None  # 0-100 percentage score
    created_at: str | None = None  # ISO-8601 datetime string



async def get_current_user(
    request: Request,
    cred: HTTPAuthorizationCredentials = Depends(auth_scheme),
):
    """
    Validate the Supabase JWT by calling Supabase's Auth API.
    Returns the user dict if valid; otherwise raises 401.
    """
    # Localhost dev bypass — skip JWT validation for local testing
    host = request.headers.get("host", "")
    if (host.startswith("localhost") or host.startswith("127.0.0.1")) and (cred is None or cred.credentials == "dev"):
        return {"id": "local-dev", "email": "dev@localhost"}

    if cred is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing Authorization header",
        )

    token = cred.credentials

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise HTTPException(
            status_code=500,
            detail="Supabase config missing on server",
        )

    auth_url = f"{SUPABASE_URL}/auth/v1/user"

    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.get(
            auth_url,
            headers={
                "apikey": SUPABASE_ANON_KEY,
                "Authorization": f"Bearer {token}",
            },
        )

    if resp.status_code != 200:
        # Invalid or expired token
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired token",
        )

    return resp.json()


# ── API Key Authentication (B2B licensing) ───────────────────────────

async def _lookup_api_key(raw_key: str) -> dict | None:
    """Look up an API key by its SHA-256 hash. Returns the api_keys row or None."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return None
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    url = f"{SUPABASE_URL}/rest/v1/api_keys?key_hash=eq.{key_hash}&select=*"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })
    if resp.status_code != 200:
        return None
    rows = resp.json()
    return rows[0] if rows else None


async def get_api_client(request: Request) -> dict | None:
    """
    Check if the request carries a valid X-API-Key header.
    Returns an api_keys row dict if valid, None if no key present.
    Raises 401/403 if the key is present but invalid/expired/revoked.
    """
    raw_key = request.headers.get("X-API-Key")
    if not raw_key:
        return None

    api_key = await _lookup_api_key(raw_key)

    if not api_key:
        raise HTTPException(status_code=401, detail="Invalid API key.")

    if not api_key.get("is_active", False):
        raise HTTPException(status_code=403, detail="API key has been revoked.")

    # Check expiry
    expires = api_key.get("expires_at")
    if expires:
        from datetime import datetime, timezone
        try:
            exp_dt = datetime.fromisoformat(expires.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) > exp_dt:
                raise HTTPException(status_code=403, detail="API key has expired.")
        except (ValueError, TypeError):
            pass

    # Check IP allowlist
    allowed_ips = api_key.get("allowed_ips")
    if allowed_ips:
        client_ip = get_remote_address(request)
        if client_ip not in allowed_ips:
            raise HTTPException(
                status_code=403,
                detail="Request from unauthorized IP address.",
            )

    # Check monthly quota
    monthly_quota = api_key.get("monthly_quota")
    if monthly_quota is not None:
        usage_count = await _count_api_usage_this_month(api_key["id"])
        if usage_count >= monthly_quota:
            raise HTTPException(
                status_code=429,
                detail=f"Monthly quota of {monthly_quota} requests exceeded.",
            )

    return api_key


async def _count_api_usage_this_month(api_key_id: str) -> int:
    """Count API usage for the current calendar month."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return 0
    now = datetime.datetime.now(datetime.timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    url = (
        f"{SUPABASE_URL}/rest/v1/api_usage"
        f"?api_key_id=eq.{api_key_id}"
        f"&created_at=gte.{month_start}"
        f"&select=id"
    )
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Prefer": "count=exact",
                "Range-Unit": "items",
                "Range": "0-0",
            })
        content_range = resp.headers.get("content-range", "")
        return int(content_range.split("/")[-1])
    except (ValueError, IndexError, Exception):
        return 0


async def _log_api_usage(
    api_key_id: str,
    endpoint: str,
    status_code: int,
    chars_processed: int,
    response_ms: int,
    client_ip: str,
    metadata: dict | None = None,
):
    """Log an API-key request to the api_usage table (best-effort)."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return
    try:
        url = f"{SUPABASE_URL}/rest/v1/api_usage"
        payload = {
            "api_key_id": api_key_id,
            "endpoint": endpoint,
            "status_code": status_code,
            "chars_processed": chars_processed,
            "response_ms": response_ms,
            "client_ip": client_ip,
            "metadata": metadata,
        }
        async with httpx.AsyncClient(timeout=5) as client:
            await client.post(url, json=payload, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type": "application/json",
                "Prefer": "return=minimal",
            })
    except Exception as e:
        print("Failed to log api_usage:", repr(e))


def require_api_product(*products: str):
    """
    FastAPI dependency that accepts EITHER a Supabase JWT user OR an API key.
    API key clients bypass product checks (their access is controlled by
    allowed_endpoints in the api_keys table).
    Existing Supabase users go through the normal product check.
    """
    async def _check(
        request: Request,
        cred: HTTPAuthorizationCredentials = Depends(auth_scheme),
    ):
        # 1. Try API key first
        api_client = await get_api_client(request)
        if api_client:
            # Check endpoint access
            endpoint = request.url.path
            allowed = api_client.get("allowed_endpoints", [])
            if allowed and endpoint not in allowed:
                raise HTTPException(
                    status_code=403,
                    detail=f"API key does not have access to {endpoint}.",
                )
            # Return a synthetic user dict that marks this as an API client
            return {
                "id": f"apikey:{api_client['id']}",
                "email": api_client.get("contact_email", ""),
                "_is_api_client": True,
                "_api_key_id": api_client["id"],
                "_api_key_rate_limit": api_client.get("rate_limit", 20),
                "_client_name": api_client.get("client_name", ""),
            }

        # 2. Fall back to normal Supabase JWT auth
        user = await get_current_user(request, cred)
        user_id = user.get("id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Could not determine user")
        if user_id == "local-dev":
            return user
        profile = await get_user_profile(user_id)
        if not profile:
            raise HTTPException(
                status_code=403,
                detail="No profile found. Please complete onboarding.",
            )
        product_map = {
            "mark": profile.get("has_mark", False),
            "revise": profile.get("has_revise", False),
            "write": profile.get("has_write", False),
        }
        if not any(product_map.get(p, False) for p in products):
            raise HTTPException(
                status_code=403,
                detail=f"This feature requires one of: {', '.join(products)}",
            )
        return user
    return _check


async def _revoke_expired_coupon_access(user_id: str, profile: dict) -> dict:
    """Check if the user's access came from an expired coupon and revoke it.

    Returns the (possibly updated) profile dict.  Skips users with an
    active Stripe subscription — their access is paid, not coupon-based.
    """
    # Don't touch paying customers
    if profile.get("subscription_status") == "active":
        return profile

    # Only check users who currently have some product access
    if not profile.get("has_mark") and not profile.get("has_revise"):
        return profile

    try:
        # Fetch this user's coupon redemptions with coupon details
        url = (
            f"{SUPABASE_URL}/rest/v1/coupon_redemptions"
            f"?user_id=eq.{user_id}"
            f"&select=coupon_id,coupon_codes(expires_at,grants_mark,grants_revise)"
        )
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            })
        if resp.status_code != 200:
            return profile  # query failed — don't block the user

        redemptions = resp.json()
        if not redemptions:
            return profile  # no coupons — nothing to revoke

        from datetime import datetime, timezone
        now = datetime.now(timezone.utc)
        revoke_mark = False
        revoke_revise = False

        for r in redemptions:
            coupon = r.get("coupon_codes") or {}
            expires_at = coupon.get("expires_at")
            if not expires_at:
                continue  # no expiry — coupon is perpetual
            try:
                exp = datetime.fromisoformat(expires_at.replace("Z", "+00:00"))
            except (ValueError, TypeError):
                continue
            if now <= exp:
                continue  # coupon still valid

            # This coupon has expired — flag its grants for revocation
            if coupon.get("grants_mark"):
                revoke_mark = True
            if coupon.get("grants_revise"):
                revoke_revise = True

        if not revoke_mark and not revoke_revise:
            return profile

        # Build the revocation patch
        patch = {}
        if revoke_mark and profile.get("has_mark"):
            patch["has_mark"] = False
        if revoke_revise and profile.get("has_revise"):
            patch["has_revise"] = False

        if not patch:
            return profile

        # If revoking all products, also reset tier to free
        new_mark = patch.get("has_mark", profile.get("has_mark", False))
        new_revise = patch.get("has_revise", profile.get("has_revise", False))
        if not new_mark and not new_revise:
            patch["subscription_tier"] = "free"
            patch["subscription_status"] = "none"

        await _update_profile_fields(user_id, patch)

        # Return updated profile
        profile = {**profile, **patch}
        print(f"Revoked expired coupon access for user {user_id}: {patch}")
        return profile
    except Exception as e:
        print(f"Coupon expiry check failed for {user_id}: {repr(e)}")
        return profile  # fail open — don't block the user


async def _enforce_product_for_mode(user: dict, student_mode: bool) -> None:
    """Raise 403 if the user lacks the product required for the request context.

    • student_mode=True  (Revise / Write) → requires has_revise
    • student_mode=False (Mark / teacher)  → requires has_mark

    API-key clients and local-dev bypass this check (they are gated by
    allowed_endpoints in the api_keys table).
    """
    if user.get("_is_api_client"):
        return
    user_id = user.get("id")
    if not user_id or user_id == "local-dev":
        return
    profile = await get_user_profile(user_id)
    if not profile:
        return  # no profile → require_api_product already raised 403

    # Check for expired coupon access and revoke if needed
    profile = await _revoke_expired_coupon_access(user_id, profile)

    if student_mode:
        if not profile.get("has_revise", False):
            raise HTTPException(
                status_code=403,
                detail="This feature requires access to Revise.",
            )
    else:
        if not profile.get("has_mark", False):
            raise HTTPException(
                status_code=403,
                detail="This feature requires access to Mark.",
            )


# ── Profile helpers & endpoints ──────────────────────────────────────

async def get_user_profile(user_id: str) -> dict | None:
    """Fetch the user's profile row from Supabase. Returns dict or None."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return None
    url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}&select=*"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })
    if resp.status_code != 200:
        return None
    rows = resp.json()
    return rows[0] if rows else None


async def count_user_marks(user_id: str) -> int:
    """Count total mark_events for a user (for free tier usage tracking)."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return 0
    url = f"{SUPABASE_URL}/rest/v1/mark_events?user_id=eq.{user_id}&select=id"
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Prefer": "count=exact",
                "Range-Unit": "items",
                "Range": "0-0",
            })
        content_range = resp.headers.get("content-range", "")
        return int(content_range.split("/")[-1])
    except (ValueError, IndexError, Exception):
        return 0


_FREE_TIER_MARK_LIMIT = 3
_MOBILE_MARK_LIMIT = 5      # Total mobile marks before paywall
_MOBILE_DAILY_LIMIT = 2     # Max mobile marks per day
_MOBILE_PAGE_LIMIT = 15     # Max pages per mobile OCR upload


async def count_mobile_marks(user_id: str) -> int:
    """Count total mobile mark_events for a user."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return 0
    url = (
        f"{SUPABASE_URL}/rest/v1/mark_events"
        f"?user_id=eq.{user_id}&source=eq.mobile&select=id"
    )
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Prefer": "count=exact",
                "Range-Unit": "items",
                "Range": "0-0",
            })
        content_range = resp.headers.get("content-range", "")
        return int(content_range.split("/")[-1])
    except (ValueError, IndexError, Exception):
        return 0


async def count_mobile_marks_today(user_id: str) -> int:
    """Count mobile mark_events created today for daily rate limiting."""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return 0
    from datetime import datetime, timezone
    today = datetime.now(timezone.utc).strftime("%Y-%m-%dT00:00:00Z")
    url = (
        f"{SUPABASE_URL}/rest/v1/mark_events"
        f"?user_id=eq.{user_id}&source=eq.mobile"
        f"&created_at=gte.{today}&select=id"
    )
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Prefer": "count=exact",
                "Range-Unit": "items",
                "Range": "0-0",
            })
        content_range = resp.headers.get("content-range", "")
        return int(content_range.split("/")[-1])
    except (ValueError, IndexError, Exception):
        return 0


class ProfileUpdateRequest(BaseModel):
    has_mark: bool | None = None
    has_revise: bool | None = None
    has_write: bool | None = None
    display_name: str | None = None
    date_of_birth: str | None = None
    subscription_tier: str | None = None


@app.get("/api/profile")
@limiter.limit("30/minute")
async def get_profile(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Return the authenticated user's profile. Auto-creates if missing."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    profile = await get_user_profile(user_id)

    if not profile:
        email = user.get("email", "")
        meta = user.get("user_metadata") or {}
        display_name = meta.get("display_name") or meta.get("full_name", "")
        dob = meta.get("date_of_birth")
        payload = {
            "id": user_id,
            "email": email,
            "display_name": display_name,
            "has_mark": True,
            "has_revise": True,
            "has_write": False,
            "subscription_status": "none",
            "subscription_tier": "free",
        }
        if dob:
            payload["date_of_birth"] = dob
        url = f"{SUPABASE_URL}/rest/v1/profiles"
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.post(url, json=payload, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type": "application/json",
                "Prefer": "return=representation",
            })
        if 200 <= resp.status_code < 300:
            rows = resp.json()
            profile = rows[0] if rows else payload
        else:
            profile = payload

    # Enrich profile with usage data for frontend entitlement logic
    marks_used = await count_user_marks(user_id)
    profile["marks_used"] = marks_used
    mobile_marks_used = await count_mobile_marks(user_id)
    profile["mobile_marks_used"] = mobile_marks_used
    profile["mobile_marks_limit"] = _MOBILE_MARK_LIMIT
    if "subscription_tier" not in profile:
        profile["subscription_tier"] = "free"

    return profile


@app.patch("/api/profile")
@limiter.limit("20/minute")
async def update_profile(
    request: Request,
    body: ProfileUpdateRequest,
    user: dict = Depends(get_current_user),
):
    """Update product flags and display name on the user's profile."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    patch = {}
    if body.has_mark is not None:
        patch["has_mark"] = body.has_mark
    if body.has_revise is not None:
        patch["has_revise"] = body.has_revise
    if body.has_write is not None:
        patch["has_write"] = body.has_write
    if body.display_name is not None:
        patch["display_name"] = body.display_name[:200]
    if body.date_of_birth is not None:
        patch["date_of_birth"] = body.date_of_birth
    # NOTE: subscription_tier is NOT user-settable — only the Stripe webhook
    # may change it.  Accepting it here would let any user bypass the paywall.

    if not patch:
        raise HTTPException(status_code=400, detail="No fields to update")

    # Set onboarded_at on first product selection
    profile = await get_user_profile(user_id)
    if profile and not profile.get("onboarded_at"):
        if any(patch.get(k) for k in ("has_mark", "has_revise", "has_write")):
            patch["onboarded_at"] = datetime.datetime.utcnow().isoformat()

    # Set subscription_status to active on product selection, but never
    # downgrade an existing paid tier back to free.
    if any(patch.get(k) for k in ("has_mark", "has_revise", "has_write")):
        patch.setdefault("subscription_status", "active")
        existing_tier = (profile or {}).get("subscription_tier", "free")
        if existing_tier != "paid":
            patch.setdefault("subscription_tier", "free")

    url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.patch(url, json=patch, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
            "Prefer": "return=representation",
        })

    if resp.status_code not in (200, 204):
        raise HTTPException(status_code=502, detail="Failed to update profile")

    # 204 has no body; only parse JSON from 200
    if resp.status_code == 200:
        rows = resp.json()
        if rows:
            return rows[0]

    # PATCH matched 0 rows — profile doesn't exist yet. Create it.
    email = user.get("email", "")
    meta = user.get("user_metadata") or {}
    insert_payload = {
        "id": user_id,
        "email": email,
        "display_name": meta.get("display_name") or meta.get("full_name", ""),
        "has_mark": False,
        "has_revise": False,
        "has_write": False,
        "subscription_status": "none",
        "subscription_tier": "free",
    }
    insert_payload.update(patch)
    insert_url = f"{SUPABASE_URL}/rest/v1/profiles"
    async with httpx.AsyncClient(timeout=5) as client:
        ins_resp = await client.post(insert_url, json=insert_payload, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
            "Prefer": "return=representation",
        })
    if 200 <= ins_resp.status_code < 300:
        ins_rows = ins_resp.json()
        if ins_rows:
            return ins_rows[0]

    # Last fallback: return the patch data so the frontend can proceed
    return insert_payload


_AVATAR_MAX_BYTES = 2 * 1024 * 1024  # 2 MB
_AVATAR_ALLOWED_TYPES = {
    "image/jpeg": "jpg",
    "image/png": "png",
    "image/gif": "gif",
    "image/webp": "webp",
}


@app.post("/api/avatar")
@limiter.limit("10/minute")
async def upload_avatar(
    request: Request,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
):
    """Upload a custom profile avatar image."""
    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    content_type = (file.content_type or "").lower()
    ext = _AVATAR_ALLOWED_TYPES.get(content_type)
    if not ext:
        raise HTTPException(
            status_code=400,
            detail="Only JPEG, PNG, GIF, and WebP images are allowed.",
        )

    image_bytes = await file.read()
    if len(image_bytes) > _AVATAR_MAX_BYTES:
        raise HTTPException(status_code=400, detail="Image must be under 2 MB.")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Storage not configured")

    storage_path = f"{user_id}/avatar.{ext}"
    storage_url = f"{SUPABASE_URL}/storage/v1/object/avatars/{storage_path}"

    async with httpx.AsyncClient(timeout=15) as client:
        resp = await client.post(
            storage_url,
            content=image_bytes,
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type": content_type,
                "x-upsert": "true",
            },
        )
    if resp.status_code not in (200, 201):
        raise HTTPException(status_code=502, detail="Failed to upload avatar")

    public_url = (
        f"{SUPABASE_URL}/storage/v1/object/public/avatars/{storage_path}"
    )

    # Save URL to profiles table
    profile_url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}"
    async with httpx.AsyncClient(timeout=5) as client:
        await client.patch(
            profile_url,
            json={"avatar_url": public_url},
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type": "application/json",
            },
        )

    return {"avatar_url": public_url}


def require_product(*products: str):
    """
    FastAPI dependency factory that checks the user's profile
    for the required product(s). Raises 403 if not authorized.
    Usage: user=Depends(require_product("mark"))
    """
    async def _check(user: dict = Depends(get_current_user)):
        user_id = user.get("id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Could not determine user")
        # Localhost dev bypass — skip product check
        if user_id == "local-dev":
            return user
        profile = await get_user_profile(user_id)
        if not profile:
            raise HTTPException(
                status_code=403,
                detail="No profile found. Please complete onboarding.",
            )
        product_map = {
            "mark": profile.get("has_mark", False),
            "revise": profile.get("has_revise", False),
            "write": profile.get("has_write", False),
        }
        if not any(product_map.get(p, False) for p in products):
            raise HTTPException(
                status_code=403,
                detail=f"This feature requires one of: {', '.join(products)}",
            )
        return user
    return _check


# ===== Dev-only: reset user for testing (localhost only) =====

@app.post("/api/dev/reset-user")
async def dev_reset_user(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Reset the current user's profile to 'brand-new' state.
    Only works on localhost — blocked in production."""
    host = request.headers.get("host", "")
    if not host.startswith("localhost") and not host.startswith("127.0.0.1"):
        raise HTTPException(status_code=404, detail="Not found")

    user_id = user.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    # Reset profile to fresh signup state
    reset_fields = {
        "has_mark": False,
        "has_revise": False,
        "has_write": False,
        "subscription_status": "none",
        "subscription_tier": "free",
        "onboarded_at": None,
        "stripe_customer_id": None,
        "subscription_plan": None,
        "subscription_ends_at": None,
        "trial_ends_at": None,
    }
    url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.patch(url, json=reset_fields, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
        })

    # Delete all mark_events so usage counter resets
    events_url = f"{SUPABASE_URL}/rest/v1/mark_events?user_id=eq.{user_id}"
    async with httpx.AsyncClient(timeout=5) as client:
        await client.delete(events_url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })

    return {"ok": True, "message": "Profile reset to new-user state. Clear localStorage and refresh."}


# ===== Error reporting endpoints =====

@app.post("/api/report-error")
@limiter.limit("10/minute")
async def report_error(
    request: Request,
    body: ErrorReportRequest,
    user=Depends(get_current_user),
):
    """Accept a user-submitted error/issue report."""
    user_id = user.get("id") if isinstance(user, dict) else None
    email = user.get("email") if isinstance(user, dict) else None

    if not body.message or not body.message.strip():
        raise HTTPException(status_code=400, detail="Message is required")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        print(f"[ERROR_REPORT] user={user_id} msg={body.message[:200]}")
        return {"ok": True}

    payload = {
        "user_id": user_id,
        "email": email,
        "message": body.message[:2000],
        "debug_info": body.debug_info,
        "page_url": body.page_url,
    }

    try:
        db_url = f"{SUPABASE_URL}/rest/v1/error_reports"
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.post(
                db_url,
                json=payload,
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
            )
            if resp.status_code >= 300:
                print(f"[ERROR_REPORT] Supabase insert failed: {resp.status_code}")
    except Exception as e:
        print(f"[ERROR_REPORT] Insert error: {repr(e)}")

    return {"ok": True}


@app.post("/api/log-error")
@limiter.limit("20/minute")
async def log_error_endpoint(
    request: Request,
    body: AutoErrorLogRequest,
    user=Depends(get_current_user),
):
    """Auto-log critical frontend errors to Supabase."""
    user_id = user.get("id") if isinstance(user, dict) else None

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        print(f"[AUTO_ERROR] type={body.error_type} msg={body.message[:200]}")
        return {"ok": True}

    payload = {
        "user_id": user_id,
        "error_type": body.error_type,
        "message": body.message[:2000],
        "details": body.details,
        "page_url": body.page_url,
        "build_id": body.build_id,
    }

    try:
        db_url = f"{SUPABASE_URL}/rest/v1/error_logs"
        async with httpx.AsyncClient(timeout=5) as client:
            await client.post(
                db_url,
                json=payload,
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
            )
    except Exception as e:
        print(f"[AUTO_ERROR] Insert error: {repr(e)}")

    return {"ok": True}


# ===== Stripe endpoints =====

# Map price IDs to product flags
def _price_to_products(price_id: str) -> dict:
    """Return the product flags for a given Stripe price ID.
    Always returns BOTH flags so the patch resets access to exactly
    what was purchased (prevents leftover free-tier flags).
    """
    if price_id == STRIPE_PRICE_BOTH:
        return {"has_mark": True, "has_revise": True}
    elif price_id == STRIPE_PRICE_MARK:
        return {"has_mark": True, "has_revise": False}
    elif price_id == STRIPE_PRICE_REVISE:
        return {"has_mark": False, "has_revise": True}
    return {}


async def _update_profile_fields(user_id: str, fields: dict):
    """Patch profile fields via Supabase REST API (service key)."""
    url = f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}"
    async with httpx.AsyncClient(timeout=5) as client:
        await client.patch(url, json=fields, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
        })


async def _find_user_by_stripe_customer(customer_id: str) -> str | None:
    """Look up a Supabase user_id by their stripe_customer_id."""
    url = (
        f"{SUPABASE_URL}/rest/v1/profiles"
        f"?stripe_customer_id=eq.{customer_id}&select=id"
    )
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })
    if resp.status_code == 200:
        rows = resp.json()
        if rows:
            return rows[0]["id"]
    return None


class CheckoutRequest(BaseModel):
    price_id: str | None = None
    product: str | None = None  # "mark" | "revise" | "both"
    return_path: str | None = None


_PRODUCT_TO_PRICE = {
    "mark": STRIPE_PRICE_MARK,
    "revise": STRIPE_PRICE_REVISE,
    "both": STRIPE_PRICE_BOTH,
}


@app.post("/api/stripe/checkout")
@limiter.limit("10/minute")
async def create_checkout_session(
    request: Request,
    body: CheckoutRequest,
    user: dict = Depends(get_current_user),
):
    """Create a Stripe Checkout Session for a subscription."""
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=503, detail="Payments not configured")

    user_id = user.get("id")
    email = user.get("email", "")

    # Look up or create Stripe customer
    profile = await get_user_profile(user_id)
    customer_id = (profile or {}).get("stripe_customer_id")

    if not customer_id:
        try:
            customer = stripe.Customer.create(
                email=email,
                metadata={"supabase_user_id": user_id},
            )
        except stripe.StripeError as e:
            raise HTTPException(status_code=502, detail=f"Stripe customer creation failed: {e}")
        customer_id = customer.id
        await _update_profile_fields(user_id, {
            "stripe_customer_id": customer_id,
        })

    price = body.price_id or _PRODUCT_TO_PRICE.get(body.product or "") or STRIPE_PRICE_BOTH
    if not price:
        raise HTTPException(status_code=400, detail="No pricing configured")

    # Determine where Stripe redirects after checkout.
    # Only allow known local paths to prevent open-redirect attacks.
    allowed_return_pages = {
        "/teacher_react.html", "/student_react.html",
        "/profile_react.html", "/write_react.html",
    }
    return_page = body.return_path if body.return_path in allowed_return_pages else "/profile_react.html"
    base = str(request.base_url).rstrip("/")

    try:
        session = stripe.checkout.Session.create(
            customer=customer_id,
            mode="subscription",
            line_items=[{"price": price, "quantity": 1}],
            allow_promotion_codes=True,
            success_url=f"{base}{return_page}?checkout=success",
            cancel_url=f"{base}{return_page}?checkout=cancelled",
            metadata={"supabase_user_id": user_id},
        )
    except stripe.StripeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return {"checkout_url": session.url}


@app.post("/api/stripe/portal")
@limiter.limit("10/minute")
async def create_portal_session(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """Create a Stripe Customer Portal session for managing billing."""
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=503, detail="Payments not configured")

    profile = await get_user_profile(user.get("id"))
    customer_id = (profile or {}).get("stripe_customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="No billing account found")

    session = stripe.billing_portal.Session.create(
        customer=customer_id,
        return_url=str(request.base_url).rstrip("/") + "/profile_react.html",
    )
    return {"portal_url": session.url}


# ── Course-code enrollment ──────────────────────────────────────────

class EnrollRequest(BaseModel):
    code: str


@app.post("/api/enroll")
@limiter.limit("10/minute")
async def enroll_with_course_code(
    request: Request,
    body: EnrollRequest,
    user: dict = Depends(get_current_user),
):
    """Redeem a course code: validate, enroll, and grant paid access."""
    code = body.code.strip().upper()
    if not code:
        raise HTTPException(status_code=400, detail="Course code is required")

    user_id = user.get("id")

    # 1. Look up the course code
    url = (
        f"{SUPABASE_URL}/rest/v1/course_codes"
        f"?code=eq.{urllib.parse.quote(code, safe='')}&select=id,name,max_students,expires_at"
    )
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })
    courses = resp.json() if resp.status_code == 200 else []
    if not courses:
        raise HTTPException(status_code=404, detail="Invalid course code")

    course = courses[0]

    # 2. Check expiry
    from datetime import datetime, timezone
    if course.get("expires_at"):
        expires = datetime.fromisoformat(course["expires_at"].replace("Z", "+00:00"))
        if expires < datetime.now(timezone.utc):
            raise HTTPException(status_code=410, detail="This course code has expired")

    # 3. Check capacity
    count_url = (
        f"{SUPABASE_URL}/rest/v1/enrollments"
        f"?course_code_id=eq.{course['id']}&select=id"
    )
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(count_url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Prefer": "count=exact",
            "Range-Unit": "items",
            "Range": "0-0",
        })
    count = 0
    cr = resp.headers.get("content-range", "")
    if "/" in cr:
        try:
            count = int(cr.split("/")[1])
        except (ValueError, IndexError):
            pass
    if course.get("max_students") and count >= course["max_students"]:
        raise HTTPException(status_code=409, detail="This course is full")

    # 4. Create enrollment (unique constraint prevents duplicates)
    enroll_url = f"{SUPABASE_URL}/rest/v1/enrollments"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.post(enroll_url, json={
            "user_id": user_id,
            "course_code_id": course["id"],
        }, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
            "Prefer": "return=minimal",
        })
    if resp.status_code == 409 or (resp.status_code >= 400 and "23505" in resp.text):
        raise HTTPException(status_code=409, detail="You are already enrolled in this course")
    if resp.status_code >= 400:
        raise HTTPException(status_code=502, detail="Enrollment failed")

    # 5. Grant paid access
    await _update_profile_fields(user_id, {
        "subscription_tier": "paid",
        "subscription_status": "active",
        "has_mark": True,
        "has_revise": True,
    })

    return {"enrolled": True, "course_name": course.get("name", code)}


# ── Coupon code redemption ──────────────────────────────────────────

class RedeemCouponRequest(BaseModel):
    code: str


@app.post("/api/redeem-coupon")
@limiter.limit("10/minute")
async def redeem_coupon(
    request: Request,
    body: RedeemCouponRequest,
    user: dict = Depends(get_current_user),
):
    """Redeem a coupon code: validate, record, and grant access."""
    code = body.code.strip().upper()
    if not code:
        raise HTTPException(status_code=400, detail="Coupon code is required")

    user_id = user.get("id")

    # 1. Look up the coupon code
    url = (
        f"{SUPABASE_URL}/rest/v1/coupon_codes"
        f"?code=eq.{urllib.parse.quote(code, safe='')}"
        f"&is_active=eq.true"
        f"&select=id,description,max_redemptions,grants_tier,grants_mark,grants_revise,expires_at"
    )
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.get(url, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        })
    coupons = resp.json() if resp.status_code == 200 else []
    if not coupons:
        raise HTTPException(status_code=404, detail="Invalid coupon code")

    coupon = coupons[0]

    # 2. Check expiry
    from datetime import datetime, timezone
    if coupon.get("expires_at"):
        expires = datetime.fromisoformat(coupon["expires_at"].replace("Z", "+00:00"))
        if expires < datetime.now(timezone.utc):
            raise HTTPException(status_code=410, detail="This coupon code has expired")

    # 3. Check redemption capacity
    if coupon.get("max_redemptions"):
        count_url = (
            f"{SUPABASE_URL}/rest/v1/coupon_redemptions"
            f"?coupon_id=eq.{coupon['id']}&select=id"
        )
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.get(count_url, headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Prefer": "count=exact",
                "Range-Unit": "items",
                "Range": "0-0",
            })
        count = 0
        cr = resp.headers.get("content-range", "")
        if "/" in cr:
            try:
                count = int(cr.split("/")[1])
            except (ValueError, IndexError):
                pass
        if count >= coupon["max_redemptions"]:
            raise HTTPException(status_code=409, detail="This coupon has reached its redemption limit")

    # 4. Record redemption (unique constraint prevents duplicates)
    redeem_url = f"{SUPABASE_URL}/rest/v1/coupon_redemptions"
    async with httpx.AsyncClient(timeout=5) as client:
        resp = await client.post(redeem_url, json={
            "coupon_id": coupon["id"],
            "user_id": user_id,
        }, headers={
            "apikey": SUPABASE_SERVICE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            "Content-Type": "application/json",
            "Prefer": "return=minimal",
        })
    if resp.status_code == 409 or (resp.status_code >= 400 and "23505" in resp.text):
        raise HTTPException(status_code=409, detail="You have already redeemed this coupon")
    if resp.status_code >= 400:
        raise HTTPException(status_code=502, detail="Redemption failed")

    # 5. Grant access
    grants_tier = coupon.get("grants_tier", "paid")
    update_fields = {
        "subscription_tier": grants_tier,
        "subscription_status": "active",
    }
    # If coupon specifies product grants, set them on the profile
    if coupon.get("grants_mark") is not None:
        update_fields["has_mark"] = coupon["grants_mark"]
    if coupon.get("grants_revise") is not None:
        update_fields["has_revise"] = coupon["grants_revise"]
    await _update_profile_fields(user_id, update_fields)

    return {"redeemed": True, "description": coupon.get("description", "")}


@app.post("/api/delete-account")
@limiter.limit("5/minute")
async def delete_account(
    request: Request,
    user: dict = Depends(get_current_user),
):
    """
    Permanently delete the authenticated user's account:
    0. Store deletion feedback (reason / details)
    1. Cancel active Stripe subscriptions
    2. Delete all data from app tables
    3. Delete uploaded files from Supabase Storage
    4. Delete the auth user record
    """
    user_id = user.get("id") if isinstance(user, dict) else None
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Server configuration missing")

    # 0. Store deletion feedback before we delete anything
    try:
        body = await request.json()
    except Exception:
        body = {}
    feedback_reason = (body.get("reason") or "")[:500]
    feedback_details = (body.get("details") or "")[:2000]
    if feedback_reason:
        try:
            async with httpx.AsyncClient(timeout=10) as fb_client:
                await fb_client.post(
                    f"{SUPABASE_URL}/rest/v1/deletion_feedback",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=minimal",
                    },
                    json={
                        "user_id": user_id,
                        "reason": feedback_reason,
                        "details": feedback_details or None,
                    },
                )
        except Exception as exc:
            print(f"[delete-account] Feedback save error (non-fatal): {exc}")

    # 1. Cancel active Stripe subscriptions
    profile = await get_user_profile(user_id)
    stripe_customer_id = (profile or {}).get("stripe_customer_id")
    if stripe_customer_id and STRIPE_SECRET_KEY:
        try:
            subs = stripe.Subscription.list(customer=stripe_customer_id, status="active")
            for sub in subs.auto_paging_iter():
                stripe.Subscription.cancel(sub.id)
            # Also cancel past_due / trialing
            for status in ("past_due", "trialing"):
                subs = stripe.Subscription.list(customer=stripe_customer_id, status=status)
                for sub in subs.auto_paging_iter():
                    stripe.Subscription.cancel(sub.id)
        except Exception as exc:
            print(f"[delete-account] Stripe cancellation error (non-fatal): {exc}")

    headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
    }

    async with httpx.AsyncClient(timeout=30) as client:
        # 2. Delete data from app tables (child tables first)
        for table in [
            "issue_examples",
            "dismissed_issue_feedback",
            "revision_drafts",
            "mark_events",
        ]:
            try:
                await client.delete(
                    f"{SUPABASE_URL}/rest/v1/{table}",
                    headers=headers,
                    params={"user_id": f"eq.{user_id}"},
                )
            except Exception as exc:
                print(f"[delete-account] Error deleting {table}: {exc}")

        # Delete profile row
        try:
            await client.delete(
                f"{SUPABASE_URL}/rest/v1/profiles",
                headers=headers,
                params={"id": f"eq.{user_id}"},
            )
        except Exception as exc:
            print(f"[delete-account] Error deleting profile: {exc}")

        # 3. Delete uploaded files from Supabase Storage (best-effort)
        try:
            list_url = f"{SUPABASE_URL}/storage/v1/object/list/originals"
            resp = await client.post(
                list_url,
                headers={**headers, "Content-Type": "application/json"},
                json={"prefix": f"{user_id}/", "limit": 1000},
            )
            if resp.status_code == 200:
                files = resp.json()
                for f in files:
                    fname = f.get("name", "")
                    if fname:
                        del_url = f"{SUPABASE_URL}/storage/v1/object/originals/{user_id}/{fname}"
                        await client.delete(del_url, headers=headers)
        except Exception as exc:
            print(f"[delete-account] Storage cleanup error (non-fatal): {exc}")

        # 4. Delete the auth user via Supabase Admin API
        try:
            auth_url = f"{SUPABASE_URL}/auth/v1/admin/users/{user_id}"
            resp = await client.delete(auth_url, headers=headers)
            if resp.status_code not in (200, 204):
                print(f"[delete-account] Auth user deletion returned {resp.status_code}: {resp.text[:200]}")
        except Exception as exc:
            print(f"[delete-account] Auth user deletion error: {exc}")

    return {"deleted": True}


@app.post("/api/stripe/webhook")
async def stripe_webhook(request: Request):
    """Handle incoming Stripe webhook events."""
    if not STRIPE_WEBHOOK_SECRET:
        raise HTTPException(status_code=503, detail="Webhooks not configured")

    payload = await request.body()
    sig = request.headers.get("stripe-signature")

    try:
        event = stripe.Webhook.construct_event(
            payload, sig, STRIPE_WEBHOOK_SECRET,
        )
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid payload")
    except stripe.error.SignatureVerificationError:
        raise HTTPException(status_code=400, detail="Invalid signature")

    event_type = event["type"]
    data = event["data"]["object"]

    if event_type == "checkout.session.completed":
        customer_id = data.get("customer")
        user_id = data.get("metadata", {}).get("supabase_user_id")
        subscription_id = data.get("subscription")

        if not user_id and customer_id:
            user_id = await _find_user_by_stripe_customer(customer_id)

        if user_id and subscription_id:
            # Fetch subscription to determine which price was purchased
            sub = stripe.Subscription.retrieve(subscription_id)
            products = {}
            for item in sub["items"]["data"]:
                price_id = item["price"]["id"]
                products.update(_price_to_products(price_id))

            patch = {
                **products,
                "subscription_status": "active",
                "subscription_tier": "paid",
                "stripe_customer_id": customer_id,
                "onboarded_at": datetime.datetime.utcnow().isoformat(),
            }
            await _update_profile_fields(user_id, patch)

    elif event_type == "customer.subscription.updated":
        customer_id = data.get("customer")
        status_val = data.get("status")
        user_id = await _find_user_by_stripe_customer(customer_id)

        if user_id:
            status_map = {
                "active": "active",
                "past_due": "past_due",
                "canceled": "cancelled",
                "unpaid": "past_due",
                "incomplete": "none",
                "incomplete_expired": "cancelled",
                "trialing": "trial",
                "paused": "cancelled",
            }
            new_status = status_map.get(status_val, "none")
            patch = {"subscription_status": new_status}

            # If active (reactivation), restore paid tier
            if new_status == "active":
                patch["subscription_tier"] = "paid"

            # If cancelled, remove product access
            if new_status == "cancelled":
                patch["has_mark"] = False
                patch["has_revise"] = False
                patch["has_write"] = False
                patch["subscription_tier"] = "free"

            await _update_profile_fields(user_id, patch)

    elif event_type == "customer.subscription.deleted":
        customer_id = data.get("customer")
        user_id = await _find_user_by_stripe_customer(customer_id)
        if user_id:
            await _update_profile_fields(user_id, {
                "subscription_status": "cancelled",
                "subscription_tier": "free",
                "has_mark": False,
                "has_revise": False,
                "has_write": False,
            })

    elif event_type == "invoice.payment_failed":
        customer_id = data.get("customer")
        user_id = await _find_user_by_stripe_customer(customer_id)
        if user_id:
            await _update_profile_fields(user_id, {
                "subscription_status": "past_due",
                "subscription_tier": "free",
                "has_mark": False,
                "has_revise": False,
                "has_write": False,
            })

    return {"received": True}


@app.get("/")
def read_root():
    return RedirectResponse(url="/signin.html")


@app.get("/student_react.html")
def serve_student_react():
    """Serve the React student interface"""
    return FileResponse("student_react.html")


@app.get("/write_react.html")
def serve_write_react():
    """Serve the React write interface"""
    return FileResponse("write_react.html")


@app.get("/teacher_react.html")
def serve_teacher_react():
    """Serve the React teacher interface"""
    return FileResponse("teacher_react.html")


@app.get("/profile_react.html")
def serve_profile_react():
    """Serve the React profile page"""
    return FileResponse("profile_react.html")


@app.get("/practice.html")
def serve_practice():
    """Serve the Practice page (no auth required)"""
    return FileResponse("practice.html")


@app.get("/manifest.json")
def serve_manifest():
    """Serve PWA manifest"""
    return FileResponse("manifest.json", media_type="application/manifest+json")


@app.get("/sw.js")
def serve_service_worker():
    """Serve PWA service worker (must be at root scope)"""
    return FileResponse("sw.js", media_type="application/javascript",
                        headers={"Service-Worker-Allowed": "/"})


@app.get("/student-react-config.json")
def serve_student_react_config():
    """Serve the React student config"""
    return FileResponse("student-react-config.json")


@app.get("/power_verbs_2025.json")
def serve_power_verbs():
    """Serve power verbs lexicon (static asset, no auth needed)"""
    return FileResponse("power_verbs_2025.json")


@app.get("/thesis_devices.txt")
def serve_thesis_devices():
    """Serve thesis devices lexicon (static asset, no auth needed)"""
    return FileResponse("thesis_devices.txt")


@app.get("/role.html")
def serve_role():
    """Serve the product selection page"""
    return FileResponse("role.html")


@app.get("/signin.html")
def serve_signin():
    """Serve the signin page"""
    return FileResponse("signin.html")


@app.get("/student_progress.html")
def serve_student_progress():
    """Serve the student progress page"""
    return FileResponse("student_progress.html")


@app.get("/progress.html")
def serve_progress():
    """Serve the teacher progress page"""
    return FileResponse("progress.html")


@app.get("/terms.html")
def serve_terms():
    """Serve the Terms of Service page"""
    return FileResponse("terms.html")


@app.get("/privacy.html")
def serve_privacy():
    """Serve the Privacy Policy page"""
    return FileResponse("privacy.html")


@app.get("/api/lexis")
@limiter.limit("60/minute")
async def get_all_lexis(request: Request, user: dict = Depends(get_current_user)):
    """Return the full lexis database for A-Z dictionary browsing (auth required).

    Returns a compact list: [{term, term_norm, focus_type, definition,
    part_of_speech, tags}, ...] sorted alphabetically by term.
    Full detail for any term can be fetched via GET /api/lexis/{term_norm}.
    """
    from marker import load_lexis_database
    import math

    lexis_df = load_lexis_database()
    if lexis_df.empty:
        return JSONResponse({"error": "Lexis database not loaded"}, status_code=503)

    # Only active terms
    active_df = lexis_df[lexis_df["active"] == True]  # noqa: E712
    # Sort alphabetically (case-insensitive)
    active_df = active_df.copy()
    active_df["_sort"] = active_df["term"].str.lower()
    active_df = active_df.sort_values("_sort")

    # Return compact fields for the A-Z list (keeps payload small)
    compact_cols = [
        "term", "term_norm", "focus_type", "definition",
        "part_of_speech", "tags", "etymology", "application",
    ]
    result = []
    for _, row in active_df.iterrows():
        entry = {}
        for col in compact_cols:
            val = row.get(col)
            if val is None:
                continue
            if hasattr(val, "item"):
                val = val.item()
            if isinstance(val, float) and math.isnan(val):
                continue
            entry[col] = val
        if entry.get("term"):
            result.append(entry)

    return JSONResponse({"terms": result, "count": len(result)})


@app.get("/api/lexis/{term_norm}")
@limiter.limit("60/minute")
async def get_lexis_term(request: Request, term_norm: str, user: dict = Depends(get_current_user)):
    """Look up a single lexis term by its normalised name (auth required)."""
    from marker import load_lexis_database
    import math

    lexis_df = load_lexis_database()
    if lexis_df.empty:
        return JSONResponse({"error": "Lexis database not loaded"}, status_code=503)

    import re
    query = re.sub(r"['\"\[\]]+", "", term_norm).lower().strip()
    # Exact match first
    matches = lexis_df[lexis_df["term_norm"] == query]
    # Fallback: term_norm starts with query or query starts with term_norm
    # (handles plural/singular mismatches like leitmotif vs leitmotifs)
    if matches.empty:
        matches = lexis_df[
            lexis_df["term_norm"].apply(
                lambda t: t.startswith(query) or query.startswith(t)
            )
        ]
    if matches.empty:
        return JSONResponse({"error": "Term not found"}, status_code=404)

    row = matches.iloc[0]
    result = {}
    for col in matches.columns:
        val = row.get(col)
        if val is None:
            continue
        # Convert numpy/pandas types to native Python
        if hasattr(val, "item"):
            val = val.item()
        if isinstance(val, float) and math.isnan(val):
            continue
        result[col] = val
    return JSONResponse(result)


@app.post("/mark")
@limiter.limit("40/minute")
async def mark_essay(
    request: Request,
    file: UploadFile = File(...),
    mode: str = Form("textual_analysis"),
    user: dict = Depends(require_product("mark", "revise")),
    include_summary_table: bool = Form(True),
    student_mode: bool | None = Form(None),
    return_metadata: bool = Form(False),  # NEW: Return JSON with metadata

    # Primary work
    author: str | None = Form(None),
    title: str | None = Form(None),
    text_is_minor_work: bool | None = Form(None),

    # New metadata (optional)
    student_name: str | None = Form(None),
    assignment_name: str | None = Form(None),
    class_id: str | None = Form(None),

    # Second work (optional)
    author2: str | None = Form(None),
    title2: str | None = Form(None),
    text_is_minor_work_2: bool | None = Form(None),

    # Third work (optional)
    author3: str | None = Form(None),
    title3: str | None = Form(None),
    text_is_minor_work_3: bool | None = Form(None),

    # Rule toggles (optional)
    forbid_personal_pronouns: bool | None = Form(None),
    forbid_audience_reference: bool | None = Form(None),
    enforce_closed_thesis: bool | None = Form(None),
    require_body_evidence: bool | None = Form(None),
    allow_intro_summary_quotes: bool | None = Form(None),
    enforce_intro_quote_rule: bool | None = Form(None),
    enforce_long_quote_rule: bool | None = Form(None),
    enforce_contractions_rule: bool | None = Form(None),
    enforce_which_rule: bool | None = Form(None),
    enforce_weak_verbs_rule: bool | None = Form(None),
    enforce_fact_proof_rule: bool | None = Form(None),
    enforce_human_people_rule: bool | None = Form(None),
    enforce_vague_terms_rule: bool | None = Form(None),
    enforce_sva_rule: bool | None = Form(None),
    enforce_spelling_rule: bool | None = Form(None),
    enforce_confused_words_rule: bool | None = Form(None),
    enforce_intro_comma_rule: bool | None = Form(None),
    enforce_apostrophe_rule: bool | None = Form(None),
    enforce_present_tense_rule: bool | None = Form(None),
    highlight_thesis_devices: bool | None = Form(None),
):
    """
    Mark a .docx essay using the Vysti engine.

    Accepts up to three works (author/title) plus:
      - text_is_minor_work, text_is_minor_work_2, text_is_minor_work_3
      - forbid_personal_pronouns
      - enforce_closed_thesis

    These map directly onto MarkerConfig in marker.py.
    """
    # 0. Product-level access check based on calling context
    await _enforce_product_for_mode(user, bool(student_mode))

    # 1. Basic validation
    _fname_lower = file.filename.lower() if file.filename else ""
    _is_pdf = _fname_lower.endswith(".pdf")
    _is_docx = _fname_lower.endswith(".docx")
    if not _is_docx and not _is_pdf:
        return JSONResponse(
            status_code=400,
            content={"error": "Please upload a .docx or .pdf file."},
        )

    # 1b. File size limit (10 MB)
    _MAX_UPLOAD_BYTES = 10 * 1024 * 1024
    contents = await file.read()
    if len(contents) > _MAX_UPLOAD_BYTES:
        return JSONResponse(
            status_code=400,
            content={"error": "File exceeds the 10 MB size limit."},
        )
    await file.seek(0)  # reset for downstream readers

    # 1c. PDF → docx conversion (extract text, build synthetic docx)
    _pdf_converted = False
    if _is_pdf:
        try:
            _pdf_text = extract_text_from_pdf(contents)
        except PDFExtractionError as exc:
            # Scanned/image PDF → try OCR transcription
            if "scanned" in str(exc).lower() or "image-based" in str(exc).lower():
                try:
                    _pdf_text = await transcribe_scanned_pdf(contents)
                except Exception as ocr_exc:
                    return JSONResponse(
                        status_code=400,
                        content={"error": str(ocr_exc)},
                    )
            else:
                return JSONResponse(status_code=400, content={"error": str(exc)})
        contents = build_doc_from_text(_pdf_text)
        _pdf_converted = True

    # 1c. Free-tier usage check
    _mark_user_id = user.get("id") if isinstance(user, dict) else None
    if _mark_user_id:
        _profile = await get_user_profile(_mark_user_id)
        _tier = (_profile or {}).get("subscription_tier", "free")
        if _tier == "free":
            _marks_used = await count_user_marks(_mark_user_id)
            if _marks_used >= _FREE_TIER_MARK_LIMIT:
                raise HTTPException(
                    status_code=402,
                    detail={
                        "message": "Subscribe for unlimited essay uploads.",
                        "code": "USAGE_LIMIT",
                    },
                )

    # 2. Validate class_id if provided
    class_id_validated = None
    if class_id:
        user_id = user.get("id") if isinstance(user, dict) else None
        if not user_id:
            return JSONResponse(
                status_code=400,
                content={"error": "Invalid user"},
            )
        
        # Verify class exists and belongs to user
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            db_url = f"{SUPABASE_URL}/rest/v1/classes"
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.get(
                    f"{db_url}?id=eq.{class_id}&user_id=eq.{user_id}&archived=eq.false",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if not data or len(data) == 0:
                        return JSONResponse(
                            status_code=400,
                            content={"error": "Class not found or does not belong to user"},
                        )
                    class_id_validated = class_id
                else:
                    return JSONResponse(
                        status_code=400,
                        content={"error": "Failed to validate class"},
                    )

    # 2b. Word count check (hard cap rejects, soft cap warns in metadata)
    _wc = 0
    try:
        _tmp_doc = Document(BytesIO(contents))
        _essay_words = " ".join(p.text.strip() for p in _tmp_doc.paragraphs if p.text.strip()).split()
        _wc = len(_essay_words)
        if _wc > _HARD_WORD_LIMIT:
            return JSONResponse(
                status_code=400,
                content={
                    "error": f"Essay exceeds the {_HARD_WORD_LIMIT:,} word limit ({_wc:,} words). "
                             "Please check that you've uploaded the correct file."
                },
            )
        _word_count_warning = (
            f"This essay is {_wc:,} words — above the typical range. Marking may take longer."
            if _wc > _SOFT_WORD_LIMIT else None
        )
    except Exception:
        _word_count_warning = None  # if we can't parse it, let the marker handle the error

    # 2c. Concurrency guard — one active mark per user
    _mark_user_id = user.get("id") if isinstance(user, dict) else None
    if _mark_user_id and _mark_user_id in _active_marks:
        return JSONResponse(
            status_code=429,
            content={"error": "A marking request is already in progress. Please wait for it to finish."},
        )
    if _mark_user_id:
        _active_marks.add(_mark_user_id)

    # 3. Read file bytes
    if _pdf_converted:
        docx_bytes = contents  # already converted to docx bytes above
    else:
        docx_bytes = await file.read()

    # 3b. Upload original to Supabase Storage (best-effort)
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            upload_user_id = user.get("id") if isinstance(user, dict) else None
            if upload_user_id and file.filename:
                safe_upload_name = _sanitize_filename(file.filename)
                storage_path = f"{upload_user_id}/{safe_upload_name}"
                storage_url = f"{SUPABASE_URL}/storage/v1/object/originals/{storage_path}"
                _upload_ct = (
                    "application/pdf" if _is_pdf
                    else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                async with httpx.AsyncClient(timeout=15) as client:
                    resp = await client.post(
                        storage_url,
                        content=docx_bytes,
                        headers={
                            "apikey": SUPABASE_SERVICE_KEY,
                            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                            "Content-Type": _upload_ct,
                            "x-upsert": "true",
                        },
                    )
                    if _DEBUG: print(f"[DEBUG] Original upload: status={resp.status_code}, path={storage_path}")
    except Exception as e:
        print(f"Failed to upload original: {repr(e)}")

    # 4. Build teacher_config from form fields (matches MarkerConfig)
    teacher_config: dict = {}

    # --- Works / titles ---
    if author:
        teacher_config["author_name"] = author
    if title:
        teacher_config["text_title"] = title
    if text_is_minor_work is not None:
        teacher_config["text_is_minor_work"] = text_is_minor_work

    if author2:
        teacher_config["author_name_2"] = author2
    if title2:
        teacher_config["text_title_2"] = title2
    if text_is_minor_work_2 is not None:
        teacher_config["text_is_minor_work_2"] = text_is_minor_work_2

    if author3:
        teacher_config["author_name_3"] = author3
    if title3:
        teacher_config["text_title_3"] = title3
    if text_is_minor_work_3 is not None:
        teacher_config["text_is_minor_work_3"] = text_is_minor_work_3

    # --- Rule overrides ---
    # These override the defaults chosen by get_preset_config(mode)
    if forbid_personal_pronouns is not None:
        teacher_config["forbid_personal_pronouns"] = forbid_personal_pronouns
    if forbid_audience_reference is not None:
        teacher_config["forbid_audience_reference"] = forbid_audience_reference
    if enforce_closed_thesis is not None:
        teacher_config["enforce_closed_thesis"] = enforce_closed_thesis
    if require_body_evidence is not None:
        teacher_config["require_body_evidence"] = require_body_evidence
    if allow_intro_summary_quotes is not None:
        teacher_config["allow_intro_summary_quotes"] = allow_intro_summary_quotes
    if enforce_intro_quote_rule is not None:
        teacher_config["enforce_intro_quote_rule"] = enforce_intro_quote_rule
    if enforce_long_quote_rule is not None:
        teacher_config["enforce_long_quote_rule"] = enforce_long_quote_rule
    if enforce_contractions_rule is not None:
        teacher_config["enforce_contractions_rule"] = enforce_contractions_rule
    if enforce_which_rule is not None:
        teacher_config["enforce_which_rule"] = enforce_which_rule
    if enforce_weak_verbs_rule is not None:
        teacher_config["enforce_weak_verbs_rule"] = enforce_weak_verbs_rule
    if enforce_fact_proof_rule is not None:
        teacher_config["enforce_fact_proof_rule"] = enforce_fact_proof_rule
    if enforce_human_people_rule is not None:
        teacher_config["enforce_human_people_rule"] = enforce_human_people_rule
    if enforce_vague_terms_rule is not None:
        teacher_config["enforce_vague_terms_rule"] = enforce_vague_terms_rule
    if enforce_sva_rule is not None:
        teacher_config["enforce_sva_rule"] = enforce_sva_rule
    if enforce_spelling_rule is not None:
        teacher_config["enforce_spelling_rule"] = enforce_spelling_rule
    if enforce_confused_words_rule is not None:
        teacher_config["enforce_confused_words_rule"] = enforce_confused_words_rule
    if enforce_intro_comma_rule is not None:
        teacher_config["enforce_intro_comma_rule"] = enforce_intro_comma_rule
    if enforce_apostrophe_rule is not None:
        teacher_config["enforce_apostrophe_rule"] = enforce_apostrophe_rule
    if enforce_present_tense_rule is not None:
        teacher_config["enforce_present_tense_rule"] = enforce_present_tense_rule
    if highlight_thesis_devices is not None:
        teacher_config["highlight_thesis_devices"] = highlight_thesis_devices
    if student_mode is True:
        teacher_config["student_mode"] = True

    # 5. Call your engine (release concurrency guard when done)
    try:
        mark_docx_bytes, _ = get_engine()
        marked_bytes, metadata = mark_docx_bytes(
            docx_bytes,
            mode=mode,
            teacher_config=teacher_config if teacher_config else None,
            include_summary_table=include_summary_table,
        )
    finally:
        _active_marks.discard(_mark_user_id or "")

    if _DEBUG:
        print("Vysti metadata:", metadata)
        print("Teacher config used:", teacher_config)

    # ----- Extract examples from metadata -----
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []

    # ----- Count yellow labels from metadata -----
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []

    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl or lbl.startswith("__"):
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        label_counter[lbl] += (cnt_i if cnt_i > 0 else 1)

    total_labels = sum(label_counter.values())

    # Compute word_count + scores before mark_events insert so they can be persisted
    _meta_word_count = metadata.get("word_count") if isinstance(metadata, dict) else None
    _meta_scores = None
    try:
        orig_doc = Document(BytesIO(docx_bytes))
        essay_text = "\n\n".join(p.text.strip() for p in orig_doc.paragraphs if p.text.strip())
        _meta_scores = _compute_scores(
            essay_text,
            mode=mode,
            label_counts=dict(label_counter),
            mark_event_id=None,
            sentence_types=(metadata.get("sentence_types", {}) if isinstance(metadata, dict) else {}),
            repeated_nouns=(metadata.get("repeated_nouns", []) if isinstance(metadata, dict) else []),
        )
    except Exception as e:
        if _DEBUG: print(f"[SCORE] Pre-insert _compute_scores failed: {repr(e)}")

    # Log usage in Supabase mark_events (best-effort; do not break marking if this fails)
    mark_event_id = None
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            user_id = user.get("id") if isinstance(user, dict) else None

            # Clear old mark_events for this user/file to ensure fresh start
            encoded_filename = urllib.parse.quote(file.filename)
            delete_url = f"{SUPABASE_URL}/rest/v1/mark_events"
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.delete(
                    f"{delete_url}?user_id=eq.{user_id}&file_name=eq.{encoded_filename}",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )
                if _DEBUG: print(f"[DEBUG] Deleted old mark_events: status={resp.status_code}, filename={file.filename}")

            db_url = f"{SUPABASE_URL}/rest/v1/mark_events?select=id"
            payload = {
                "user_id": user_id,
                "file_name": file.filename,
                "mode": mode,
                "bytes": len(docx_bytes),
                "student_name": student_name,
                "assignment_name": assignment_name,
                "class_id": class_id_validated,
                "total_labels": total_labels,
                "label_counts": dict(label_counter),
                "issues": issues,
                "review_status": "pending",
                "word_count": _meta_word_count,
                "scores": _sanitize_for_json(_meta_scores) if _meta_scores else None,
            }

            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.post(
                    db_url,
                    json=payload,
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=representation",
                    },
                )
                # Capture mark_event_id from response
                if resp.status_code >= 200 and resp.status_code < 300:
                    resp_data = resp.json()
                    if resp_data and isinstance(resp_data, list) and len(resp_data) > 0:
                        mark_event_id = resp_data[0].get("id")
    except Exception as e:
        print("Failed to log mark_event:", repr(e))

    # Prune old mark_events beyond retention cap (best-effort)
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and mark_event_id:
            user_id = user.get("id") if isinstance(user, dict) else None
            if user_id:
                # Fetch the Nth oldest event ID to use as a cutoff
                prune_url = (
                    f"{SUPABASE_URL}/rest/v1/mark_events"
                    f"?user_id=eq.{user_id}&select=id&order=created_at.desc"
                    f"&offset={_MAX_MARK_EVENTS_PER_USER}&limit=1"
                )
                async with httpx.AsyncClient(timeout=5) as client:
                    prune_resp = await client.get(
                        prune_url,
                        headers={
                            "apikey": SUPABASE_SERVICE_KEY,
                            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        },
                    )
                    if prune_resp.status_code == 200:
                        prune_data = prune_resp.json()
                        if prune_data and len(prune_data) > 0:
                            cutoff_id = prune_data[0]["id"]
                            # Get IDs of all events older than the cutoff
                            old_url = (
                                f"{SUPABASE_URL}/rest/v1/mark_events"
                                f"?user_id=eq.{user_id}&select=id&order=created_at.desc"
                                f"&offset={_MAX_MARK_EVENTS_PER_USER}"
                            )
                            old_resp = await client.get(
                                old_url,
                                headers={
                                    "apikey": SUPABASE_SERVICE_KEY,
                                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                },
                            )
                            if old_resp.status_code == 200:
                                old_ids = [r["id"] for r in old_resp.json() if "id" in r]
                                if old_ids:
                                    # Delete old events and their examples
                                    for old_id in old_ids:
                                        await client.delete(
                                            f"{SUPABASE_URL}/rest/v1/issue_examples?mark_event_id=eq.{old_id}",
                                            headers={
                                                "apikey": SUPABASE_SERVICE_KEY,
                                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                            },
                                        )
                                        await client.delete(
                                            f"{SUPABASE_URL}/rest/v1/mark_events?id=eq.{old_id}",
                                            headers={
                                                "apikey": SUPABASE_SERVICE_KEY,
                                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                            },
                                        )
                                    if _DEBUG:
                                        print(f"[DEBUG] Pruned {len(old_ids)} old mark_events for user {user_id}")
    except Exception as e:
        if _DEBUG:
            print(f"[DEBUG] Mark event pruning failed: {repr(e)}")

    # Log examples to Supabase issue_examples (best-effort; do not break marking if this fails)
    try:
        if _DEBUG: print(f"[DEBUG] Examples from marker: count={len(examples) if examples else 0}")
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and examples:
            user_id = user.get("id") if isinstance(user, dict) else None
            if _DEBUG: print(f"[DEBUG] Saving examples: user_id={user_id}, file_name={file.filename}")
            if user_id:
                # Clear old cached examples for this user/file to ensure fresh start
                encoded_filename = urllib.parse.quote(file.filename)
                delete_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                async with httpx.AsyncClient(timeout=5) as client:
                    resp = await client.delete(
                        f"{delete_url}?user_id=eq.{user_id}&file_name=eq.{encoded_filename}",
                        headers={
                            "apikey": SUPABASE_SERVICE_KEY,
                            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        },
                    )
                    if _DEBUG: print(f"[DEBUG] Deleted old examples: status={resp.status_code}")

                example_rows = []
                for ex in examples:
                    if not isinstance(ex, dict):
                        continue
                    label = ex.get("label")
                    sentence = ex.get("sentence")
                    paragraph_index = ex.get("paragraph_index")
                    if not label or not sentence:
                        continue
                    example_row = {
                        "user_id": user_id,
                        "class_id": class_id_validated,
                        "assignment_name": assignment_name,
                        "student_name": student_name,
                        "mode": mode,
                        "file_name": file.filename,
                        "label": label,
                        "sentence": sentence,
                        "paragraph_index": paragraph_index,
                        "mark_event_id": mark_event_id,
                        # Include context fields for dynamic guidance (use None if missing to ensure all rows have same keys)
                        "found_value": ex.get("found_value"),
                        "topics": ex.get("topics"),
                        "thesis": ex.get("thesis"),
                        "confidence": ex.get("confidence"),
                        "original_phrase": ex.get("original_phrase"),
                    }

                    example_rows.append(example_row)

                if _DEBUG: print(f"[DEBUG] Created {len(example_rows)} example rows to insert")
                if example_rows:
                    if _DEBUG: print(f"[DEBUG] First example row: {example_rows[0]}")
                    db_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                    async with httpx.AsyncClient(timeout=5) as client:
                        post_resp = await client.post(
                            db_url,
                            json=example_rows,
                            headers={
                                "apikey": SUPABASE_SERVICE_KEY,
                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                "Content-Type": "application/json",
                                "Prefer": "return=minimal",
                            },
                        )
                        if _DEBUG: print(f"[DEBUG] Inserted examples: status={post_resp.status_code}, response={post_resp.text[:200]}")
                else:
                    if _DEBUG: print("[DEBUG] No valid example rows to insert")
        else:
            if _DEBUG: print(f"[DEBUG] Skipping examples insert: SUPABASE_URL={bool(SUPABASE_URL)}, SUPABASE_SERVICE_KEY={bool(SUPABASE_SERVICE_KEY)}, examples={bool(examples)}")
    except Exception as e:
        print("Failed to log issue_examples:", repr(e))

    # 5. Return response - JSON with metadata or stream the marked .docx
    clean_name = _sanitize_filename(file.filename or "essay.docx")
    base_name = clean_name.rsplit(".", 1)[0] if clean_name else "essay"
    output_filename = f"{base_name}_marked.docx"

    # NEW: Return JSON if requested
    if return_metadata:
        # Enrich metadata with computed values the frontend needs
        enriched = dict(metadata) if isinstance(metadata, dict) else {}
        enriched["total_labels"] = total_labels
        enriched["label_counts"] = dict(label_counter)
        enriched["mark_event_id"] = mark_event_id
        # Reuse pre-computed scores (computed before mark_events insert)
        if _meta_scores:
            enriched["scores"] = _meta_scores
        # Strip proprietary fields before sending to client
        if "issues" in enriched:
            enriched["issues"] = _strip_ip_from_issues(enriched["issues"])
        if "examples" in enriched:
            enriched["examples"] = _strip_ip_from_examples(enriched["examples"])
        if _word_count_warning:
            enriched["word_count_warning"] = _word_count_warning
        return JSONResponse({
            "document": base64.b64encode(marked_bytes).decode('utf-8'),
            "filename": output_filename,
            "metadata": _sanitize_for_json(enriched),
        })

    # Existing: Stream the document
    return StreamingResponse(
        io.BytesIO(marked_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
    )


@app.post("/export_docx")
@limiter.limit("30/minute")
async def export_docx(
    request: Request,
    body: ExportDocxRequest,
    user: dict = Depends(get_current_user),
):
    """
    Export a clean .docx from plain text.
    This is for Student mode "Download revised essay" — no Vysti marks, no summary table.
    """
    # Free-tier students cannot download
    _exp_user_id = user.get("id") if isinstance(user, dict) else None
    if _exp_user_id:
        _exp_profile = await get_user_profile(_exp_user_id)
        _exp_tier = (_exp_profile or {}).get("subscription_tier", "free")
        if _exp_tier == "free":
            raise HTTPException(
                status_code=402,
                detail={
                    "message": "Subscribe to download your essay.",
                    "code": "DOWNLOAD_BLOCKED",
                },
            )

    if not body.text or not body.text.strip():
        raise HTTPException(status_code=400, detail="Missing text")

    # Build a clean document (NO marking)
    docx_bytes = build_doc_from_text(body.text)

    safe_name = _sanitize_filename(body.file_name.strip() if body.file_name else "essay_revised.docx")
    if not safe_name.lower().endswith(".docx"):
        safe_name += ".docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@app.post("/export_teacher_docx")
@limiter.limit("30/minute")
async def export_teacher_docx(
    request: Request,
    body: ExportTeacherDocxRequest,
    user: dict = Depends(get_current_user),
):
    """
    Export a teacher-marked .docx preserving Vysti labels (yellow highlight)
    and teacher annotations (green highlight).  Used by Teacher mode
    'Download Marked Essay'.
    """
    if not body.text or not body.text.strip():
        raise HTTPException(status_code=400, detail="Missing text")

    # No free-tier guard here: marking itself is already gated by the mark
    # limit, so once a teacher has marked an essay they should always be
    # able to download the result — regardless of tier.

    docx_bytes = build_teacher_doc_from_text(body.text, body.comment or "")

    safe_name = _sanitize_filename(body.file_name.strip() if body.file_name else "essay_marked.docx")
    if not safe_name.lower().endswith(".docx"):
        safe_name += ".docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@app.post("/delete_mark_events")
@limiter.limit("20/minute")
async def delete_mark_events(
    request: Request,
    body: DeleteMarkEventsRequest,
    user: dict = Depends(get_current_user),
):
    """
    Delete mark_events (and related issue_examples / dismissed_issue_feedback)
    for the authenticated user by file_name list.  Uses the service key so
    RLS policies don't block the operation.
    """
    if not body.file_names:
        raise HTTPException(status_code=400, detail="No file names provided")

    user_id = user.get("id") if isinstance(user, dict) else None
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase config missing on server")

    headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
    }

    deleted = 0
    tables = ["issue_examples", "dismissed_issue_feedback", "mark_events"]

    async with httpx.AsyncClient(timeout=30) as client:
        for chunk_start in range(0, len(body.file_names), 20):
            chunk = body.file_names[chunk_start : chunk_start + 20]
            # Build PostgREST in-filter value with quoted strings
            in_value = "(" + ",".join(f'"{fn}"' for fn in chunk) + ")"
            for table in tables:
                try:
                    resp = await client.delete(
                        f"{SUPABASE_URL}/rest/v1/{table}",
                        headers={**headers, "Prefer": "return=representation"},
                        params={
                            "user_id": f"eq.{user_id}",
                            "file_name": f"in.{in_value}",
                        },
                    )
                    if resp.status_code in (200, 204):
                        try:
                            resp_body = resp.json()
                            if isinstance(resp_body, list):
                                deleted += len(resp_body)
                        except Exception:
                            pass
                    else:
                        if _DEBUG: print(f"[delete_mark_events] {table}: status={resp.status_code} body={resp.text[:300]}")
                except Exception as exc:
                    if _DEBUG: print(f"[delete_mark_events] Error deleting from {table}: {exc}")

            # Also delete originals from Supabase Storage (best-effort)
            for fn in chunk:
                try:
                    storage_path = f"{user_id}/{fn}"
                    storage_url = f"{SUPABASE_URL}/storage/v1/object/originals/{storage_path}"
                    await client.delete(storage_url, headers=headers)
                except Exception:
                    pass

    return {"deleted": deleted}


@app.patch("/update_mark_event")
@limiter.limit("60/minute")
async def update_mark_event(
    request: Request,
    body: UpdateMarkEventRequest,
    user: dict = Depends(get_current_user),
):
    """
    Update editable fields on a mark_event for the authenticated user.
    Uses the service key so RLS policies don't block the operation.
    """
    if not body.file_name:
        raise HTTPException(status_code=400, detail="file_name is required")

    user_id = user.get("id") if isinstance(user, dict) else None
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase config missing on server")

    # Build PATCH payload from non-None fields
    patch_body = {}
    if body.assignment_name is not None:
        patch_body["assignment_name"] = body.assignment_name
    if body.essay_title is not None:
        patch_body["essay_title"] = body.essay_title
    if body.source_works is not None:
        patch_body["source_works"] = [w.model_dump() for w in body.source_works]
    if body.notes is not None:
        patch_body["notes"] = body.notes
    if body.student_name is not None:
        patch_body["student_name"] = body.student_name
    if body.class_id is not None:
        patch_body["class_id"] = body.class_id if body.class_id else None
    if body.teacher_comment is not None:
        patch_body["teacher_comment"] = body.teacher_comment
    if body.review_status is not None:
        if body.review_status in ("pending", "in_progress", "completed", "archived"):
            patch_body["review_status"] = body.review_status
    if body.score is not None:
        patch_body["score"] = max(0, min(100, body.score))
    if body.created_at is not None:
        patch_body["created_at"] = body.created_at

    if not patch_body:
        return {"updated": 0}

    headers = {
        "apikey": SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }

    # Target by mark_event_id (precise) or fall back to file_name (legacy)
    if body.mark_event_id:
        target_params = {"user_id": f"eq.{user_id}", "id": f"eq.{body.mark_event_id}"}
    else:
        target_params = {"user_id": f"eq.{user_id}", "file_name": f"eq.{body.file_name}"}

    updated = 0
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.patch(
            f"{SUPABASE_URL}/rest/v1/mark_events",
            headers=headers,
            params=target_params,
            json=patch_body,
        )
        if resp.status_code in (200, 204):
            try:
                result = resp.json()
                if isinstance(result, list):
                    updated = len(result)
            except Exception:
                updated = 1
        else:
            if _DEBUG: print(f"[update_mark_event] mark_events: status={resp.status_code} body={resp.text[:300]}")
            raise HTTPException(status_code=502, detail="Failed to update mark event")

        # If assignment_name changed, also update issue_examples for consistency
        if body.assignment_name is not None:
            try:
                ie_params = {"user_id": f"eq.{user_id}"}
                if body.mark_event_id:
                    ie_params["mark_event_id"] = f"eq.{body.mark_event_id}"
                else:
                    ie_params["file_name"] = f"eq.{body.file_name}"
                await client.patch(
                    f"{SUPABASE_URL}/rest/v1/issue_examples",
                    headers=headers,
                    params=ie_params,
                    json={"assignment_name": body.assignment_name},
                )
            except Exception as exc:
                if _DEBUG: print(f"[update_mark_event] issue_examples sync: {exc}")

    return {"updated": updated}


@app.get("/mark_event")
@limiter.limit("60/minute")
async def get_mark_event(
    request: Request,
    event_id: str,
    user: dict = Depends(get_current_user),
):
    """Fetch a single mark_event by ID for the authenticated user."""
    user_id = user.get("id") if isinstance(user, dict) else None
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")
    if not event_id:
        raise HTTPException(status_code=400, detail="event_id is required")
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase config missing")

    url = (
        f"{SUPABASE_URL}/rest/v1/mark_events"
        f"?id=eq.{event_id}&user_id=eq.{user_id}&select=*"
    )
    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.get(
            url,
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            },
        )
        if resp.status_code != 200:
            raise HTTPException(status_code=502, detail="Failed to fetch mark event")
        rows = resp.json()
        if not rows:
            raise HTTPException(status_code=404, detail="Mark event not found")
        return rows[0]


@app.get("/download_original")
@limiter.limit("30/minute")
async def download_original(
    request: Request,
    file_name: str,
    user: dict = Depends(get_current_user),
):
    """Download the original uploaded .docx from Supabase Storage."""
    user_id = user.get("id") if isinstance(user, dict) else None
    if not user_id:
        raise HTTPException(status_code=401, detail="Could not determine user")
    if not file_name:
        raise HTTPException(status_code=400, detail="file_name is required")
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=500, detail="Supabase config missing")

    safe_name = _sanitize_filename(file_name)
    storage_path = f"{user_id}/{safe_name}"
    storage_url = f"{SUPABASE_URL}/storage/v1/object/originals/{storage_path}"

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(
            storage_url,
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
            },
        )
        if resp.status_code != 200:
            raise HTTPException(status_code=404, detail="Original not found")

        return StreamingResponse(
            io.BytesIO(resp.content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )


def _make_teacher_label_run(paragraph, text):
    """Create a yellow-highlighted bold run for a teacher annotation label,
    matching the style marker.py uses for Vysti labels (→ Label)."""
    lbl = paragraph.add_run(text)
    lbl.font.name = "Times New Roman"
    lbl._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    lbl.font.size = Pt(12)
    lbl.font.bold = True
    lbl.font.highlight_color = WD_COLOR_INDEX.YELLOW
    lbl.font.color.rgb = RGBColor(0, 0, 0)
    lbl.font.underline = False
    return lbl


def _make_annotation_run_xml(label_text):
    """Build a <w:r> element for a yellow-highlighted teacher annotation."""
    ann_run = OxmlElement("w:r")
    ann_rPr = OxmlElement("w:rPr")
    for tag, val in [
        ("w:b", None),
        ("w:highlight", "yellow"),
        ("w:color", "000000"),
    ]:
        el = OxmlElement(tag)
        if val is not None:
            el.set(qn("w:val"), val)
        ann_rPr.append(el)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    ann_rPr.append(rFonts)
    for sz_tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(sz_tag)
        sz.set(qn("w:val"), "24")  # 12pt = 24 half-points
        ann_rPr.append(sz)
    ann_run.append(ann_rPr)
    ann_t = OxmlElement("w:t")
    ann_t.text = f" \u2192 {label_text}"
    ann_t.set(qn("xml:space"), "preserve")
    ann_run.append(ann_t)
    return ann_run


def _insert_annotation_inline(doc, wrapped_text, label):
    """Find `wrapped_text` in a .docx paragraph and insert ' → label' right after it.

    Walks all paragraphs, finds the first whose full text contains `wrapped_text`,
    then splits the run at the match boundary and inserts a yellow label run.
    Returns True if inserted, False if the text was not found.
    """
    from copy import deepcopy

    if not wrapped_text:
        return False

    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        pos = full_text.find(wrapped_text)
        if pos < 0:
            continue

        target_end = pos + len(wrapped_text)

        # Walk runs to find where target_end falls
        char_offset = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_end = char_offset + run_len

            if target_end <= run_end:
                split_at = target_end - char_offset
                ann_run = _make_annotation_run_xml(label)

                if split_at < run_len:
                    # Split this run: keep text[:split_at], remainder into a new run
                    remainder_text = run.text[split_at:]
                    run.text = run.text[:split_at]

                    # Build a CLEAN remainder run — only copy <w:rPr> (font/style),
                    # not the full XML (which may carry <w:tab/>, highlights, etc.)
                    remainder_elem = OxmlElement("w:r")
                    orig_rPr = run._element.find(qn("w:rPr"))
                    if orig_rPr is not None:
                        remainder_elem.append(deepcopy(orig_rPr))
                    rem_t = OxmlElement("w:t")
                    rem_t.text = remainder_text
                    if remainder_text and remainder_text[0] == " ":
                        rem_t.set(qn("xml:space"), "preserve")
                    remainder_elem.append(rem_t)

                    # Insert: current_run → annotation → remainder
                    run._element.addnext(ann_run)
                    ann_run.addnext(remainder_elem)
                else:
                    # Perfect alignment — just insert after this run
                    run._element.addnext(ann_run)

                return True

            char_offset = run_end

    return False


@app.post("/annotate_docx")
@limiter.limit("20/minute")
async def annotate_docx(
    request: Request,
    file: UploadFile = File(...),
    annotations: str = Form("[]"),
    revised_text: str = Form(""),
    teacher_comment: str = Form(""),
    user: dict = Depends(get_current_user),
):
    """
    Take an already-marked .docx and insert teacher annotations INLINE,
    right after the highlighted phrase — matching how marker.py places
    its own yellow labels.  Falls back to appending at the end if the
    wrapped text can't be found.

    If ``revised_text`` is provided (teacher edited the preview), the
    document body is rebuilt from that text so teacher edits (typed
    corrections, sp marks, etc.) appear in the download.
    """
    import json as _json

    try:
        ann_list = _json.loads(annotations)
    except Exception:
        ann_list = []

    revised = (revised_text or "").strip()
    comment = (teacher_comment or "").strip()

    if not ann_list and not revised and not comment:
        # No annotations and no edits — just return the original file
        content = await file.read()
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{_sanitize_filename(file.filename or "essay.docx")}"'},
        )

    # Read original .docx
    docx_bytes = await file.read()

    # If teacher provided revised text, rebuild the document from it
    # so typed corrections, sp marks, etc. appear in the download.
    if revised:
        doc_bytes = build_doc_from_text(revised)
        doc = Document(io.BytesIO(doc_bytes))
    else:
        doc = Document(io.BytesIO(docx_bytes))

    # Insert each annotation inline after its wrapped text
    fallback_labels = []
    for ann in ann_list:
        # ann is either a string (legacy) or {label, wrappedText} object
        if isinstance(ann, dict):
            label = (ann.get("label") or "").strip()
            wrapped = (ann.get("wrappedText") or "").strip()
        elif isinstance(ann, str):
            label = ann.strip()
            wrapped = ""
        else:
            continue
        if not label:
            continue

        if wrapped and _insert_annotation_inline(doc, wrapped, label):
            continue  # successfully inserted inline
        # Fallback: couldn't find the wrapped text — collect for end-of-doc
        fallback_labels.append(label)

    # Any annotations that couldn't be placed inline go at the end
    if fallback_labels:
        doc.add_paragraph("")
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("Teacher Comments")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = "Times New Roman"
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        for fb_label in fallback_labels:
            _make_teacher_label_run(
                doc.add_paragraph(),
                f"\u2192 {fb_label}",
            )

    # Append teacher's comment notebook (structured comment block)
    if comment:
        doc.add_paragraph("")
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("Teacher\u2019s Comment")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = "Times New Roman"
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        for line in comment.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            # Bold section headers
            if line.startswith("Score:") or line in (
                "Strengths",
                "Areas for Growth",
                "Next Steps",
            ):
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
            else:
                run = p.add_run(line)
                run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _sanitize_filename(file.filename or "essay_annotated.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@app.post("/ingest_marked")
@limiter.limit("20/minute")
async def ingest_marked_essay(
    request: Request,
    file: UploadFile = File(...),
    student_name: str | None = Form(None),
    assignment_name: str | None = Form(None),
    mode: str = Form("imported_marked"),
    class_id: str | None = Form(None),
    user: dict = Depends(require_product("mark")),
):
    """
    Ingest an already-marked .docx file by extracting label counts from the summary table
    and logging it to mark_events.
    """
    # 1. Basic validation
    if not file.filename.lower().endswith(".docx"):
        return JSONResponse(
            status_code=400,
            content={"error": "Please upload a .docx file"},
        )

    # 2. Read file bytes
    docx_bytes = await file.read()

    # 3. Parse the document to extract summary metadata
    try:
        doc = Document(BytesIO(docx_bytes))
        _, extract_summary_metadata = get_engine()
        metadata = extract_summary_metadata(doc)
    except Exception as e:
        print(f"[ERROR] Failed to parse document: {repr(e)}")
        return JSONResponse(
            status_code=400,
            content={"error": "Failed to parse document. Please ensure this is a valid Vysti-marked .docx file."},
        )

    # 4. Validate that this is a marked document
    issues = metadata.get("issues", [])
    if not issues:
        return JSONResponse(
            status_code=400,
            content={"error": "This doc doesn't appear to be a Vysti-marked file"},
        )

    # 5. Build label_counts from issues
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl or lbl.startswith("__"):
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        # Ensure count is at least 1
        label_counter[lbl] += max(cnt_i, 1)

    total_labels = sum(label_counter.values())

    # 5b. Compute word_count and scores for persistence
    _ingest_word_count = None
    _ingest_scores = None
    try:
        essay_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        _ingest_word_count = len(essay_text.split()) if essay_text else 0
        _ingest_scores = _compute_scores(
            essay_text,
            mode=mode if mode != "imported_marked" else "textual_analysis",
            label_counts=dict(label_counter),
            mark_event_id=None,
            sentence_types={},
            repeated_nouns=[],
        )
    except Exception:
        pass

    # 6. Validate class_id if provided
    class_id_validated = None
    if class_id:
        user_id = user.get("id") if isinstance(user, dict) else None
        if not user_id:
            raise HTTPException(
                status_code=400,
                detail="Invalid user",
            )
        
        # Verify class exists and belongs to user
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            db_url = f"{SUPABASE_URL}/rest/v1/classes"
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.get(
                    f"{db_url}?id=eq.{class_id}&user_id=eq.{user_id}&archived=eq.false",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if not data or len(data) == 0:
                        raise HTTPException(
                            status_code=400,
                            detail="Class not found or does not belong to user",
                        )
                    class_id_validated = class_id
                else:
                    raise HTTPException(
                        status_code=400,
                        detail="Failed to validate class",
                    )

    # 7. Log to Supabase mark_events
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(
            status_code=500,
            detail="Supabase configuration missing on server",
        )

    user_id = user.get("id") if isinstance(user, dict) else None
    db_url = f"{SUPABASE_URL}/rest/v1/mark_events"
    payload = {
        "user_id": user_id,
        "file_name": file.filename,
        "mode": mode,
        "bytes": len(docx_bytes),
        "student_name": student_name,
        "assignment_name": assignment_name,
        "class_id": class_id_validated,
        "total_labels": total_labels,
        "label_counts": dict(label_counter),
        "issues": issues,
        "review_status": "pending",
        "word_count": _ingest_word_count,
        "scores": _sanitize_for_json(_ingest_scores) if _ingest_scores else None,
    }

    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.post(
                db_url,
                json=payload,
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
            )
            # Check if the insert was successful
            if resp.status_code < 200 or resp.status_code >= 300:
                print(f"[ERROR] Supabase insert failed: status={resp.status_code}, body={resp.text[:200]}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to log mark event to database.",
                )
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Failed to log mark event: {repr(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to log mark event.",
        )

    # 8. Return success response
    return JSONResponse(
        content={
            "ok": True,
            "total_labels": total_labels,
            "label_counts": dict(label_counter),
            "issues": _strip_ip_from_issues(issues),
        }
    )


def normalize_label(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip().lower()


def normalize_text(value: str | None) -> str:
    if not value:
        return ""
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    normalized = (
        normalized.replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201C", "\"")
        .replace("\u201D", "\"")
    )
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def split_into_paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    rewrite_pattern = r"\s*\*\s*Rewrite this paragraph for practice\s*\*\s*"
    normalized = re.sub(rewrite_pattern, "", normalized, flags=re.IGNORECASE)
    para_chunks = re.split(r"\n{2,}", normalized)
    paragraphs: list[str] = []
    for chunk in para_chunks:
        para_text = re.sub(r"\n+", " ", chunk).strip()
        if para_text:
            paragraphs.append(para_text)
    return paragraphs


def apply_rewrite_to_paragraph(
    paragraphs: list[str],
    paragraph_index: int | None,
    original_sentence: str,
    rewrite: str,
) -> tuple[list[str], bool, int | None]:
    if not paragraphs or not original_sentence:
        return paragraphs, False, paragraph_index

    if paragraph_index is not None and 0 <= paragraph_index < len(paragraphs):
        para_text = paragraphs[paragraph_index]
        if original_sentence in para_text:
            paragraphs[paragraph_index] = para_text.replace(original_sentence, rewrite, 1)
            return paragraphs, True, paragraph_index

    for idx, para_text in enumerate(paragraphs):
        if original_sentence in para_text:
            paragraphs[idx] = para_text.replace(original_sentence, rewrite, 1)
            return paragraphs, True, idx

    return paragraphs, False, paragraph_index


def build_doc_from_text(text: str) -> bytes:
    # Normalize newlines to \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Safety net: Remove rewrite-practice tag if it appears in the text
    rewrite_pattern = r"\s*\*\s*Rewrite this paragraph for practice\s*\*\s*"
    text = re.sub(rewrite_pattern, "", text, flags=re.IGNORECASE)

    # Split paragraphs on 2+ newlines
    para_chunks = re.split(r"\n{2,}", text)

    # Create document
    doc = Document()

    # Set default style to Times New Roman 12pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    # Set eastAsia font too
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Helper function to detect header-like lines
    def is_header_line(line_text: str) -> bool:
        """Detect if a line is likely a header (teacher name, date, course, etc.)"""
        text_lower = line_text.lower().strip()
        # Check for teacher titles
        if re.match(r"^(mr|ms|mrs|dr|prof)\.?\s+", text_lower):
            return True
        # Check for date patterns (e.g., "January 1, 2024" or "1/1/2024")
        if re.search(r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d+", text_lower):
            return True
        if re.search(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}", line_text):
            return True
        # Check for course/class keywords
        if "course" in text_lower or "class" in text_lower:
            return True
        # Short name-like lines (2-3 words, no sentence-ending punctuation)
        words = line_text.split()
        if len(words) <= 3 and not re.search(r"[.!?]$", line_text):
            return True
        return False

    # Helper function to check if text is a sentence (ends with .?!)
    def is_sentence(line_text: str) -> bool:
        """Check if text appears to be a sentence"""
        return bool(re.search(r"[.!?]$", line_text.strip()))

    # Track if we've found the essay title (first non-header short non-sentence line)
    title_found = False

    # Add paragraphs (collapsing single newlines within paragraphs to spaces)
    for para_chunk in para_chunks:
        # Collapse single newlines within paragraph to spaces
        para_text = re.sub(r"\n+", " ", para_chunk).strip()
        if not para_text:  # Skip empty paragraphs
            continue

        para = doc.add_paragraph(para_text)

        # Check if this is a header line
        is_header = is_header_line(para_text)
        is_sent = is_sentence(para_text)

        # Apply formatting based on paragraph type
        if is_header:
            # Header lines: no indentation, left-aligned
            para.paragraph_format.first_line_indent = Inches(0)
        elif not title_found and not is_header and not is_sent and len(para_text.split()) <= 10:
            # First non-header short non-sentence line: likely essay title - center it
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Inches(0)
            title_found = True
        else:
            # Prose paragraphs: apply MLA first-line indent (0.5")
            para.paragraph_format.first_line_indent = Inches(0.5)

    # If no paragraphs were created, add at least one
    if len(doc.paragraphs) == 0:
        para = doc.add_paragraph(text.strip() or "Empty document")
        para.paragraph_format.first_line_indent = Inches(0.5)

    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()
    return docx_bytes


# ── Label-delimiter regex (matches «→ Label text» from frontend extraction) ──
_LABEL_RE = re.compile(r"\u00AB(\u2192[^\u00BB]+)\u00BB")         # «→ Label»
_TEACHER_RE = re.compile(r"\[Teacher:\s*([^\]]+)\]")               # [Teacher: ...]
_SQUIGGLY_RE = re.compile(r"\{~([^~]+)~\}")                        # {~text~}
_STRIKE_RE = re.compile(r"\{x:([^}]+)\}")                          # {x:text}
_HL_AQUA_RE = re.compile(r"\{hl:([^}]+)\}")                        # {hl:text}
_HL_GRAY_RE = re.compile(r"\{g:([^}]+)\}")                         # {g:text}
_SP_RE = re.compile(r"\{sp\}")                                      # {sp}
_WC_RE = re.compile(r"\{wc\}")                                      # {wc}
_HL_GREEN_RE = re.compile(r"\{gr:([^}]+)\}")                        # {gr:text}  green highlight
_CONFUSION_RE = re.compile(r"\{\?\?\?\}")                            # {???}     confusion mark
_POSITIVE_RE = re.compile(r"\{\+([^}]+)\}")                          # {+✓} or {+☺} or {+★}
_NEGATIVE_RE = re.compile(r"\{-([^}]+)\}")                           # {-☹} negative indicator
_UNHAPPY_RE = re.compile(r"\{unhappy:([^}]+)\}")                     # {unhappy:text} unhappy highlight
_COMMENT_RE = re.compile(r"\{c\|([^|]*)\|([^}]+)\}")               # {c|anchor|comment} (anchor may be empty)
_BOLD_RE = re.compile(r"\{b:([^}]+)\}")                              # {b:text}  teacher bold
_ITALIC_RE = re.compile(r"\{i:([^}]+)\}")                            # {i:text}  italic text
_UNDERLINE_RE = re.compile(r"\{u:([^}]+)\}")                         # {u:text}  solid underline
_CARET_RE = re.compile(r"\{\^\}")                                    # {^}       missing element
_STAR_RE = re.compile(r"\{star:([^}]+)\}")                           # {star:text} exemplary highlight
_INSERT_RE = re.compile(r"\{ins:([^}]+)\}")                          # {ins:text}  teacher insert
_CUSTOM_SUP_RE = re.compile(r"\{sup:([^}]+)\}")                     # {sup:label} custom superscript
_TAGGED_HL_RE = re.compile(r"\{tag:(\w+):([^:]*):([^}]+)\}")       # {tag:hl:label:text}
_PARA_RE = re.compile(r"\{para:([^}]+)\}")                          # {para:label}
_REORDER_RE = re.compile(r"\{reorder:([^}]+)\}")                    # {reorder:①}
_ARROW_MARK_RE = re.compile(r"\{arrow\}")                           # {arrow}   inline arrow mark
_COMBINED_RE = re.compile(
    r"("
    r"\u00AB\u2192[^\u00BB]+\u00BB"           # «→ Label»
    r"|\[Teacher:\s*[^\]]+\]"                  # [Teacher: ...]
    r"|\{c\|[^|]*\|[^}]+\}"                   # {c|anchor|comment} (anchor may be empty)
    r"|\{~[^~]+~\}"                            # {~squiggly~}
    r"|\{x:[^}]+\}"                            # {x:strikethrough}
    r"|\{tag:\w+:[^:]*:[^}]+\}"               # {tag:hl:label:text} tagged highlight (before {hl:})
    r"|\{hl:[^}]+\}"                           # {hl:aqua highlight}
    r"|\{star:[^}]+\}"                         # {star:text} exemplary (before {sp})
    r"|\{g:[^}]+\}"                            # {g:gray highlight}
    r"|\{gr:[^}]+\}"                           # {gr:green highlight}
    r"|\{b:[^}]+\}"                            # {b:text} teacher bold
    r"|\{i:[^}]+\}"                            # {i:text} italic text
    r"|\{u:[^}]+\}"                            # {u:text} solid underline
    r"|\{ins:[^}]+\}"                          # {ins:text} teacher insert
    r"|\{sup:[^}]+\}"                          # {sup:label} custom superscript
    r"|\{para:[^}]+\}"                         # {para:label} paragraph note
    r"|\{reorder:[^}]+\}"                      # {reorder:①} reorder marker
    r"|\{arrow\}"                               # {arrow} inline arrow mark
    r"|\{sp\}"                                  # {sp} spelling
    r"|\{wc\}"                                  # {wc} word choice
    r"|\{\^\}"                                  # {^} missing element caret
    r"|\{\?\?\?\}"                              # {???} confusion
    r"|\{\+[^}]+\}"                             # {+✓} or {+☺} or {+★} positive indicator
    r"|\{-[^}]+\}"                              # {-☹} negative indicator
    r"|\{unhappy:[^}]+\}"                       # {unhappy:text} unhappy highlight
    r")"
)


def _build_comments_part(doc, comments):
    """Create a Word comments.xml part and add it to the document.

    comments: list of dicts with keys: id (int), text (str), author (str)
    """
    if not comments:
        return

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.Element(qn("w:comments"), nsmap={"w": W})
    now = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    for c in comments:
        cel = etree.SubElement(root, qn("w:comment"))
        cel.set(qn("w:id"), str(c["id"]))
        cel.set(qn("w:author"), c.get("author", "Teacher"))
        cel.set(qn("w:date"), now)
        p_el = etree.SubElement(cel, qn("w:p"))
        r_el = etree.SubElement(p_el, qn("w:r"))
        t_el = etree.SubElement(r_el, qn("w:t"))
        t_el.text = c["text"]
        t_el.set(qn("xml:space"), "preserve")

    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = Part(
        partname=PackURI("/word/comments.xml"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        blob=blob,
        package=doc.part.package,
    )
    doc.part.relate_to(
        part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )


def _apply_comment_shading(run):
    """Apply soft maroon background shading (#F4D4D8) to a comment anchor run."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4D4D8")  # ~15% opacity of brand maroon #A90D22
    run._element.get_or_add_rPr().append(shd)


def _wrap_run_with_comment(run, comment_id):
    """Add commentRangeStart/End around a run and a commentReference after it."""
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    run._element.addprevious(start)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    run._element.addnext(end)

    ref_run = OxmlElement("w:r")
    ref_rPr = OxmlElement("w:rPr")
    ref_style = OxmlElement("w:rStyle")
    ref_style.set(qn("w:val"), "CommentReference")
    ref_rPr.append(ref_style)
    ref_run.append(ref_rPr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    end.addnext(ref_run)


def build_teacher_doc_from_text(text: str, comment: str = "") -> bytes:
    """Build a .docx for teacher 'Download Marked Essay'.

    Vysti labels (wrapped in «→ Label» by the frontend) are rendered with
    yellow highlighting and bold. Teacher comments ({c|anchor|comment}) become
    Word margin comments. An optional teacher comment section is appended at
    the end.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Strip null bytes and XML-incompatible control characters (U+0000–U+0008,
    # U+000B–U+000C, U+000E–U+001F) that can leak from docx-preview DOM text.
    # lxml/python-docx raises ValueError if these are present.
    _ctrl_re = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
    text = _ctrl_re.sub("", text)
    comment = _ctrl_re.sub("", comment)

    # Safety net: strip rewrite-practice tags
    text = re.sub(
        r"\s*\*\s*Rewrite this paragraph for practice\s*\*\s*",
        "", text, flags=re.IGNORECASE,
    )

    para_chunks = re.split(r"\n{2,}", text)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    pending_comments = []  # Word margin comments to finalize

    for para_chunk in para_chunks:
        para_text = re.sub(r"\n+", " ", para_chunk).strip()
        if not para_text:
            continue

        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.5)

        # Split into segments: normal text, «→ label», {c|...|...}, marks
        segments = _COMBINED_RE.split(para_text)

        for segment in segments:
            if not segment:
                continue

            label_m = _LABEL_RE.match(segment)
            teacher_m = _TEACHER_RE.match(segment)
            comment_m = _COMMENT_RE.match(segment)
            squiggly_m = _SQUIGGLY_RE.match(segment)
            strike_m = _STRIKE_RE.match(segment)
            hl_aqua_m = _HL_AQUA_RE.match(segment)
            hl_gray_m = _HL_GRAY_RE.match(segment)
            hl_green_m = _HL_GREEN_RE.match(segment)
            sp_m = _SP_RE.match(segment)
            wc_m = _WC_RE.match(segment)
            confusion_m = _CONFUSION_RE.match(segment)
            positive_m = _POSITIVE_RE.match(segment)
            negative_m = _NEGATIVE_RE.match(segment)
            unhappy_m = _UNHAPPY_RE.match(segment)
            bold_m = _BOLD_RE.match(segment)
            italic_m = _ITALIC_RE.match(segment)
            underline_m = _UNDERLINE_RE.match(segment)
            caret_m = _CARET_RE.match(segment)
            star_m = _STAR_RE.match(segment)
            insert_m = _INSERT_RE.match(segment)
            custom_sup_m = _CUSTOM_SUP_RE.match(segment)
            tagged_hl_m = _TAGGED_HL_RE.match(segment)
            para_m = _PARA_RE.match(segment)
            reorder_m = _REORDER_RE.match(segment)
            arrow_mark_m = _ARROW_MARK_RE.match(segment)

            if label_m:
                # Vysti label → yellow highlight, bold
                run = para.add_run(label_m.group(1))
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif comment_m:
                # Teacher comment → Word margin comment
                anchor_text = comment_m.group(1).strip()
                comment_text = comment_m.group(2)
                cid = len(pending_comments)
                pending_comments.append({"id": cid, "text": comment_text, "author": "Teacher"})
                if anchor_text:
                    # Normal case: anchor text wraps the commented word(s)
                    run = para.add_run(anchor_text)
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"
                    _apply_comment_shading(run)
                    _wrap_run_with_comment(run, cid)
                else:
                    # Empty anchor (cross-boundary selection): point comment
                    # Insert a zero-width space so commentRangeStart/End have a run
                    run = para.add_run("\u200B")
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"
                    _wrap_run_with_comment(run, cid)
            elif teacher_m:
                # Old palette annotation → bold red text
                run = para.add_run(f"[Teacher: {teacher_m.group(1)}]")
                run.bold = True
                run.font.color.rgb = RGBColor(211, 47, 47)  # #D32F2F
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif squiggly_m:
                # Squiggly underline → red wavy underline
                run = para.add_run(squiggly_m.group(1))
                run.font.underline = WD_UNDERLINE.WAVY
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif strike_m:
                # Red strikethrough
                run = para.add_run(strike_m.group(1))
                run.font.strike = True
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif hl_aqua_m:
                # Aqua highlight
                run = para.add_run(hl_aqua_m.group(1))
                run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif hl_gray_m:
                # Gray highlight
                run = para.add_run(hl_gray_m.group(1))
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif hl_green_m:
                # Green highlight
                run = para.add_run(hl_green_m.group(1))
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif bold_m:
                # Teacher inline comment → red bold on yellow highlight
                run = para.add_run(bold_m.group(1))
                run.bold = True
                run.font.color.rgb = RGBColor(211, 47, 47)  # #D32F2F
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif italic_m:
                # Italic text (preserved from original Word doc or teacher edit)
                run = para.add_run(italic_m.group(1))
                run.italic = True
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif underline_m:
                # Solid underline
                run = para.add_run(underline_m.group(1))
                run.font.underline = True
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif star_m:
                # Exemplary highlight → yellow highlight (closest Word color to amber)
                run = para.add_run(star_m.group(1))
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif insert_m:
                # Teacher insert → bold red inline text
                run = para.add_run(insert_m.group(1))
                run.bold = True
                run.font.color.rgb = RGBColor(211, 47, 47)  # #D32F2F
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif custom_sup_m:
                # Custom superscript label → red bold superscript
                run = para.add_run(custom_sup_m.group(1))
                run.font.superscript = True
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif tagged_hl_m:
                # Tagged highlight → color-highlighted text + bold superscript label
                color_code = tagged_hl_m.group(1)
                tag_label = tagged_hl_m.group(2)
                tag_text = tagged_hl_m.group(3)
                color_map = {"hl": WD_COLOR_INDEX.TURQUOISE, "g": WD_COLOR_INDEX.GRAY_25, "gr": WD_COLOR_INDEX.BRIGHT_GREEN}
                hl_color = color_map.get(color_code, WD_COLOR_INDEX.TURQUOISE)
                run = para.add_run(tag_text)
                run.font.highlight_color = hl_color
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                if tag_label:
                    label_run = para.add_run(tag_label)
                    label_run.font.superscript = True
                    label_run.font.color.rgb = RGBColor(211, 47, 47)
                    label_run.bold = True
                    label_run.font.size = Pt(12)
                    label_run.font.name = "Times New Roman"
            elif para_m:
                # Paragraph note — silently consumed (teachers use margin comments instead)
                pass
            elif reorder_m:
                # Reorder marker → purple bold superscript
                run = para.add_run(reorder_m.group(1))
                run.font.superscript = True
                run.font.color.rgb = RGBColor(124, 58, 237)  # #7C3AED
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            elif arrow_mark_m:
                # Arrow mark → bold maroon "→" inline
                run = para.add_run(" \u2192 ")
                run.bold = True
                run.font.color.rgb = RGBColor(169, 13, 34)  # #A90D22 maroon
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            elif sp_m:
                # Spelling mark → red superscript "sp"
                run = para.add_run("sp")
                run.font.superscript = True
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            elif wc_m:
                # Word choice mark → red superscript "wc"
                run = para.add_run("wc")
                run.font.superscript = True
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            elif caret_m:
                # Missing element caret → red superscript "^"
                run = para.add_run("^")
                run.font.superscript = True
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            elif confusion_m:
                # Confusion mark → red bold inline "???" with yellow highlight
                run = para.add_run("???")
                run.font.superscript = False
                run.font.color.rgb = RGBColor(211, 47, 47)
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif positive_m:
                # Positive indicator (✓, ☺, or ★) → green/amber superscript
                icon = positive_m.group(1)
                if icon == "\u2605":
                    # Star → Wingdings 2 six-pointed star for Word compatibility
                    run = para.add_run("\u00AB")  # renders as star in Wingdings 2
                    run.font.name = "Wingdings 2"
                    run.font.color.rgb = RGBColor(217, 119, 6)  # #D97706 amber
                elif icon == "\u2713":
                    # Checkmark → Wingdings checkmark
                    run = para.add_run(chr(252))
                    run.font.name = "Wingdings"
                    run.font.color.rgb = RGBColor(22, 163, 74)  # #16A34A green
                elif icon == "\u263A":
                    # Smiley → Wingdings smiley
                    run = para.add_run("J")
                    run.font.name = "Wingdings"
                    run.font.color.rgb = RGBColor(22, 163, 74)  # #16A34A green
                else:
                    run = para.add_run(icon)
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(22, 163, 74)
                if icon == "\u263A":
                    # Smiley sits inline at body size (not superscript)
                    run.font.superscript = False
                    run.font.size = Pt(12)
                else:
                    run.font.superscript = True
                    run.font.size = Pt(10)
                run.bold = True
            elif negative_m:
                # Negative indicator (☹) → red, inline at body size
                icon = negative_m.group(1)
                if icon == "\u2639":
                    # Frowny → Wingdings frown (letter L)
                    run = para.add_run("L")
                    run.font.name = "Wingdings"
                else:
                    run = para.add_run(icon)
                    run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(220, 38, 38)  # #DC2626 red
                run.font.superscript = False
                run.font.size = Pt(12)
                run.bold = True
            elif unhappy_m:
                # Unhappy highlight → red-tinted text (body of the mark)
                run = para.add_run(unhappy_m.group(1))
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            else:
                # Normal essay text
                run = para.add_run(segment)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

    if len(doc.paragraphs) == 0:
        para = doc.add_paragraph(text.strip() or "Empty document")
        para.paragraph_format.first_line_indent = Inches(0.5)

    # ── Finalize Word margin comments ──
    if pending_comments:
        _build_comments_part(doc, pending_comments)

    # ── Teacher comment section ──
    if comment and comment.strip():
        # Separator line
        sep = doc.add_paragraph()
        sep.paragraph_format.first_line_indent = Inches(0)
        sep_run = sep.add_run("\u2500" * 40)
        sep_run.font.size = Pt(10)
        sep_run.font.color.rgb = RGBColor(160, 160, 160)

        # Header
        header_para = doc.add_paragraph()
        header_para.paragraph_format.first_line_indent = Inches(0)
        header_run = header_para.add_run("Teacher Comment")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = "Times New Roman"

        # Comment body
        for line in comment.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            r = p.add_run(line)
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()
    return docx_bytes


def build_teacher_config_from_titles(titles: list[TitleInfo] | None) -> dict | None:
    if not titles:
        return {"highlight_thesis_devices": False}

    teacher_config: dict = {"highlight_thesis_devices": False}
    trimmed = titles[:3]
    if len(trimmed) > 0:
        t1 = trimmed[0]
        teacher_config["author_name"] = t1.author
        teacher_config["text_title"] = t1.title
        teacher_config["text_is_minor_work"] = t1.is_minor
    if len(trimmed) > 1:
        t2 = trimmed[1]
        teacher_config["author_name_2"] = t2.author
        teacher_config["text_title_2"] = t2.title
        teacher_config["text_is_minor_work_2"] = t2.is_minor
    if len(trimmed) > 2:
        t3 = trimmed[2]
        teacher_config["author_name_3"] = t3.author
        teacher_config["text_title_3"] = t3.title
        teacher_config["text_is_minor_work_3"] = t3.is_minor
    return teacher_config


def find_matching_examples(
    examples: list,
    label_value: str,
    paragraph_index: int | None,
    original_sentence: str | None,
) -> tuple[list[dict], int | None]:
    normalized_label = normalize_label(label_value)
    label_matches = [
        ex
        for ex in examples
        if isinstance(ex, dict)
        and normalize_label(ex.get("label")) == normalized_label
    ]

    if paragraph_index is not None:
        matches = [ex for ex in label_matches if ex.get("paragraph_index") == paragraph_index]
        return matches, paragraph_index

    if not original_sentence:
        return [], None

    normalized_original = normalize_text(original_sentence)
    best_ratio = 0.0
    best_example = None
    for ex in label_matches:
        sentence = ex.get("sentence") or ""
        ratio = difflib.SequenceMatcher(
            None,
            normalize_text(sentence),
            normalized_original,
        ).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_example = ex

    if best_example and best_ratio >= 0.7:
        target_index = best_example.get("paragraph_index")
        if target_index is None:
            return [best_example], None
        matches = [ex for ex in label_matches if ex.get("paragraph_index") == target_index]
        return matches, target_index

    return [], None


@app.post("/revision/check")
@limiter.limit("40/minute")
async def check_revision(
    request: Request,
    body: RevisionCheckRequest,
    user: dict = Depends(require_product("revise")),
):
    """
    Check if a rewritten sentence still triggers a specific issue label.
    Validates the rewrite text in isolation without requiring Preview context.
    """
    # Validate rewrite: reject empty/whitespace
    if not body.rewrite or not body.rewrite.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite cannot be empty or whitespace only",
        )

    # Cap rewrite length (2000 chars)
    if len(body.rewrite) > 2000:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite exceeds maximum length of 2000 characters",
        )

    if not body.original_sentence:
        return JSONResponse(
            content={
                "approved": False,
                "message": "Original sentence is required for validation.",
            }
        )

    if normalize_text(body.rewrite) == normalize_text(body.original_sentence):
        return JSONResponse(
            content={
                "approved": False,
                "message": "No changes detected — edit the example before checking.",
            }
        )

    label_value = body.label_trimmed or body.label
    mode = body.mode or "textual_analysis"
    teacher_config = build_teacher_config_from_titles(body.titles)

    mark_docx_bytes, _ = get_engine()
    normalized_label = normalize_label(label_value)

    # Mark the rewrite in isolation to see if the issue still triggers
    doc_rewrite = build_doc_from_text(body.rewrite.strip())
    _, metadata_rewrite = mark_docx_bytes(
        doc_rewrite,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
    )

    examples_rewrite = metadata_rewrite.get("examples", []) if isinstance(metadata_rewrite, dict) else []

    # Debug: log all labels found in the rewrite
    rewrite_labels = [ex.get("label", "") for ex in examples_rewrite if isinstance(ex, dict)]
    if _DEBUG:
        print(f"[REVISION CHECK] Target label: '{label_value}' (normalized: '{normalized_label}')")
        print(f"[REVISION CHECK] Labels found in rewrite: {rewrite_labels}")

    # Count how many times the target label appears in the rewrite
    rewrite_count = sum(
        1 for ex in examples_rewrite
        if isinstance(ex, dict) and normalize_label(ex.get("label", "")) == normalized_label
    )

    if _DEBUG: print(f"[REVISION CHECK] Rewrite count for target label: {rewrite_count}")

    # Approved if the rewrite no longer triggers the issue
    if rewrite_count == 0:
        return JSONResponse(
            content={
                "approved": True,
                "message": "Looks good! Revision approved.",
                "after_count": rewrite_count,
            }
        )

    return JSONResponse(
        content={
            "approved": False,
            "message": "Still needs revision — the issue is still triggering here.",
            "after_count": rewrite_count,
        }
    )


class PracticeAnalyzeRequest(BaseModel):
    sentence: str
    mode: str | None = None


# Labels that require full-essay context and cannot be fixed in single-sentence Practice
_PRACTICE_SKIP_LABELS = {
    "the first sentence should state the author",
    "closed thesis statement",
    "close thesis statement",
    "use a closed thesis",
    "the thesis should close",
    "topic sentence should state",
    "topic sentences stating",
    "the topic sentence should clearly",
    "final sentence of a body paragraph",
    "final sentence should link",
    "explain your evidence",
    "explain evidence",
    "introduce quotations",
    "quotation should be introduced",
    "embed quotations",
    "clarify pronoun",
    "review how you refer to the author",
    "refer to the author by last name",
    "check title formatting",
    "paragraph is too short",
    "paragraph is too long",
    "noun repetition",
    "no quotations in thesis",
    "essay title format",
    "capitalize the words in the title",
    "topics in the thesis statement",
}


def _is_practice_relevant(label: str) -> bool:
    """Return True if this label is fixable in single-sentence Practice."""
    norm = label.lower().strip()
    for skip in _PRACTICE_SKIP_LABELS:
        if skip in norm:
            return False
    return True


def _extract_practice_issues(metadata: dict) -> list[dict]:
    """
    Extract issues from marker metadata for the Practice page.
    Returns a list of {label, found_value, explanation, student_guidance}
    using the *examples* list (one entry per flagged word) enriched with
    explanation/guidance from the *issues* list (one entry per unique label).
    Filters out structural/context-dependent labels.
    """
    issues_meta = metadata.get("issues", []) if isinstance(metadata, dict) else []
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []

    # Build lookup: label → {explanation, student_guidance} from issues metadata
    guidance_by_label = {}
    for iss in issues_meta:
        if not isinstance(iss, dict):
            continue
        lbl = iss.get("label", "")
        if lbl and lbl not in guidance_by_label:
            guidance_by_label[lbl] = {
                "explanation": iss.get("explanation", ""),
                "student_guidance": iss.get("student_guidance", ""),
            }

    # Build practice issues from examples (each flagged word = one issue)
    result = []
    for ex in examples:
        if not isinstance(ex, dict):
            continue
        label = ex.get("label", "")
        if not label or not _is_practice_relevant(label):
            continue
        info = guidance_by_label.get(label, {})
        result.append({
            "label": label,
            "found_value": ex.get("found_value", ""),
            "explanation": info.get("explanation", ""),
            "student_guidance": info.get("student_guidance", ""),
        })

    return result


@app.post("/practice/analyze")
@limiter.limit("20/minute")
async def analyze_practice(
    request: Request,
    body: PracticeAnalyzeRequest,
):
    """
    Unauthenticated endpoint that runs the marker on a sentence and returns
    all issues found, with explanation and guidance for each.
    Used by the Practice page to dynamically discover issues at load time.
    """
    if not body.sentence or not body.sentence.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Sentence cannot be empty",
        )

    if len(body.sentence) > 2000:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Sentence exceeds maximum length of 2000 characters",
        )

    mode = body.mode or "textual_analysis"
    mark_docx_bytes, _ = get_engine()

    doc = build_doc_from_text(body.sentence.strip())
    _, metadata = mark_docx_bytes(doc, mode=mode)

    issues = _extract_practice_issues(metadata)

    return JSONResponse(content={"issues": issues})


@app.post("/practice/check-all")
@limiter.limit("20/minute")
async def check_practice_all(
    request: Request,
    body: RevisionCheckRequest,
):
    """
    Unauthenticated endpoint that runs the marker on a rewrite and returns
    ALL issues found (with explanation and guidance).  The frontend compares
    against the original issue list to decide resolved / still-present / new.
    """
    if not body.rewrite or not body.rewrite.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite cannot be empty or whitespace only",
        )

    if len(body.rewrite) > 2000:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite exceeds maximum length of 2000 characters",
        )

    if not body.original_sentence:
        return JSONResponse(content={"issues": [], "error": "Original sentence is required."})

    if normalize_text(body.rewrite) == normalize_text(body.original_sentence):
        return JSONResponse(content={"issues": [], "unchanged": True})

    mode = body.mode or "textual_analysis"
    teacher_config = build_teacher_config_from_titles(body.titles)

    mark_docx_bytes, _ = get_engine()

    doc_rewrite = build_doc_from_text(body.rewrite.strip())
    _, metadata_rewrite = mark_docx_bytes(
        doc_rewrite,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
    )

    issues = _extract_practice_issues(metadata_rewrite)

    return JSONResponse(content={"issues": issues})


# ── Practice essay directory ──
_PRACTICE_ESSAYS_DIR = pathlib.Path(__file__).parent / "Test Essays"


@app.get("/practice/random-essay")
@limiter.limit("10/minute")
async def practice_random_essay(request: Request):
    """
    Unauthenticated endpoint that picks a random Test Essay .docx,
    runs it through the marker, computes scores, and returns the marked
    document + metadata in the same JSON format as /mark (return_metadata).
    """
    if not _PRACTICE_ESSAYS_DIR.is_dir():
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Practice essays directory not found",
        )

    docx_files = list(_PRACTICE_ESSAYS_DIR.glob("*.docx"))
    if not docx_files:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="No practice essays available",
        )

    chosen = random.choice(docx_files)
    docx_bytes = chosen.read_bytes()
    mode = "textual_analysis"

    mark_docx_bytes_fn, _ = get_engine()
    marked_bytes, metadata = mark_docx_bytes_fn(
        docx_bytes,
        mode=mode,
        include_summary_table=True,
    )

    # Count labels
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label", "")
        cnt = issue.get("count")
        if lbl:
            label_counter[lbl] = cnt if isinstance(cnt, int) else 1

    # Compute scores
    scores = None
    try:
        from docx import Document as _Document
        orig_doc = _Document(BytesIO(docx_bytes))
        essay_text = "\n\n".join(
            p.text.strip() for p in orig_doc.paragraphs if p.text.strip()
        )
        scores = _compute_scores(
            essay_text,
            mode=mode,
            label_counts=dict(label_counter),
            mark_event_id=None,
            sentence_types=(
                metadata.get("sentence_types", {})
                if isinstance(metadata, dict) else {}
            ),
            repeated_nouns=(
                metadata.get("repeated_nouns", [])
                if isinstance(metadata, dict) else []
            ),
        )
    except Exception as e:
        if _DEBUG:
            print(f"[PRACTICE] Score computation failed: {repr(e)}")

    # Build enriched metadata (same shape as /mark return_metadata)
    enriched = dict(metadata) if isinstance(metadata, dict) else {}
    enriched["label_counts"] = dict(label_counter)
    if scores:
        enriched["scores"] = scores
    if "issues" in enriched:
        enriched["issues"] = _strip_ip_from_issues(enriched["issues"])
    if "examples" in enriched:
        enriched["examples"] = _strip_ip_from_examples(enriched["examples"])

    return JSONResponse({
        "document": base64.b64encode(marked_bytes).decode("utf-8"),
        "filename": chosen.name,
        "metadata": _sanitize_for_json(enriched),
    })


@app.post("/practice/revision-check")
@limiter.limit("40/minute")
async def practice_revision_check(
    request: Request,
    body: RevisionCheckRequest,
):
    """
    Unauthenticated version of /revision/check for Practice page.
    Checks if a rewritten sentence still triggers a specific issue label.
    """
    if not body.rewrite or not body.rewrite.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite cannot be empty or whitespace only",
        )

    if len(body.rewrite) > 2000:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite exceeds maximum length of 2000 characters",
        )

    if not body.original_sentence:
        return JSONResponse(
            content={
                "approved": False,
                "message": "Original sentence is required for validation.",
            }
        )

    if normalize_text(body.rewrite) == normalize_text(body.original_sentence):
        return JSONResponse(
            content={
                "approved": False,
                "message": "No changes detected — edit the example before checking.",
            }
        )

    label_value = body.label_trimmed or body.label
    mode = body.mode or "textual_analysis"
    teacher_config = build_teacher_config_from_titles(body.titles)

    mark_docx_bytes_fn, _ = get_engine()
    normalized_label_val = normalize_label(label_value)

    doc_rewrite = build_doc_from_text(body.rewrite.strip())
    _, metadata_rewrite = mark_docx_bytes_fn(
        doc_rewrite,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
    )

    examples_rewrite = (
        metadata_rewrite.get("examples", [])
        if isinstance(metadata_rewrite, dict)
        else []
    )

    rewrite_count = sum(
        1
        for ex in examples_rewrite
        if isinstance(ex, dict)
        and normalize_label(ex.get("label", "")) == normalized_label_val
    )

    if rewrite_count == 0:
        return JSONResponse(
            content={
                "approved": True,
                "message": "Looks good! Revision approved.",
                "after_count": rewrite_count,
            }
        )

    return JSONResponse(
        content={
            "approved": False,
            "message": "Still needs revision — the issue is still triggering here.",
            "after_count": rewrite_count,
        }
    )


@app.post("/mark_text")
@limiter.limit("20/minute")
async def mark_text(
    request: Request,
    body: MarkTextRequest,
    user: dict = Depends(require_api_product("mark", "revise")),
):
    """
    Mark text content by creating a .docx in memory and running the marking pipeline.

    Request JSON:
    {
      "file_name": "OriginalFileName.docx",
      "text": "Full essay text with paragraphs",
      "mode": "student"
    }

    Returns the marked .docx bytes (same as /mark).
    """
    _is_api_client = user.get("_is_api_client", False) if isinstance(user, dict) else False
    _api_start_time = time.time()

    # 0. Product-level access check based on calling context
    await _enforce_product_for_mode(user, body.student_mode)

    # 0a. Free-tier usage check — skip for API key clients (quota checked in middleware)
    if not _is_api_client:
        _mt_user_id = user.get("id") if isinstance(user, dict) else None
        if _mt_user_id:
            _mt_profile = await get_user_profile(_mt_user_id)
            _mt_tier = (_mt_profile or {}).get("subscription_tier", "free")
            if _mt_tier == "free":
                _mt_used = await count_user_marks(_mt_user_id)
                if _mt_used >= _FREE_TIER_MARK_LIMIT:
                    raise HTTPException(
                        status_code=402,
                        detail={
                            "message": "Subscribe for unlimited essay marking.",
                            "code": "USAGE_LIMIT",
                        },
                    )

    # 0b. Mobile-specific usage limits (separate from desktop free tier)
    _is_mobile = body.source == "mobile"
    if _is_mobile and not _is_api_client:
        _mt_user_id_mob = user.get("id") if isinstance(user, dict) else None
        if _mt_user_id_mob:
            _mob_total = await count_mobile_marks(_mt_user_id_mob)
            if _mob_total >= _MOBILE_MARK_LIMIT:
                raise HTTPException(
                    status_code=402,
                    detail={
                        "message": "You've used all your free mobile marks. Subscribe for unlimited marks.",
                        "code": "MOBILE_LIMIT",
                    },
                )
            _mob_today = await count_mobile_marks_today(_mt_user_id_mob)
            if _mob_today >= _MOBILE_DAILY_LIMIT:
                raise HTTPException(
                    status_code=429,
                    detail={
                        "message": "Daily mobile limit reached. Try again tomorrow, or use the full desktop version.",
                        "code": "MOBILE_DAILY_LIMIT",
                    },
                )

    # 0c. Text length limit (50,000 chars ≈ 8,000 words)
    _MAX_TEXT_CHARS = 50_000
    if body.text and len(body.text) > _MAX_TEXT_CHARS:
        raise HTTPException(status_code=400, detail=f"Text exceeds {_MAX_TEXT_CHARS} character limit.")

    # 1. Create .docx from text
    docx_bytes = build_doc_from_text(body.text)

    # 2. Build teacher_config from body.titles + optional rule overrides
    teacher_config = build_teacher_config_from_titles(body.titles) or {}
    teacher_config["student_mode"] = body.student_mode
    # Apply teacher rule overrides when present (sent by teacher recheck)
    _rule_fields = [
        "forbid_personal_pronouns", "forbid_audience_reference",
        "enforce_closed_thesis", "require_body_evidence",
        "allow_intro_summary_quotes", "enforce_intro_quote_rule",
        "enforce_long_quote_rule", "enforce_contractions_rule",
        "enforce_which_rule", "enforce_weak_verbs_rule",
        "enforce_fact_proof_rule", "enforce_human_people_rule",
        "enforce_vague_terms_rule", "highlight_thesis_devices",
    ]
    for _rf in _rule_fields:
        _rv = getattr(body, _rf, None)
        if _rv is not None:
            teacher_config[_rf] = _rv

    # 3. Call mark_docx_bytes (same pipeline as /mark)
    mark_docx_bytes, _ = get_engine()
    mode = body.mode or "textual_analysis"
    marked_bytes, metadata = mark_docx_bytes(
        docx_bytes,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
        include_summary_table=bool(body.include_summary_table),
    )
    
    # 4. Extract examples and issues from metadata
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []
    
    # 4. Count labels
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl or lbl.startswith("__"):
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        label_counter[lbl] += (cnt_i if cnt_i > 0 else 1)
    
    total_labels = sum(label_counter.values())
    
    # 5. Log to Supabase mark_events (best-effort) — skip for API key clients
    mark_event_id = None
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and not _is_api_client:
            user_id = user.get("id") if isinstance(user, dict) else None

            # Clear old mark_events for this user/file to ensure fresh start
            delete_url = f"{SUPABASE_URL}/rest/v1/mark_events"
            encoded_fn = urllib.parse.quote(body.file_name or "", safe="")
            async with httpx.AsyncClient(timeout=5) as client:
                await client.delete(
                    f"{delete_url}?user_id=eq.{user_id}&file_name=eq.{encoded_fn}",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )

            db_url = f"{SUPABASE_URL}/rest/v1/mark_events?select=id"
            payload = {
                "user_id": user_id,
                "file_name": body.file_name,
                "mode": mode,
                "bytes": len(docx_bytes),
                "total_labels": total_labels,
                "label_counts": dict(label_counter),
                "issues": issues,
                "review_status": "pending",
                "source": body.source or "desktop",
            }

            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.post(
                    db_url,
                    json=payload,
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=representation",
                    },
                )
                if resp.status_code >= 200 and resp.status_code < 300:
                    resp_data = resp.json()
                    if resp_data and isinstance(resp_data, list) and len(resp_data) > 0:
                        mark_event_id = resp_data[0].get("id")
    except Exception as e:
        print("Failed to log mark_event:", repr(e))

    # 6. Log examples to Supabase issue_examples (best-effort)
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and examples:
            user_id = user.get("id") if isinstance(user, dict) else None
            if user_id:
                # Clear old cached examples for this user/file to ensure fresh start
                delete_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                encoded_fn2 = urllib.parse.quote(body.file_name or "", safe="")
                async with httpx.AsyncClient(timeout=5) as client:
                    await client.delete(
                        f"{delete_url}?user_id=eq.{user_id}&file_name=eq.{encoded_fn2}",
                        headers={
                            "apikey": SUPABASE_SERVICE_KEY,
                            "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        },
                    )

                example_rows = []
                for ex in examples:
                    if not isinstance(ex, dict):
                        continue
                    label = ex.get("label")
                    sentence = ex.get("sentence")
                    paragraph_index = ex.get("paragraph_index")
                    if not label or not sentence:
                        continue
                    example_row = {
                        "user_id": user_id,
                        "file_name": body.file_name,
                        "mode": mode,
                        "label": label,
                        "sentence": sentence,
                        "paragraph_index": paragraph_index,
                        "mark_event_id": mark_event_id,
                        # Include context fields for dynamic guidance (use None if missing to ensure all rows have same keys)
                        "found_value": ex.get("found_value"),
                        "topics": ex.get("topics"),
                        "thesis": ex.get("thesis"),
                        "confidence": ex.get("confidence"),
                        "original_phrase": ex.get("original_phrase"),
                    }

                    example_rows.append(example_row)
                
                if example_rows:
                    db_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                    async with httpx.AsyncClient(timeout=5) as client:
                        await client.post(
                            db_url,
                            json=example_rows,
                            headers={
                                "apikey": SUPABASE_SERVICE_KEY,
                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                "Content-Type": "application/json",
                                "Prefer": "return=minimal",
                            },
                        )
    except Exception as e:
        print("Failed to log issue_examples:", repr(e))
    
    # 7. Return marked .docx bytes
    clean_name = _sanitize_filename(body.file_name or "essay.docx")
    base_name = clean_name.rsplit(".", 1)[0] if clean_name else "essay"
    output_filename = f"{base_name}_marked.docx"

    # Log API key usage (best-effort, non-blocking)
    if _is_api_client:
        _api_elapsed = int((time.time() - _api_start_time) * 1000)
        await _log_api_usage(
            api_key_id=user.get("_api_key_id", ""),
            endpoint="/mark_text",
            status_code=200,
            chars_processed=len(body.text or ""),
            response_ms=_api_elapsed,
            client_ip=get_remote_address(request),
            metadata={"mode": mode, "total_labels": total_labels},
        )

    if body.return_metadata:
        import base64
        enriched = dict(metadata) if isinstance(metadata, dict) else {}
        enriched["total_labels"] = total_labels
        enriched["label_counts"] = dict(label_counter)
        enriched["mark_event_id"] = mark_event_id
        # Compute scores server-side
        try:
            enriched["scores"] = _compute_scores(
                body.text,
                mode=mode,
                label_counts=dict(label_counter),
                mark_event_id=mark_event_id,
                sentence_types=enriched.get("sentence_types", {}),
                repeated_nouns=enriched.get("repeated_nouns", []),
            )
        except Exception:
            pass
        # Strip proprietary fields before sending to client
        if "issues" in enriched:
            enriched["issues"] = _strip_ip_from_issues(enriched["issues"])
        if "examples" in enriched:
            enriched["examples"] = _strip_ip_from_examples(enriched["examples"])
        # For API clients, strip mark_event_id (internal tracking)
        if _is_api_client:
            enriched.pop("mark_event_id", None)
        return JSONResponse({
            "document": base64.b64encode(marked_bytes).decode('utf-8'),
            "filename": output_filename,
            "metadata": _sanitize_for_json(enriched),
        })

    return StreamingResponse(
        io.BytesIO(marked_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
    )


@app.post("/check_text")
@limiter.limit("60/minute")
async def check_text(
    request: Request,
    body: MarkTextRequest,
    user: dict = Depends(require_api_product("mark", "revise")),
):
    """
    Analyse text and return JSON feedback (issues, examples, label counts).
    Same marking pipeline as /mark_text but returns structured JSON instead
    of a .docx binary — designed for the live Write page.
    """
    _is_api_client = user.get("_is_api_client", False) if isinstance(user, dict) else False
    _api_start_time = time.time()

    # 0. Product-level access check based on calling context
    await _enforce_product_for_mode(user, body.student_mode)

    # 0a. Free-tier usage check — skip for API key clients and student rechecks.
    #     Students can recheck freely; their paywall is on download only.
    if not _is_api_client and not body.student_mode:
        _ct_user_id = user.get("id") if isinstance(user, dict) else None
        if _ct_user_id and _ct_user_id != "local-dev":
            _ct_profile = await get_user_profile(_ct_user_id)
            _ct_tier = (_ct_profile or {}).get("subscription_tier", "free")
            if _ct_tier == "free":
                _ct_used = await count_user_marks(_ct_user_id)
                if _ct_used >= _FREE_TIER_MARK_LIMIT:
                    raise HTTPException(
                        status_code=402,
                        detail={
                            "message": "Subscribe for unlimited text checking.",
                            "code": "USAGE_LIMIT",
                        },
                    )

    # 0b. Text length limit (50,000 chars ≈ 8,000 words)
    _MAX_TEXT_CHARS = 50_000
    if body.text and len(body.text) > _MAX_TEXT_CHARS:
        raise HTTPException(status_code=400, detail=f"Text exceeds {_MAX_TEXT_CHARS} character limit.")

    # 1. Create .docx from text
    docx_bytes = build_doc_from_text(body.text)

    # 2. Build teacher_config from body.titles
    teacher_config = build_teacher_config_from_titles(body.titles) or {}
    teacher_config["student_mode"] = body.student_mode

    # 3. Call mark_docx_bytes (same pipeline as /mark and /mark_text)
    mark_docx_bytes_fn, _ = get_engine()
    mode = body.mode or "textual_analysis"
    _marked_bytes, metadata = mark_docx_bytes_fn(
        docx_bytes,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
        include_summary_table=False,
    )

    # 4. Extract issues, examples, detected_lexis from metadata
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []
    detected_lexis = metadata.get("detected_lexis", []) if isinstance(metadata, dict) else []
    techniques_discussed = metadata.get("techniques_discussed", []) if isinstance(metadata, dict) else []
    sentence_types = metadata.get("sentence_types", {}) if isinstance(metadata, dict) else {}
    first_sentence_components = metadata.get("first_sentence_components", {}) if isinstance(metadata, dict) else {}

    # 5. Count labels
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl or lbl.startswith("__"):
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        label_counter[lbl] += (cnt_i if cnt_i > 0 else 1)

    total_labels = sum(label_counter.values())

    # 6. Word count
    cleaned = (body.text or "").strip()
    word_count = len(cleaned.split()) if cleaned else 0

    # 6b. Compute scores before mark_events insert so they can be persisted
    scores = None
    try:
        scores = _compute_scores(
            body.text,
            mode=mode,
            label_counts=dict(label_counter),
            mark_event_id=None,
            sentence_types={str(k): v for k, v in sentence_types.items()},
            repeated_nouns=metadata.get("repeated_nouns", []) if isinstance(metadata, dict) else [],
        )
    except Exception:
        pass

    # 7. Log to Supabase mark_events (best-effort) — skip for API key clients
    mark_event_id = None
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and not _is_api_client:
            user_id = user.get("id") if isinstance(user, dict) else None
            db_url = f"{SUPABASE_URL}/rest/v1/mark_events?select=id"
            payload = {
                "user_id": user_id,
                "file_name": body.file_name or "write_session",
                "mode": mode,
                "bytes": len(docx_bytes),
                "total_labels": total_labels,
                "label_counts": dict(label_counter),
                "issues": issues,
                "review_status": "pending",
                "word_count": word_count,
                "scores": _sanitize_for_json(scores) if scores else None,
            }
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.post(
                    db_url,
                    json=payload,
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=representation",
                    },
                )
                if 200 <= resp.status_code < 300:
                    resp_data = resp.json()
                    if resp_data and isinstance(resp_data, list) and len(resp_data) > 0:
                        mark_event_id = resp_data[0].get("id")
    except Exception as e:
        print("Failed to log mark_event (check_text):", repr(e))

    # 8. Log examples to Supabase issue_examples (best-effort) — skip for API clients
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and examples and not _is_api_client:
            user_id = user.get("id") if isinstance(user, dict) else None
            if user_id:
                example_rows = []
                for ex in examples:
                    if not isinstance(ex, dict):
                        continue
                    label = ex.get("label")
                    sentence = ex.get("sentence")
                    if not label or not sentence:
                        continue
                    example_rows.append({
                        "user_id": user_id,
                        "file_name": body.file_name or "write_session",
                        "mode": mode,
                        "label": label,
                        "sentence": sentence,
                        "paragraph_index": ex.get("paragraph_index"),
                        "mark_event_id": mark_event_id,
                        "found_value": ex.get("found_value"),
                        "topics": ex.get("topics"),
                        "thesis": ex.get("thesis"),
                        "confidence": ex.get("confidence"),
                        "original_phrase": ex.get("original_phrase"),
                    })
                if example_rows:
                    db_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                    async with httpx.AsyncClient(timeout=5) as client:
                        await client.post(
                            db_url,
                            json=example_rows,
                            headers={
                                "apikey": SUPABASE_SERVICE_KEY,
                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                "Content-Type": "application/json",
                                "Prefer": "return=minimal",
                            },
                        )
    except Exception as e:
        print("Failed to log issue_examples (check_text):", repr(e))

    # 9. Log API key usage (best-effort)
    if _is_api_client:
        _api_elapsed = int((time.time() - _api_start_time) * 1000)
        await _log_api_usage(
            api_key_id=user.get("_api_key_id", ""),
            endpoint="/check_text",
            status_code=200,
            chars_processed=len(body.text or ""),
            response_ms=_api_elapsed,
            client_ip=get_remote_address(request),
            metadata={"mode": mode, "total_labels": total_labels, "word_count": word_count},
        )

    # 10. Return JSON response (strip proprietary fields)
    _response_data = {
        "issues": _strip_ip_from_issues(issues),
        "examples": _strip_ip_from_examples(examples),
        "label_counts": dict(label_counter),
        "detected_lexis": detected_lexis,
        "techniques_discussed": techniques_discussed,
        "total_labels": total_labels,
        "word_count": word_count,
        "sentence_types": {str(k): v for k, v in sentence_types.items()},
        "first_sentence_components": first_sentence_components,
        "scores": scores,
    }
    # For regular users, include mark_event_id; strip it for API clients
    if not _is_api_client:
        _response_data["mark_event_id"] = mark_event_id
    return JSONResponse(_sanitize_for_json(_response_data))
