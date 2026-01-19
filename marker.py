#!/usr/bin/env python3
# ============================================================
#   VYSTI MARKER — CLEAN ENGINE
# ============================================================

import os
import sys
import re
import tempfile
import hashlib
from io import BytesIO
import pandas as pd
import docx  # type: ignore
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn  # type: ignore[attr-defined]
from docx.oxml import OxmlElement  # type: ignore[attr-defined]
from docx.opc.constants import RELATIONSHIP_TYPE
import spacy

# Custom logical color for grammar issues (implemented via shading)
GRAMMAR_ORANGE = "GRAMMAR_ORANGE"

try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    # Fallback: download the model at runtime if it's not installed
    from spacy.cli import download
    print("Downloading spaCy model 'en_core_web_sm'...")
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

from dataclasses import dataclass
from difflib import SequenceMatcher
from typing import Dict, Tuple, NamedTuple


# ============================================================
# FOUNDATION ASSIGNMENT 1 — GLOBAL LABEL TRACKING
# ============================================================
# For Foundation Assignment 1, we track the position of the last piece of
# extra content (extra sentence or extra paragraph) so we can attach a single
# yellow label at the very end of all red-struck content.
# Format: (paragraph_index, char_position_in_flat_text) or None
FOUNDATION1_LABEL_TARGET = None


def normalize_title_key(s: str) -> str:
    """
    Normalize a title or quoted string so we can match titles
    even if students add extra punctuation or use different capitalization.
    Removes all punctuation, collapses whitespace, and lowercases.
    """
    s = s.strip()
    # Remove all punctuation characters (keep only letters, digits, and spaces)
    s = re.sub(r"[^\w\s]", "", s)
    s = " ".join(s.split())
    return s.lower()


def title_similarity(a: str, b: str) -> float:
    """
    Fuzzy similarity between two title-like strings in [0, 1].
    We:
      - normalize both with normalize_title_key (strip punctuation/case)
      - compute SequenceMatcher ratio on the normalized forms
    """
    ak = normalize_title_key(a)
    bk = normalize_title_key(b)
    if not ak or not bk:
        return 0.0
    return SequenceMatcher(None, ak, bk).ratio()


def is_title_case_like(snippet: str) -> bool:
    """
    Heuristic: return True if `snippet` looks like a title.

    We treat something as title-like if:
      - it has at least a few words
      - at least ~half of those words start with an uppercase letter
    This lets us catch Title Case like
        "Facebook Multiplies Genders But Offers Users..."
    while ignoring ordinary clauses like
        "Facebook offered 58 genders but only allowed users..."
    """
    words = re.findall(r"[A-Za-z]+", snippet)
    if len(words) < 3:
        return False

    caps = sum(1 for w in words if w[0].isupper())
    if caps < 2:
        return False

    ratio = caps / len(words)
    return ratio >= 0.5


def get_config_title_keys(config) -> set[str]:
    """
    Return a set containing the normalized title keys from all teacher-supplied works,
    or an empty set if no titles are configured.
    """
    keys: set[str] = set()
    for work in iter_teacher_works(config):
        key = normalize_title_key(work.title)
        if key:
            keys.add(key)
    return keys


def normalize_title_for_exact_match(s: str) -> str:
    """
    Normalize a title or fragment just enough to compare the teacher-supplied
    text with the student's text "as written", while still smoothing out
    Word-specific quirks (curly quotes, non-breaking spaces, long dashes).
    This does NOT strip ordinary punctuation or change case.
    """
    if not s:
        return ""
    out = s.replace("\u00A0", " ")

    quote_map = {
        # Double quotes (curly variants) → "
        "\u201C": '"',  # LEFT DOUBLE QUOTATION MARK
        "\u201D": '"',  # RIGHT DOUBLE QUOTATION MARK
        "\u201E": '"',  # DOUBLE LOW-9 QUOTATION MARK
        "\u201F": '"',  # DOUBLE HIGH-REVERSED-9 QUOTATION MARK

        # Single quotes / apostrophes (curly variants) → '
        "\u2018": "'",  # LEFT SINGLE QUOTATION MARK / opening apostrophe
        "\u2019": "'",  # RIGHT SINGLE QUOTATION MARK / apostrophe
        "\u201A": "'",  # SINGLE LOW-9 QUOTATION MARK
        "\u201B": "'",  # SINGLE HIGH-REVERSED-9 QUOTATION MARK

        # Angle quote characters
        "\u00AB": '"',  # LEFT-POINTING DOUBLE ANGLE QUOTATION MARK
        "\u00BB": '"',  # RIGHT-POINTING DOUBLE ANGLE QUOTATION MARK
        "\u2039": "'",  # SINGLE LEFT-POINTING ANGLE QUOTATION MARK
        "\u203A": "'",  # SINGLE RIGHT-POINTING ANGLE QUOTATION MARK
    }
    for bad, good in quote_map.items():
        out = out.replace(bad, good)

    # Normalize dashes
    out = out.replace("—", "-").replace("–", "-")
    return out.strip()


def normalize_author_key(s: str) -> str:
    """
    Normalize an author name so we can match the teacher-supplied author_name
    against the student's first sentence even when punctuation and spacing
    differ.

    Rules:
      - normalize curly quotes, NBSP, and dashes via normalize_title_for_exact_match
      - strip a possessive "'s" (Fisher's -> Fisher)
      - remove all non-letter characters (periods between initials, commas, digits)
        but keep spaces
      - collapse whitespace and lowercase

    This makes "M. F. K. Fisher", "M.F.K. Fisher", and "M.F.K. Fisher's"
    all normalize to the same key while still requiring all name parts
    to be present.
    """
    if not s:
        return ""
    # Normalize quotes, NBSP, dashes, etc.
    s = normalize_title_for_exact_match(s)
    # Treat possessive 's as part of the base name: Fisher's -> Fisher
    s = re.sub(r"'s\b", "", s)
    # Keep only letters and spaces (periods, commas, etc. → spaces)
    s = re.sub(r"[^A-Za-z\s]", " ", s)
    # Collapse multiple spaces and lowercase
    s = " ".join(s.split())
    return s.lower()


def author_full_name_present(author_name: str, first_sentence_text: str) -> bool:
    """
    Return True if the first sentence appears to contain the *full* author name
    (all given names / initials + last name), allowing for differences in
    spacing, periods, and a trailing possessive 's.

    Examples that should count as present:
        author_name = "M. F. K. Fisher"
        first sentence: "M. F. K. Fisher's self-reflective essay..."
        first sentence: "In her essay 'Young Hunger,' M. F. K. Fisher describes..."

    Examples that should NOT count as present:
        first sentence: "Fisher's essay..."
        first sentence: "M. Fisher's essay..."
    """
    import re

    # Extract alphabetic tokens from the teacher name
    author_words = re.findall(r"[A-Za-z]+", author_name or "")
    if not author_words:
        return False

    # Last token is treated as the surname; everything before it are given names / initials
    last = author_words[-1].lower()
    given = [w.lower() for w in author_words[:-1]]

    # Extract alphabetic tokens from the first sentence
    sent_words = re.findall(r"[A-Za-z]+", first_sentence_text or "")
    if not sent_words:
        return False

    # Find all positions where the surname appears (or as the start of a word, e.g. Fisher/Fishers)
    candidate_idxs = [
        i for i, w in enumerate(sent_words)
        if w.lower().startswith(last)
    ]
    if not candidate_idxs:
        return False

    # For each surname occurrence, look backward for the given names / initials
    for idx in candidate_idxs:
        # Look back up to len(given) tokens before the surname
        window_start = max(0, idx - len(given))
        preceding = sent_words[window_start:idx]
        prec_low = [w.lower() for w in preceding]

        # Full name is present only if *all* given-name tokens appear before the surname
        if given and all(g in prec_low for g in given):
            return True
        # If there are no given names (single-word author), just seeing the surname is enough
        if not given:
            return True

    return False


def find_title_span_in_first_sentence(flat_text: str, sentences, config) -> tuple[tuple[int, int] | None, bool]:
    """
    Try to locate the teacher-supplied title inside the FIRST sentence of the
    paragraph.

    Returns:
        (span, is_exact)

        span     -> (start, end) character offsets into flat_text, or None
        is_exact -> True  if the student text matches the teacher input
                           (after normalizing curly quotes/dashes),
                    False if it is only a fuzzy match.
    """
    if not config or not getattr(config, "text_title", None) or not sentences:
        return None, False

    # Target for fuzzy matching (ignore punctuation/case)
    norm_target = normalize_title_key(config.text_title)
    if not norm_target:
        return None, False

    # Target for "exact" matching (preserve punctuation/case, but normalize quotes/dashes)
    exact_target = normalize_title_for_exact_match(config.text_title)

    first_start, first_end = sentences[0]
    sentence_text = flat_text[first_start:first_end]
    n = len(sentence_text)

    best_span: tuple[int, int] | None = None
    best_sim: float = 0.0
    best_exact = False

    for i in range(n):
        for j in range(i + 1, n + 1):
            snippet = sentence_text[i:j]

            # Cheap reject: must contain at least one letter
            if not any(ch.isalpha() for ch in snippet):
                continue

            # Normalize snippet for fuzzy matching
            norm_snip = normalize_title_key(snippet)
            if not norm_snip:
                continue

            # NEW: if the normalized snippet equals the normalized teacher title,
            # accept it even if it's only 1–2 words.
            if norm_snip != norm_target:
                # Only fuzzy candidates need to "look like a title"
                # This filters out ordinary clauses like
                # "Facebook offered 58 genders but only allowed users..."
                # while keeping real/attempted titles like
                # "Facebook's New Gender Options: Multiple Choices, Same Three Tired Pronouns"
                if not is_title_case_like(snippet):
                    continue

            # Ignore substrings that are far shorter than the full title;
            # this prevents short evidence like "three tired pronouns" from
            # being mistaken for the title.
            if len(norm_snip) < max(5, int(0.5 * len(norm_target))):
                continue

            # Exact fuzzy-normalized match
            if norm_snip == norm_target:
                sim = 1.0
            else:
                # Real fuzzy: allow small word changes ("Fenders" vs "Genders",
                # or Ham's alternate title)
                sim = title_similarity(snippet, config.text_title)

            # Require reasonably high similarity so we don't confuse ordinary
            # quotations with the title.
            if sim < 0.6:
                continue

            span = (first_start + i, first_start + j)

            # Check if this candidate is also an exact match "as written"
            is_exact_here = (normalize_title_for_exact_match(snippet) == exact_target)

            # If we ever find an exact match, take it immediately.
            if is_exact_here:
                return span, True

            # Otherwise keep the best fuzzy match so far.
            if sim > best_sim:
                best_sim = sim
                best_span = span
                best_exact = False

    if best_span is not None:
        return best_span, best_exact

    return None, False


# Bookmark-related constants and helpers for internal hyperlinks
BOOKMARK_PREFIX = "vysti_issue_"

BOOKMARK_MAX_LEN = 40  # Word's limit for bookmark names

ARTICLE_ERROR_LABEL = "Article error"
ARTICLE_ERROR_EXPLANATION = "Use a before consonants and an before vowels."
ARTICLE_ERROR_GUIDANCE = "Swap the article so it matches the next word (a + consonant, an + vowel)."

# Global counter so bookmark IDs are unique in the document
BOOKMARK_ID_COUNTER = 1


def bookmark_name_for_label(label: str) -> str:
    """
    Turn an issue label into a safe Word bookmark name.

    Word bookmark names:
      - must start with a letter
      - cannot contain spaces or most punctuation
      - must be <= 40 characters total
    """
    base = re.sub(r"\W+", "_", label).strip("_")

    if not base:
        base = "generic"

    available = BOOKMARK_MAX_LEN - len(BOOKMARK_PREFIX)
    if available < 1:
        available = 1

    if len(base) > available:
        base = base[:available]

    return BOOKMARK_PREFIX + base


def add_bookmark_to_paragraph(paragraph, bookmark_name: str):
    """
    Wrap the given paragraph in a Word bookmark so we can hyperlink to it.
    We don't care about the exact span; starting at the paragraph start is fine.
    """
    global BOOKMARK_ID_COUNTER

    p = paragraph._p
    # Create bookmarkStart
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(BOOKMARK_ID_COUNTER))
    start.set(qn("w:name"), bookmark_name)

    # Create bookmarkEnd
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(BOOKMARK_ID_COUNTER))

    BOOKMARK_ID_COUNTER += 1

    # Insert at beginning & end of paragraph
    # Put bookmarkStart before the first paragraph child
    p.insert(0, start)
    # Append bookmarkEnd at the end
    p.append(end)


def wrap_run_in_internal_link(paragraph, run, bookmark_name: str):
    """
    Take an existing Run object, remove it from the paragraph, and wrap it in
    a Word internal hyperlink pointing at the given bookmark.
    """
    p = paragraph._p
    r_el = run._element

    # Remove the run from the paragraph's children
    p.remove(r_el)

    # Build <w:hyperlink w:anchor="bookmark_name">
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    # history="1" tells Word to treat it like a visited link in nav history
    hyperlink.set(qn("w:history"), "1")

    # Attach the run inside the hyperlink
    hyperlink.append(r_el)

    # Append the hyperlink element back to the paragraph
    p.append(hyperlink)


def wrap_run_in_external_link(paragraph, run, url: str):
    """
    Turn an existing run into an external hyperlink pointing at `url`.
    The run keeps its font formatting; it just becomes clickable.
    """
    p = paragraph._p
    r_el = run._element

    # Remove the run from the paragraph
    p.remove(r_el)

    # Build <w:hyperlink r:id="...">
    hyperlink = OxmlElement("w:hyperlink")

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    # Put the original run inside the hyperlink
    hyperlink.append(r_el)

    # Append the hyperlink back into the paragraph
    p.append(hyperlink)


# Pattern for creative essay titles like:
#   "Stepping stones of every creative mind": Topic in "The Nature of Scientific Reasoning"
# or the looser form:
#   "Stepping stones of every creative mind": Topic "The Nature of Scientific Reasoning"
#
# We only require:
#   - a leading quoted phrase (student's quotation)
#   - a colon and some topic text
#   - a second quoted phrase (the known source title)
# The connector word "in" is now optional and just treated as part of the topic text.
TITLE_PATTERN = re.compile(
    r'^\s*"([^"]+)"\s*:\s*(.+?)\s+"([^"]+)"\s*\.?\s*$',
    re.IGNORECASE,
)

# Pattern for creative essay titles that omit the colon, like:
#   "I did not have the sense to explain to them how starved I was" Desire of "Young Hunger"
#
# This matches the same structure as TITLE_PATTERN but without requiring a colon
# between the first quotation and the topic segment. The final quoted segment
# should still be a known work title from KNOWN_TITLES_MINOR or KNOWN_TITLES_MAJOR.
TITLE_PATTERN_NO_COLON = re.compile(
    r'^\s*"([^"]+)"\s+(.+?)\s+"([^"]+)"\s*\.?\s*$',
    re.IGNORECASE,
)
# Prefix patterns for creative titles at the *start* of a paragraph,
# even when students continue the intro in the same line.
TITLE_PREFIX_PATTERN = re.compile(
    r'^\s*"([^"]+)"\s*:\s*(.+?)\s+"([^"]+)"\s*\.?\s*',
    re.IGNORECASE,
)

TITLE_PREFIX_NO_COLON_PATTERN = re.compile(
    r'^\s*"([^"]+)"\s+(.+?)\s+"([^"]+)"\s*\.?\s*',
    re.IGNORECASE,
)

# Device/strategy keywords to detect in thesis statements.
# This set is deliberately lemma-based: we check token.lemma_.lower() against it,
# so "metaphors", "metaphorical", etc. still count as a "metaphor" device.
THESIS_DEVICE_WORDS: set[str] = set()

# Synonyms for thesis device words to improve semantic alignment
# Maps alternative terms students might use to the canonical device names
THESIS_DEVICE_SYNONYMS: dict[str, str] = {}

# Multi-word thesis device synonyms
# Maps tuples of token strings to canonical device names
# Example: ("rhetorical", "question") -> "rhetorical question"
THESIS_MULTIWORD_SYNONYMS: Dict[Tuple[str, ...], str] = {}

# Maximum length of multi-word phrases we need to check
THESIS_MULTIWORD_MAX_LEN: int = 1


@dataclass
class MarkerConfig:
    """
    Configuration switches for the Vysti marker.

    mode:
        'textual_analysis' (default)
        'intertextual_analysis' (supports up to three works)
        'reader_response'
        'no_title'
        'image_analysis'
        'argumentation'
        'foundation_1' ... 'foundation_6'
        (others can be added later)

    author_name / text_title:
        Provided by the teacher via GUI later. For now they are
        stored but not yet used, just wired through.
    """
    mode: str = "textual_analysis"

    # Teacher-provided metadata
    author_name: str | None = None
    text_title: str | None = None
    text_is_minor_work: bool = True  # True = quoted, False = italic (for future use)

    # NEW: up to two additional works for intertextual analysis
    author_name_2: str | None = None
    text_title_2: str | None = None
    text_is_minor_work_2: bool | None = None

    author_name_3: str | None = None
    text_title_3: str | None = None
    text_is_minor_work_3: bool | None = None

    # Rule toggles (default = current behavior)
    enforce_closed_thesis: bool = True
    enforce_specific_thesis_topics: bool = True
    enforce_thesis_organization: bool = True  # "Organization of thesis statement"
    enforce_topic_thesis_alignment: bool = True  # "Follow the organization..." etc.
    enforce_off_topic: bool = True

    # Student mode toggle (default = teacher mode)
    student_mode: bool = False

    require_body_evidence: bool = True  # "Every paragraph needs evidence"
# Controls whether the generic contractions rule is enforced

    enforce_contractions_rule: bool = True
    enforce_long_quote_rule: bool = True
    # Controls whether we enforce the "Avoid the word 'which'" rule
    enforce_which_rule: bool = True
    enforce_weak_verbs_rule: bool = True
    enforce_fact_proof_rule: bool = True
    enforce_human_people_rule: bool = True
    # Controls whether we enforce the "society / universe / reality / life / truth" vague-terms rule
    enforce_vague_terms_rule: bool = True
    enforce_sva_rule: bool = False
    enforce_present_tense_rule: bool = False


    forbid_audience_reference: bool = True  # "Avoid referring to the reader or audience..."
    forbid_personal_pronouns: bool = True  # "No 'I', 'we', 'us', 'our' or 'you'..."

    enforce_essay_title_format: bool = True  # "Essay title format"
    enforce_essay_title_capitalization: bool = True  # "Capitalize the words in titles"

    # Intro quotation behavior
    # If True, allow direct quotations in *introductory summary* sentences
    # (between the first sentence and the thesis), but still forbid them
    # in the first sentence and in the thesis statement.
    allow_intro_summary_quotes: bool = False

    # If False, we do NOT use the generic "Avoid quotations in the introduction"
    # rule at all. We still use "No quotations in thesis statements".
    enforce_intro_quote_rule: bool = True
        # Controls BRIGHT_GREEN highlighting of thesis devices/rhetorical strategies
    highlight_thesis_devices: bool = True



class TeacherWork(NamedTuple):
    """Represents a teacher-supplied work (author + title + minor/major flag)."""
    author: str | None
    title: str
    is_minor: bool


def iter_teacher_works(config: MarkerConfig | None) -> list[TeacherWork]:
    """
    Return all teacher-supplied works (author + title + minor/major) as
    a flat list. Always includes the legacy single work if present.
    """
    if not config:
        return []

    works: list[TeacherWork] = []

    def add(author_field: str | None, title_field: str | None, minor_flag: bool | None):
        if not title_field:
            return
        is_minor = config.text_is_minor_work if minor_flag is None else minor_flag
        works.append(TeacherWork(author_field, title_field, is_minor))

    add(config.author_name, config.text_title, config.text_is_minor_work)
    add(config.author_name_2, config.text_title_2, config.text_is_minor_work_2)
    add(config.author_name_3, config.text_title_3, config.text_is_minor_work_3)

    # Deduplicate by normalized title in case the same work is entered twice
    seen: set[str] = set()
    unique: list[TeacherWork] = []
    for w in works:
        key = normalize_title_key(w.title)
        if key and key not in seen:
            seen.add(key)
            unique.append(w)
    return unique


def iter_author_names(config: MarkerConfig | None) -> list[str]:
    """
    Return all configured author names as a list of normalized strings.
    """
    if not config:
        return []
    out: list[str] = []
    for name in (config.author_name, config.author_name_2, config.author_name_3):
        if name:
            norm = normalize_title_key(name)
            if norm:
                out.append(norm)
    return out


def get_preset_config(mode: str = "textual_analysis") -> MarkerConfig:
    """
    Return a MarkerConfig preconfigured for a given assignment mode.

    Modes:

    - textual_analysis: full rule set (current default behavior)
    - intertextual_analysis: same as textual_analysis, but supports up to three works
    - reader_response: allow I/you/reader; keep most academic rules
    - no_title: students are not required to have an essay title
    - image_analysis: no quotation/evidence requirement per paragraph
    - argumentation: disable device-based closed-thesis structure checks
    - foundation_1: Foundation Assignment 1 – first sentence only
    - foundation_2: Foundation Assignment 2 – first sentence + closed thesis
    - foundation_3: Foundation Assignment 3 – full introduction
    - foundation_4: Foundation Assignment 4 – intro + first body topic sentence
    - foundation_5: Foundation Assignment 5 – intro + full body paragraphs
    - foundation_6: Foundation Assignment 6 – full essay
    """
    cfg = MarkerConfig(mode=mode)

    if mode == "reader_response":
        # Allow I/you/reader/audience; still formal in other ways
        cfg.forbid_audience_reference = False
        cfg.forbid_personal_pronouns = False
        cfg.enforce_contractions_rule = False
    elif mode == "no_title":
        # No essay title required; source-title rules still apply
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False

    elif mode == "image_analysis":
        # No quotation-based evidence requirement
        cfg.require_body_evidence = False

    elif mode == "argumentation":
        # Argument essays may use open thesis and non-device topics.
        # Disable device-based closed-thesis rules and alignment.
        cfg.enforce_closed_thesis = False
        cfg.enforce_specific_thesis_topics = False
        cfg.enforce_thesis_organization = False
        cfg.enforce_topic_thesis_alignment = False
        cfg.require_body_evidence = False
        cfg.enforce_off_topic = False
        # Allow direct address to the reader/audience in argumentation
        # ("Avoid referring to the reader or audience..." is off),
        # but keep the ban on first-person pronouns by default.
        cfg.forbid_audience_reference = False
        cfg.highlight_thesis_devices = False


    elif mode == "analytic_frame":
        # Analytic frame essays:
        # - Use the full textual-analysis rule set
        # - Support multiple works via author_name_2 / text_title_2 / author_name_3 / text_title_3
        # - Allow direct quotations in the *introductory summary* (between first sentence and thesis)
        # - Turn OFF the generic "Avoid quotations in the introduction" rule
        #   (but still forbid quotations in the thesis sentence itself).
        cfg.allow_intro_summary_quotes = True
        cfg.enforce_intro_quote_rule = False


    # ---------- Foundation modes ----------

    elif mode == "foundation_1":
        # ============================================================
        # Foundation Assignment 1: First Sentence Only
        # ============================================================
        # Students write ONLY the first sentence of an analytical essay.
        # This sentence must include:
        #   - Author's full name
        #   - Text title (properly formatted)
        #   - Genre
        #   - Brief summary
        #
        # Special behavior:
        #   - The first sentence is marked normally with all first-sentence rules
        #   - Any extra sentences in the intro paragraph → red highlight + strikethrough
        #   - Any extra paragraphs (body, conclusion) → red highlight + strikethrough
        #   - A SINGLE yellow label appears at the END of all extra content:
        #     "The assignment is to write the first sentence"
        #   - Extra content gets no other issue labels (only the assignment violation)
        #
        # Implementation: See analyze_text (extra sentences in intro), run_marker
        # (extra paragraphs), and label insertion code at end of run_marker.
        #
        # Disable all rules that require multi-paragraph structure:
        cfg.enforce_closed_thesis = False
        cfg.enforce_specific_thesis_topics = False
        cfg.enforce_thesis_organization = False
        cfg.enforce_topic_thesis_alignment = False
        cfg.enforce_off_topic = False
        cfg.require_body_evidence = False
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False

    elif mode in ("foundation_2", "foundation_3"):
        # Foundation 2: first sentence + closed thesis.
        # Foundation 3: full introduction (first sentence, summary, thesis).
        # Thesis rules stay ON, but we do not require body paragraphs or essay titles yet.
        cfg.require_body_evidence = False
        cfg.enforce_topic_thesis_alignment = False
        cfg.enforce_off_topic = False
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False

    elif mode == "foundation_4":
        # Foundation 4: introduction + 1st body paragraph topic sentence.
        # Keep thesis and alignment rules, but do NOT require evidence yet.
        cfg.require_body_evidence = False
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False

    elif mode == "foundation_5":
        # Foundation 5: introduction + full body paragraphs (no conclusion).
        # Full thesis + body behavior, but still do NOT require a creative essay title.
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False

    # elif mode == "foundation_6":
        # Foundation 6: full essay, like textual_analysis,
        # plus experimental grammar checks (subject–verb agreement).
        # cfg.enforce_sva_rule = True
        # cfg.enforce_present_tense_rule = True

    elif mode == "peel_paragraph":
        # Single PEEL paragraph: Point–Evidence–Explanation–Link.
        # We want body-paragraph evidence rules, but we *do not* want
        # multi-paragraph thesis organization / off-topic logic or essay-title requirements.
        cfg.enforce_thesis_organization = False
        cfg.enforce_topic_thesis_alignment = False
        cfg.enforce_off_topic = False
        cfg.enforce_essay_title_format = False
        cfg.enforce_essay_title_capitalization = False
        # Keep:
        #   - enforce_closed_thesis = True  (we'll use it for a simple "no questions" check)
        #   - enforce_specific_thesis_topics = True (harmless; we won't run the big thesis logic)
        #   - require_body_evidence = True (PEEL needs Evidence in the middle)

    # textual_analysis and any unknown mode fall back to defaults (all True)
    return cfg


def load_thesis_devices(path: str = "thesis_devices.txt") -> None:
    """
    Load thesis device words and synonyms from a text file.
    
    The file format is:
    - Lines starting with # are comments
    - Every other non-empty line has the format: term,canonical_device
    - If term == canonical_device, it declares a canonical device
    - If term != canonical_device, it's a synonym for the canonical device
    """
    global THESIS_MULTIWORD_MAX_LEN
    
    # Get the directory where marker.py is located
    file_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(file_dir, path)
    
    if not os.path.exists(file_path):
        raise RuntimeError(f"thesis_devices.txt not found at {file_path}")
    
    with open(file_path, encoding="utf-8") as f:
        for line in f:
            # Skip empty lines and comments
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            
            # Split on first comma
            if "," not in line:
                continue
            
            term, canonical_device = line.split(",", 1)
            term = term.strip().lower()
            canonical_device = canonical_device.strip().lower()
            
            # Split term on whitespace to check for multi-word terms
            tokens = term.split()
            
            # Always add canonical device to THESIS_DEVICE_WORDS (existing behavior)
            THESIS_DEVICE_WORDS.add(canonical_device)
            
            if len(tokens) == 1:
                # Single-word term: keep existing synonym behavior
                if term != canonical_device:
                    THESIS_DEVICE_SYNONYMS[term] = canonical_device
            else:
                # Multi-word term (len(tokens) > 1): add to multi-word synonyms
                # This handles both canonical multi-word devices (e.g., "rhetorical question")
                # and multi-word synonyms (e.g., "rhetorical questions" -> "rhetorical question")
                key = tuple(tokens)
                THESIS_MULTIWORD_SYNONYMS[key] = canonical_device
                THESIS_MULTIWORD_MAX_LEN = max(THESIS_MULTIWORD_MAX_LEN, len(tokens))
            
            # Also handle the canonical_device if it's multi-word (regardless of term length)
            # This ensures canonical multi-word devices are always in the multi-word map
            canonical_tokens = canonical_device.split()
            if len(canonical_tokens) > 1:
                canonical_key = tuple(canonical_tokens)
                THESIS_MULTIWORD_SYNONYMS[canonical_key] = canonical_device
                THESIS_MULTIWORD_MAX_LEN = max(THESIS_MULTIWORD_MAX_LEN, len(canonical_tokens))


# Load devices at module import time
load_thesis_devices()


def canonical_device_key(tok):
    """
    Return the canonical device/strategy key for this token, or None
    if the token is not one of our known devices (including synonyms).
    
    This function uses lemma-first matching to ensure that inflected forms
    (e.g., "contrasts", "contrasted", "contrasting") map to the same canonical
    key (e.g., "contrast") as their base form, even if the inflected forms
    are not explicitly listed in thesis_devices.txt.
    
    Also supports multi-word devices (e.g., "rhetorical question") by checking
    phrases starting at the current token before falling back to single-token matching.
    """
    # Check multi-word phrases first (before single-token logic)
    # This allows detection of phrases like "rhetorical question" starting at "rhetorical"
    doc = tok.doc
    i = tok.i
    
    # Iterate from max length down to 2 (minimum multi-word phrase)
    for length in range(THESIS_MULTIWORD_MAX_LEN, 1, -1):
        # Guard against running past the end of the doc
        if i + length > len(doc):
            continue
        
        # Build a tuple of lowercased text from doc[i:i+length]
        phrase_tokens = tuple(doc[j].text.lower() for j in range(i, i + length))
        
        # Check if this phrase exists in our multi-word synonyms
        if phrase_tokens in THESIS_MULTIWORD_SYNONYMS:
            return THESIS_MULTIWORD_SYNONYMS[phrase_tokens]
    
    # Fall back to existing single-token matching logic
    lemma = tok.lemma_.lower()
    lower = tok.text.lower()

    # Lemma-first: this ensures "contrasts" -> "contrast" via lemma mapping
    if lemma in THESIS_DEVICE_WORDS:
        return lemma
    # Fallback to exact word form match
    if lower in THESIS_DEVICE_WORDS:
        return lower
    # Check synonyms (exact word form)
    if lower in THESIS_DEVICE_SYNONYMS:
        return THESIS_DEVICE_SYNONYMS[lower]
    # Also check if lemma maps to a synonym's canonical form
    if lemma in THESIS_DEVICE_SYNONYMS:
        return THESIS_DEVICE_SYNONYMS[lemma]
    return None


def iter_device_spans(doc, start_char: int | None = None, end_char: int | None = None):
    """
    Yield (canonical_device_key, span_start, span_end) for each thesis device
    in the given character range of `doc`.

    - Uses THESIS_MULTIWORD_SYNONYMS and THESIS_MULTIWORD_MAX_LEN to detect
      multi-word devices like 'false dilemma' and 'rhetorical question' as
      single units.

    - Falls back to single-token detection via canonical_device_key for
      simple devices.

    - `start_char` / `end_char` are inclusive/exclusive bounds in doc text;
      if None, use the full doc.
    """
    text = doc.text
    n = len(doc)
    if start_char is None:
        start_char = 0
    if end_char is None:
        end_char = len(text)

    i = 0
    while i < n:
        tok = doc[i]
        # Skip tokens outside the requested char-range
        if tok.idx >= end_char:
            break
        if tok.idx + len(tok.text) <= start_char:
            i += 1
            continue

        # Try multi-word match first, from longest to shortest
        matched = False
        for length in range(THESIS_MULTIWORD_MAX_LEN, 1, -1):
            if i + length > n:
                continue
            phrase_tokens = doc[i:i+length]
            span_start_char = phrase_tokens[0].idx
            span_end_char = phrase_tokens[-1].idx + len(phrase_tokens[-1].text)
            # If phrase lies outside requested range, skip
            if span_end_char <= start_char or span_start_char >= end_char:
                continue
            phrase_key = tuple(t.text.lower() for t in phrase_tokens)
            if phrase_key in THESIS_MULTIWORD_SYNONYMS:
                canonical = THESIS_MULTIWORD_SYNONYMS[phrase_key]
                yield canonical, span_start_char, span_end_char
                i += length
                matched = True
                break

        if matched:
            continue

        # Single-token fallback
        canonical = canonical_device_key(tok)
        if canonical is not None:
            span_start_char = tok.idx
            span_end_char = tok.idx + len(tok.text)
            if span_end_char > start_char and span_start_char < end_char:
                yield canonical, span_start_char, span_end_char

        i += 1


THESIS_TOPIC_CONNECTOR_PHRASES = [
    ("and",),
    ("alongside",),
    ("with",),
    ("as", "well", "as"),
    ("in", "conjunction", "with"),
    ("in", "addition", "to"),
    ("in", "concert", "with"),
    ("in", "combination", "with"),
    ("along", "with"),
    ("together", "with"),
    ("in", "tandem", "with"),
    ("coupled", "with"),
    ("combined", "with"),
    ("plus",),
]
THESIS_TOPIC_CONNECTOR_MAX_LEN = max(len(phrase) for phrase in THESIS_TOPIC_CONNECTOR_PHRASES)


def _find_thesis_connector_span(thesis_tokens, index):
    n = len(thesis_tokens)
    if n == 0:
        return None
    start = max(0, index - THESIS_TOPIC_CONNECTOR_MAX_LEN + 1)
    for i in range(start, index + 1):
        for phrase in THESIS_TOPIC_CONNECTOR_PHRASES:
            length = len(phrase)
            if i + length > n:
                continue
            if all(thesis_tokens[i + k].text.lower() == phrase[k] for k in range(length)):
                if i <= index <= i + length - 1:
                    return (i, i + length - 1)
    return None


def _has_device_on_side(thesis_tokens, start_idx, step):
    i = start_idx
    while 0 <= i < len(thesis_tokens):
        tok = thesis_tokens[i]
        if tok.text in {",", ";"}:
            break
        if canonical_device_key(tok) is not None:
            return True
        i += step
    return False


def is_thesis_device_separator(thesis_tokens, index):
    span = _find_thesis_connector_span(thesis_tokens, index)
    if not span:
        return False
    left_has_device = _has_device_on_side(thesis_tokens, span[0] - 1, -1)
    if not left_has_device:
        return False
    right_has_device = _has_device_on_side(thesis_tokens, span[1] + 1, 1)
    return right_has_device


def is_embedded_device(thesis_tokens, i):
    """
    Heuristic: return True if the device token at index i is embedded
    inside the same phrase as an *earlier* device token, with no comma,
    semicolon, or device-connector between them.

    Examples that should be treated as ONE topic:
      - 'sexually charged diction on description of food'
        => 'description' is embedded in the 'diction' phrase.
      - 'metaphorical image'
        => 'image' is embedded in the 'metaphorical' device.

    Examples that should remain SEPARATE topics:
      - 'imagery of food and hyperbole about hunger'
      - 'comparison and contrast of X and Y'
    """
    # If this token is not a device at all, it's not "embedded".
    if canonical_device_key(thesis_tokens[i]) is None:
        return False

    j = i - 1
    while j >= 0:
        prev = thesis_tokens[j]
        txt = prev.text

        # Hard boundaries: if we hit punctuation or a device connector,
        # we are starting a new list item / phrase.
        if txt in {",", ";"} or is_thesis_device_separator(thesis_tokens, j):
            break

        prev_key = canonical_device_key(prev)
        if prev_key is not None:
            # We found an earlier device in the same comma/and chunk,
            # so this one is part of that device's clarifying phrase.
            return True

        j -= 1

    return False


def normalize_quote_text(s: str) -> str:
    """
    Normalize quotation content so repeated evidence like 'orgy', 'orgy.' or
    '"orgy"' are detected as the same. Used for the 'Only cite a quotation once'
    rule inside body paragraphs.
    """
    s = s.strip().lower()
    # Strip common punctuation and quote marks at both ends
    s = s.strip(".,!?;:\"'()[]{}")
    return s


# Subjective evaluation words we want to strip from intro/conclusion
SUBJECTIVE_EVAL_WORDS = {
    # Core ones Mark specified
    "successful",
    "successfully",
    "poignant",    
    "insightful",
    "thought-provoking",
    "insightfully",
    "sophisticated",
    "wonderful",
    "persuasive",
    # Other common empty evaluators
    "powerful",
    "effective",
    "interesting",
    "compelling",
    "moving",
    "beautiful",
    "brilliant",
    "amazing",
    "masterful",
    "great",
    "excellent",
}

THESIS_VERB_LEMMAS = {
    "argue", "argues", "arguing",
    "claim", "claims", "claiming",
    "suggest", "suggests", "suggesting",
    "show", "shows", "showing",
    "demonstrate", "demonstrates", "demonstrating",
    "reveal", "reveals", "revealing",
    "explore", "explores", "exploring",
    "emphasize", "emphasizes", "emphasizing",
    "illustrate", "illustrates", "illustrating",
    "highlight", "highlights", "highlighting",
    "contend", "contends", "contending",
    "assert", "asserts", "asserting",
    "imply", "implies", "implying",
    "maintain", "maintains", "maintaining",
    "propose", "proposes", "proposing",
    "present", "presents", "presenting",
    "explain", "explains", "explaining",
    "convey", "conveys", "conveying",
    "portray",  # e.g. "Bradbury successfully portrays..."
}

# Ordered device/strategy lemmas extracted from the thesis sentence
THESIS_DEVICE_SEQUENCE = []

# Will hold the ordered list of thesis topics (device/strategy lemmas)
THESIS_TOPIC_ORDER = []

# All device/strategy lemmas that appear anywhere in the thesis sentence
THESIS_ALL_DEVICE_KEYS = set()

# New: raw thesis paragraph text (lowercased) for simple substring checks
THESIS_TEXT_LOWER: str = ""

# New global state for body paragraph indexing and bridge detection
BODY_PARAGRAPH_COUNT = 0
BRIDGE_PARAGRAPHS = set()
BRIDGE_DEVICE_KEYS: dict[int, set[str]] = {}

# Foundation Assignment 4 — Track thesis location for assignment-completion rule
THESIS_PARAGRAPH_INDEX = None
THESIS_ANCHOR_POS = None

# Global state for collecting example sentences per label
DOC_EXAMPLES = []
DOC_EXAMPLE_COUNTS = {}  # dict: label -> int (count of examples stored for this label)
DOC_EXAMPLE_SENT_HASHES = set()  # set of tuples: (label, md5(sentence)) for deduplication
MAX_EXAMPLES_PER_LABEL = 10


def extract_thesis_topics(thesis_tokens):
    """
    Extract an ordered list of thesis 'topics' (devices/strategies) from the thesis sentence.

    We define a 'topic' as a device/strategy word from THESIS_DEVICE_WORDS that has at least
    one non-trivial clarifying content word in its local phrase (within the same comma/and chunk).

    Returns a list of canonical lemmas (strings) in order of appearance.
    """
    topics = []

    STOPLIKE = {
        "a", "an", "the",
        "of", "to", "for", "on", "in", "at", "by", "from", "with", "as",
        "that", "this", "these", "those",
        "his", "her", "their", "its",
        "use",
    }

    n = len(thesis_tokens)
    i = 0
    while i < n:
        tok = thesis_tokens[i]
        device_key = canonical_device_key(tok)
        if device_key is None:
            i += 1
            continue

        # Skip devices embedded in the phrase of an earlier device
        if is_embedded_device(thesis_tokens, i):
            i += 1
            continue

        clarifier_count = 0
        j = i + 1
        while j < n:
            nxt = thesis_tokens[j]
            txt = nxt.text
            lower2 = txt.lower()

            # Stop at list boundaries
            if txt in {",", ";"} or is_thesis_device_separator(thesis_tokens, j):
                break

            if any(ch.isalpha() for ch in txt) and lower2 not in STOPLIKE:
                clarifier_count += 1
            j += 1

        if clarifier_count > 0:
            topics.append(device_key)

        i += 1

    return topics


def compute_quote_spans(text: str):
    """
    Returns a list of (start, end) spans of all text inside quotes (exclusive of the quote marks).
    Handles nested quotes, mismatches, labels added after marking, and arbitrary run splitting.
    """
    spans = []
    in_quote = False
    start = None
    i = 0
    while i < len(text):
        ch = text[i]
        # Entering quote
        if ch in ['"', '"', '"'] and not in_quote:
            in_quote = True
            start = i + 1
        # Exiting quote
        elif ch in ['"', '"', '"'] and in_quote:
            end = i
            if start is not None and end > start:
                spans.append((start, end))
            in_quote = False
            start = None
        i += 1
    return spans


def pos_in_spans(pos, spans):
    """
    Check if a position is inside any of the given spans.
    Returns True if pos is within [start, end] (inclusive) of any span.
    """
    return any(start <= pos <= end for start, end in spans)


def compute_topic_sentence_span(flat_text: str, quote_spans: list) -> tuple[int, int]:
    """
    Compute the topic sentence span for a body paragraph using character offsets.
    
    The topic sentence is defined as:
    - Starting at the first non-space character of the paragraph
    - Ending at the first sentence-ending punctuation (. ? !) that is OUTSIDE any quote spans
    - If no such punctuation is found, fall back to the first spaCy sentence
    
    Args:
        flat_text: The flattened paragraph text
        quote_spans: List of (start, end) tuples for quote spans (exclusive of quote marks)
    
    Returns:
        (topic_start, topic_end) character offsets
    """
    # Find first non-space character
    topic_start = 0
    while topic_start < len(flat_text) and flat_text[topic_start].isspace():
        topic_start += 1
    
    if topic_start >= len(flat_text):
        # Empty paragraph
        return (0, len(flat_text))
    
    # Scan forward for first . ? ! outside quotes
    topic_end = None
    for i in range(topic_start, len(flat_text)):
        ch = flat_text[i]
        if ch in {'.', '?', '!'}:
            # Check if this punctuation is inside any quote span
            if not pos_in_spans(i, quote_spans):
                topic_end = i + 1  # Include the punctuation
                break
    
    # Fallback: if no sentence-ending punctuation found outside quotes,
    # use the first spaCy sentence as a safety net
    if topic_end is None:
        doc = nlp(flat_text)
        if doc.sents:
            first_sent = next(doc.sents)
            topic_end = first_sent.end_char
        else:
            topic_end = len(flat_text)
    
    return (topic_start, topic_end)


def get_paragraph_role(paragraph_index, intro_idx, total_paragraphs, config: MarkerConfig | None = None):
    """
    Classify the current paragraph as 'intro', 'body', 'conclusion', or 'other'
    based on its index, the intro index, and the total number of real paragraphs.
    
    Special case: in foundation_4 and foundation_5, we *do not* treat the final
    paragraph as a conclusion. Those assignments do not require a conclusion,
    so every paragraph after the introduction is treated as a body paragraph.
    """
    if paragraph_index is None or total_paragraphs is None:
        return "other"

    mode = getattr(config, "mode", None)

    # PEEL paragraph – treat the entire assignment as ONE logical paragraph.
    # Only the first content paragraph receives analysis; all others are ignored.
    if mode == "peel_paragraph":
        if paragraph_index == intro_idx:
            return "body"   # The only real PEEL paragraph
        return "other"       # Ignore all subsequent paragraphs completely

    # Introduction paragraph
    if paragraph_index == intro_idx:
        return "intro"

    foundation_no_conclusion = mode in ("foundation_4", "foundation_5")

    if foundation_no_conclusion:
        # For these modes, anything after the intro is just a body paragraph.
        if intro_idx is not None and paragraph_index > intro_idx:
            return "body"
        return "other"

    # Conclusion paragraph: last real paragraph AFTER the intro
    if total_paragraphs > 1 and paragraph_index == total_paragraphs - 1 and paragraph_index > intro_idx:
        return "conclusion"

    # Body paragraphs: anything between intro and conclusion
    if intro_idx is not None and intro_idx < paragraph_index < total_paragraphs - 1:
        return "body"

    return "other"


def get_sentence_index_for_pos(pos, sentences):
    """
    Given a character position `pos` in flat_text, return the index of the
    sentence in `sentences` that contains it, or None if none match.
    """
    for idx, (s_start, s_end) in enumerate(sentences):
        if s_start <= pos < s_end:
            return idx
    return None


def enforce_font(run):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def run_is_italic(run) -> bool:
    """
    Return True if this run should be treated as italic, either because
    its direct font formatting is italic or because its character style
    is italic (e.g., Word's Emphasis style).
    """
    # Direct formatting
    if bool(run.font.italic):
        return True

    # Style-level formatting (e.g. Emphasis)
    style = getattr(run, "style", None)
    if style is not None:
        try:
            if bool(style.font.italic):
                return True
        except AttributeError:
            pass

    return False


def load_rules(excel_path):
    df = pd.read_excel(excel_path, header=None)
    df = df.dropna(subset=[0, 1])
    df[0] = df[0].astype(str).str.strip()
    df[1] = df[1].astype(str).str.strip()
    return dict(zip(df[0], df[1]))


def load_student_guidance(excel_path) -> dict[str, str]:
    df = pd.read_excel(excel_path, header=None)
    if df.shape[1] < 3:
        return {}
    df = df.dropna(subset=[0, 2])
    if df.empty:
        return {}
    df[0] = df[0].astype(str).str.strip()
    df[2] = df[2].astype(str).str.strip()
    df = df[(df[0] != "") & (df[2] != "")]
    if df.empty:
        return {}

    # Optional safety: drop a header row if someone adds one
    df = df[~df[0].str.lower().isin(["issue", "label"])]

    return dict(zip(df[0], df[2]))


def is_intro_paragraph(idx, intro_idx):
    return idx == intro_idx


def normalize_leading_whitespace(text: str, *, strip_all_tabs_for_title: bool = False) -> str:
    """
    Normalizes leading whitespace for student paragraphs.

    - If strip_all_tabs_for_title=True:
         remove ALL leading tabs entirely and keep at most ONE leading space.
         (Titles should not carry indentation semantics.)
    - Otherwise (normal paragraph mode):
         collapse multiple leading tabs to ONE tab,
         collapse any number of leading spaces to ONE space.

    Only affects the **leading margin**; interior whitespace is untouched.
    """
    if not text:
        return text

    # Extract leading whitespace (spaces and tabs only)
    i = 0
    while i < len(text) and text[i] in (" ", "\t"):
        i += 1

    lead = text[:i]
    rest = text[i:]

    if strip_all_tabs_for_title:
        # Titles: drop all tabs; keep at most one space if they tried to indent with spaces
        had_space = " " in lead
        lead = "" if not had_space else " "
        return lead + rest

    # Normal paragraph mode:
    # Any number of leading tabs -> exactly one tab (wipe out any leading spaces).
    tab_count = lead.count("\t")
    if tab_count >= 1:
        lead = "\t"
    # If there are no tabs but there are spaces, collapse to a single space.
    elif " " in lead:
        lead = " "

    return lead + rest


def flatten_paragraph(paragraph, *, skip_vysti: bool = True):
    """
    Convert Word's fragmented runs into:
      • flat_text (single unified string)
      • segments (list of run position mappings)
    
    Args:
        paragraph: The paragraph to flatten
        skip_vysti: If True, skip Vysti-generated runs. If False, include all runs.
    """
    flat = []
    segments = []
    offset = 0

    for i, run in enumerate(paragraph.runs):
        # Skip Vysti-generated runs if skip_vysti is True
        if skip_vysti and run._element.get("data-vysti") == "yes":
            continue

        clean = extract_run_text(run)
        flat.append(clean)

        length = len(clean)
        segments.append((i, offset, offset + length))
        offset += length

    flat_text = "".join(flat)
    flat_text = normalize_leading_whitespace(flat_text)
    return flat_text, segments


def flatten_paragraph_without_labels(paragraph):
    """
    Similar to flatten_paragraph, but specifically skips Vysti label runs
    (yellow/green highlighted runs with data-vysti="yes").
    
    This is used for title paragraphs to ignore label text while keeping
    highlighted student words.
    
    Returns:
        (flat_text, segments) - same format as flatten_paragraph
    """
    flat = []
    segments = []
    offset = 0

    for i, run in enumerate(paragraph.runs):
        # Skip Vysti label arrows (yellow/green highlighted runs with data-vysti="yes")
        if run._element.get("data-vysti") == "yes":
            highlight = run.font.highlight_color
            if highlight in (WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BRIGHT_GREEN):
                continue

        clean = extract_run_text(run)
        flat.append(clean)

        length = len(clean)
        segments.append((i, offset, offset + length))
        offset += length

    flat_text = "".join(flat)
    flat_text = normalize_leading_whitespace(flat_text)
    return flat_text, segments


def spacy_parse(text):
    """
    Returns:
        doc  -> spaCy Doc object
        tokens -> list of (token.text, start_char, end_char)
        sentences -> list of (sent.start_char, sent.end_char)
    """
    doc = nlp(text)
    tokens = [(t.text, t.idx, t.idx + len(t.text)) for t in doc]

    raw_sentences = [(s.start_char, s.end_char) for s in doc.sents]

    # NEW: merge sentence breaks caused by . ? ! that occur INSIDE double quotes
    quote_spans = compute_quote_spans(text)

    def _quote_span_for_pos(pos: int):
        # compute_quote_spans returns interior spans (start..end) where end is the closing quote index
        for qs, qe in quote_spans:
            if qs <= pos <= (qe - 1):
                return (qs, qe)
        return None

    def _should_merge_at_boundary(boundary_pos: int) -> bool:
        if boundary_pos < 0 or boundary_pos >= len(text):
            return False
        if text[boundary_pos] not in ".?!":
            return False

        span = _quote_span_for_pos(boundary_pos)
        if not span:
            return False

        qs, qe = span

        # If punctuation occurs anywhere inside the quoted interior, merge.
        if boundary_pos < (qe - 1):
            return True

        # If punctuation is at the end of a quoted chunk, only merge when it’s clearly mid-sentence
        # e.g. ... "What is Love?", the author ...
        j = qe + 1
        while j < len(text) and text[j].isspace():
            j += 1
        if j >= len(text):
            return False

        nxt = text[j]
        return (nxt in {",", ";", ":", ")", "]"} or nxt.islower() or nxt.isdigit())

    # Merge pass (only if we have at least 2 sentences)
    if len(raw_sentences) >= 2:
        merged = []
        cur_start, cur_end = raw_sentences[0]
        for ns, ne in raw_sentences[1:]:
            boundary_pos = cur_end - 1
            if _should_merge_at_boundary(boundary_pos):
                cur_end = ne
            else:
                merged.append((cur_start, cur_end))
                cur_start, cur_end = ns, ne
        merged.append((cur_start, cur_end))
        raw_sentences = merged

    sentences: list[tuple[int, int]] = []

    for s_start, s_end in raw_sentences:
        cur_start = s_start
        i = s_start
        # Look for . ? ! immediately followed by an uppercase letter with no space
        while i < s_end - 1:
            ch = text[i]
            nxt = text[i + 1]
            if ch in ".?!" and not nxt.isspace() and nxt.isalpha() and nxt.isupper():
                cut_end = i + 1
                if cur_start < cut_end:
                    sentences.append((cur_start, cut_end))
                cur_start = cut_end
            i += 1
        if cur_start < s_end:
            sentences.append((cur_start, s_end))

    return doc, tokens, sentences


WEAK_TRANSITIONS_MULTI = [
    "to begin",
    "to start",
    "in addition",
    "on the other hand",
    "in summary",
    "to summarize",
    "in conclusion",
    "to conclude",
    "in fact",
    "in contrast",
    "by contrast",
    "even so",
]

WEAK_TRANSITIONS_SINGLE = [
    "first",
    "second",
    "third",
    "next",
    "then",
    "afterward",
    "subsequently",
    "later",
    "moreover",
    "furthermore",
    "additionally",
    "also",
    "besides",
    "overall",
    "finally",
    "ultimately",
    "however",
    "but",
    "yet",
    "nevertheless",
    "nonetheless",
    "still",
    "indeed",
]


def _build_weak_transition_lemmas() -> set[str]:
    lemmas: set[str] = set()
    for phrase in WEAK_TRANSITIONS_MULTI + WEAK_TRANSITIONS_SINGLE:
        for tok in nlp(phrase):
            if tok.is_space or tok.is_punct:
                continue
            lemma = tok.lemma_.lower().strip()
            if lemma:
                lemmas.add(lemma)
    return lemmas


WEAK_TRANSITION_LEMMAS = _build_weak_transition_lemmas()
CONTENT_POS = {"NOUN", "PROPN", "VERB", "ADJ", "ADV"}


def extract_content_lemmas(doc, start_char, end_char, extra_exclude=None) -> set[str]:
    exclude = {"be", "have", "do"} | WEAK_TRANSITION_LEMMAS
    if extra_exclude:
        exclude |= {lemma for lemma in extra_exclude if lemma}

    lemmas: set[str] = set()
    for tok in doc:
        if tok.idx < start_char:
            continue
        if tok.idx >= end_char:
            break
        if tok.is_stop or tok.is_punct or tok.is_space:
            continue
        if tok.pos_ not in CONTENT_POS:
            continue
        lemma = tok.lemma_.lower().strip()
        if not lemma or lemma in exclude:
            continue
        lemmas.add(lemma)
    return lemmas



def extract_run_text(run):
    """
    C1 version:
    Lightweight, safe text extraction.
    Only normalizes:
      • NBSP → space
      • curly quotes → straight quotes
      • en/em dashes → hyphen
    """
    raw = run.text or ""
    out = raw

    # Detect Word tab and break nodes
    xml = run._element.xml
    # Treat Word tabs/breaks as normal spaces in the analytic string.
    # Layout belongs to Word; analysis wants clean prose.
    if "<w:tab" in xml:
        out = " " + out
    if "<w:br" in xml or "<w:cr" in xml:
        out = " " + out

    # Non-breaking spaces → normal spaces
    out = out.replace("\u00A0", " ")

    # Curly → straight quotes (comprehensive map for all quote variants)
    quote_map = {
        # Double quotes (curly variants) → "
        "\u201C": '"',  # LEFT DOUBLE QUOTATION MARK
        "\u201D": '"',  # RIGHT DOUBLE QUOTATION MARK
        "\u201E": '"',  # DOUBLE LOW-9 QUOTATION MARK
        "\u201F": '"',  # DOUBLE HIGH-REVERSED-9 QUOTATION MARK

        # Single quotes / apostrophes (curly variants) → '
        # Do NOT map these to " — Word uses these for apostrophes.
        "\u2018": "'",  # LEFT SINGLE QUOTATION MARK / opening apostrophe
        "\u2019": "'",  # RIGHT SINGLE QUOTATION MARK / apostrophe
        "\u201A": "'",  # SINGLE LOW-9 QUOTATION MARK
        "\u201B": "'",  # SINGLE HIGH-REVERSED-9 QUOTATION MARK

        # Angle quote characters
        "\u00AB": '"',  # LEFT-POINTING DOUBLE ANGLE QUOTATION MARK
        "\u00BB": '"',  # RIGHT-POINTING DOUBLE ANGLE QUOTATION MARK
        "\u2039": "'",  # SINGLE LEFT-POINTING ANGLE QUOTATION MARK
        "\u203A": "'",  # SINGLE RIGHT-POINTING ANGLE QUOTATION MARK
    }

    for bad, good in quote_map.items():
        out = out.replace(bad, good)

    # Normalize dashes
    out = out.replace("—", "-").replace("–", "-")

    return out


def compute_teacher_title_spans(flat_text: str, config: MarkerConfig | None) -> list[tuple[int, int]]:
    """
    Return a list of (start, end) character spans in flat_text where any
    teacher-supplied work title appears *as written* (case-sensitive).

    We reuse the same logic as collect_text_title_format_marks so that:
      - only properly-cased matches to the teacher title count as titles
      - generic phrases like 'being and time' (all lowercase) are NOT
        treated as titles

    These spans are used to exempt parts of titles (e.g. the 'and' in
    'Being and Time') from rules like 'Avoid using the word "and"...'.
    """
    spans: list[tuple[int, int]] = []
    if not config:
        return spans

    for work in iter_teacher_works(config):
        title_text = work.title
        if not title_text:
            continue

        # Match the teacher title text in a case-insensitive way
        pattern = re.compile(re.escape(title_text), re.IGNORECASE)

        for m in pattern.finditer(flat_text):
            start, end = m.start(), m.end()
            matched_text = flat_text[start:end]

            if work.is_minor:
                # Preserve the same guards as collect_text_title_format_marks:
                #   - skip all-lowercase matches
                #   - require exact-case match to the teacher title
                if matched_text == matched_text.lower():
                    continue
                if matched_text != title_text:
                    continue
            else:
                # Major works: only treat exact-case matches as the title
                if matched_text != title_text:
                    continue

            spans.append((start, end))

    return spans


def collect_text_title_format_marks(
    paragraph,
    flat_text: str,
    segments,
    spans,
    config: MarkerConfig | None,
    labels_used: list[str],
    paragraph_role: str | None = None,
    sentences: list[tuple[int, int]] | None = None,
) -> list[dict]:
    """
    Return mark dicts that enforce the teacher-supplied text_title formatting
    (minor works in double quotes, not italic; major works italicized) for
    the given paragraph text.

    This function now supports multiple works (for intertextual analysis mode).
    It loops over all configured teacher works and applies formatting rules to each.

    This is the logic that used to live under the "TITLE FORMATTING RULES"
    block inside analyze_text.
    """
    marks: list[dict] = []

    if not config:
        return marks

    works = iter_teacher_works(config)
    if not works:
        return marks

    def is_span_italic(start: int, end: int) -> bool:
        has_any = False
        for run_idx, seg_start, seg_end in segments:
            if seg_end <= start or seg_start >= end:
                continue
            has_any = True
            run = paragraph.runs[run_idx]
            if not run_is_italic(run):
                return False
        return has_any

    # Loop over all teacher-supplied works
    for work in works:
        title_text = work.title
        pattern = re.compile(re.escape(title_text), re.IGNORECASE)

        # IMPORTANT: we want to skip clearly non-title uses where the matched
        # substring is entirely lowercase (e.g. "young hunger" used as evidence).
        # That behavior should be preserved from the original code.

        if work.is_minor:
            note = "The title of minor works should be inside double quotation marks"
            for m in pattern.finditer(flat_text):
                start, end = m.start(), m.end()
                matched_text = flat_text[start:end]

                # Preserve original guard: skip all-lowercase matches
                if matched_text == matched_text.lower():
                    continue

                # NEW: only treat exact-case matches as actual references to the title.
                # This prevents phrases like "The knife" from being flagged when the
                # configured title is "The Knife".
                if matched_text != title_text:
                    continue

                inside_double = pos_in_spans(start, spans) and pos_in_spans(end - 1, spans)
                italic = is_span_italic(start, end)

                # Correct formatting: inside double quotes and not italic
                if inside_double and not italic:
                    continue

                first_time = note not in labels_used
                mark = {
                    "start": start,
                    "end": end,
                    "note": note,
                    "color": WD_COLOR_INDEX.GRAY_25,
                }
                if first_time:
                    mark["label"] = True
                    labels_used.append(note)
                marks.append(mark)
        else:
            note = "The title of major works should be italicized"
            is_single_word = " " not in title_text
            
            # For single-word titles, apply special heuristic to avoid false positives
            # from character-name uses (e.g., "Antigone" as character vs. *Antigone* as title)
            if is_single_word and paragraph_role == "intro" and sentences:
                # Only enforce in intro paragraph's first sentence
                first_start, first_end = sentences[0]
                
                # Collect all exact-case matches within the first sentence
                first_sentence_matches = []
                for m in pattern.finditer(flat_text):
                    start, end = m.start(), m.end()
                    matched_text = flat_text[start:end]
                    
                    # Only exact-case matches
                    if matched_text != title_text:
                        continue
                    
                    # Skip possessive forms
                    next_two = flat_text[end:end+2] if end + 2 <= len(flat_text) else ""
                    if next_two in ("'s", "\u2019s"):
                        continue
                    
                    # Only consider matches within the first sentence
                    if start < first_start or start >= first_end:
                        continue
                    
                    first_sentence_matches.append((start, end))
                
                # If there are matches in the first sentence, check if any are italicized
                if first_sentence_matches:
                    has_italicized = any(is_span_italic(start, end) for start, end in first_sentence_matches)
                    
                    if has_italicized:
                        # At least one italicized occurrence exists → don't flag any unitalicized ones
                        # (they're likely character-name uses)
                        continue
                    else:
                        # No italicized occurrences → flag the first unitalicized one
                        match_start, match_end = first_sentence_matches[0]
                        first_time = note not in labels_used
                        mark = {
                            "start": match_start,
                            "end": match_end,
                            "note": note,
                            "color": WD_COLOR_INDEX.GRAY_25,
                        }
                        if first_time:
                            mark["label"] = True
                            labels_used.append(note)
                        marks.append(mark)
                # If no matches in first sentence, skip (title not mentioned in first sentence)
            elif is_single_word and paragraph_role in ("body", "conclusion"):
                # Single-word title in body/conclusion paragraphs → skip enforcement
                # (don't flag unitalicized single-word titles to avoid character-name false positives)
                continue
            elif is_single_word:
                # Single-word title in other contexts (e.g., title lines, or intro without sentences)
                # → use standard enforcement (no special heuristic needed)
                for m in pattern.finditer(flat_text):
                    start, end = m.start(), m.end()
                    matched_text = flat_text[start:end]

                    # Only exact-case matches
                    if matched_text != title_text:
                        continue

                    # Skip possessive forms
                    next_two = flat_text[end:end+2] if end + 2 <= len(flat_text) else ""
                    if next_two in ("'s", "\u2019s"):
                        continue

                    # If the span is fully italicized, it's correct
                    if is_span_italic(start, end):
                        continue

                    first_time = note not in labels_used
                    mark = {
                        "start": start,
                        "end": end,
                        "note": note,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    }
                    if first_time:
                        mark["label"] = True
                        labels_used.append(note)
                    marks.append(mark)
            else:
                # Multi-word title: keep current behavior (enforce italics everywhere)
                for m in pattern.finditer(flat_text):
                    start, end = m.start(), m.end()
                    matched_text = flat_text[start:end]

                    # NEW: only enforce formatting for exact-case matches of the title.
                    if matched_text != title_text:
                        continue

                    # If the span is fully italicized, it's correct
                    if is_span_italic(start, end):
                        continue

                    first_time = note not in labels_used
                    mark = {
                        "start": start,
                        "end": end,
                        "note": note,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    }
                    if first_time:
                        mark["label"] = True
                        labels_used.append(note)
                    marks.append(mark)

    return marks


def analyze_text(
    paragraph,
    paragraph_index=None,
    total_paragraphs=None,
    labels_used=None,
    intro_idx=0,
    config: MarkerConfig | None = None,
    prev_body_last_sentence_content_words: set[str] | None = None,
):
    """
    Phase 1 — Forbidden Words
    Phase 5A — Delete-phrases
    Phase 5B — Text-as-text rule
    Phase 6 — Weak Verbs
    Phase 7 — Number Rule
    Phase 8 — Weak Transitions
    
    IMPORTANT: This function recomputes flat_text from the paragraph to ensure
    it reflects any mutations (e.g., intro quotation marks) that occurred before
    this function was called. Uses flatten_paragraph_without_labels to ignore previous Vysti labels.
    """
    global THESIS_TOPIC_ORDER
    global THESIS_DEVICE_SEQUENCE
    global THESIS_ALL_DEVICE_KEYS
    global BODY_PARAGRAPH_COUNT
    global BRIDGE_PARAGRAPHS
    global BRIDGE_DEVICE_KEYS
    global FOUNDATION1_LABEL_TARGET
    global THESIS_TEXT_LOWER

    if labels_used is None:
        labels_used = []
    
    if config is None:
        config = get_preset_config("textual_analysis")

    marks = []

    # -----------------------
    # RECOMPUTE flat_text from paragraph, but ignore previous Vysti labels
    # -----------------------
    flat_text, segments = flatten_paragraph_without_labels(paragraph)

    # -----------------------
    # SPACY PROCESSING
    # -----------------------
    doc, tokens, sentences = spacy_parse(flat_text)
    last_sentence_content_words: set[str] = set()
    if sentences:
        last_start, last_end = sentences[-1]
        last_sentence_content_words = extract_content_lemmas(doc, last_start, last_end)

    # -----------------------
    # Always compute spans fresh from the updated flat_text
    # This happens AFTER any mutations, so spans are accurate
    # -----------------------
    spans = compute_quote_spans(flat_text)

    # -----------------------
    # PEEL PARAGRAPH: First sentence behaves like
    # "First Sentence of analysis" + a simple closed thesis
    # -----------------------
    if config and getattr(config, "mode", None) == "peel_paragraph":
        # First content paragraph in the essay (after MLA header / title)
        is_first_content_para = (
            paragraph_index is not None
            and intro_idx is not None
            and paragraph_index == intro_idx
        )

        if is_first_content_para and sentences:
            # Keep a copy of the paragraph text in lowercase (harmless here, but
            # consistent with how intros store THESIS_TEXT_LOWER)
            THESIS_TEXT_LOWER = flat_text.lower()

            # ---------- First Sentence of analysis: author + genre + title + summary ----------
            if config.text_title:
                title_presence_note = (
                    "The first sentence must state the author's full name, genre and title of the text, "
                    "and present a concrete and general summary"
                )

                # Re-use the same fuzzy title finder we use in intros
                title_span, title_is_exact = find_title_span_in_first_sentence(flat_text, sentences, config)

                # Check author full name presence against ANY configured author
                missing_author = False
                author_fields: list[str] = []
                for name in (config.author_name, config.author_name_2, config.author_name_3):
                    if name:
                        author_fields.append(name)

                if author_fields:
                    first_start, first_end = sentences[0]
                    first_sentence_text = flat_text[first_start:first_end]

                    has_any_full_author = any(
                        author_full_name_present(author_name, first_sentence_text)
                        for author_name in author_fields
                    )
                    missing_author = not has_any_full_author

                def find_author_like_span_in_first_sentence() -> tuple[int, int] | None:
                    """
                    Try to locate the student's attempted author name in the FIRST sentence.
                    Copied from the intro logic.
                    """
                    if not getattr(config, "author_name", None):
                        return None
                    if not sentences:
                        return None

                    first_start, first_end = sentences[0]
                    author_target = config.author_name

                    # Collect tokens in the first sentence
                    first_tokens = [t for t in doc if first_start <= t.idx < first_end]

                    def looks_name_like(tok):
                        txt = tok.text
                        return txt and txt[0].isupper() and any(ch.isalpha() for ch in txt)

                    best_span = None
                    best_sim = 0.0

                    n = len(first_tokens)
                    for i in range(n):
                        if not looks_name_like(first_tokens[i]):
                            continue
                        for length in (1, 2, 3):
                            j = i + length
                            if j > n:
                                break
                            window = first_tokens[i:j]
                            snippet_text = flat_text[window[0].idx: window[-1].idx + len(window[-1].text)]
                            sim = SequenceMatcher(None, snippet_text, author_target).ratio()
                            if sim > best_sim and sim >= 0.4:
                                best_sim = sim
                                best_span = (window[0].idx, window[-1].idx + len(window[-1].text))
                    return best_span

                # Decide whether to attach the yellow "First sentence must…" label
                should_label_first_sentence = False
                if title_span is None:
                    should_label_first_sentence = True
                if missing_author:
                    should_label_first_sentence = True

                if should_label_first_sentence:
                    first_start, first_end = sentences[0]
                    anchor_pos = first_end
                    if title_presence_note not in labels_used:
                        marks.append({
                            "start": anchor_pos,
                            "end": anchor_pos,
                            "note": title_presence_note,
                            "color": None,
                            "label": True,
                        })
                        labels_used.append(title_presence_note)
                    else:
                        marks.append({
                            "start": anchor_pos,
                            "end": anchor_pos,
                            "note": title_presence_note,
                            "color": None,
                        })

                # If the author is missing, gray-highlight the guessed author and add local label
                if missing_author:
                    author_span = find_author_like_span_in_first_sentence()
                    if author_span is not None:
                        a_start, a_end = author_span
                        marks.append({
                            "start": a_start,
                            "end": a_end,
                            "color": WD_COLOR_INDEX.GRAY_25,
                        })
                        marks.append({
                            "start": a_end,
                            "end": a_end,
                            "note": "Is this the author's full name?",
                            "color": None,
                            "label": True,
                        })

                # If the student's title is fuzzy but not exact, gray-highlight and label it
                if title_span is not None and not title_is_exact:
                    ts, te = title_span
                    marks.append({
                        "start": ts,
                        "end": te,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    })
                    marks.append({
                        "start": te,
                        "end": te,
                        "note": "Is this the correct title?",
                        "color": None,
                        "label": True,
                    })

            # ---------- Simple "closed thesis" behavior for the Point sentence ----------
            # For PEEL, your Point sentence shouldn't be a question.
            first_start, first_end = sentences[0]
            first_text = flat_text[first_start:first_end].strip()
            if config.enforce_closed_thesis and "?" in first_text:
                anchor_pos = first_end
                closed_note = "Use a closed thesis statement"
                if closed_note not in labels_used:
                    marks.append({
                        "start": anchor_pos,
                        "end": anchor_pos,
                        "note": closed_note,
                        "color": None,
                        "label": True,
                    })
                    labels_used.append(closed_note)
                else:
                    marks.append({
                        "start": anchor_pos,
                        "end": anchor_pos,
                        "note": closed_note,
                        "color": None,
                    })

    # -----------------------
    # Teacher-supplied title keys for fuzzy matching
    # -----------------------
    # Build a list of all teacher-supplied titles for structural quote checks
    teacher_titles = [w.title for w in iter_teacher_works(config)]
    teacher_title_norm_keys = {
        normalize_title_key(t) for t in teacher_titles if t
    }

    # Precompute spans for all teacher-supplied titles as they appear in this
    # paragraph's text. We use these spans to exempt title words (for example,
    # the "and" in Being and Time) from global rules like the repeated-"and"
    # check.
    teacher_title_spans: list[tuple[int, int]] = []
    for title_text in teacher_titles:
        if not title_text:
            continue
        norm_title = normalize_title_for_exact_match(title_text)
        if not norm_title:
            continue
        pattern = re.compile(re.escape(norm_title), re.IGNORECASE)
        for m in pattern.finditer(flat_text):
            teacher_title_spans.append((m.start(), m.end()))

    def in_teacher_title(pos: int) -> bool:
        """Return True iff `pos` lies inside any teacher-supplied title span."""
        return pos_in_spans(pos, teacher_title_spans)

    def is_teacher_title_interior(interior: str) -> bool:
        """
        Return True if `interior` is one of the teacher-supplied text titles,
        allowing small deviations (typos, paraphrasing) but *not* short
        fragments like "three tired pronouns".

        Used to exempt titles from structural quotation bans
        (intro/body/conclusion).

        IMPORTANT for intertextual mode:
        This checks **all** configured works (text_title, text_title_2, text_title_3),
        not just the first one.
        """
        if not teacher_titles:
            return False

        key = normalize_title_key(interior)
        if not key:
            return False

        # Fast path: exact normalized match to any teacher title
        if key in teacher_title_norm_keys:
            return True

        # Fuzzy path: treat near-misses as the title if similarity is high enough.
        # This keeps the old behavior (e.g. slight Ham/Facebook title variants)
        # while still rejecting short fragments.
        for title_text in teacher_titles:
            norm_title = normalize_title_key(title_text)
            if not norm_title:
                continue

            # Ignore fragments that are much shorter than the full title
            if len(key) < max(5, int(0.5 * len(norm_title))):
                continue

            sim = title_similarity(interior, title_text)
            if sim >= 0.6:
                return True

        return False

    # -----------------------
    # STRUCTURAL QUOTATION RULES (intro/body/conclusion)
    # -----------------------
    def add_structural_mark(start, end, note, color=WD_COLOR_INDEX.GRAY_25):
        """
        Helper to add a mark for structural quotation & thesis rules, ensuring we only
        create a labeled yellow comment once per note across the document.
        Allows zero-length spans so we can attach label-only comments at the end
        of a paragraph without highlighting any characters.
        """
        if start is None or end is None or start > end:
            return

        # Always show a yellow label for Undeveloped paragraph in every paragraph
        # and for the body-evidence rule in every paragraph that lacks quotations.
        if note in {"Undeveloped paragraph", "Every paragraph needs evidence"}:
            marks.append({
                "start": start,
                "end": end,
                "note": note,
                "color": color,
                "label": True,
            })
            return

        if note not in labels_used:
            marks.append({
                "start": start,
                "end": end,
                "note": note,
                "color": color,
                "label": True,
            })
            labels_used.append(note)
        else:
            marks.append({
                "start": start,
                "end": end,
                "note": note,
                "color": color,
            })

    paragraph_role = get_paragraph_role(paragraph_index, intro_idx, total_paragraphs, config=config)
    num_sentences = len(sentences)

    # Detect creative essay title lines like:
    #   "Quote from text": Topic in "Text Title"
    stripped = flat_text.strip()
    is_essay_title_line = bool(
        TITLE_PATTERN.match(stripped) or TITLE_PATTERN_NO_COLON.match(stripped)
    )

    # -----------------------
    # Foundation Assignment 3: treat ALL content paragraphs as part of a
    # single logical introduction, regardless of paragraph breaks.
    # -----------------------
    is_foundation3 = bool(config and config.mode == "foundation_3")

    if is_foundation3:
        paragraph_role = "intro"

    # Convenience flags: first and last content paragraph indices
    is_first_intro_para = (
        paragraph_index is not None
        and intro_idx is not None
        and paragraph_index == intro_idx
    )
    is_last_intro_para = (
        paragraph_index is not None
        and total_paragraphs is not None
        and paragraph_index == total_paragraphs - 1
    )

    # =====================================================================
    # FOUNDATION ASSIGNMENT 1 — EXTRA SENTENCES IN INTRO PARAGRAPH
    # =====================================================================
    # Foundation Assignment 1 requires ONLY the first sentence of the essay.
    # 
    # If the student writes multiple sentences in the intro paragraph:
    #   1. The first sentence is marked normally with all first-sentence rules
    #   2. All extra sentences get:
    #      - Red highlight + strikethrough (matching deleted-word style)
    #   3. The position of the last extra sentence is tracked globally so we
    #      can attach a single yellow label at the very end (see run_marker)
    #   4. No other issue labels are attached to extra sentences (handled by
    #      the filtering logic later in this function)
    #
    # This handles cases like Foundation_HW1_Ben where extra sentences appear
    # in the same paragraph as the first sentence.
    # =====================================================================
    if config.mode == "foundation_1" and paragraph_role == "intro" and num_sentences > 1:
        for idx, (s_start, s_end) in enumerate(sentences):
            if idx == 0:
                continue  # Skip the first sentence - it gets normal marking below

            # Apply red highlight + strikethrough to the entire extra sentence
            # This matches the "deleted word" style (see lines ~2411-2418)
            marks.append(
                {
                    "start": s_start,
                    "end": s_end,
                    "note": None,
                    "color": WD_COLOR_INDEX.RED,  # Red highlight
                    "strike": True,  # Strikethrough
                }
            )

            # Track this extra sentence's end position for label placement
            # Each extra sentence overwrites the previous target, so we end up
            # with the position of the LAST extra sentence in the intro paragraph
            FOUNDATION1_LABEL_TARGET = (paragraph_index, s_end)

    # Map each quote interior span to the sentence index that contains it
    quote_sentence_indices = []
    for (q_start, q_end) in spans:
        if q_start >= len(flat_text):
            quote_sentence_indices.append(None)
            continue

        # Prefer the character just before the interior if possible (opening quote),
        # otherwise the interior start, otherwise the last interior char.
        pos = q_start - 1 if q_start > 0 else q_start
        sent_idx = get_sentence_index_for_pos(pos, sentences)
        if sent_idx is None and q_end > q_start:
            pos2 = q_end - 1
            sent_idx = get_sentence_index_for_pos(pos2, sentences)
        
        # Adjust for quotes immediately following a colon in body paragraphs
        # This handles cases where spaCy splits the topic sentence at a colon,
        # placing the quote in the next sentence, but conceptually it's part of the topic sentence.
        if paragraph_role == "body" and sent_idx is not None and sent_idx > 0:
            # Find the last non-space character before the opening quote
            before = pos - 1
            while before >= 0 and flat_text[before].isspace():
                before -= 1
            
            # If the quote starts right after a colon, treat it as part of the previous sentence
            if before >= 0 and flat_text[before] == ":":
                # Check if the colon is at the end of the previous sentence
                prev_sent_idx = sent_idx - 1
                if 0 <= prev_sent_idx < len(sentences):
                    prev_s_start, prev_s_end = sentences[prev_sent_idx]
                    # If the colon is in the previous sentence, treat the quote as part of that sentence
                    if before >= prev_s_start and before < prev_s_end:
                        # Colon is at the end of the previous sentence, so move quote to that sentence
                        sent_idx = prev_sent_idx
        
        quote_sentence_indices.append(sent_idx)

    # -----------------------
    # SUBJECTIVE EVALUATION WORDS (INTRO + CONCLUSION ONLY)
    # -----------------------
    # We want to delete empty evaluative language like "insightful essay",
    # "successful essay", "successfully argues", etc. in the introduction
    # and conclusion. These should be red strikethrough only — no yellow label,
    # no entry in the Issues/Explanation summary.
    if paragraph_role in ("intro", "conclusion"):
        for token in doc:
            lower = token.text.lower()
            if lower not in SUBJECTIVE_EVAL_WORDS:
                continue

            start = token.idx
            end = token.idx + len(token.text)

            # Do not touch text inside direct quotations
            if pos_in_spans(start, spans) or pos_in_spans(end - 1, spans):
                continue

            # Do not touch text inside teacher-supplied titles
            if in_teacher_title(start) or in_teacher_title(end - 1):
                continue

            marks.append({
                "start": start,
                "end": end,
                "note": "",  # no inline comment text
                "color": WD_COLOR_INDEX.RED,
                "strike": True,
                # critically: no "label" key → no yellow label and
                # no addition to labels_used/summary table
            })

    # Apply structural quotation rules based on paragraph role
    if paragraph_role == "intro":
        # Remember the full introduction paragraph text for on-thesis checks
        THESIS_TEXT_LOWER = flat_text.lower()

        # ---------------------------------------------
        # First-sentence TITLE check (teacher-supplied text_title)
        # ---------------------------------------------
        if (
            config
            and config.text_title
            and sentences
            and (not is_foundation3 or is_first_intro_para)
        ):
            title_presence_note = (
                "The first sentence must state the author's full name, genre and title of the text, "
                "and present a concrete and general summary"
            )
            title_span, title_is_exact = find_title_span_in_first_sentence(flat_text, sentences, config)

            # ---------- Check author full name whenever the teacher supplied one ----------
            missing_author = False
            # Collect all raw teacher-supplied author names (do NOT normalize here;
            # author_full_name_present expects the original strings).
            author_fields: list[str] = []
            if config is not None:
                for name in (config.author_name, config.author_name_2, config.author_name_3):
                    if name:
                        author_fields.append(name)

            if author_fields and sentences:
                first_start, first_end = sentences[0]
                first_sentence_text = flat_text[first_start:first_end]

                # The first sentence is OK if it contains the *full* name of at least
                # ONE configured author (intertextual essays may start with any work).
                has_any_full_author = any(
                    author_full_name_present(author_name, first_sentence_text)
                    for author_name in author_fields
                )
                missing_author = not has_any_full_author

            def find_author_like_span_in_first_sentence() -> tuple[int, int] | None:
                """
                Try to locate the student's attempted author name in the FIRST sentence.

                We:
                  - Restrict to the first sentence span [first_start, first_end)
                  - Consider 1–3-token windows that look like a name (capitalized words)
                  - Use fuzzy similarity vs config.author_name to pick the best candidate
                  - Return (start, end) char offsets in flat_text, or None if no good candidate
                """
                if not getattr(config, "author_name", None):
                    return None
                if not sentences:
                    return None

                first_start, first_end = sentences[0]
                author_target = config.author_name

                # Collect tokens in the first sentence
                first_tokens = [t for t in doc if first_start <= t.idx < first_end]

                def looks_name_like(tok):
                    txt = tok.text
                    # Simple heuristic: alphabetic and capitalized (e.g. Dennis, Baron, Brown)
                    return txt and txt[0].isupper() and any(ch.isalpha() for ch in txt)

                best_span = None
                best_sim = 0.0

                n = len(first_tokens)
                for i in range(n):
                    # Only start on name-like tokens
                    if not looks_name_like(first_tokens[i]):
                        continue

                    for length in (1, 2, 3):
                        j = i + length
                        if j > n:
                            break
                        window = first_tokens[i:j]
                        snippet_text = flat_text[window[0].idx: window[-1].idx + len(window[-1].text)]

                        # Compute fuzzy similarity against the teacher-supplied author_name
                        sim = SequenceMatcher(None, snippet_text, author_target).ratio()
                        if sim > best_sim and sim >= 0.4:
                            best_sim = sim
                            best_span = (window[0].idx, window[-1].idx + len(window[-1].text))

                return best_span

            # Decide whether to attach the yellow label
            should_label_first_sentence = False
            # Original behavior: if we never see a fuzzy title at all, complain.
            if title_span is None:
                should_label_first_sentence = True
            # New behavior: for foundation assignments, also complain if author is missing,
            # even when a fuzzy title was found.
            if missing_author:
                should_label_first_sentence = True

            if should_label_first_sentence:
                first_start, first_end = sentences[0]
                anchor_pos = first_end
                if title_presence_note not in labels_used:
                    marks.append({
                        "start": anchor_pos,
                        "end": anchor_pos,
                        "note": title_presence_note,
                        "color": None,  # label-only, no highlight
                        "label": True,
                    })
                    labels_used.append(title_presence_note)
                else:
                    # Extra safeguard if this ever runs more than once
                    marks.append({
                        "start": anchor_pos,
                        "end": anchor_pos,
                        "note": title_presence_note,
                        "color": None,
                    })

            # If the author is missing (e.g. misspelled or wrong name),
            # gray-highlight the student's attempted author name and add
            # a local yellow label that does NOT appear in the summary.
            if missing_author:
                author_span = find_author_like_span_in_first_sentence()
                if author_span is not None:
                    a_start, a_end = author_span

                    # Gray highlight on the name itself
                    marks.append({
                        "start": a_start,
                        "end": a_end,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    })

                    # Local yellow label: "Is this the author's full name?"
                    # Note: we do NOT add this note to labels_used, so it
                    # will NOT appear in the Issues/Explanation summary.
                    marks.append({
                        "start": a_end,
                        "end": a_end,
                        "note": "Is this the author's full name?",
                        "color": None,   # label only, no extra highlight
                        "label": True,
                    })

            # If we found a fuzzy-but-not-exact match for the title in the
            # FIRST sentence, gray-highlight it and add a local yellow label
            # that does NOT appear in the Issues/Explanation summary.
            if title_span is not None and not title_is_exact:
                ts, te = title_span

                # Gray highlight on the attempted title
                marks.append({
                    "start": ts,
                    "end": te,
                    "color": WD_COLOR_INDEX.GRAY_25,
                })

                # Local yellow label: "Is this the correct title?"
                marks.append({
                    "start": te,
                    "end": te,
                    "note": "Is this the correct title?",
                    "color": None,   # label only
                    "label": True,
                })

        # ---------------------------------------------
        # Existing Rule 1: Avoid quotations in the introduction
        # (first sentence, summary sentences), EXCEPT:
        #   - quotes in the THESIS sentence should use:
        #       "No quotations in thesis statements"
        #   - and we always exempt the teacher-supplied title.
        # ---------------------------------------------
        first_idx = 0
        allow_summary_quotes = getattr(config, "allow_intro_summary_quotes", False)

        def strike_commas_around_title(q_start: int, q_end: int) -> None:
            """
            For a teacher-supplied title in quotes in the FIRST sentence,
            strike commas that incorrectly set the title off on both sides.

            We look for:
              - A comma before the opening quote (allowing for spaces)
              - A comma either:
                    * just before the closing quote ("Title,"), OR
                    * just after the closing quote ("Title", Author)
            When and only when we find commas on BOTH sides, we apply
            red highlight + strikethrough to those two commas, with
            no label text.
            """
            # LEFT side: comma before the opening quote
            # spans are interior-only, so the opening quote is at q_start - 1
            i = q_start - 1

            # Step left past any spaces (defensive)
            while i >= 0 and flat_text[i].isspace():
                i -= 1

            # Step left past the opening quote itself and any spaces before it
            if i >= 0 and flat_text[i] == '"':
                i -= 1
                while i >= 0 and flat_text[i].isspace():
                    i -= 1

            left_comma = i if i >= 0 and flat_text[i] == "," else None

            # RIGHT side: comma either inside before the closing quote,
            # or just after the closing quote (skipping spaces).
            right_comma = None

            # Inside: last interior character (q_end - 1)
            inside = q_end - 1
            if 0 <= inside < len(flat_text) and flat_text[inside] == ",":
                right_comma = inside
            else:
                # After the closing quote: q_end is the closing quote
                j = q_end + 1
                while j < len(flat_text) and flat_text[j].isspace():
                    j += 1
                if j < len(flat_text) and flat_text[j] == ",":
                    right_comma = j

            # Only strike if we have commas on BOTH sides
            if left_comma is not None and right_comma is not None:
                for idx in (left_comma, right_comma):
                    marks.append({
                        "start": idx,
                        "end": idx + 1,
                        "note": "",
                        "color": WD_COLOR_INDEX.RED,
                        "strike": True,
                    })

        # Compute which sentence index is the thesis sentence, using the same
        # logic as the closed-thesis block below:
        thesis_sent_idx = None
        if sentences:
            # For most modes, thesis = last sentence of this intro paragraph.
            # For foundation_3, the actual thesis is the last sentence of the
            # LAST intro paragraph, so we only treat the last intro paragraph's
            # final sentence as the thesis.
            if not is_foundation3 or is_last_intro_para:
                thesis_sent_idx = len(sentences) - 1

        for (q_start, q_end), sent_idx in zip(spans, quote_sentence_indices):
            if sent_idx is None:
                continue

            interior = flat_text[q_start:q_end].strip()
            ...
            # If this quotation is the teacher-supplied title for any configured work,
            # do NOT treat it as a real intro quotation.
            if (
                is_teacher_title_interior(interior)
                or (in_teacher_title(q_start) and in_teacher_title(q_end - 1))
            ):
                continue
            # Identify thesis vs summary vs first-sentence quotes
            is_thesis_sentence = (
                thesis_sent_idx is not None and sent_idx == thesis_sent_idx
            )

            is_summary_sentence = (
                allow_summary_quotes
                and thesis_sent_idx is not None
                and sent_idx > first_idx
                and sent_idx < thesis_sent_idx
            )

            # In Analytic frame mode or when the toggle is on, allow
            # quotations in the *summary* sentences between first and thesis.
            if is_summary_sentence:
                continue

            # Thesis sentences are still forbidden to contain quotations,
            # unless closed-thesis rules are turned off.
            if is_thesis_sentence:
                if not config.enforce_closed_thesis:
                    # When the teacher disables closed-thesis rules,
                    # we also allow quotations in the thesis sentence.
                    continue
                quote_note = "No quotations in thesis statements"
            else:
                # First sentence (and any other non-summary, non-thesis sentence)
                # normally uses the standard intro quotation rule.
                quote_note = "Avoid quotations in the introduction"


            # NEW: if the generic intro-quote rule is disabled, skip
            # non-thesis intro-quotation marks entirely.
            if (
                quote_note == "Avoid quotations in the introduction"
                and not getattr(config, "enforce_intro_quote_rule", True)
            ):
                continue

            add_structural_mark(
                q_start,
                q_end,
                quote_note,
            )

        # -----------------------
        # THESIS STATEMENT CHECK (CLOSED THESIS)
        # Thesis = final sentence of the introduction.
        # For Foundation 3, treat the final sentence of the LAST content
        # paragraph as the thesis, even if students break the intro into
        # multiple paragraphs.
        # -----------------------
        if sentences and (not is_foundation3 or is_last_intro_para):
            thesis_start, thesis_end = sentences[-1]
            thesis_text = flat_text[thesis_start:thesis_end].strip()

            # Anchor for label-only comment placed "after" the paragraph.
            # For Foundation 2, anchor at the end of the thesis sentence instead.
            anchor_pos = thesis_end if config and config.mode == "foundation_2" else len(flat_text) + 1

            # New: flag an introduction that has only a one-sentence summary
            # before the thesis. With two sentences total in the intro, there is
            # exactly one summary sentence and one thesis sentence, which is
            # considered insufficient.
            # EXCEPTION: Foundation 2 expects exactly 2 sentences (first sentence + thesis),
            # so we skip this rule for that mode.
            if len(sentences) == 2 and not (config and config.mode == "foundation_2"):
                one_sentence_summary_note = "A one-sentence summary is always insufficient"

                # Attach the label-only comment directly after the FIRST sentence
                # (the introductory summary), not at the end of the paragraph.
                first_start, first_end = sentences[0]

                marks.append({
                    "start": first_end,
                    "end": first_end,
                    "note": one_sentence_summary_note,
                    "color": None,   # label-only, no highlight
                    "label": True,
                })

            # If the final sentence is a question, treat this as "no thesis".
            # (Thesis statements in this system cannot be questions.)
            ends_with_question = "?" in thesis_text
            if config.enforce_closed_thesis and ends_with_question:
                add_structural_mark(
                    anchor_pos,
                    anchor_pos,
                    "Use a closed thesis statement",
                    color=None,  # label only, no highlighting
                )
            else:
                # Foundation 4: Record thesis location for assignment-completion rule
                if config.mode == "foundation_4":
                    global THESIS_PARAGRAPH_INDEX, THESIS_ANCHOR_POS
                    THESIS_PARAGRAPH_INDEX = paragraph_index
                    THESIS_ANCHOR_POS = thesis_end
                
                # Collect spaCy tokens that lie in the thesis sentence span
                thesis_tokens = [t for t in doc if thesis_start <= t.idx < thesis_end]

                device_count = 0
                clarifier_devices = 0
                thesis_devices_in_order = []

                # Very small stoplist so we don't count glue words as "clarifying"
                STOPLIKE = {
                    "a", "an", "the",
                    "of", "to", "for", "on", "in", "at", "by", "from", "with", "as",
                    "that", "this", "these", "those",
                    "his", "her", "their", "its",
                    "use",
                }

                # Track where devices and the main argumentative verb appear
                thesis_verb_index = None
                device_positions = []

                for i, tok in enumerate(thesis_tokens):
                    lemma = tok.lemma_.lower()
                    lower = tok.text.lower()

                    # Rough position of the main argumentative verb (argues, suggests, shows, etc.)
                    if thesis_verb_index is None and lemma in THESIS_VERB_LEMMAS:
                        thesis_verb_index = i

                    device_key = canonical_device_key(tok)
                    if device_key is None:
                        # Not a device/strategy token – keep going.
                        continue

                    # We only get here for real device tokens
                    device_count += 1
                    device_positions.append(i)

                    # Only treat this as a NEW topic if it is not embedded in the phrase
                    # of an earlier device (e.g. 'diction on description of food').
                    if not is_embedded_device(thesis_tokens, i):
                        # Keep every non-embedded device in order, even if the canonical keys match.
                        # This allows a thesis like "zoomorphication ... and imagery ..." to
                        # support two separate body paragraphs, even if synonyms collapse them.
                        thesis_devices_in_order.append(device_key)

                    # -----------------------------
                    # How specific is THIS device?
                    # -----------------------------
                    has_clarifier = False

                    # (a) adjective immediately BEFORE the device, e.g. "sexual metaphors", "quantitative data"
                    # Also check 1-2 tokens before to handle hyphenated compounds like "religiously-charged"
                    for lookback in [1, 2]:
                        if i >= lookback:
                            prev = thesis_tokens[i - lookback]
                            prev_lower = prev.text.lower()
                            # Check for ADJ or ADV (for compounds like "religiously-charged")
                            if (
                                prev.pos_ in {"ADJ", "ADV"}
                                and prev_lower not in STOPLIKE
                                and prev.lemma_.lower() not in THESIS_DEVICE_WORDS
                            ):
                                has_clarifier = True
                                break
                    
                    # (b) OR some content phrase AFTER the device in the same comma/and chunk,
                    #     e.g. "metaphors of hunger", "symbolism of the chocolate orgy"
                    if not has_clarifier:
                        j = i + 1
                        while j < len(thesis_tokens):
                            nxt = thesis_tokens[j]
                            txt = nxt.text
                            nxt_lower = txt.lower()

                            # Stop at list boundaries
                            if txt in {",", ";"} or is_thesis_device_separator(thesis_tokens, j):
                                break

                            # Count as clarifier if it's a content word (NOUN, PROPN, ADJ)
                            # and not a stopword or device word
                            # This handles cases like "quantitative data", "disorienting imagery"
                            # where the adjective comes before, but also cases where content follows
                            if (
                                any(ch.isalpha() for ch in txt)
                                and nxt_lower not in STOPLIKE
                                and nxt.lemma_.lower() not in THESIS_DEVICE_WORDS
                                and nxt.pos_ in {"NOUN", "PROPN", "ADJ"}
                            ):
                                has_clarifier = True
                                break

                            j += 1

                    if has_clarifier:
                            clarifier_devices += 1

                # Closed / specificity checks
                if config.enforce_closed_thesis and device_count == 0:
                    # No devices at all in the thesis → needs to be closed.
                    add_structural_mark(
                        anchor_pos,
                        anchor_pos,
                        "Use a closed thesis statement",
                        color=None,  # label only, no highlighting
                    )
                else:
                    # Require all devices to be specific in a one-topic thesis,
                    # but allow at most ONE bare device once there are 2+ topics.
                    #
                    # Examples:
                    #   device_count = 1 → required_specific = 1
                    #   device_count = 2 → required_specific = 1
                    #   device_count = 3 → required_specific = 2
                    #   device_count = 4 → required_specific = 3
                    required_specific = max(1, device_count - 1)

                    if (
                        config.enforce_closed_thesis
                        and config.enforce_specific_thesis_topics
                        and clarifier_devices < required_specific
                    ):
                        add_structural_mark(
                            anchor_pos,
                            anchor_pos,
                            "The topics in the thesis statement should be specific devices or strategies",
                            color=None,  # label only, no highlighting
                        )


                # Persist the ordered thesis devices for later body-paragraph checks
                if thesis_devices_in_order:
                    THESIS_DEVICE_SEQUENCE = thesis_devices_in_order
                else:
                    THESIS_DEVICE_SEQUENCE = []

                # Build a set of all device lemmas that appear anywhere in the thesis
                all_thesis_devices = set()
                for tok in thesis_tokens:
                    key = canonical_device_key(tok)
                    if key is not None:
                        all_thesis_devices.add(key)

                # Persist for body-paragraph checks
                THESIS_ALL_DEVICE_KEYS = all_thesis_devices

                # Only extract thesis topics when we have a valid, non-question, closed thesis
                if device_count > 0 and clarifier_devices == device_count and not ends_with_question:
                    THESIS_TOPIC_ORDER = extract_thesis_topics(thesis_tokens)
                else:
                    THESIS_TOPIC_ORDER = []

                # --- Organization of thesis statement: devices/strategies should come
                #     before the main argumentative verb in the thesis.
                verb_idx = None
                device_positions = []

                for i, tok in enumerate(thesis_tokens):
                    lemma = tok.lemma_.lower()
                    lower = tok.text.lower()

                    # First argumentative verb in the thesis
                    if verb_idx is None and lemma in THESIS_VERB_LEMMAS:
                        verb_idx = i

                    # Record positions of any device/strategy tokens
                    if lemma in THESIS_DEVICE_WORDS or lower in THESIS_DEVICE_WORDS:
                        device_positions.append(i)

                if verb_idx is not None and device_positions:
                    has_before = any(pos < verb_idx for pos in device_positions)
                    has_after = any(pos > verb_idx for pos in device_positions)

                    # We want at least one device/topic before the argumentative verb/claim.
                    # Only flag when ALL devices come after the verb.
                    if (
                        config.enforce_closed_thesis
                        and config.enforce_thesis_organization
                        and not has_before
                    ):
                        add_structural_mark(
                            anchor_pos,
                            anchor_pos,
                            "Organization of thesis statement",
                            color=None,  # label-only, no gray highlight
                        )

    elif paragraph_role == "body" and num_sentences > 0:
        # Rule 2 and Rule 4: body paragraphs
        first_idx = 0
        last_idx = num_sentences - 1
        has_interior_quote = False

        # Determine whether this is a one-sentence "bridge" line:
        # e.g. "By relying upon zoomorphication of the craving for meals,"
        trimmed = flat_text.strip()
        is_topic_only_bridge = (
            num_sentences == 1
            and bool(trimmed)
            and trimmed[-1] in {",", ":"}
        )

        # Remember bridge paragraphs by index so run_marker can avoid praising them
        if is_topic_only_bridge and paragraph_index is not None:
            BRIDGE_PARAGRAPHS.add(paragraph_index)

            # Record any thesis-device keywords that appear in this bridge line
            bridge_keys: set[str] = set()
            for tok in doc:
                key = canonical_device_key(tok)
                if key is not None:
                    bridge_keys.add(key)
            if bridge_keys:
                BRIDGE_DEVICE_KEYS[paragraph_index] = bridge_keys

        # Compute topic sentence span using character offsets (not spaCy sentence boundaries)
        # This handles cases like "When describing the grandeur of the mall, Guterson contrasts..."
        # where spaCy might split at a colon, but we want the entire first sentence.
        topic_start, topic_end = compute_topic_sentence_span(flat_text, spans)

        for (q_start, q_end), sent_idx in zip(spans, quote_sentence_indices):
            if sent_idx is None:
                continue

            # Check if quote overlaps with topic sentence span (excluding known titles)
            quote_overlaps_topic = (
                q_start < topic_end and q_end > topic_start
            )
            
            if quote_overlaps_topic:
                # Topic sentence – skip if the quotation is the teacher-supplied title
                interior = flat_text[q_start:q_end].strip()
                if is_teacher_title_interior(interior):
                    continue

                add_structural_mark(
                    q_start,
                    q_end,
                    "No quotations in topic sentences",
                )
            elif sent_idx == last_idx:
                # Skip if this quotation is the teacher-supplied title
                interior = flat_text[q_start:q_end].strip()
                if is_teacher_title_interior(interior):
                    continue

                add_structural_mark(
                    q_start,
                    q_end,
                    "No quotations in the final sentence of a body paragraph",
                )
            else:
                # Interior sentence: counts as a valid body quote
                has_interior_quote = True

        # Rule 4: Require at least one direct quotation in the body
        # Only meaningful if there is at least one interior sentence (3+ sentences total)
        # In PEEL mode, only apply to the first content paragraph (intro_idx)
        if (
            (config.mode != "peel_paragraph" or paragraph_index == intro_idx)
            and config.require_body_evidence
            and num_sentences >= 3
            and not has_interior_quote
        ):
            # Attach a label-only mark at the very end of the paragraph.
            # We use an anchor past the end of the text so it won't overlap
            # with other marks or cause any highlighting.
            anchor_pos = len(flat_text) + 1
            add_structural_mark(
                anchor_pos,
                anchor_pos,
                "Every paragraph needs evidence",
                color=None,  # explicitly no highlight
            )

        # ---------------------------------------------
        # Undeveloped paragraph: too few interior sentences
        # ---------------------------------------------
        # We treat the first sentence as the topic sentence
        # and the last sentence as the final sentence.
        # Everything in between is the "body" of the paragraph.
        body_sentence_count = max(num_sentences - 2, 0)

        # Skip undeveloped-paragraph rule for Foundation 4 (topic sentence only)
        # Skip undeveloped-paragraph rule for PEEL mode (treats entire submission as one paragraph)
        if config.mode == "foundation_4":
            pass
        elif config.mode != "peel_paragraph":
            if body_sentence_count <= 2 and not is_topic_only_bridge:
                # Place label at the END of the entire paragraph
                anchor_pos = len(flat_text)

                add_structural_mark(
                    anchor_pos,
                    anchor_pos,
                    "Undeveloped paragraph",
                    color=None
                )

                # ---------------------------------------------
        # Thesis/topic alignment for body paragraphs
        # ---------------------------------------------
        if (
            intro_idx is not None
            and config.enforce_closed_thesis
            and config.enforce_topic_thesis_alignment
        ):

            # Only count non-bridge body paragraphs toward thesis organization.
            # One-sentence lines ending with ',' or ':' (like "By relying upon ... ,")
            # should NOT advance the body index.
            #
            # Special case: In foundation_4 and foundation_5, there is no conclusion,
            # so the last paragraph should be treated as a body paragraph.
            body_idx = None
            is_foundation_no_conclusion = config and config.mode in ("foundation_4", "foundation_5")
            
            # Determine if this paragraph should be counted as a body paragraph
            is_body_for_alignment = (
                not is_topic_only_bridge
                and paragraph_index is not None
                and paragraph_index > intro_idx
            )
            
            # In normal essays, exclude the conclusion (last paragraph)
            # In foundation_4/5, include all paragraphs after the intro
            if not is_foundation_no_conclusion:
                is_body_for_alignment = is_body_for_alignment and paragraph_index < (total_paragraphs - 1)
            
            if is_body_for_alignment:
                BODY_PARAGRAPH_COUNT += 1
                body_idx = BODY_PARAGRAPH_COUNT

            # If this paragraph doesn't map to a thesis "slot", skip thesis-topic labels entirely.
            if body_idx is not None:
                expected_idx = body_idx - 1

                # Topic sentence span computed using character offsets (handles "When..." style sentences)
                # This is already computed above for the quotation rule, so reuse it
                # topic_start, topic_end already defined above

                # Collect device lemmas actually used in the topic sentence
                # Use iter_device_spans to handle multi-word devices as single atomic units
                topic_devices = set()
                topic_device_positions: dict[str, tuple[int, int]] = {}

                for key, dev_start, dev_end in iter_device_spans(doc, topic_start, topic_end):
                    topic_devices.add(key)
                    # Record first occurrence for highlighting labels
                    if key not in topic_device_positions:
                        topic_device_positions[key] = (dev_start, dev_end)

                # OUTSIDE loop → build thesis device set once
                thesis_device_set = set(THESIS_DEVICE_SEQUENCE)

                # Also treat any device that appears anywhere in the thesis as "in thesis"
                # This prevents mislabeling cases like BP2 "symbolism" when the thesis contains
                # "allegorical symbolism" but that device was treated as embedded.
                if THESIS_ALL_DEVICE_KEYS:
                    thesis_device_set |= THESIS_ALL_DEVICE_KEYS

                anchor_pos = topic_end  # default anchor for label-only comments

                def mark_topic_label(label_text: str):
                    """
                    If the label is 'Put this topic in the thesis statement', attach the label
                    directly to the first off-thesis device token and gray-highlight that device.
                    Otherwise, attach a label-only comment at the end of the topic sentence.
                    """
                    # Respect configuration: skip some labels entirely in certain modes
                    if (
                        label_text == "Follow the organization of the thesis"
                        and not config.enforce_topic_thesis_alignment
                    ):
                        return
                    if label_text == "Off-topic" and not config.enforce_off_topic:
                        return

                    if "Put this topic in the thesis statement" in label_text or label_text == "Off-topic":
                        thesis_device_set = set(THESIS_DEVICE_SEQUENCE)
                        if THESIS_ALL_DEVICE_KEYS:
                            thesis_device_set |= THESIS_ALL_DEVICE_KEYS

                        # NEW: If every device in the topic sentence already appears
                        # textually in the thesis paragraph, do NOT treat it as a new topic.
                        if THESIS_TEXT_LOWER and topic_devices:
                            all_in_thesis_text = True
                            for d in topic_devices:
                                token = d.lower()
                                if token and token not in THESIS_TEXT_LOWER:
                                    all_in_thesis_text = False
                                    break
                            if all_in_thesis_text:
                                return  # The thesis already literally mentions these devices.

                        # Only treat *new* devices as off-thesis. If every device in the topic
                        # sentence already appears somewhere in the thesis device set, do NOT label it.
                        off_thesis_devices = [d for d in topic_devices if d not in thesis_device_set]
                        if not off_thesis_devices:
                            return  # Nothing truly new here; skip this label.

                        # Otherwise, highlight the first genuinely new device.
                        target_device = off_thesis_devices[0]

                        start, end = topic_device_positions.get(
                            target_device,
                            (anchor_pos, anchor_pos)
                        )
                        add_structural_mark(start, end, label_text)

                    else:
                        # Default: attach label at end of topic sentence
                        add_structural_mark(
                            anchor_pos,
                            anchor_pos,
                            label_text,
                            color=None
                        )

                # ---------------------------------------------
                # Evaluate topic alignment
                # ---------------------------------------------
                # Expected behavior: If the topic sentence contains an inflected form of the
                # expected thesis device (e.g., "contrasts" when expected_device is "contrast"),
                # canonical_device_key should map it to the same canonical key via lemma matching,
                # so expected_device in topic_devices should be True and no "Move this to the
                # topic sentence" label should be added.
                # Example: In Foundation_HW7_Darren, BP1's topic sentence contains "contrasts"
                # and the first thesis device is "contrast". These should match, so no label.
                if 0 <= expected_idx < len(THESIS_DEVICE_SEQUENCE):
                    expected_device = THESIS_DEVICE_SEQUENCE[expected_idx]

                    # NEW: treat the topic as present in the topic sentence if EITHER:
                    #   - the canonical device key appears in topic_devices, OR
                    #   - the raw topic sentence text contains the device word as a substring
                    #
                    # This prevents cases where we miss the device via canonical_device_key
                    # (e.g. thesis device "image" but topic sentence uses "imagery"), yet
                    # the topic sentence clearly mentions the same topic.
                    topic_sentence_text = flat_text[topic_start:topic_end].lower()
                    device_token = expected_device.lower() if expected_device else ""
                    topic_mentions_expected = bool(device_token) and device_token in topic_sentence_text

                    # Also treat the expected device as "already introduced" when it
                    # appears in a one-line bridge paragraph immediately before this one.
                    expected_in_bridge = False
                    if paragraph_index is not None:
                        bridge_keys = BRIDGE_DEVICE_KEYS.get(paragraph_index - 1)
                        if bridge_keys and expected_device in bridge_keys:
                            expected_in_bridge = True

                    if expected_device in topic_devices or topic_mentions_expected or expected_in_bridge:
                        # Correct topic in correct order, or at least clearly mentioned
                        # in the topic sentence / bridge. Do NOT flag later mentions with
                        # "Move this to the topic sentence".
                        pass
                    else:
                        # Look for expected topic elsewhere in paragraph body (after topic sentence)
                        # Use canonical_device_key consistently to handle inflected forms and synonyms
                        expected_start = None
                        expected_end = None

                        for tok in doc:
                            if tok.idx < topic_end:
                                continue
                            
                            # Use canonical_device_key to check if this token matches expected_device
                            key = canonical_device_key(tok)
                            if key == expected_device:
                                # Avoid evidence: skip inside quotes
                                if pos_in_spans(tok.idx, spans) or pos_in_spans(
                                    tok.idx + len(tok.text) - 1, spans
                                ):
                                    continue

                                expected_start = tok.idx
                                expected_end = tok.idx + len(tok.text)
                                break

                        if expected_start is not None:
                            # Misplaced topic — highlight and instruct student to move it
                            marks.append(
                                {
                                    "start": expected_start,
                                    "end": expected_end,
                                    "note": "Move this to the topic sentence",
                                    "color": WD_COLOR_INDEX.GRAY_25,
                                    "label": True,
                                }
                            )
                        else:
                            # No occurrence anywhere else — classify based on topic sentence
                            if topic_devices & thesis_device_set:
                                label_text = "Follow the organization of the thesis"
                            elif topic_devices:
                                label_text = "Put this topic in the thesis statement"
                            else:
                                label_text = "Off-topic"

                            mark_topic_label(label_text)
                else:
                    # More body paragraphs than thesis topics
                    if topic_devices & thesis_device_set:
                        label_text = "Follow the organization of the thesis"
                    elif topic_devices:
                        label_text = "Put this topic in the thesis statement"
                    else:
                        label_text = "Off-topic"

                    mark_topic_label(label_text)




        # ---------------------------------------------
        # Additional body-only quotation rules
        # 1) Floating quotations (entire sentence quoted)
        # 2) Long quotations (> 5 words) in interior sentences
        # ---------------------------------------------
        floating_sentence_indices = set()

        # Detect floating quotations in interior sentences.
        # NEW definition (much stricter and matches the teaching rule):
        #   A floating quotation is when the sentence *begins* with a double quote
        #   (ignoring leading spaces) and *ends* with a double quote or with a
        #   double quote immediately followed by sentence punctuation (".", "?!").
        for idx_sent, (s_start, s_end) in enumerate(sentences):
            # Only consider interior sentences (ignore topic + final)
            if idx_sent == first_idx or idx_sent == last_idx:
                continue

            # Extract the raw sentence substring
            s_text = flat_text[s_start:s_end]
            if not s_text.strip():
                continue

            # Leading/trailing whitespace-trimmed view for boundary checks
            s_trim = s_text.strip()
            if not s_trim:
                continue

            # Sentence must start with a double quote to even be a candidate
            if not s_trim.startswith('"'):
                continue

            # Find the last non-space character
            j = len(s_trim) - 1
            while j >= 0 and s_trim[j].isspace():
                j -= 1
            if j < 0:
                continue

            last_char = s_trim[j]
            prev_char = s_trim[j - 1] if j - 1 >= 0 else ""

            # Valid endings:
            #   ... "word"
            #   ... "word".
            #   ... "word"?
            #   ... "word"!
            ends_with_quote = last_char == '"'
            ends_with_quote_and_punct = (
                last_char in ".?!"
                and prev_char == '"'
            )

            if not (ends_with_quote or ends_with_quote_and_punct):
                continue

            # At this point, the entire sentence is syntactically framed by quotes
            # (ignoring leading/trailing spaces). This is a true floating quotation.
            floating_sentence_indices.add(idx_sent)

            # Highlight from first to last alphabetic token for clarity.
            sentence_word_tokens = []
            for tok_text, tok_start, tok_end in tokens:
                if tok_start < s_start or tok_start >= s_end:
                    continue
                if not any(ch.isalpha() for ch in tok_text):
                    continue
                sentence_word_tokens.append((tok_text, tok_start, tok_end))

            if sentence_word_tokens:
                float_start = sentence_word_tokens[0][1]
                float_end = sentence_word_tokens[-1][2]
                add_structural_mark(
                    float_start,
                    float_end,
                    "Floating quotation",
                )

        # Long quotations (> 5 words) in interior, non-floating sentences
        if getattr(config, "enforce_long_quote_rule", True):
            for (q_start, q_end), sent_idx in zip(spans, quote_sentence_indices):
                if sent_idx is None:
                    continue
                # Only interior sentences
                if sent_idx == first_idx or sent_idx == last_idx:
                    continue
                # Ignore quotations in floating-quotation sentences
                if sent_idx in floating_sentence_indices:
                    continue

                # Count alphabetic word tokens inside this quotation span
                word_count = 0
                for tok_text, tok_start, tok_end in tokens:
                    if tok_start < q_start or tok_start >= q_end:
                        continue
                    if any(ch.isalpha() for ch in tok_text):
                        word_count += 1

                if word_count > 5:
                    add_structural_mark(
                        q_start,
                        q_end,
                        "Shorten, modify, and integrate quotations",
                    )

        # ---------------------------------------------
        # Evidence-context rule:
        # Flag sentences where a quotation is introduced
        # with only a very short opening phrase before
        # the first comma.
        #
        # Heuristic:
        #   - Only interior sentences (not topic/final).
        #   - Sentence must contain at least one quote.
        #   - There must be a comma before the first quote.
        #   - The opening phrase (from sentence start up to
        #     that comma) has 1–5 alphabetic words.
        #   - The region between that comma and the quote
        #     does NOT already contain 4+ content words
        #     (NOUN/PROPN/ADJ/VERB).
        #
        # If all of the above hold, attach a label right
        # after the opening phrase:
        #   "Follow the process for inserting evidence"
        # ---------------------------------------------
        rule_note_evidence = "Follow the process for inserting evidence"
        rule_note_repeat_quote = "Only cite a quotation once"

        # Track which evidence quotations have already been introduced
        # earlier in this paragraph.
        used_evidence_quotes = set()

        # Track which *content words* (by lemma) have already appeared inside quotes
        # earlier in this paragraph. This lets us catch cases like
        # "giggly black ghosts" → later "ghosts" as repeated evidence.
        used_quoted_lemmas = set()

        for idx_sent, (s_start, s_end) in enumerate(sentences):
            # Only interior sentences: evidence should live here
            if idx_sent == first_idx or idx_sent == last_idx:
                continue

            # Collect all quotation spans in this sentence
            sent_quote_spans = [
                (q_start, q_end)
                for (q_start, q_end), sent_idx in zip(spans, quote_sentence_indices)
                if sent_idx == idx_sent
            ]
            if not sent_quote_spans:
                continue

            # Use the first quotation in the sentence as the evidence anchor
            q_start, q_end = min(sent_quote_spans, key=lambda x: x[0])

            # Collect content-word lemmas inside this quotation
            quote_lemmas = set()
            for tok in doc:
                if tok.idx < q_start or tok.idx >= q_end:
                    continue
                # Only count content words as "evidence" tokens
                if tok.pos_ in {"NOUN", "PROPN", "VERB", "ADJ", "ADV"}:
                    lemma = tok.lemma_.lower().strip()
                    if lemma:
                        quote_lemmas.add(lemma)

            # Normalize this quotation text so we can detect repeated evidence.
            quote_text = normalize_quote_text(flat_text[q_start:q_end])

            # Check for repetition in two ways BEFORE updating tracking sets:
            #  1) Entire quote string repeated
            #  2) Any content-word lemma repeated from an earlier quote
            is_repeated_string = bool(quote_text) and quote_text in used_evidence_quotes
            is_repeated_lemma = bool(quote_lemmas & used_quoted_lemmas)
            is_repeated = is_repeated_string or is_repeated_lemma

            if is_repeated:
                # Repeated evidence: label it with 'Only cite a quotation once'
                # and skip the evidence-context rule entirely.
                if rule_note_repeat_quote not in labels_used:
                    marks.append({
                        "start": q_start,
                        "end": q_end,
                        "note": rule_note_repeat_quote,
                        "color": WD_COLOR_INDEX.GRAY_25,
                        "label": True,
                    })
                    labels_used.append(rule_note_repeat_quote)
                else:
                    marks.append({
                        "start": q_start,
                        "end": q_end,
                        "note": rule_note_repeat_quote,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    })
                # Update tracking sets after labeling
                if quote_text:
                    used_evidence_quotes.add(quote_text)
                used_quoted_lemmas.update(quote_lemmas)
                # Skip the normal context-checking for this repeated quote
                continue
            else:
                # First time this piece of evidence appears in the paragraph:
                # update tracking sets but do NOT label it.
                if quote_text:
                    used_evidence_quotes.add(quote_text)
                used_quoted_lemmas.update(quote_lemmas)

            # Find the first comma before this quotation, inside the same sentence
            first_comma_pos = None
            for tok in doc:
                if tok.idx < s_start or tok.idx >= q_start:
                    continue
                if tok.text == ",":
                    first_comma_pos = tok.idx
                    break

            # If there is no comma before the quote, this sentence
            # does not follow the simple "opening phrase, ... "quote""
            # pattern we want to constrain.
            if first_comma_pos is None:
                continue

            # Opening phrase: from sentence start up to that comma
            opening_tokens = []
            for tok in doc:
                if tok.idx < s_start or tok.idx >= first_comma_pos:
                    continue
                if not any(ch.isalpha() for ch in tok.text):
                    continue
                # Ignore anything inside direct quotations (paranoid, but cheap)
                tok_start = tok.idx
                tok_end = tok.idx + len(tok.text)
                if pos_in_spans(tok_start, spans) or pos_in_spans(tok_end - 1, spans):
                    continue
                opening_tokens.append(tok)

            opening_word_count = len(opening_tokens)
            # If there is no opening phrase or it is already fairly long,
            # we assume the student has done enough setup.
            if opening_word_count == 0 or opening_word_count > 5:
                continue

            # NEW: treat temporal clauses like "When interviewing Arab women,"
            # as sufficient context before a quotation. If the short opening
            # phrase begins with "When" followed by a gerund (VBG), do not
            # attach the "Follow the process for inserting evidence" label.
            if (
                opening_tokens
                and opening_tokens[0].text.lower() == "when"
                and len(opening_tokens) >= 2
                and opening_tokens[1].tag_ == "VBG"
            ):
                continue

            # Check the region between the comma and the quotation for
            # richer context. If there are already several content words
            # (nouns/verbs/adjectives/proper nouns) there, we treat that
            # as adequate context and do NOT flag.
            content_tokens = []
            for tok in doc:
                if tok.idx <= first_comma_pos or tok.idx >= q_start:
                    continue
                if not any(ch.isalpha() for ch in tok.text):
                    continue
                if tok.pos_ in {"NOUN", "PROPN", "ADJ", "VERB"}:
                    tok_start = tok.idx
                    tok_end = tok.idx + len(tok.text)
                    if pos_in_spans(tok_start, spans) or pos_in_spans(tok_end - 1, spans):
                        continue
                    content_tokens.append(tok)

            if len(content_tokens) >= 4:
                # There is already a reasonably detailed pre-quote setup; skip.
                continue

            # Highlight the opening phrase in GRAY_25
            opening_phrase_start = s_start
            opening_phrase_end = first_comma_pos
            marks.append({
                "start": opening_phrase_start,
                "end": opening_phrase_end,
                "color": WD_COLOR_INDEX.GRAY_25,
            })

            # Anchor the label right after the opening phrase (just before the comma).
            anchor_pos = first_comma_pos

            if rule_note_evidence not in labels_used:
                marks.append({
                    "start": anchor_pos,
                    "end": anchor_pos,
                    "note": rule_note_evidence,
                    "color": WD_COLOR_INDEX.GRAY_25,
                    "label": True,
                })
                labels_used.append(rule_note_evidence)
            else:
                # Subsequent instances: gray highlight only, no new yellow arrow
                marks.append({
                    "start": anchor_pos,
                    "end": anchor_pos,
                    "note": rule_note_evidence,
                    "color": WD_COLOR_INDEX.GRAY_25,
                })

    elif paragraph_role == "conclusion":
        # New rule: Incomplete conclusion if there is only one sentence.
        # Attach a label-only comment after the paragraph (no highlight).
        if num_sentences == 1:
            anchor_pos = len(flat_text) + 1
            add_structural_mark(
                anchor_pos,
                anchor_pos,
                "Incomplete conclusion",
                color=None,  # label-only, no highlighting
            )

        # Rule 3: No quotations anywhere in the conclusion,
        # EXCEPT allow known titles quoted in the conclusion.
        for (q_start, q_end), sent_idx in zip(spans, quote_sentence_indices):
            if sent_idx is None:
                continue

            # Interior of the quotation
            interior = flat_text[q_start:q_end].strip()

            # If this quoted string is the teacher-supplied title, do NOT mark it.
            if is_teacher_title_interior(interior):
                continue

            add_structural_mark(
                q_start,
                q_end,
                "Avoid quotations in the conclusion",
            )

    # -----------------------
    # -----------------------
    # TITLE FORMATTING RULES (teacher-supplied title from config)
    # -----------------------
    # Apply wherever the teacher's text_title appears in student prose
    # (intro, title line, body, conclusion).
    title_marks = collect_text_title_format_marks(
        paragraph=paragraph,
        flat_text=flat_text,
        segments=segments,
        spans=spans,
        config=config,
        labels_used=labels_used,
        paragraph_role=paragraph_role,
        sentences=sentences,
    )
    if title_marks:
        marks.extend(title_marks)

    # -----------------------
    # CLARIFY PRONOUNS (He/She at sentence start)
    # -----------------------
    rule_note_pronoun_antecedent = "Clarify pronouns and antecedents"

    # For each sentence, find the first meaningful token and check for He/She
    for (s_start, s_end) in sentences:
        first_token = None

        for tok_text, tok_start, tok_end in tokens:
            # Skip tokens that are before this sentence
            if tok_start < s_start:
                continue
            # Stop once we've moved past this sentence
            if tok_start >= s_end:
                break

            # Skip tokens that are purely punctuation (no letters)
            if not any(ch.isalpha() for ch in tok_text):
                continue

            first_token = (tok_text, tok_start, tok_end)
            break

        if first_token is None:
            continue

        tok_text, tok_start, tok_end = first_token
        lower = tok_text.lower()

        # Skip if this first word is inside a direct quotation
        if pos_in_spans(tok_start, spans) or pos_in_spans(tok_end - 1, spans):
            continue

        if lower in ("he", "she"):
            # He/She at sentence start → Clarify pronouns and antecedents (TURQUOISE + label)
            if rule_note_pronoun_antecedent not in labels_used:
                marks.append({
                    "start": tok_start,
                    "end": tok_end,
                    "note": rule_note_pronoun_antecedent,
                    "color": WD_COLOR_INDEX.TURQUOISE,
                    "label": True,
                })
                labels_used.append(rule_note_pronoun_antecedent)
            else:
                marks.append({
                    "start": tok_start,
                    "end": tok_end,
                    "note": rule_note_pronoun_antecedent,
                    "color": WD_COLOR_INDEX.TURQUOISE,
                })

    # -----------------------
    # AVOID BEGINNING A SENTENCE WITH A QUOTATION
    # -----------------------
    rule_note_quotation_start = "Avoid beginning a sentence with a quotation"

    # Only apply this rule to intro/body/conclusion paragraphs,
    # and never to the creative essay title line.
    is_content_paragraph = paragraph_role in ("intro", "body", "conclusion")

    if is_content_paragraph and not is_essay_title_line:
        # For each sentence, find the first meaningful token and check if it starts with a quotation mark
        for (s_start, s_end) in sentences:
            first_token = None

            for tok_text, tok_start, tok_end in tokens:
                # Skip tokens that are before this sentence
                if tok_start < s_start:
                    continue
                # Stop once we've moved past this sentence
                if tok_start >= s_end:
                    break

                # Skip tokens that are purely punctuation (no letters)
                if not any(ch.isalpha() for ch in tok_text):
                    continue

                first_token = (tok_text, tok_start, tok_end)
                break

            if first_token is None:
                continue

            tok_text, tok_start, tok_end = first_token

            # Check if the first meaningful token begins with a double quotation mark
            quote_start = None
            if flat_text[tok_start] == '"':
                quote_start = tok_start
            else:
                # Check if there's a quote token immediately before the first meaningful token
                for prev_tok_text, prev_tok_start, prev_tok_end in tokens:
                    if prev_tok_start < s_start:
                        continue
                    if prev_tok_end == tok_start and prev_tok_text == '"':
                        quote_start = prev_tok_start
                        break
                    if prev_tok_start >= tok_start:
                        break

            if quote_start is None:
                continue

            # Extract the interior text by scanning forward until the next "
            quote_end = None
            for i in range(quote_start + 1, min(s_end, len(flat_text))):
                if flat_text[i] == '"':
                    quote_end = i
                    break

            if quote_end is None:
                continue

            # Extract interior text (between the quotes)
            interior = flat_text[quote_start + 1:quote_end].strip()

            # Skip if the quoted text is the teacher-supplied title
            if is_teacher_title_interior(interior):
                continue

            # Highlight the entire offending quotation (from quote_start to quote_end, exclusive of closing quote)
            marks.append({
                "start": quote_start,
                "end": quote_end,
                "color": WD_COLOR_INDEX.GRAY_25,
            })

            # Insert a yellow label-only mark directly after the closing quotation mark
            if rule_note_quotation_start not in labels_used:
                marks.append({
                    "start": quote_end,
                    "end": quote_end,
                    "note": rule_note_quotation_start,
                    "color": None,
                    "label": True,
                })
                labels_used.append(rule_note_quotation_start)

    # -----------------------
    # HIGHLIGHT ALL "this"/"This" OUTSIDE DIRECT QUOTATIONS
    # -----------------------
    for tok_text, tok_start, tok_end in tokens:
        if tok_text.lower() == "this":
            # Skip any 'this' that appears inside direct quotations
            if pos_in_spans(tok_start, spans) or pos_in_spans(tok_end - 1, spans):
                continue

            # Turquoise highlight only, no label or summary entry
            marks.append({
                "start": tok_start,
                "end": tok_end,
                "color": WD_COLOR_INDEX.TURQUOISE,
            })

    # -----------------------
    # REPEATED "AND" IN A SENTENCE
    # -----------------------
    rule_note_and = "Avoid using the word 'and' more than once in a sentence"

    # Have we already attached a yellow label for this rule anywhere in the document?
    and_label_attached = rule_note_and in labels_used

    for (s_start, s_end) in sentences:
        and_tokens = []

        for tok_text, tok_start, tok_end in tokens:
            if tok_start < s_start or tok_start >= s_end:
                continue
            if tok_text.lower() != "and":
                continue

            # Ignore "and" inside direct quotations or teacher-supplied titles
            if (
                pos_in_spans(tok_start, spans)
                or pos_in_spans(tok_end - 1, spans)
                or in_teacher_title(tok_start)
                or in_teacher_title(tok_end - 1)
            ):
                continue

            and_tokens.append((tok_start, tok_end))

        if len(and_tokens) > 1:
            # Highlight ALL 'and's in this sentence in TURQUOISE
            for idx, (tok_start, tok_end) in enumerate(and_tokens):
                mark = {
                    "start": tok_start,
                    "end": tok_end,
                    "color": WD_COLOR_INDEX.TURQUOISE,
                }

                # Attach a SINGLE yellow label to the LAST "and"
                # in the FIRST offending sentence only
                if not and_label_attached and idx == len(and_tokens) - 1:
                    mark["note"] = rule_note_and
                    mark["label"] = True
                    labels_used.append(rule_note_and)
                    and_label_attached = True

                marks.append(mark)

    # -----------------------
    # HIGHLIGHT THESIS DEVICE WORDS (from thesis_devices.txt)
    # -----------------------
    # Any token/phrase that resolves to a canonical thesis device (including synonyms
    # and inflected forms via canonical_device_key and THESIS_MULTIWORD_SYNONYMS)
    # gets a simple BRIGHT_GREEN highlight, as long as it is outside of direct quotations.
    # NOTE: Do NOT highlight devices in essay title lines (matching TITLE_PATTERN or TITLE_PATTERN_NO_COLON)
    if getattr(config, "highlight_thesis_devices", True) and not is_essay_title_line:
        for device_key, start, end in iter_device_spans(doc):
            # Skip any device words/phrases that appear inside direct quotations
            if pos_in_spans(start, spans) or pos_in_spans(end - 1, spans):
                continue

            marks.append({
                "start": start,
                "end": end,
                "color": WD_COLOR_INDEX.BRIGHT_GREEN,
                "device_highlight": True,  # mark as non-issue, just a visual aid
            })

    # -----------------------
    # PHASE 1 — FORBIDDEN WORDS
    # -----------------------
    forbidden = {
        "i": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "you": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "we": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "us": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "our": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "your": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "yours": "No 'I', 'we', 'us', 'our' or 'you' in academic writing",
        "ethos": "Avoid the words 'ethos', 'pathos', and 'logos'",
        "pathos": "Avoid the words 'ethos', 'pathos', and 'logos'",
        "logos": "Avoid the words 'ethos', 'pathos', and 'logos'",
        "very": "Avoid the words 'very' and 'a lot'",
        "a lot": "Avoid the words 'very' and 'a lot'",
        "which": "Avoid the word 'which'",
        "human": "Avoid using the words 'human', 'people', 'everyone', or 'individual'",
        "people": "Avoid using the words 'human', 'people', 'everyone', or 'individual'",
        "everyone": "Avoid using the words 'human', 'people', 'everyone', or 'individual'",
        "individual": "Avoid using the words 'human', 'people', 'everyone', or 'individual'",
        "fact": "Avoid the words 'fact', 'proof', and 'prove'",
        "facts": "Avoid the words 'fact', 'proof', and 'prove'",
        "proof": "Avoid the words 'fact', 'proof', and 'prove'",
        "prove": "Avoid the words 'fact', 'proof', and 'prove'",
        "proves": "Avoid the words 'fact', 'proof', and 'prove'",
        "society": "Avoid overly general words like 'society', 'universe', 'reality', 'life', and 'truth'",
        "universe": "Avoid overly general words like 'society', 'universe', 'reality', 'life', and 'truth'",
        "life": "Avoid overly general words like 'society', 'universe', 'reality', 'life', and 'truth'",
        "truth": "Avoid overly general words like 'society', 'universe', 'reality', 'life', and 'truth'",
        "reality": "Avoid overly general words like 'society', 'universe', 'reality', 'life', and 'truth'",
        "etc": "Avoid 'etc.' in academic writing",
        "audience": "Avoid referring to the reader or audience unless necessary",
        "audiences": "Avoid referring to the reader or audience unless necessary",
        "reader": "Avoid referring to the reader or audience unless necessary",
        "readers": "Avoid referring to the reader or audience unless necessary",
    }

    # Apply configuration: optionally allow I/you/reader/audience in some modes
    if not config.forbid_personal_pronouns:
        for key in ["i", "you", "we", "us", "our", "your", "yours"]:
            forbidden.pop(key, None)

    if not config.forbid_audience_reference:
        for key in ["audience", "audiences", "reader", "readers"]:
            forbidden.pop(key, None)
    # Optionally allow "which"
    if not getattr(config, "enforce_which_rule", True):
        forbidden.pop("which", None)
        # Optionally allow 'fact', 'proof', 'prove'
    if not getattr(config, "enforce_fact_proof_rule", True):
        for key in ["fact", "facts", "proof", "prove", "proves"]:
            forbidden.pop(key, None)
    # Optionally allow 'human', 'people', 'everyone', 'individual'
    if not getattr(config, "enforce_human_people_rule", True):
        for key in ["human", "people", "everyone", "individual"]:
            forbidden.pop(key, None)

    # Optionally allow vague general nouns 'society', 'universe', 'reality', 'life', 'truth'
    if not getattr(config, "enforce_vague_terms_rule", True):
        for key in ["society", "universe", "reality", "life", "truth"]:
            forbidden.pop(key, None)

    # Exceptions for technical / idiomatic uses of general words like
    # 'reality', 'truth', 'life', 'society', and 'universe'.
    # We record the character positions of the keyword in allowed phrases
    # and skip those specific tokens in the forbidden-word loop.
    general_allowed_positions: dict[str, set[int]] = {}

    # Only build exception maps for keys that are actually forbidden
    general_keys = ["reality", "truth", "life", "society", "universe"]
    active_general_keys = [k for k in general_keys if k in forbidden]
    if active_general_keys:
        general_allowed_positions = {k: set() for k in active_general_keys}
        abstract_general_nouns = {
            "life",
            "truth",
            "reality",
            "existence",
            "world",
            "universe",
            "everything",
        }

        # --- REALITY exceptions ---
        if "reality" in forbidden:
            reality_patterns = [
                re.compile(r"\b(?P<kw>reality)\s+principle\b", re.IGNORECASE),
                re.compile(r"\bobjective\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\bsocial\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\bmaterial\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\bvirtual\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\bconstructed\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\blived\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(r"\bon-the-ground\s+(?P<kw>reality)\b", re.IGNORECASE),
                re.compile(
                    r"\bappearance\s+and\s+(?P<kw>reality)\b",
                    re.IGNORECASE,
                ),
                # the reality of X, but avoid cosmic X like "life", "existence", etc.
                # No required "the" so we allow "reality of adolescent hunger" etc.
                re.compile(
                    r"\b(?P<kw>reality)\s+of\s+(?P<obj>[A-Za-z][A-Za-z-]*)\b",
                    re.IGNORECASE,
                ),
            ]
            for pat in reality_patterns:
                for m in pat.finditer(flat_text):
                    if "obj" in m.groupdict():
                        obj = m.group("obj").lower()
                        if obj in abstract_general_nouns:
                            continue
                    general_allowed_positions["reality"].add(m.start("kw"))

        # --- TRUTH exceptions ---
        if "truth" in forbidden:
            truth_patterns = [
                re.compile(
                    r"\bthe\s+(?P<kw>truth)\s+of\s+the\s+matter\b",
                    re.IGNORECASE,
                ),
                re.compile(r"\b(?P<kw>truth)\s+value\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>truth)\s+claims\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>truth)\s+conditions\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>truth)\s+content\b", re.IGNORECASE),
                re.compile(r"\btell[- ]tale\s+(?P<kw>truth)\b", re.IGNORECASE),
                re.compile(r"\bhard\s+(?P<kw>truth)\b", re.IGNORECASE),
                re.compile(
                    r"\b(?P<kw>truth)\s+and\s+reconciliation\b",
                    re.IGNORECASE,
                ),
                # the truth behind X, again with a concrete-ish X
                re.compile(
                    r"\bthe\s+(?P<kw>truth)\s+behind\s+(?P<obj>[A-Za-z][A-Za-z-]*)\b",
                    re.IGNORECASE,
                ),
            ]
            for pat in truth_patterns:
                for m in pat.finditer(flat_text):
                    if "obj" in m.groupdict():
                        obj = m.group("obj").lower()
                        if obj in abstract_general_nouns:
                            continue
                    general_allowed_positions["truth"].add(m.start("kw"))

        # --- LIFE exceptions ---
        if "life" in forbidden:
            life_patterns = [
                re.compile(r"\beveryday\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\binner\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\bprivate\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\bpublic\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\bpolitical\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\bsocial\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>life)\s+cycle\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>life)\s+expectancy\b", re.IGNORECASE),
                re.compile(r"\b(?P<kw>life)\s+stages\b", re.IGNORECASE),
                re.compile(r"\bway\s+of\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(r"\bstages\s+of\s+(?P<kw>life)\b", re.IGNORECASE),
                re.compile(
                    r"\b(?P<kw>life)\s*[- ]and[- ]death\b",
                    re.IGNORECASE,
                ),
                re.compile(
                    r"\bthe\s+(?P<kw>life)\s+of\s+the\s+mind\b",
                    re.IGNORECASE,
                ),
            ]
            for pat in life_patterns:
                for m in pat.finditer(flat_text):
                    general_allowed_positions["life"].add(m.start("kw"))

        # --- SOCIETY exceptions ---
        if "society" in forbidden:
            society_patterns = [
                re.compile(r"\bpatriarchal\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bindustrial\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bconsumer\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bmodern\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bcontemporary\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bcivil\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(r"\bhigh\s+(?P<kw>society)\b", re.IGNORECASE),
                re.compile(
                    r"\b(?P<kw>society)\s+at\s+large\b",
                    re.IGNORECASE,
                ),
            ]
            for pat in society_patterns:
                for m in pat.finditer(flat_text):
                    general_allowed_positions["society"].add(m.start("kw"))

            # Generic "specific society" patterns:
            #   - adjective + society  ("American society")
            #   - society of X         ("the society of Shanghai in the 1920s")
            adj_society_pattern = re.compile(
                r"\b(?P<adj>[A-Za-z][A-Za-z-]*)\s+(?P<kw>society)\b",
                re.IGNORECASE,
            )
            for m in adj_society_pattern.finditer(flat_text):
                adj = m.group("adj").lower()
                # Skip bare determiners/pronouns like "the society", "our society"
                if adj in {
                    "a", "an", "the",
                    "this", "that", "these", "those",
                    "any", "some", "each", "every",
                    "many", "few", "several", "no",
                    "our", "my", "your", "their", "his", "her", "its",
                }:
                    continue
                general_allowed_positions["society"].add(m.start("kw"))

            society_of_pattern = re.compile(
                r"\b(?P<kw>society)\s+of\s+(?P<obj>[A-Za-z][A-Za-z-]*)\b",
                re.IGNORECASE,
            )
            for m in society_of_pattern.finditer(flat_text):
                general_allowed_positions["society"].add(m.start("kw"))

        # --- UNIVERSE exceptions ---
        if "universe" in forbidden:
            universe_patterns = [
                re.compile(r"\bobservable\s+(?P<kw>universe)\b", re.IGNORECASE),
                re.compile(r"\bknown\s+(?P<kw>universe)\b", re.IGNORECASE),
                re.compile(r"\bnarrative\s+(?P<kw>universe)\b", re.IGNORECASE),
                re.compile(r"\bdiegetic\s+(?P<kw>universe)\b", re.IGNORECASE),
                re.compile(r"\bmarvel\s+(?P<kw>universe)\b", re.IGNORECASE),
                re.compile(
                    r"\b(?P<kw>universe)\s+of\s+discourse\b",
                    re.IGNORECASE,
                ),
            ]
            for pat in universe_patterns:
                for m in pat.finditer(flat_text):
                    general_allowed_positions["universe"].add(m.start("kw"))

            # Generic "specific universe" patterns:
            #   - adjective + universe ("fictional universe", "Marvel universe")
            #   - universe of X        ("universe of the NBA")
            adj_universe_pattern = re.compile(
                r"\b(?P<adj>[A-Za-z][A-Za-z-]*)\s+(?P<kw>universe)\b",
                re.IGNORECASE,
            )
            for m in adj_universe_pattern.finditer(flat_text):
                adj = m.group("adj").lower()
                if adj in {
                    "a", "an", "the",
                    "this", "that", "these", "those",
                    "any", "some", "each", "every",
                    "many", "few", "several", "no",
                    "our", "my", "your", "their", "his", "her", "its",
                }:
                    continue
                general_allowed_positions["universe"].add(m.start("kw"))

            universe_of_pattern = re.compile(
                r"\b(?P<kw>universe)\s+of\s+(?P<obj>[A-Za-z][A-Za-z-]*)\b",
                re.IGNORECASE,
            )
            for m in universe_of_pattern.finditer(flat_text):
                general_allowed_positions["universe"].add(m.start("kw"))

        # --- HUMAN exceptions ---
        # Allow the fixed phrase "human rights" even though "human" is normally forbidden.
        if "human" in forbidden:
            human_rights_pattern = re.compile(
                r"\b(?P<kw>human)\s+rights\b",
                re.IGNORECASE,
            )
            for m in human_rights_pattern.finditer(flat_text):
                general_allowed_positions.setdefault("human", set()).add(m.start("kw"))

        # Precompute positions where 'very' is allowed in fixed idioms like
    # "the very beginning", "the very end", "the very fact that", etc.
    very_ok_pattern = re.compile(
        r"\b(?:the\s+)?(?P<very>very)\s+("
        r"outset|beginning|end|moment|instant|idea|thought|"
        r"fact\s+that|same|heart\s+of|center|core|essence|"
        r"reason|point|place|man|person|thing|process"
        r")\b",
        re.IGNORECASE,
    )
    allowed_very_positions = {
        m.start("very") for m in very_ok_pattern.finditer(flat_text)
    }

    # Also allow "very" when it directly modifies a noun:
    # e.g. "the very process", "this very idea".
    for tok in doc:
        if tok.text.lower() == "very":
            # Skip 'very' inside direct quotations – those are already exempted earlier
            if pos_in_spans(tok.idx, spans) or pos_in_spans(tok.idx + len(tok.text) - 1, spans):
                continue

            # Look at the next token; if it's a noun or proper noun, allow this "very"
            if tok.i + 1 < len(doc):
                nxt = doc[tok.i + 1]
                if nxt.pos_ in {"NOUN", "PROPN"}:
                    allowed_very_positions.add(tok.idx)


    fw_pattern = r"\b(" + "|".join(map(re.escape, forbidden.keys())) + r")\b"
    fw_regex = re.compile(fw_pattern, re.IGNORECASE)

    for match in fw_regex.finditer(flat_text):
        match_start = match.start()
        match_end = match.end()
        # Skip forbidden-term marking inside ANY quotation (BEFORE any other checks)
        if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
            continue

        raw = match.group(0)          # preserve original casing
        word = raw.lower()

        # Treat all‑caps "US" as an acronym, not the pronoun "us"
        if word == "us" and raw.isupper():
            continue

        # Skip technical / idiomatic uses of general words like
        # 'reality', 'truth', 'life', 'society', and 'universe'
        if word in general_allowed_positions and match_start in general_allowed_positions[word]:
            continue

        # Allow fixed idioms like "the very beginning", "the very end",
        # "the very fact that", etc. Do NOT flag those uses of "very".
        if word == "very" and match_start in allowed_very_positions:
            continue

        rule_note = forbidden[word]
        if rule_note not in labels_used:
            marks.append({
                "start": match_start,
                "end": match_end,
                "note": rule_note,
                "color": WD_COLOR_INDEX.GRAY_25,
                "label": True,
            })
            labels_used.append(rule_note)
        else:
            marks.append({
                "start": match_start,
                "end": match_end,
                "note": rule_note,
                "color": WD_COLOR_INDEX.GRAY_25,
            })
     
        # -----------------------
    # PHASE 1.5 — SUBJECT–VERB AGREEMENT (experimental)
    # -----------------------
    """
    if getattr(config, "enforce_sva_rule", False):
        rule_note_sva = "Check subject–verb agreement"
        rule_note_sva_short = "s-v"

        def classify_subject_number(tok):
            # Pronouns first
            if tok.pos_ == "PRON":
                lower = tok.text.lower()
                if lower in {"he", "she", "it", "this", "that"}:
                    return "sing"
                if lower in {"they", "we", "these", "those"}:
                    return "plur"
            # Nouns by tag
            if tok.tag_ in {"NN", "NNP"}:
                return "sing"
            if tok.tag_ in {"NNS", "NNPS"}:
                return "plur"
            return None

        for sent in doc.sents:
            # Only check sentences that actually contain a verb
            if not any(t.pos_ in {"VERB", "AUX"} for t in sent):
                continue

            for tok in sent:
                # Look for nominal subjects
                if tok.dep_ not in {"nsubj", "nsubjpass"}:
                    continue
                subj = tok
                verb = subj.head

                if verb.pos_ not in {"VERB", "AUX"}:
                    continue

                # Don't flag things inside quotations
                subj_start = subj.idx
                subj_end = subj.idx + len(subj.text)
                verb_start = verb.idx
                verb_end = verb.idx + len(verb.text)
                if (
                    pos_in_spans(subj_start, spans)
                    or pos_in_spans(subj_end - 1, spans)
                    or pos_in_spans(verb_start, spans)
                    or pos_in_spans(verb_end - 1, spans)
                ):
                    continue

                # 1) Subject number from tags/pronouns
                subj_num = classify_subject_number(subj)

                # 2) Subject & verb number from morphology (if available)
                subj_num_feat = subj.morph.get("Number")
                subj_num_morph = subj_num_feat[0].lower() if subj_num_feat else None
                if subj_num_morph in {"sing", "plur"}:
                    # Prefer explicit morph number if present
                    subj_num = subj_num or subj_num_morph

                verb_num_feat = verb.morph.get("Number")
                verb_num = verb_num_feat[0].lower() if verb_num_feat else None

                verb_tag = verb.tag_
                lemma = verb.lemma_.lower()

                mismatch = False

                # Primary check: if both subject and verb have a Number feature, use that
                if subj_num in {"sing", "plur"} and verb_num in {"sing", "plur"}:
                    if subj_num != verb_num:
                        mismatch = True
                else:
                    # Fallback heuristic for present-tense main verbs
                    if subj_num == "sing":
                        if lemma != "be" and verb_tag == "VBP":
                            mismatch = True
                        elif lemma == "be" and verb_num == "plur":
                            mismatch = True
                    elif subj_num == "plur":
                        if lemma != "be" and verb_tag == "VBZ":
                            mismatch = True
                        elif lemma == "be" and verb_num == "sing":
                            mismatch = True

                if not mismatch:
                    continue

                # Highlight the verb as the error locus
                start = verb_start
                end = verb_end

                if rule_note_sva not in labels_used:
                    # First occurrence: full label text
                    marks.append({
                        "start": start,
                        "end": end,
                        "note": rule_note_sva,
                        "color": GRAMMAR_ORANGE,
                        "label": True,
                    })
                    labels_used.append(rule_note_sva)
                else:
                    # Subsequent occurrences: short display text "s-v",
                    # but still use the full note for linking/summary.
                    marks.append({
                        "start": start,
                        "end": end,
                        "note": rule_note_sva,
                        "display_note": rule_note_sva_short,
                        "color": GRAMMAR_ORANGE,
                        "label": True,
                    })

     if getattr(config, "enforce_present_tense_rule", False):
     rule_note_tense = "Write in the present tense"
     rule_note_tense_short = "tense"

+    def preceded_by_that(token, max_back: int = 3) -> bool:
+        # Simple heuristic: if "that" appears immediately before the verb
+        # (or a couple tokens back, e.g., "that he was"), ignore it.
+        for j in range(1, max_back + 1):
+            k = token.i - j
+            if k < 0:
+                break
+            prev = doc[k]
+            if prev.is_punct:
+                break
+            if prev.lower_ == "that":
+                return True
+        return False

     for tok in doc:
         # VBD = simple past (walked, said, was, were, had, did...)
         if tok.tag_ != "VBD":
             continue
         if tok.pos_ not in ("VERB", "AUX"):
             continue
+
+        # NEW: Ignore VBD verbs preceded by "that"
+        if preceded_by_that(tok):
+            continue

         start = tok.idx
         end = tok.idx + len(tok.text)

         # Ignore past tense inside direct quotations
         if pos_in_spans(start, spans) or pos_in_spans(end - 1, spans):
             continue

         ...
"""


    # -----------------------
    # PHASE 2 — CONTRACTIONS
    # -----------------------
    if getattr(config, "enforce_contractions_rule", True):
        contractions = {
            "don't", "doesn't", "didn't",
            "can't", "couldn't", "won't", "wouldn't", "shouldn't",
            "isn't", "aren't", "wasn't", "weren't",
            "hasn't", "haven't", "hadn't",
            "mustn't", "mightn't", "shan't",
            "it's", "that's", "there's", "what's", "who's", "where's",
            "when's", "why's", "how's",
            "i'm", "you're", "we're", "they're", "he's", "she's",
            "i've", "you've", "we've", "they've",
            "i'd", "you'd", "he'd", "she'd", "we'd", "they'd",
            "i'll", "you'll", "he'll", "she'll", "we'll", "they'll",
            "let's",
            "could've", "would've", "should've", "must've",
            "ain't",
        }

        contractions_note = "No contractions in academic writing"

        contr_pattern = r"\b(" + "|".join(map(re.escape, contractions)) + r")\b"
        contr_regex = re.compile(contr_pattern, re.IGNORECASE)

        for match in contr_regex.finditer(flat_text):
            match_start = match.start()
            match_end = match.end()

            # Ignore contractions inside direct quotations
            if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
                continue

            if contractions_note not in labels_used:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": contractions_note,
                    "color": WD_COLOR_INDEX.GRAY_25,
                    "label": True,
                    })
                labels_used.append(contractions_note)
            else:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": contractions_note,
                    "color": WD_COLOR_INDEX.GRAY_25,
                })

    # -----------------------
    # PHASE 2.1 — Article errors (a/an)
    # -----------------------
    article_regex = re.compile(r"\b(a|an)\s+([A-Za-z])", re.IGNORECASE)
    for match in article_regex.finditer(flat_text):
        article = match.group(1) or ""
        next_letter = match.group(2) or ""
        if not article or not next_letter:
            continue

        article_lower = article.lower()
        next_lower = next_letter.lower()
        should_be_an = next_lower in {"a", "e", "i", "o", "u"}
        is_error = (article_lower == "a" and should_be_an) or (article_lower == "an" and not should_be_an)
        if not is_error:
            continue

        start = match.start(1)
        end = match.end(1)

        # Ignore article mistakes inside direct quotations
        if pos_in_spans(start, spans) or pos_in_spans(end - 1, spans):
            continue

        if ARTICLE_ERROR_LABEL not in labels_used:
            marks.append({
                "start": start,
                "end": end,
                "note": ARTICLE_ERROR_LABEL,
                "color": WD_COLOR_INDEX.GRAY_25,
                "label": True,
            })
            labels_used.append(ARTICLE_ERROR_LABEL)
        else:
            marks.append({
                "start": start,
                "end": end,
                "note": ARTICLE_ERROR_LABEL,
                "color": WD_COLOR_INDEX.GRAY_25,
            })

    # -----------------------
    # PHASE 5A — Delete-phrases
    # -----------------------
    delete_phrases = [
        "vividly",
        "vivid",
        "in conclusion",
        "all in all",
        "in summary",
        "to conclude",
        "to summarize",
        "the use of",
    ]
    delete_phrases = sorted(
        set(p.strip() for p in delete_phrases if p.strip()),
        key=len,
        reverse=True,
    )

    delete_pattern = re.compile(
        r"\b(" + "|".join(re.escape(p) for p in delete_phrases) + r")\b",
        re.IGNORECASE,
    )

    # These should only be deleted in the conclusion paragraph
    conclusion_only_delete_phrases = {
        "in conclusion",
        "all in all",
        "in summary",
        "to conclude",
        "to summarize",
    }

    rule_description = ""

    for match in delete_pattern.finditer(flat_text):
        match_start, match_end = match.start(1), match.end(1)

        # Phrase text in lowercase so we can compare
        phrase_text = match.group(1).lower()

        # Only delete these in the conclusion paragraph
        if phrase_text in conclusion_only_delete_phrases and paragraph_role != "conclusion":
            continue

        # Skip phrases that fall inside direct quotation spans
        if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
            continue

        # Ensure this issue appears in the summary table, but we *don't*
        # print a yellow inline label in the text.
        if rule_description not in labels_used:
            labels_used.append(rule_description)

        marks.append({
            "start": match_start,
            "end": match_end,
            "note": rule_description,
            "color": WD_COLOR_INDEX.RED,
            "strike": True,
            # no "label": this prevents a yellow arrow comment
        })

    # -----------------------
    # PHASE 5A.1 — Logical connectors: therefore/thereby/hence/thus
    # -----------------------
    rule_note_logical = "Avoid the words 'therefore', 'thereby', 'hence', and 'thus'"
    logical_regex = re.compile(r"\b(therefore|thereby|hence|thus)\b", re.IGNORECASE)

    for match in logical_regex.finditer(flat_text):
        match_start, match_end = match.start(1), match.end(1)

        # Skip matches inside direct quotations
        if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
            continue

        # Ensure this rule appears in the summary table, but do NOT create a yellow inline label
        if rule_note_logical not in labels_used:
            labels_used.append(rule_note_logical)

        marks.append({
            "start": match_start,
            "end": match_end,
            "note": rule_note_logical,
            "color": WD_COLOR_INDEX.RED,
            "strike": True,
            # no "label": key here – we only want red strikethrough in the text
        })

    # -----------------------
    # PHASE 5B — TEXT-AS-TEXT RULES
    # -----------------------
    text_as_text_phrases = [
        "in this paragraph",
        "in the paragraph",
        "this paragraph",
        "in this sentence",
        "in the sentence",
        "this sentence",
        "in this quotation",
        "in the quotation",
        "this quotation",
        "in this passage",
        "in the passage",
        "this passage",
        "in this essay",
        "in the essay",
        "within the reading",
        "throughout the essay",
        "throughout the article",
        "throughout the short story",
        "throughout the novel",
        "throughout the story",
        "throughout the poem",
        "throughout the narrative",
        "throughout the passage",
        "through this essay",
        "throughout the text",
        "in the text",
        "in this quote",
        "the quote",
        "the text",
        "the paragraph",
        "the passage",
        "quote",
        "quotation",
        "paragraphs",
    ]

    rule_note_text_as_text = "Do not refer to the text as a text; refer to context instead"

    for phrase in text_as_text_phrases:
        pattern = r'\b' + re.escape(phrase) + r'\b'
        phrase_regex = re.compile(pattern, re.IGNORECASE)

        for match in phrase_regex.finditer(flat_text):
            match_start = match.start()
            match_end = match.end()

            # Skip inside quotations
            if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
                continue

            # Allow references to the essay in the FIRST sentence only
            # (students often write "In this essay..." as part of genre naming)
            if sentences:
                sent_idx = get_sentence_index_for_pos(match_start, sentences)
                if sent_idx == 0 and "essay" in phrase.lower():
                    continue

            # Skip the text-as-text rule in the THESIS sentence for
            # any phrase containing "quote" or "quotation", and also
            # for phrases that are thesis devices.
            if paragraph_role == "intro" and sentences:
                thesis_start, thesis_end = sentences[-1]
                if thesis_start <= match_start < thesis_end:
                    lower_phrase = phrase.lower()

                    # Allow "quote"/"quotation" words in the thesis sentence
                    # (e.g. "a quote of X", "a key quotation of Y") without
                    # triggering the text-as-text rule.
                    if "quote" in lower_phrase or "quotation" in lower_phrase:
                        continue

                    # Preserve previous behaviour: skip if phrase is also
                    # a thesis device word.
                    if lower_phrase in THESIS_DEVICE_WORDS:
                        continue

            if rule_note_text_as_text not in labels_used:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": rule_note_text_as_text,
                    "color": WD_COLOR_INDEX.GRAY_25,
                    "label": True,
                })
                labels_used.append(rule_note_text_as_text)
            else:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": rule_note_text_as_text,
                    "color": WD_COLOR_INDEX.GRAY_25,
                })



    # -----------------------
    # PHASE 6 — WEAK VERBS
    # -----------------------
    if getattr(config, "enforce_weak_verbs_rule", True):
        weak_verbs_regex = re.compile(r"\b(show|shows|showing|use|uses|using)\b", re.IGNORECASE)

        rule_note_weak_verbs = "Avoid weak verbs"

        for match in weak_verbs_regex.finditer(flat_text):
            match_start = match.start()
            match_end = match.end()

            # Special case: "the use of" is handled as a delete-phrase with
            # red strikethrough, so we do NOT also flag "use" as a weak verb there.
            word_lower = match.group(0).lower()
            if word_lower.startswith("use"):
                phrase_start = match_start - 4  # position of "the "
                phrase_end = match_end + 3      # position after " of"
                if phrase_start >= 0 and phrase_end <= len(flat_text):
                    if flat_text[phrase_start:phrase_end].lower() == "the use of":
                        continue

            # Skip forbidden-term marking inside ANY quotation (BEFORE any other checks)
            if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
                continue

            if rule_note_weak_verbs not in labels_used:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": rule_note_weak_verbs,
                    "color": WD_COLOR_INDEX.TURQUOISE,
                    "label": True
                })
                labels_used.append(rule_note_weak_verbs)
            else:
                marks.append({
                    "start": match_start,
                    "end": match_end,
                    "note": rule_note_weak_verbs,
                    "color": WD_COLOR_INDEX.TURQUOISE
                })

    # -----------------------
    # PHASE 7 — NUMBER RULE (1–10)
    # -----------------------
    number_regex = re.compile(r"\b(1|2|3|4|5|6|7|8|9|10)\b")

    rule_note_number = "Write out the numbers one through ten"

    def is_parenthetical_citation(flat: str, start: int, end: int) -> bool:
        """
        Return True if the number at [start:end) is part of a parenthetical
        citation, including MLA forms like:

            (1), (1-3), (1, 2, 3)
            (Kristof 4), (Baron 12-13), (Kristof and Smith 4-5)

        Any number that lives between matching parentheses and is preceded
        (inside the same parentheses) by an author-ish chunk but not followed
        by more letters is treated as a citation and exempt from the
        1–10 spelling rule.
        """
        n = len(flat)

        # Find the '(' that starts this parenthetical
        left = flat.rfind("(", 0, start)
        if left == -1:
            return False

        # Make sure there isn't a ')' between '(' and the number
        if flat.rfind(")", left, start) != -1:
            return False

        # Find the closing ')'
        right = flat.find(")", end)
        if right == -1:
            return False

        inside = flat[left + 1:right].strip()

        # Must contain at least one digit
        if not re.search(r"\d", inside):
            return False

        # Purely numeric parenthetical: (1), (1-3), (1, 2, 3)
        if re.fullmatch(r"[0-9][0-9\s,.-]*", inside):
            return True

        # MLA-style: some non-digit content (author name, etc.) followed by
        # a page number, with no letters after the page number.
        first_digit = re.search(r"\d", inside)
        if not first_digit:
            return False

        before = inside[: first_digit.start()]
        after = inside[first_digit.start():]

        # Require at least one letter before the digits (author-ish chunk)
        if not re.search(r"[A-Za-z]", before):
            return False

        # After the page number we only allow digits and basic separators,
        # not more letters.
        if re.search(r"[A-Za-z]", after):
            return False

        return True

    for match in number_regex.finditer(flat_text):
        match_start = match.start()
        match_end = match.end()

        # Skip numbers inside direct quotations
        if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
            continue

        # Skip parenthetical citations like (1) or (1-3)
        if is_parenthetical_citation(flat_text, match_start, match_end):
            continue

        if rule_note_number not in labels_used:
            marks.append({
                "start": match_start,
                "end": match_end,
                "note": rule_note_number,
                "color": WD_COLOR_INDEX.GRAY_25,
                "label": True
            })
            labels_used.append(rule_note_number)
        else:
            marks.append({
                "start": match_start,
                "end": match_end,
                "note": rule_note_number,
                "color": WD_COLOR_INDEX.GRAY_25
            })

    # -----------------------
    # PHASE 9 — UNCOUNTABLE NOUNS
    # -----------------------
    # Conservative detection of a few high‑value uncountable nouns being treated as countable.
    # We only flag:
    #   - plural forms (e.g. "evidences")
    #   - or use with "counting" determiners/numbers (many/two/a/an),
    # and we skip anything inside direct quotations.
    uncountable_lemmas = {"evidence", "imagery", "research", "information", "advice", "diction", "jargon"}
    uncountable_note = "Uncountable noun"

    for token in doc:
        lemma = token.lemma_.lower()
        if lemma not in uncountable_lemmas:
            continue

        # Calculate character span of this token in flat_text
        tok_start = token.idx
        tok_end = token.idx + len(token.text)

        # Ignore anything inside direct quotations
        if pos_in_spans(tok_start, spans) or pos_in_spans(tok_end - 1, spans):
            continue

        is_error = False

        # 1) Plural form of an uncountable noun (e.g. "evidences", "researches")
        if token.tag_ in ("NNS", "NNPS"):
            is_error = True

        # Helper: safe previous token
        prev_token = doc[token.i - 1] if token.i > 0 else None

        # 2) Count determiners like "many", "few", "several" directly attached or just before
        if not is_error:
            # Check dependency children
            for child in token.children:
                if child.dep_ == "det" and child.text.lower() in {"many", "few", "several"}:
                    is_error = True
                    break
            # Check immediate left neighbor
            if not is_error and prev_token is not None and prev_token.text.lower() in {"many", "few", "several"}:
                is_error = True

        # 3) Numeric determiners: e.g. "two evidence", "3 research"
        if not is_error and prev_token is not None and prev_token.like_num:
            is_error = True

        # 4) Indefinite article "a"/"an" directly attached or just before
        if not is_error:
            for child in token.children:
                if child.dep_ == "det" and child.text.lower() in {"a", "an"}:
                    is_error = True
                    break
            if not is_error and prev_token is not None and prev_token.text.lower() in {"a", "an"}:
                is_error = True

        if not is_error:
            continue

        # Add the mark, with a yellow label only on the first occurrence
        if uncountable_note not in labels_used:
            marks.append({
                "start": tok_start,
                "end": tok_end,
                "note": uncountable_note,
                "color": WD_COLOR_INDEX.GRAY_25,
                "label": True,
            })
            labels_used.append(uncountable_note)
        else:
            marks.append({
                "start": tok_start,
                "end": tok_end,
                "note": uncountable_note,
                "color": WD_COLOR_INDEX.GRAY_25,
            })

    # -----------------------
    # PHASE 8 — WEAK TRANSITION DETECTION
    # -----------------------
    rule_note_weak_transition = "Use a boundary statement when transitioning between paragraphs"

    if paragraph_role == "body" and sentences:
        # Check multi-word transitions first (longer phrases first to avoid partial matches)
        weak_transitions_multi_sorted = sorted(WEAK_TRANSITIONS_MULTI, key=len, reverse=True)

        matched = False
        match_start = None
        match_end = None
        matched_phrase = None

        # Check if topic sentence starts with any multi-word transition
        topic_start, topic_end = sentences[0]
        topic_text = flat_text[topic_start:topic_end]
        trimmed_text = topic_text.lstrip()
        leading_whitespace_len = len(topic_text) - len(trimmed_text)

        for transition in weak_transitions_multi_sorted:
            transition_lower = transition.lower()
            if trimmed_text.lower().startswith(transition_lower):
                transition_len = len(transition)
                if transition_len >= len(trimmed_text) or not trimmed_text[transition_len].isalnum():
                    match_start = topic_start + leading_whitespace_len
                    match_end = match_start + transition_len
                    matched = True
                    matched_phrase = transition
                    break

        # If no multi-word match, check single-word transitions against first token
        if not matched and tokens:
            first_token = None
            for tok_text, tok_start, tok_end in tokens:
                if tok_start < topic_start:
                    continue
                if tok_start >= topic_end:
                    break
                first_token = (tok_text, tok_start, tok_end)
                break

            if first_token:
                first_text = first_token[0].lower()
                if first_text in WEAK_TRANSITIONS_SINGLE:
                    match_start = first_token[1]
                    match_end = first_token[2]
                    matched = True
                    matched_phrase = first_token[0]

        # If a match was found at the start, mark it (unless boundary overlap exists)
        if matched and match_start is not None and match_end is not None:
            # Skip forbidden-term marking inside ANY quotation (BEFORE any other checks)
            if pos_in_spans(match_start, spans) or pos_in_spans(match_end - 1, spans):
                pass  # Skip marking
            else:
                weak_transition_lemmas = set()
                if matched_phrase:
                    for tok in nlp(matched_phrase):
                        if tok.is_space or tok.is_punct:
                            continue
                        lemma = tok.lemma_.lower().strip()
                        if lemma:
                            weak_transition_lemmas.add(lemma)

                topic_content = extract_content_lemmas(
                    doc,
                    topic_start,
                    topic_end,
                    extra_exclude=weak_transition_lemmas,
                )
                has_overlap = False
                if prev_body_last_sentence_content_words:
                    has_overlap = bool(prev_body_last_sentence_content_words & topic_content)

                if not has_overlap:
                    if rule_note_weak_transition not in labels_used:
                        marks.append({
                            "start": match_start,
                            "end": match_end,
                            "note": rule_note_weak_transition,
                            "color": WD_COLOR_INDEX.GRAY_25,
                            "label": True,
                        })
                        labels_used.append(rule_note_weak_transition)
                    else:
                        marks.append({
                            "start": match_start,
                            "end": match_end,
                            "note": rule_note_weak_transition,
                            "color": WD_COLOR_INDEX.GRAY_25,
                        })

    # =====================================================================
    # FOUNDATION ASSIGNMENT 1 — FILTER MARKS IN EXTRA SENTENCES
    # =====================================================================
    # For Foundation Assignment 1, we must prevent normal analytic issue labels
    # from appearing on extra sentences in the intro paragraph.
    #
    # Strategy:
    #   - Keep ALL marks that touch the first sentence (normal analysis)
    #   - Keep our Foundation 1 red-strike marks (added earlier in this function)
    #   - Drop any other marks that only appear in extra sentences
    #     (e.g., "Avoid the word 'which'", "Off-topic", etc.)
    #
    # Note: The single yellow label for Foundation 1 is added later in run_marker,
    # after all paragraphs are processed, so there's no label to filter here.
    # =====================================================================
    if config.mode == "foundation_1" and paragraph_role == "intro" and sentences:
        first_start, first_end = sentences[0]

        filtered_marks: list[dict] = []

        for m in marks:
            start = m.get("start", 0)
            end = m.get("end", 0)

            # Identify our Foundation 1 red-strike marks by their characteristics:
            # red highlight + strikethrough, with no note attached
            is_assignment_strike = (
                m.get("strike")
                and m.get("note") is None
                and m.get("color") == WD_COLOR_INDEX.RED
            )

            if end <= first_end:
                # Mark touches the first sentence → keep for normal analysis
                filtered_marks.append(m)
            elif is_assignment_strike:
                # Our Foundation 1 red-strike marks → keep (they're in extra sentences)
                filtered_marks.append(m)
            else:
                # Regular analytic mark in an extra sentence → drop it
                # (e.g., weak verb, off-topic, quotation rule, etc.)
                continue

        marks = filtered_marks

    # =====================================================================
    # FOUNDATION ASSIGNMENT 2 — ONLY FIRST SENTENCE + THESIS ARE MARKED
    # =====================================================================
    # For Foundation 2, the intro should contain:
    #   - the first sentence (context, author, title, genre, etc.)
    #   - a closed thesis (last sentence of the intro)
    #
    # We still allow students to write extra sentences, but we do not want
    # normal analytic rules firing on those middle sentences. So:
    #   - keep marks that touch the first sentence or the thesis sentence
    #   - drop marks that live entirely inside any "middle" sentences.
    if config.mode == "foundation_2" and paragraph_role == "intro" and sentences:
        first_start, first_end = sentences[0]
        thesis_start, thesis_end = sentences[-1]

        allowed_ranges = [
            (first_start, first_end),
            (thesis_start, thesis_end),
        ]

        def overlaps_allowed(start: int, end: int) -> bool:
            # Anchor-only marks (start == end) are treated as overlapping if
            # their anchor lies inside either the first sentence or the thesis.
            if start == end:
                return any(a_start <= start <= a_end for (a_start, a_end) in allowed_ranges)
            # Normal span overlap check
            return any(start < a_end and end > a_start for (a_start, a_end) in allowed_ranges)

        filtered_marks: list[dict] = []

        for m in marks:
            start = m.get("start", 0)
            end = m.get("end", 0)

            if overlaps_allowed(start, end):
                filtered_marks.append(m)

        marks = filtered_marks

    def ensure_student_labels_per_sentence(marks_in, sentence_spans):
        """
        In student mode, ensure each sentence shows at least one yellow label
        per issue note, without repeating the same label within the sentence.
        """
        if not sentence_spans or not marks_in:
            return marks_in

        groups: dict[tuple[int, str], list[dict]] = {}

        for m in marks_in:
            note = m.get("note")
            if not note:
                continue
            if m.get("praise") or m.get("device_highlight"):
                continue

            start = m.get("start", 0)
            sent_idx = get_sentence_index_for_pos(start, sentence_spans)
            if sent_idx is None:
                if start < sentence_spans[0][0]:
                    sent_idx = 0
                elif start >= sentence_spans[-1][1]:
                    sent_idx = len(sentence_spans) - 1
                else:
                    best_idx = 0
                    best_dist = None
                    for idx, (s_start, s_end) in enumerate(sentence_spans):
                        if start < s_start:
                            dist = s_start - start
                        elif start >= s_end:
                            dist = start - s_end
                        else:
                            dist = 0
                        if best_dist is None or dist < best_dist:
                            best_dist = dist
                            best_idx = idx
                    sent_idx = best_idx

            key = (sent_idx, str(note))
            groups.setdefault(key, []).append(m)

        for group_marks in groups.values():
            if any(mark.get("label") for mark in group_marks):
                continue
            def mark_sort_key(mark):
                end = mark.get("end", mark.get("start", 0))
                start = mark.get("start", 0)
                return (end, start)
            best_mark = max(group_marks, key=mark_sort_key)
            best_mark["label"] = True

        return marks_in

    if getattr(config, "student_mode", False):
        marks = ensure_student_labels_per_sentence(marks, sentences)

    return marks, flat_text, segments, sentences, last_sentence_content_words


# Constants for quotation label insertion
CLOSING_QUOTE_CHARS = {'"', '"', "'", "'"}
TRAILING_QUOTE_PUNCT = {",", ".", "!", "?", ";", ":"}


def apply_marks(paragraph, flat_text, segments, marks, sentences=None, paragraph_index=None):
    """
    Rebuild `paragraph` from `flat_text` and a list of `marks`.

    Each mark is a dict with at least:

        - "start": start index in flat_text

        - "end": end index in flat_text

    Optional keys:

        - "note": label text

        - "label": bool, whether to attach a yellow label arrow

        - "color": a WD_COLOR_INDEX value (or None for no highlight)

        - "strike": bool, whether to apply strikethrough

    This version:

      * Safely handles anchor-only marks that sit past the end of the text.

      * Keeps Times New Roman 12pt via enforce_font.

      * Moves quotation-rule labels to AFTER the closing quote.

      * Preserves italic formatting from the original student text character-by-character.

    """
    global DOC_EXAMPLES, DOC_EXAMPLE_COUNTS, DOC_EXAMPLE_SENT_HASHES
    
    # Collect example sentences from marks before mutating paragraph
    sentences_for_examples = sentences
    if sentences_for_examples is None and paragraph_index is not None:
        # Title paragraphs pass sentences=None; treat whole paragraph as one sentence
        if flat_text and len(flat_text) > 0:
            sentences_for_examples = [(0, len(flat_text))]

    if sentences_for_examples is not None and paragraph_index is not None:
        for mark in marks:
            # Collect examples for ANY mark that has a real, non-empty note string
            note = mark.get("note")
            if not note:
                continue
            
            # Skip the "The title of major works should be italicized" label
            if note == "The title of major works should be italicized":
                continue
            
            # If text is empty, don't try to clamp / index
            if not flat_text:
                continue
            
            # Find the sentence containing the mark
            mark_start = mark.get("start", 0)
            # Clamp to valid range
            mark_start = max(0, min(mark_start, len(flat_text) - 1))

            # Avoid trailing whitespace: prefer a safe position inside text
            trimmed_len = len(flat_text.rstrip())
            if trimmed_len <= 0:
                continue
            safe_pos = min(mark_start, max(0, trimmed_len - 1))
            while safe_pos > 0 and flat_text[safe_pos].isspace():
                safe_pos -= 1

            # Use get_sentence_index_for_pos to locate sentence span
            sent_idx = get_sentence_index_for_pos(safe_pos, sentences_for_examples)
            if sent_idx is None:
                fallback_pos = max(0, trimmed_len - 1)
                if fallback_pos != safe_pos:
                    sent_idx = get_sentence_index_for_pos(fallback_pos, sentences_for_examples)
            if sent_idx is None:
                continue
            
            # Extract sentence text
            s_start, s_end = sentences_for_examples[sent_idx]
            sentence_text = flat_text[s_start:s_end].strip()
            
            # Normalize whitespace: collapse \s+ to single spaces
            sentence_text = re.sub(r'\s+', ' ', sentence_text)
            
            # Cap at 500 chars
            if len(sentence_text) > 500:
                sentence_text = sentence_text[:500]
            
            # Dedupe: compute md5(sentence_text) and use a set key (note, md5) in DOC_EXAMPLE_SENT_HASHES
            sentence_hash = hashlib.md5(sentence_text.encode('utf-8')).hexdigest()
            hash_key = (note, sentence_hash)
            if hash_key in DOC_EXAMPLE_SENT_HASHES:
                continue  # Skip if already seen
            
            # Cap: only store up to MAX_EXAMPLES_PER_LABEL per label using DOC_EXAMPLE_COUNTS
            current_count = DOC_EXAMPLE_COUNTS.get(note, 0)
            if current_count >= MAX_EXAMPLES_PER_LABEL:
                continue  # Skip if we've reached the cap for this label
            
            # Append example
            DOC_EXAMPLES.append({
                "label": note,
                "sentence": sentence_text,
                "paragraph_index": paragraph_index,
            })
            # Update counts and hashes
            DOC_EXAMPLE_COUNTS[note] = current_count + 1
            DOC_EXAMPLE_SENT_HASHES.add(hash_key)
    
    def append_text_with_italics(
        paragraph,
        flat_text: str,
        start: int,
        end: int,
        italic_spans: list,
        *,
        color=None,
        strike: bool = False,
    ):
        """
        Append runs for flat_text[start:end], splitting them so that italic
        formatting exactly matches the original italic_spans. Optionally apply a
        highlight color and/or strikethrough across the whole region.
        """
        pos = start
        while pos < end:
            # Should this position be italic?
            italic_here = any(i_start <= pos < i_end for (i_start, i_end) in italic_spans)

            # Find the next position where the italic state changes, or we hit `end`
            next_pos = end
            for (i_start, i_end) in italic_spans:
                if italic_here:
                    # We're currently in an italic span: boundary is its end
                    if pos < i_end <= next_pos:
                        next_pos = i_end
                else:
                    # We're currently non-italic: boundary is the start of the next italic span
                    if pos < i_start < next_pos:
                        next_pos = i_start

            chunk = flat_text[pos:next_pos]
            r = paragraph.add_run(chunk)
            enforce_font(r)

            # Apply highlight
            if color is not None:
                if color == GRAMMAR_ORANGE:
                    # Grammar issues: DARK BLUE highlight + WHITE font (Word-safe)
                    r.font.highlight_color = WD_COLOR_INDEX.DARK_BLUE
                    r.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    r.font.highlight_color = color

            # Preserve italics from original student text (this must NOT depend on strike)
            if italic_here:
                r.font.italic = True

            # Optional strikethrough
            if strike:
                r.font.strike = True

            # Mark Vysti-generated runs (useful for later passes)
            if color is not None or strike:
                r._element.set("data-vysti", "yes")

            pos = next_pos

    
    # STEP 1: Before clearing runs, capture which character ranges in flat_text were italic
    original_italic_spans = []
    for run_idx, seg_start, seg_end in segments:
        run = paragraph.runs[run_idx]
        if run_is_italic(run):
            original_italic_spans.append((seg_start, seg_end))
    
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # If there are no marks, just write the plain text with italics preserved
    if not marks:
        append_text_with_italics(paragraph, flat_text, 0, len(flat_text), original_italic_spans)
        return

    text_len = len(flat_text)

    def marks_overlap(a, b) -> bool:
        """
        Return True if mark spans a and b should be merged.
        We merge when:
          - they have real character overlap, OR
          - they have exactly the same [start, end] span (same anchor).
        We do NOT merge when they only touch at a boundary, e.g.
            a: [145, 146), b: [146, 173)
        or when an anchor-only mark at position k just sits at the edge
        of a longer span starting at k.
        """
        a_start, a_end = a["start"], a["end"]
        b_start, b_end = b["start"], b["end"]
        # Identical span (including anchor-only marks at the same position)
        if a_start == b_start and a_end == b_end:
            return True
        # Check for zero-length marks (anchors)
        a_zero = a_start == a_end
        b_zero = b_start == b_end
        # Zero-length marks should not merge with longer spans that only touch at a boundary
        # A zero-length mark at position P should not merge with a span [P, Q) that starts at P
        if a_zero or b_zero:
            # Only merge zero-length marks if there's actual character overlap
            # (i.e., the other mark contains the anchor position in its interior, not at its start/end)
            if a_zero:
                # a is zero-length at position a_start: b must contain it in its interior
                return b_start < a_start < b_end
            else:  # b_zero
                # b is zero-length at position b_start: a must contain it in its interior
                return a_start < b_start < a_end
        # Real overlap (non-empty intersection) for non-zero-length marks
        # Use strict < to avoid merging marks that only touch at boundaries
        return (a_start < b_end) and (b_start < a_end)

    # Sort and gently merge overlapping marks
    sorted_marks = sorted(marks, key=lambda m: (m["start"], m["end"]))
    merged_marks = []

    for m in sorted_marks:
        m = m.copy()
        if not merged_marks:
            merged_marks.append(m)
            continue
        cur = merged_marks[-1]
        if marks_overlap(cur, m):
            # Overlapping marks – merge spans and metadata
            cur["end"] = max(cur["end"], m["end"])
            # Prefer to keep any existing note/label; otherwise adopt new one
            if m.get("label"):
                cur["label"] = True
                if not cur.get("note") and m.get("note"):
                    cur["note"] = m["note"]
            elif m.get("note") and not cur.get("note"):
                cur["note"] = m["note"]
            # Strikethrough if either mark wants it
            if m.get("strike"):
                cur["strike"] = True
            # Preserve an explicit color if we don't already have one
            if cur.get("color") is None and m.get("color") is not None:
                cur["color"] = m["color"]
        else:
            merged_marks.append(m)

    cursor = 0
    for mark in merged_marks:
        start = mark["start"]
        end = mark["end"]
        note = mark.get("note")
        color = mark.get("color", None)
        is_label = bool(mark.get("label"))
        is_quote_label = bool(note and is_label and "quotation" in note.lower())

        # Clamp into the bounds of flat_text; anchors past the end become len(flat_text)
        mark_start = max(0, min(start, text_len))
        mark_end = max(0, min(end, text_len))

        # Unmarked text before this mark
        if cursor < mark_start:
            append_text_with_italics(
                paragraph,
                flat_text,
                cursor,
                mark_start,
                original_italic_spans,
                color=None,
                strike=False,
            )

        cursor = mark_start

        # Marked span (if any characters are actually covered)
        marked_color = mark.get("color", None)

        # Only allow strikethrough for "delete" style issues (red highlight),
        # never for gray structural/title highlights.
        raw_strike = bool(mark.get("strike"))
        marked_strike = raw_strike and marked_color == WD_COLOR_INDEX.RED

        if mark_start < mark_end:
            append_text_with_italics(
                paragraph,
                flat_text,
                mark_start,
                mark_end,
                original_italic_spans,
                color=marked_color,
                strike=marked_strike,
            )
            cursor = mark_end

        # Handle labels
        if note and is_label:
            # Praise labels (e.g. "Good paragraph.") get green; all others stay yellow.
            label_color = (
                WD_COLOR_INDEX.BRIGHT_GREEN
                if mark.get("praise")
                else WD_COLOR_INDEX.YELLOW
            )

            if is_quote_label:
                # For quotation rules, move the label AFTER the closing quote if it is next
                if cursor < len(flat_text) and flat_text[cursor] in CLOSING_QUOTE_CHARS:
                    append_text_with_italics(
                        paragraph,
                        flat_text,
                        cursor,
                        cursor + 1,
                        original_italic_spans,
                        color=None,
                        strike=False,
                    )
                    cursor += 1

                    # Optional but recommended: keep punctuation immediately after the closing quote attached
                    while cursor < len(flat_text) and flat_text[cursor] in TRAILING_QUOTE_PUNCT:
                        append_text_with_italics(
                            paragraph,
                            flat_text,
                            cursor,
                            cursor + 1,
                            original_italic_spans,
                            color=None,
                            strike=False,
                        )
                        cursor += 1

            # Build the label run
            # NOTE: Label runs are Vysti-generated and should NOT inherit student italics,
            # so we do NOT pass them through append_text_with_italics
            display_note = mark.get("display_note", note)
            lbl = paragraph.add_run(f" → {display_note}")

            enforce_font(lbl)
            lbl.font.bold = True
            lbl.font.highlight_color = label_color

            # Force all labels to use black text and no underline
            lbl.font.color.rgb = RGBColor(0, 0, 0)
            lbl.font.underline = False

            # For yellow issue labels, keep the hyperlink to the Issue row
            if not mark.get("praise"):
                bmk_name = bookmark_name_for_label(note)
                wrap_run_in_internal_link(paragraph, lbl, bmk_name)

            lbl._element.set("data-vysti", "yes")

    # Any remaining unmarked text after the last mark
    if cursor < text_len:
        append_text_with_italics(
            paragraph,
            flat_text,
            cursor,
            text_len,
            original_italic_spans,
            color=None,
            strike=False,
        )


def add_summary_table(doc, labels, rules, issue_counts=None):

    """
    Append a summary section of all yellow-labeled rules used in the document.

    `labels` is the list of rule-label strings that were actually used
    (labels_used), and `rules` is the dict returned by load_rules:
        short_label -> explanation
    """
    # Deduplicate while preserving order
    unique_labels = []
    for lbl in labels:
        if lbl and lbl not in unique_labels:
            unique_labels.append(lbl)

    # If there are no issue labels, do NOT add any summary block at all.
    if not unique_labels:
        return

    # Sort issues alphabetically by their first word ("Avoid", "Clarify", etc.)
    unique_labels_sorted = sorted(unique_labels, key=lambda s: s.split()[0].lower())

    # Spacer before summary
    doc.add_paragraph()

    # -----------------------
    # Summary table: Issue | Explanation
    # With a blank row after each Issue/Explanation pair for visual separation.
    # -----------------------
    num_pairs = len(unique_labels_sorted)
    # We want:
    #   header row
    #   blank spacer row after header
    #   (issue/explanation row + blank row) * N
    total_rows = num_pairs * 2 + 2
    table = doc.add_table(rows=total_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Make Issue column narrower than Explanation column
    table.autofit = False
    issue_width = Inches(1.8)
    expl_width = Inches(4.8)

    # python-docx is a little fussy: set both column and cell widths
    for col, width in zip(table.columns, (issue_width, expl_width)):
        col.width = width
        for cell in col.cells:
            cell.width = width

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Issue"
    hdr[1].text = "Explanation"

    # Make sure the Issue header is left-aligned
    for p in hdr[0].paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Data rows, with one blank row after each for spacing
    row_idx = 2   # instead of 1
    for lbl in unique_labels_sorted:
        issue_cell, expl_cell = table.rows[row_idx].cells

        # ISSUE cell: build with runs so we can attach a bookmark
        issue_p = issue_cell.paragraphs[0]
        issue_p.text = ""
        issue_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        issue_run = issue_p.add_run(lbl)
        enforce_font(issue_run)
        cnt = (issue_counts or {}).get(lbl, 0)
        cnt_run = issue_p.add_run(f" ({cnt})")
        enforce_font(cnt_run)


        # Attach a bookmark to this Issue paragraph so yellow labels can link to it
        bmk_name = bookmark_name_for_label(lbl)
        add_bookmark_to_paragraph(issue_p, bmk_name)

        # EXPLANATION cell
        explanation = rules.get(lbl, "")

        # Special case: make the Power Verbs explanation a hyperlink
        if lbl in ("Avoid weak verbs", "Refer to the Power Verbs list"):
            expl_p = expl_cell.paragraphs[0]
            expl_p.text = ""

            link_text = explanation or "Download the Power Verbs here"
            run = expl_p.add_run(link_text)
            enforce_font(run)

            wrap_run_in_external_link(
                expl_p,
                run,
                "https://www.vysti.org/resources",
            )
        else:
            expl_cell.text = explanation if explanation is not None else ""

        # Next row (row_idx + 1) is intentionally left blank as a spacer
        row_idx += 2

    # Ensure Times New Roman 12pt for everything in the table,
    # and make the header row bold.
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    enforce_font(run)
                    if r_idx == 0:
                        run.font.bold = True


def title_needs_capitalization_fix(title_str: str) -> bool:
    """
    Return True if the essay title looks like it has obvious capitalization problems.

    For creative titles of the form:
        "Stylish quotation": Topic phrase in "Young Hunger"
    we ONLY check the Topic phrase (the middle part), not:
      - the stylistic quotation at the start
      - the underlying work title in the final quotes.

    Heuristics:
      - Extract alphabetic words.
      - First word must be capitalized.
      - Interior words must be capitalized unless they are small function words
        (the, of, in, etc.).
      - We do NOT force the final word to be capitalized if it's one of those
        small function words.
    """
    if not title_str:
        return False

    # Same pattern used elsewhere
    m = TITLE_PATTERN.match(title_str)
    if m:
        # Only enforce title case on the topic segment (middle part)
        target = m.group(2).strip()
    else:
        # Fallback: apply to the whole string
        target = title_str

    words = re.findall(r"[A-Za-z]+(?:'[A-Za-z]+)?", target)
    if not words:
        return False

    ALWAYS_LOWER = {
        "a", "an", "the",
        "and", "but", "or", "nor", "for", "so", "yet",
        "of", "in", "on", "at", "by", "to", "up",
        "from", "into", "onto", "over", "under",
        "with", "within", "without", "about",
        "between", "before", "after", "around",
        "through", "during", "inside", "outside",
        "as", "than", "via", "per",
        "is", "are", "was", "were", "be", "been", "being", "am",
        "do", "does", "did",
        "has", "have", "had",
        "can", "could", "shall", "should",
        "will", "would", "may", "might", "must",
    }

    for i, w in enumerate(words):
        lower = w.lower()

        if i == 0:
            # First word: always needs to be capitalized
            if not w[0].isupper():
                return True
        else:
            # Other words: only enforce caps on content words
            if lower in ALWAYS_LOWER:
                continue
            if not w[0].isupper():
                return True

    return False


def topic_segment_is_too_thin(topic_segment: str) -> bool:
    """
    Return True if the 'topic' part of a creative title is basically empty –
    e.g. just 'in', 'of', 'the', etc., with no real content word.
    """
    if not topic_segment:
        return True

    words = re.findall(r"[A-Za-z]+", topic_segment)
    if not words:
        return True

    ALWAYS_LOWER = {
        "a", "an", "the",
        "and", "but", "or", "nor", "for", "so", "yet",
        "of", "in", "on", "at", "by", "to", "up",
        "from", "into", "onto", "over", "under",
        "with", "within", "without", "about",
        "between", "before", "after", "around",
        "through", "during", "inside", "outside",
        "as", "than", "via", "per",
        "is", "are", "was", "were", "be", "been", "being", "am",
        "do", "does", "did",
        "has", "have", "had",
        "can", "could", "shall", "should",
        "will", "would", "may", "might", "must",
    }

    # Content words = everything that is not just a tiny function word
    content_words = [w for w in words if w.lower() not in ALWAYS_LOWER]
    return not content_words


def is_probable_title_paragraph(paragraph, config: MarkerConfig | None = None):
    """
    Decide if a paragraph is likely the student's title line.

    We treat something as a title if EITHER:
      1. It matches the creative pattern `"Quote": Topic in "Title"` (with colon)
         or `"Quote" Topic "Title"` (without colon), where the
         second quoted string is the teacher-supplied title (case-insensitive), or
      2. It looks like a standalone heading such as "Homework #1":
         - short line,
         - no terminal sentence punctuation,
         - and either:
             • truly centered via Word alignment, OR
             • manually centered via 2+ leading TABs.
    """
    # Raw Word text, including true tabs at the margin
    raw = paragraph.text or ""

    # Count leading tabs before the first non-whitespace character
    leading_tabs = 0
    for ch in raw:
        if ch == "\t":
            leading_tabs += 1
        elif ch in (" ", "\u00A0"):
            # allow spaces before the first non-space char
            continue
        else:
            break

    # 1 tab = normal indent, 2+ tabs = student trying to center manually
    tab_centered = leading_tabs >= 2

    # Use flattened text that strips labels and normalizes quotes/dashes, etc.
    flat_text, _ = flatten_paragraph_without_labels(paragraph)
    # For title detection, ignore all leading tabs (titles shouldn't "look indented")
    text = normalize_leading_whitespace(flat_text, strip_all_tabs_for_title=True).strip()
    if not text:
        return False

    # Strip any existing Vysti label if we're re-running on a marked doc
    if "→" in text:
        text = text.split("→", 1)[0].strip()
        if not text:
            return False

    # Get normalized teacher titles for fuzzy matching
    config_title_keys = get_config_title_keys(config)
    if not config_title_keys:
        # No configured titles, so we can't match creative title patterns
        # but can still match heading-style titles below
        pass
    else:
        # 1) Strong match: `"Quote": Topic in "Title"` (with colon)
        m = TITLE_PATTERN.match(text)
        if m:
            # The third capture group is the underlying source title in quotes
            source_title_quoted = m.group(3)
            # Try fuzzy matching against all configured titles
            for config_title_key in config_title_keys:
                # Use title_similarity for fuzzy matching
                # First normalize both for comparison
                source_title_normalized = normalize_title_key(source_title_quoted)
                if source_title_normalized:
                    sim = title_similarity(source_title_quoted, config_title_key)
                    if sim >= 0.6:  # Reasonable threshold for fuzzy matching
                        return True

        # 1b) Strong match: `"Quote" Topic "Title"` (without colon)
        # This handles creative titles that omit the colon but still have the
        # clear structure of two quoted segments where the final one is the teacher-supplied title.
        m_no_colon = TITLE_PATTERN_NO_COLON.match(text)
        if m_no_colon:
            # The third capture group is the underlying source title in quotes
            source_title_quoted = m_no_colon.group(3)
            # Try fuzzy matching against all configured titles
            for config_title_key in config_title_keys:
                source_title_normalized = normalize_title_key(source_title_quoted)
                if source_title_normalized:
                    sim = title_similarity(source_title_quoted, config_title_key)
                    if sim >= 0.6:  # Reasonable threshold for fuzzy matching
                        return True

    # 2) Secondary heuristics for simple headings like "Homework #1"

    # 2a) Short, no sentence-ending punctuation, and either Word-centered
    #     or manually centered with 2+ leading tabs.
    if len(text) <= 80 and text[-1] not in ".!?;:":
        if (
            paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            or tab_centered
        ):
            return True

    # 2b) Obvious homework-style labels such as "Homework #1", "Essay 2", "Assignment 4"
    lowered = text.lower()
    if re.match(r"^(homework|essay|paper|assignment)\b", lowered) and len(text) <= 40:
        return True

    return False


def is_source_reference_subtitle_line(text: str, config: MarkerConfig | None = None) -> bool:
    """
    Heuristic: detect lines like
        in Toni Morrison's "Strangers"
    that sit under the main creative title.

    We still want to treat them as part of the title *block* for intro
    detection, but we don't want to enforce the essay-title colon pattern
    or title-capitalization rules on them.
    """
    if not config:
        return False

    stripped = text.strip()
    if not stripped:
        return False
    # If it reads like a real sentence, it is NOT a subtitle line.
    if looks_like_full_sentence(stripped):
        return False

    # We only treat lines that *start* with "in " as subtitle/source lines.
    # (e.g., "in Toni Morrison's \"Strangers\"")
    lower = stripped.lower()
    if not lower.startswith("in "):
        return False

    # Must mention at least one teacher-supplied title
    norm_line = normalize_title_key(stripped)
    if not norm_line:
        return False

    for work in iter_teacher_works(config):
        key = normalize_title_key(work.title)
        if key and key in norm_line:
            return True

    return False


def looks_like_full_sentence(t: str) -> bool:
    """
    Heuristic: return True if this top-of-page line looks like a real sentence
    rather than an MLA header.

    We treat it as a sentence if:
      - it has at least ~8 alphabetic words, and
      - it contains at least one sentence-ending punctuation mark (. ? !)

    This prevents long introduction paragraphs that happen to contain
    words like 'professor' or a date from being misclassified as headers.
    """
    t = t.strip()
    if not t:
        return False

    # Count alphabetic word tokens
    words = re.findall(r"[A-Za-z]+", t)
    if len(words) < 8:
        return False

    # Any sentence-ending punctuation at all makes this look like a sentence
    quote_spans = compute_quote_spans(t)
    for i, ch in enumerate(t):
        if ch in ".?!" and not pos_in_spans(i, quote_spans):
            return True
    return False



def is_probable_mla_header_line(text: str, config: MarkerConfig | None = None) -> bool:
    """
    Robustly detect MLA-style header lines (student name, teacher, class, date).
    Returns True for likely MLA header lines, False for actual title/intro/essay content.
    
    Rules:
    1. Pre-cleaning: Strip whitespace and Vysti labels (→ and after)
    2. Date lines: Detect various date formats (e.g., "8 November", "November 8, 2025", "08/11/2025")
    3. Teacher lines: Detect titles like Mr., Ms., Mrs., Dr., Prof., Professor
    4. Student name lines: 1-4 word-like tokens, capitalized, no digits, not the teacher-supplied title
    5. Do NOT detect full sentences with punctuation as MLA headers
    
    Only used on the first few paragraphs of the document.
    """
    # Pre-cleaning: strip whitespace
    t = text.strip()
    if not t:
        return False

    # Strip any existing Vysti label (anything after →) so re-runs on marked docs still work
    if "→" in t:
        t = t.split("→", 1)[0].strip()
        if not t:
            return False

    # NEW: If this line already looks like a full sentence, it should not be
    # treated as an MLA header, even if it happens to contain words like
    # 'professor', a date, or course terms. This protects real introductions.
    if looks_like_full_sentence(t):
            return False

    lower = t.lower()
    
    # Get normalized teacher title for matching
    config_title_keys = get_config_title_keys(config)
    normalized_config_title = next(iter(config_title_keys)) if config_title_keys else None

    # ====================================================================
    # DATE LINES
    # ====================================================================
    # Detect various date formats:
    #   8 November, 8th November, November 8, November 8th
    #   November 8, 2025, 8 November 2025, 8th November 2025, November 8th 2025
    #   08/11/2025, 11/08/2025
    
    month = r"(january|february|march|april|may|june|july|august|september|october|november|december)"
    day = r"\d{1,2}(st|nd|rd|th)?"
    year = r"\d{4}"

    date_patterns = [
        rf"\b{day}\s+{month}\b",                # 8 November, 8th November
        rf"\b{day}\s+{month}\s+{year}\b",       # 8 November 2025, 8th November 2025
        rf"\b{month}\s+{day}\b",                # November 8, November 8th
        rf"\b{month}\s+{day},?\s+{year}\b",     # November 8, 2025 or November 8th 2025
        r"\b\d{1,2}/\d{1,2}/\d{4}\b",           # 08/11/2025, 11/08/2025
    ]

    for pat in date_patterns:
        if re.search(pat, lower):
            return True

    # ====================================================================
    # SHORT SLASH-DATE LINES (WITHOUT 4-DIGIT YEAR)
    # ====================================================================
    # Some students write dates like "9/20" or "09/20" at the top of the document,
    # often combined with their name (e.g., "Liang Hong 9/20").
    # Detect lines containing a pattern like \d{1,2}/\d{1,2} when:
    #   - The line is reasonably short (<= 60 characters)
    #   - The line has no sentence-ending punctuation (., ?, !)
    # This ensures we don't accidentally treat a real sentence like
    # "9/11 changed the American imagination." as a header.
    if re.search(r'\b\d{1,2}/\d{1,2}\b', t):
        if len(t) <= 60 and not any(ch in t for ch in ".?!"):
            return True

    # ====================================================================
    # TERM + YEAR LINES
    # ====================================================================
    # Detect lines with both a term word and a 4-digit year (20xx)
    # Examples: "Spring 2025", "2025 Fall", "FoundB-Analytical Writing 2025 Spring"
    term_pattern = r"\b(spring|summer|fall|autumn|winter)\b"
    year_pattern = r"\b20\d{2}\b"
    
    if re.search(term_pattern, lower) and re.search(year_pattern, lower):
        # Must be short and have no sentence-ending punctuation
        if len(t) <= 60 and not any(ch in t for ch in ".?!"):
            return True

    # ====================================================================
    # COURSE/CLASS LINES
    # ====================================================================
    # Detect course/class lines with academic terms
    # Examples: "FoundB-Analytical Writing", "English 101", "AP Literature", "IB Lang HL"
    course_pattern = (
        r"\b(english|history|biology|chemistry|physics|math|algebra|geometry|calculus|"
        r"foundation|bfoundation|foundb|ap|ib|period|writing|analysis|analytical|"
        r"literature|language|lang|composition|comp|honors|hl|sl|class)\b"
    )
    
    if re.search(course_pattern, lower):
        # Must be short and have no sentence-ending punctuation
        if len(t) <= 60 and not any(ch in t for ch in ".?!"):
            return True

    # ====================================================================
    # TEACHER LINES
    # ====================================================================
    # Detect teacher name lines: Mr., Ms., Mrs., Dr., Prof., Professor
    if re.search(r"\b(mr\.|ms\.|mrs\.|dr\.|prof\.|professor)\b", lower):
        return True

    # ====================================================================
    # STUDENT NAME LINES
    # ====================================================================
    # Detect student's name as:
    #   - 1-4 "word-like" tokens
    #   - No digits
    #   - Not equal to any known work title
    #   - Each token is either:
    #       • A capitalized word (Wendi, Zhou), or
    #       • An initial (W.)
    
    words = t.split()
    if 1 <= len(words) <= 4 and not any(ch.isdigit() for ch in t):
        # Avoid treating the teacher-supplied work title as a header line
        if normalized_config_title is not None and normalize_title_key(t) == normalized_config_title:
            return False

        def is_simple_name_token(w: str) -> bool:
            """Check if a token looks like a name or initial."""
            if not w:
                return False
            # Allow initials like "W."
            if len(w) == 2 and w[0].isupper() and w[1] == ".":
                return True
            # Capitalized word: first letter uppercase, rest lowercase (or single letter)
            # Examples: "Wendi", "Zhou", "A"
            if w[0].isupper() and (len(w) == 1 or w[1:].islower()):
                return True
            return False

        # Only check alphabetic tokens (ignore punctuation-only tokens if any)
        alpha_tokens = [w for w in words if any(ch.isalpha() for ch in w)]
        if alpha_tokens and all(is_simple_name_token(w) for w in alpha_tokens):
            return True

    # ====================================================================
    # OTHERWISE
    # ====================================================================
    # Do not detect full sentences with punctuation as MLA headers.
    # Those should go through normal title/intro logic.
    return False


def detect_mla_header_indices(real_paragraphs, config: MarkerConfig | None = None):
    """
    Return a set of paragraph indices (in the real_paragraphs list) that form
    a top-of-document MLA-style header block (name, teacher, course, date, etc.).
    
    Algorithm:
    1. Only check the top of the document (new_idx <= 5)
    2. For each paragraph, get text via flatten_paragraph (includes tabs, actual Word text)
    3. Call is_probable_mla_header_line to check if it's a header line
    4. If True: add to header_indices
    5. If blank: skip (continue)
    6. Never break early - always check all top paragraphs
    
    Args:
        real_paragraphs: List of (old_idx, paragraph) pairs
        config: MarkerConfig to access teacher-supplied title
        
    Returns:
        Set of new_idx values that are part of the MLA header block
    """
    header_indices = set()

    for new_idx, (old_idx, p) in enumerate(real_paragraphs):
        # Only consider the very top of the document
        if new_idx > 5:
            break

        # Get text via flatten_paragraph so we see the actual Word text
        # (including any tabs, but normalization happens in is_probable_mla_header_line)
        flat_text, _ = flatten_paragraph(p, skip_vysti=False)
        
        # Skip blank lines
        if flat_text.strip() == "":
            continue
        
        # Check if this is a header line and add to indices if so
        if is_probable_mla_header_line(flat_text, config=config):
            header_indices.add(new_idx)

    return header_indices


def extract_summary_metadata(doc: Document) -> dict:
    """
    Read the summary table that add_summary_table() injects at the end of the document
    and return it as a simple JSON-friendly structure.

    Scans tables from the END backwards and picks the first table whose header row
    is exactly "Issue" and "Explanation".

    Returns:
        {
            "issues": [
                {"label": "...", "explanation": "...", "count": N},
                ...
            ]
        }
    """
    issues: list[dict] = []

    # No tables? No summary.
    if not doc.tables:
        return {"issues": issues}

    # Scan from the end backwards to find the summary table
    for table in reversed(doc.tables):
        # Sanity check: header row should say "Issue | Explanation"
        if len(table.rows) < 2 or len(table.rows[0].cells) < 2:
            continue

        header_issue = table.rows[0].cells[0].text.strip().lower()
        header_expl = table.rows[0].cells[1].text.strip().lower()
        if header_issue == "issue" and header_expl == "explanation":
            # Found the summary table! Parse it.
            # Data rows start at row index 2, with a blank row after each pair
            for row_idx in range(2, len(table.rows), 2):
                row = table.rows[row_idx]
                if len(row.cells) < 2:
                    continue

                raw = row.cells[0].text.strip()
                expl = row.cells[1].text.strip()

                # Parse trailing " (N)" without polluting the label
                m = re.match(r"^(.*?)(?:\s*\((\d+)\))?$", raw)
                label = (m.group(1) or "").strip() if m else raw
                count = int(m.group(2)) if (m and m.group(2)) else 0

                if label:
                    issues.append({"label": label, "explanation": expl, "count": count})

            # Return immediately after finding and parsing the first matching table
            return {"issues": issues}

    # No matching table found
    return {"issues": issues}


def strip_summary_table_in_place(doc: Document) -> bool:
    """
    Remove the trailing summary table whose header is exactly: Issue | Explanation.
    Returns True if removed, False if not found.
    """
    target = None

    # Find the summary table from the end (same logic as extract_summary_metadata)
    for table in reversed(doc.tables):
        try:
            if len(table.rows) < 1 or len(table.rows[0].cells) < 2:
                continue
            h0 = table.rows[0].cells[0].text.strip().lower()
            h1 = table.rows[0].cells[1].text.strip().lower()
            if h0 == "issue" and h1 == "explanation":
                target = table
                break
        except Exception:
            continue

    if target is None:
        return False

    # Remove the table XML element
    tbl_el = target._element
    parent = tbl_el.getparent()
    if parent is not None:
        parent.remove(tbl_el)

    # add_summary_table() also adds a spacer paragraph before the table.
    # After table removal, that spacer is usually at the end. Remove a small
    # number of trailing blank paragraphs (don't nuke the whole ending).
    for _ in range(3):
        if doc.paragraphs and not (doc.paragraphs[-1].text or "").strip():
            p = doc.paragraphs[-1]
            p_el = p._element
            p_parent = p_el.getparent()
            if p_parent is not None:
                p_parent.remove(p_el)
        else:
            break

    return True


def extract_richer_examples(doc: Document) -> list[dict]:
    """
    Post-process the marked document to extract richer example snippets.
    
    Scans the marked docx for label runs starting with " → " and builds
    context-aware snippets based on label type.
    
    Returns:
        List of { "label": str, "sentence": str } dictionaries.
        Max 1 example per label per essay.
    """
    examples_map = {}  # label -> example dict (for de-duplication)
    
    # Track previous paragraph for boundary statement logic
    prev_paragraph_text = None
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        # Build clean paragraph text excluding label runs, while recording label positions
        clean_text_parts = []
        label_offsets = []  # List of (offset, label_text) tuples
        
        current_offset = 0
        
        for run in paragraph.runs:
            run_text = run.text
            
            # Check if this run starts with " → " (label run)
            if run_text.startswith(" → "):
                # Extract label text (everything after " → ")
                label_text = run_text[3:].strip()
                if label_text:
                    label_offsets.append((current_offset, label_text))
            else:
                # Regular text - add to clean text
                clean_text_parts.append(run_text)
                current_offset += len(run_text)
        
        clean_text = "".join(clean_text_parts).strip()
        
        if not clean_text or not label_offsets:
            prev_paragraph_text = clean_text
            continue
        
        # Use spaCy to segment sentences
        doc_spacy = nlp(clean_text)
        sentences = [(sent.start_char, sent.end_char) for sent in doc_spacy.sents]
        
        # Process each label in this paragraph
        for label_offset, label_text in label_offsets:
            # Skip if we already have an example for this label
            if label_text in examples_map:
                continue
            
            # Find the sentence containing the label offset
            # Labels are usually inserted at the end of marks, so check if offset is at or after sentence end
            offending_sent_idx = None
            for idx, (s_start, s_end) in enumerate(sentences):
                if s_start <= label_offset <= s_end:
                    offending_sent_idx = idx
                    break
            
            if offending_sent_idx is None:
                # Fallback: use last sentence if offset is after all sentences, or first if before
                if sentences:
                    if label_offset >= sentences[-1][1]:
                        offending_sent_idx = len(sentences) - 1
                    else:
                        offending_sent_idx = 0
                else:
                    continue
            
            # Extract snippet based on label type
            snippet = None
            
            # Normalize label for matching (case-insensitive)
            label_lower = label_text.lower()
            
            # Boundary statement: last sentence of previous + first sentence of current
            if "boundary statement" in label_lower and "transitioning" in label_lower and "between paragraphs" in label_lower:
                if prev_paragraph_text and prev_paragraph_text.strip():
                    # Get last sentence of previous paragraph
                    prev_doc_spacy = nlp(prev_paragraph_text)
                    prev_sentences = list(prev_doc_spacy.sents)
                    if prev_sentences:
                        last_prev_sent = prev_sentences[-1].text.strip()
                        # Get first sentence of current paragraph (where the label is)
                        if sentences:
                            first_curr_sent = clean_text[sentences[0][0]:sentences[0][1]].strip()
                            snippet = f"{last_prev_sent}\n\n{first_curr_sent}"
                # Fallback: just first sentence of current paragraph
                if not snippet and sentences:
                    snippet = clean_text[sentences[0][0]:sentences[0][1]].strip()
            
            # Evidence-related: middle sentences only
            elif "follow the process for inserting evidence" in label_lower or ("process" in label_lower and "inserting" in label_lower and "evidence" in label_lower):
                if len(sentences) >= 3:
                    # Middle sentences: all except first and last
                    middle_sents = []
                    for idx in range(1, len(sentences) - 1):
                        s_start, s_end = sentences[idx]
                        middle_sents.append(clean_text[s_start:s_end].strip())
                    snippet = " ".join(middle_sents)
                else:
                    # Fallback: whole paragraph if too short
                    snippet = clean_text
            
            # Every paragraph needs evidence: same as above
            elif "every paragraph" in label_lower and "evidence" in label_lower:
                if len(sentences) >= 3:
                    middle_sents = []
                    for idx in range(1, len(sentences) - 1):
                        s_start, s_end = sentences[idx]
                        middle_sents.append(clean_text[s_start:s_end].strip())
                    snippet = " ".join(middle_sents)
                else:
                    snippet = clean_text
            
            # Pronoun/antecedent: preceding sentence + offending sentence
            elif "pronoun" in label_lower and "antecedent" in label_lower:
                if offending_sent_idx > 0:
                    # Preceding sentence in same paragraph
                    prev_s_start, prev_s_end = sentences[offending_sent_idx - 1]
                    prev_sent = clean_text[prev_s_start:prev_s_end].strip()
                    off_s_start, off_s_end = sentences[offending_sent_idx]
                    off_sent = clean_text[off_s_start:off_s_end].strip()
                    snippet = f"{prev_sent} {off_sent}"
                elif prev_paragraph_text and prev_paragraph_text.strip():
                    # Try last sentence of previous paragraph
                    prev_doc_spacy = nlp(prev_paragraph_text)
                    prev_sentences = list(prev_doc_spacy.sents)
                    if prev_sentences:
                        last_prev_sent = prev_sentences[-1].text.strip()
                        off_s_start, off_s_end = sentences[offending_sent_idx]
                        off_sent = clean_text[off_s_start:off_s_end].strip()
                        snippet = f"{last_prev_sent} {off_sent}"
                # Fallback: just offending sentence
                if not snippet and sentences:
                    off_s_start, off_s_end = sentences[offending_sent_idx]
                    snippet = clean_text[off_s_start:off_s_end].strip()
            
            # Quotation rules (topic sentence / final sentence): just offending sentence
            elif ("no quotations" in label_lower or "no quotation" in label_lower) and ("topic sentence" in label_lower or ("final sentence" in label_lower and "body paragraph" in label_lower)):
                if sentences:
                    off_s_start, off_s_end = sentences[offending_sent_idx]
                    snippet = clean_text[off_s_start:off_s_end].strip()
            
            # Default: just offending sentence
            else:
                if sentences:
                    off_s_start, off_s_end = sentences[offending_sent_idx]
                    snippet = clean_text[off_s_start:off_s_end].strip()
            
            # Store example if we got a snippet
            if snippet:
                examples_map[label_text] = {
                    "label": label_text,
                    "sentence": snippet,
                }
        
        # Update previous paragraph text for next iteration
        prev_paragraph_text = clean_text
    
    return list(examples_map.values())


def build_techniques_discussed(docx_bytes: bytes, mode: str) -> list[dict]:
    if mode == "argumentation":
        return []
    if not THESIS_ALL_DEVICE_KEYS:
        return []
    try:
        base_doc = Document(BytesIO(docx_bytes))
        paragraphs = [p.text for p in base_doc.paragraphs if p.text and p.text.strip()]
        if not paragraphs:
            return []
        full_text = "\n".join(paragraphs).strip()
        if not full_text:
            return []

        doc_spacy = nlp(full_text)
        from collections import Counter
        counts = Counter()
        for device_key, _, _ in iter_device_spans(doc_spacy):
            if device_key in THESIS_ALL_DEVICE_KEYS:
                counts[device_key] += 1

        if not counts:
            return []

        ordered_keys = []
        for key in THESIS_TOPIC_ORDER:
            if key in counts and key not in ordered_keys:
                ordered_keys.append(key)
        remaining = sorted(
            (key for key in counts.keys() if key not in ordered_keys),
            key=lambda k: (-counts[k], k),
        )
        ordered_keys.extend(remaining)

        return [
            {"name": key.replace("_", " ").title(), "count": counts[key]}
            for key in ordered_keys
            if counts[key] > 0
        ]
    except Exception:
        return []


def mark_docx_bytes(
    docx_bytes: bytes,
    mode: str = "textual_analysis",
    teacher_config: dict | None = None,
    rules_path: str = "Vysti Rules for Writing.xlsx",
    include_summary_table: bool = True,
) -> tuple[bytes, dict]:
    """
    High-level engine API for web/backend use.

    - Does not depend on Tkinter or any GUI.
    - Accepts a .docx file as raw bytes.
    - Returns (marked_docx_bytes, metadata_dict).

    'mode' should match the MarkerConfig modes, e.g.:
        "textual_analysis", "intertextual_analysis", "reader_response",
        "no_title", "image_analysis", "argumentation",
        "foundation_1" ... "foundation_6", "peel_paragraph", etc.

    'teacher_config' can override fields on MarkerConfig, for example:
        {
            "author_name": "M. F. K. Fisher",
            "text_title": "Young Hunger",
            "text_is_minor_work": True,
        }
    """
    # 1. Build a MarkerConfig
    config = get_preset_config(mode)

    if teacher_config:
        for key, value in teacher_config.items():
            if hasattr(config, key):
                setattr(config, key, value)

    config.student_mode = not include_summary_table

    # 2. Write the uploaded bytes to a temporary .docx file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
        tmp_in.write(docx_bytes)
        tmp_in_path = tmp_in.name

    tmp_out_path = None

    try:
        # 3. Call the existing path-based engine
        #    (this keeps run_marker unchanged for CLI/legacy use)
        tmp_out_path = run_marker(tmp_in_path, rules_path=rules_path, config=config)

        # 4. Read the marked .docx back into memory
        with open(tmp_out_path, "rb") as f:
            marked_bytes = f.read()

        # 5. Load the marked doc into python-docx to extract the summary metadata
        doc = Document(BytesIO(marked_bytes))
        metadata = extract_summary_metadata(doc)
        techniques_discussed = build_techniques_discussed(docx_bytes, mode)
        if isinstance(metadata, dict):
            metadata["techniques_discussed"] = techniques_discussed
        try:
            guidance_map = load_student_guidance(rules_path)
        except Exception:
            guidance_map = {}
        if ARTICLE_ERROR_LABEL not in guidance_map:
            guidance_map[ARTICLE_ERROR_LABEL] = ARTICLE_ERROR_GUIDANCE
        for issue in metadata.get("issues", []):
            if isinstance(issue, dict):
                issue["student_guidance"] = guidance_map.get(issue.get("label"), "")
        
        # 6. Always use DOC_EXAMPLES (the multi-example list collected during marking)
        # DO NOT overwrite with extract_richer_examples() which can only capture
        # one example per label (from yellow label runs " → ") and misses all
        # occurrences that don't have yellow labels.
        global DOC_EXAMPLES
        metadata["examples"] = DOC_EXAMPLES if DOC_EXAMPLES else []

        # If the caller doesn't want the summary table inside the returned document
        # (Student Mode), strip it AFTER extracting metadata.
        if not include_summary_table:
            removed = strip_summary_table_in_place(doc)
            if removed:
                out = BytesIO()
                doc.save(out)
                marked_bytes = out.getvalue()

    finally:
        # 6. Clean up temp files as best we can
        try:
            os.remove(tmp_in_path)
        except OSError:
            pass
        if tmp_out_path is not None:
            try:
                os.remove(tmp_out_path)
            except OSError:
                pass

    return marked_bytes, metadata


def run_marker(
    essay_path: str,
    rules_path: str = "Vysti Rules for Writing.xlsx",
    config: MarkerConfig | None = None,
) -> str:
    """
    Runs the Vysti marker on the given essay and returns the path
    to the saved *_marked.docx file.
    """
    # Reset global state for this document
    print("Vysti marker: audience/use-of/red-label version loaded")
    global THESIS_DEVICE_SEQUENCE, THESIS_TOPIC_ORDER, BODY_PARAGRAPH_COUNT, BRIDGE_PARAGRAPHS, BRIDGE_DEVICE_KEYS
    global BOOKMARK_ID_COUNTER, FOUNDATION1_LABEL_TARGET
    global THESIS_PARAGRAPH_INDEX, THESIS_ANCHOR_POS, THESIS_ALL_DEVICE_KEYS, THESIS_TEXT_LOWER
    global DOC_EXAMPLES, DOC_EXAMPLE_COUNTS, DOC_EXAMPLE_SENT_HASHES

    THESIS_DEVICE_SEQUENCE = []
    THESIS_TOPIC_ORDER = []
    THESIS_ALL_DEVICE_KEYS = set()
    THESIS_TEXT_LOWER = ""
    BODY_PARAGRAPH_COUNT = 0
    BRIDGE_PARAGRAPHS = set()
    BRIDGE_DEVICE_KEYS = {}
    BOOKMARK_ID_COUNTER = 1
    FOUNDATION1_LABEL_TARGET = None
    THESIS_PARAGRAPH_INDEX = None
    THESIS_ANCHOR_POS = None
    DOC_EXAMPLES = []
    DOC_EXAMPLE_COUNTS = {}
    DOC_EXAMPLE_SENT_HASHES = set()
    
    if config is None:
        # Default behavior remains the existing full analytic mode
        config = get_preset_config("textual_analysis")
    
    rules = load_rules(rules_path)
    if ARTICLE_ERROR_LABEL not in rules:
        rules[ARTICLE_ERROR_LABEL] = ARTICLE_ERROR_EXPLANATION
    doc = Document(essay_path)
    labels_used = []

    from collections import Counter
    issue_counts = Counter()


        # ------------------------------------------------------------------
    # TITLE PRE-PROCESSING:
    #   1) Split combined title+intro when the title and intro are in
    #      a single centered paragraph.
    #   2) Flatten multi-paragraph title block (title + subtitle lines).
    # ------------------------------------------------------------------
    tmp_real_paragraphs = [
        (i, p) for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]

    tmp_header_indices = detect_mla_header_indices(tmp_real_paragraphs, config=config)

    # ---------- (1) Split combined title + intro in one centered paragraph ----------
    for new_idx, (old_idx, p) in enumerate(tmp_real_paragraphs):
        if new_idx in tmp_header_indices:
            continue

        # Look only at non-empty paragraphs
        flat_text, _ = flatten_paragraph_without_labels(p)
        text = flat_text.strip()
        if not text:
            continue

        # Quick centering heuristic: true center alignment OR manual 2+ tab center
        raw = p.text or ""
        leading_tabs = 0
        for ch in raw:
            if ch == "\t":
                leading_tabs += 1
            elif ch in (" ", "\u00A0"):
                continue
            else:
                break
        tab_centered = leading_tabs >= 2
        is_centered = (
            p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            or tab_centered
        )

        # We only split when the student has effectively "made everything the title"
        # by centering it.
        if not is_centered:
            continue

        # Does this paragraph START with a creative title pattern?
        m_prefix = TITLE_PREFIX_PATTERN.match(text)
        if not m_prefix:
            m_prefix = TITLE_PREFIX_NO_COLON_PATTERN.match(text)
        if not m_prefix:
            continue

        end_idx = m_prefix.end()
        # If there's nothing after the creative title, it's a normal title line.
        if end_idx >= len(text):
            continue

        # Split into title_text (kept in this paragraph) and intro_text (new paragraph)
        title_text = text[:end_idx].rstrip()
        intro_text = text[end_idx:].lstrip()
        if not intro_text:
            continue

        # Overwrite the original paragraph so it only contains the title.
        # (We don't try to preserve fine-grained run formatting here; the
        #  marker will re-mark any text-title formatting issues anyway.)
        p.text = title_text

        # Insert a new paragraph *after* this one with the intro content.
        parent = p._element.getparent()
        new_p = OxmlElement("w:p")

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = intro_text
        r.append(t)
        new_p.append(r)

        parent.insert(parent.index(p._element) + 1, new_p)

        # We only need to fix the first combined title+intro; bail out.
        break

    # Rebuild after possible split
    tmp_real_paragraphs = [
        (i, p) for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]
    tmp_header_indices = detect_mla_header_indices(tmp_real_paragraphs, config=config)

    # ---------- (2) Flatten multi-paragraph title block (your existing logic) ----------
    title_start_idx = None
    title_end_idx = None

    for new_idx, (old_idx, p) in enumerate(tmp_real_paragraphs):
        if new_idx in tmp_header_indices:
            continue
        # Find first non-header paragraph that looks like a title
        if is_probable_title_paragraph(p, config=config):
            title_start_idx = new_idx
            title_end_idx = new_idx
            # Extend the title block through all immediately following paragraphs
            # that look like subtitle/source lines *or* begin with lowercase “in ”
            # or continue the title phrase.
            for look_ahead in range(new_idx + 1, len(tmp_real_paragraphs)):
                _, q = tmp_real_paragraphs[look_ahead]
                flat_q, _ = flatten_paragraph_without_labels(q)
                text_q = flat_q.strip()
                if not text_q:
                    continue

                if not looks_like_full_sentence(text_q):
                    title_end_idx = look_ahead
                else:
                    break

            break

    # If we found a contiguous block of title paragraphs, merge them
    if title_start_idx is not None and title_end_idx is not None and title_end_idx > title_start_idx:
        base_para = tmp_real_paragraphs[title_start_idx][1]
        parts = []
        for merge_idx in range(title_start_idx, title_end_idx + 1):
            _, para = tmp_real_paragraphs[merge_idx]
            text = para.text or ""
            if text.strip():
                # Separate merged pieces with a single space
                parts.append(text.strip())
            if merge_idx != title_start_idx:
                parent = para._element.getparent()
                if parent is not None:
                    parent.remove(para._element)
        base_para.text = " ".join(parts)

        # NEW: ensure the merged student title uses Times New Roman 12
        for run in base_para.runs:
            enforce_font(run)


    # ------------------------------------------------------------------
    # PEEL PRE-PROCESSING: FLATTEN PARAGRAPH BREAKS INTO ONE LOGICAL BODY
    # ------------------------------------------------------------------
    if config.mode == "peel_paragraph":
        # First pass: build a temporary real_paragraphs list so we can
        # detect MLA header lines and the first real content paragraph.
        tmp_real_paragraphs = [
            (i, p) for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        ]

        tmp_header_indices = detect_mla_header_indices(
            tmp_real_paragraphs,
            config=config,
        )

        # Find the first non-header, non-title paragraph.
        # This is the start of the logical PEEL paragraph.
        peel_start_idx = None
        for new_idx, (old_idx, p) in enumerate(tmp_real_paragraphs):
            if new_idx in tmp_header_indices:
                continue

            flat_text, _ = flatten_paragraph_without_labels(p)
            text = flat_text.strip()
            if not text:
                continue

            if is_probable_title_paragraph(p, config=config):
                continue

            peel_start_idx = new_idx
            break

        if peel_start_idx is not None:
            # Base paragraph that will hold the entire PEEL paragraph text
            base_para = tmp_real_paragraphs[peel_start_idx][1]

            parts = [base_para.text or ""]
            # Merge every later *content* paragraph into base_para
            for _, p in tmp_real_paragraphs[peel_start_idx + 1:]:
                text = p.text or ""
                if text.strip():
                    # Separate paragraphs with a single space so sentences
                    # don't jam together.
                    parts.append(" " + text)

                # Physically remove this paragraph node from the document
                parent = p._element.getparent()
                parent.remove(p._element)

            base_para.text = "".join(parts)

    # After optional PEEL flattening, recompute real_paragraphs as usual.
    real_paragraphs = [
        (i, p) for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]
    
    total_real_paras = len(real_paragraphs)

    # ====================================================================
    # STEP 1: DETECT AND SKIP MLA-STYLE HEADER BLOCK
    # ====================================================================
    # Detect top-of-document MLA-style header lines (name, teacher, course, date).
    # These "garbage" lines should not confuse title/intro detection.
    header_indices = detect_mla_header_indices(real_paragraphs, config=config)

    # ====================================================================
    # STEP 2: FIND THE INTRODUCTION PARAGRAPH
    # ====================================================================
    # Strategy:
    #   1. Try to find the first non-header, non-title paragraph with
    #      more than one sentence (normal full essay intros).
    #   2. If none exists (e.g., Foundation 1: one-sentence intro),
    #      fall back to the first non-header, non-title paragraph.
    intro_idx = None

    # Pass 1: prefer multi-sentence intro
    for new_idx, (old_idx, p) in enumerate(real_paragraphs):
        if new_idx in header_indices:
            continue

        flat_text, _ = flatten_paragraph_without_labels(p)
        text = flat_text.strip()
        if not text:
            continue

        if is_probable_title_paragraph(p, config=config):
            continue

        # Count true sentences based on .?! outside quotes
        quote_spans = compute_quote_spans(text)
        sentence_ending_count = 0
        for i, ch in enumerate(text):
            if ch in {".", "?", "!"} and not pos_in_spans(i, quote_spans):
                    sentence_ending_count += 1
        
        if sentence_ending_count > 1:
            intro_idx = new_idx
            break

    # Pass 2: fallback – first non-header, non-title paragraph
    if intro_idx is None:
        for new_idx, (old_idx, p) in enumerate(real_paragraphs):
            if new_idx in header_indices:
                continue
        
            flat_text, _ = flatten_paragraph_without_labels(p)
            text = flat_text.strip()
            if not text:
                continue

            if is_probable_title_paragraph(p, config=config):
                continue

            intro_idx = new_idx
            break

    # Extreme fallback: if somehow still None, stick with 0
    if intro_idx is None:
        intro_idx = 0

    # ====================================================================
    # STEP 3: PROCESS AND MARK EACH PARAGRAPH
    # ====================================================================
    # Loop over real_paragraphs to apply analytic rules.
    # Important: Skip MLA header indices completely - they should not get labels,
    # not be title-enforced, not be treated as intro/body/conclusion.
    
    # Foundation 4: Track whether student wrote a first body paragraph
    saw_body_para = False
    prev_body_last_sentence_content_words: set[str] | None = None
    
    for new_idx, (old_idx, p) in enumerate(real_paragraphs):
        # Skip MLA-style header lines entirely (name, teacher, class, date, etc.)
        # These should not be analyzed or marked in any way.
        if new_idx in header_indices:
            continue

        # If this is a title-like paragraph, enforce essay title format
        # and title capitalization, then skip other rules
        if is_probable_title_paragraph(p, config=config):
            flat_text, seg = flatten_paragraph_without_labels(p)
            title_text = flat_text.strip()

            # Strip any existing Vysti label if we're re-running on a marked doc
            if "→" in title_text:
                title_text = title_text.split("→", 1)[0].strip()

            # Special case: subtitle/source line like
            #   in Toni Morrison's "Strangers"
            # Treat it as part of the title block, but don't enforce
            # essay-title format or title-capitalization on it.
            if is_source_reference_subtitle_line(title_text, config=config):
                spans = compute_quote_spans(flat_text)
                title_marks = collect_text_title_format_marks(
                    paragraph=p,
                    flat_text=flat_text,
                    segments=seg,
                    spans=spans,
                    config=config,
                    labels_used=labels_used,
                )
                # Always rebuild the paragraph so any stale labels disappear
                title_sentences = [(0, len(flat_text))] if flat_text else None
                apply_marks(p, flat_text, seg, title_marks, sentences=title_sentences, paragraph_index=new_idx)
                continue

            title_marks = []

            # ------------------------------------
            # Rule A: Essay title format
            # ------------------------------------
            title_note = "Essay title format"
            m = TITLE_PATTERN.match(title_text)
            m_no_colon = TITLE_PATTERN_NO_COLON.match(title_text)
            topic_too_thin = False

            

            if m:
                topic_segment = m.group(2).strip()
                topic_too_thin = topic_segment_is_too_thin(topic_segment)
            elif m_no_colon:
                # Also check topic segment for no-colon pattern
                topic_segment = m_no_colon.group(2).strip()
                topic_too_thin = topic_segment_is_too_thin(topic_segment)
           

            # NEW: if the title explicitly includes any of the configured author names, allow it
            # without flagging "Essay title format", even if it doesn't follow
            # the strict colon pattern.
            author_in_title = False
            norm_title = normalize_title_key(title_text)
            if norm_title:
                for norm_author in iter_author_names(config):
                    # Full-name match
                    if norm_author in norm_title:
                        author_in_title = True
                        break
                    # Also allow a last-name-only match (e.g. "Kristof's \"Saudis in Bikinis\"")
                    parts = norm_author.split()
                    if parts:
                        last = parts[-1]
                        if last and last in norm_title:
                            author_in_title = True
                            break

            # If the title does NOT match the recommended colon pattern OR
            # the topic segment is basically empty (e.g. just "in"),
            # flag Essay title format.
            # Note: Titles matching TITLE_PATTERN_NO_COLON are recognized as titles
            # but still flagged here to suggest the colon format.
            if config.enforce_essay_title_format and (not m or topic_too_thin) and not author_in_title:
                anchor_pos = len(flat_text) + 1
                title_marks.append({
                    "start": anchor_pos,
                    "end": anchor_pos,
                    "note": title_note,
                    "color": None,
                    "label": True,
                })
                if title_note not in labels_used:
                    labels_used.append(title_note)

            # ------------------------------------
            # Rule B: Capitalize the words in titles
            # ------------------------------------
            cap_note = "Capitalize the words in titles"

            # Only bother with capitalization if there IS a real topic segment.
            if (
                config.enforce_essay_title_capitalization
                and not topic_too_thin
                and title_needs_capitalization_fix(title_text)
            ):
                # We want to gray-highlight the specific words that violate
                # our capitalization heuristics (e.g. "reversal" in
                #   "A Fearsome Entity": Ironic reversal in "Black Men...")
                # so the student can see exactly what to fix.
                #
                # Use the same logic as title_needs_capitalization_fix:
                #   - for creative titles, work on the topic segment (middle part)
                #   - otherwise, work on the whole title_text.
                m_cap = TITLE_PATTERN.match(title_text)
                m_cap_no_colon = TITLE_PATTERN_NO_COLON.match(title_text)

                if m_cap:
                    target = m_cap.group(2)
                    target_offset_in_title = m_cap.start(2)
                elif m_cap_no_colon:
                    target = m_cap_no_colon.group(2)
                    target_offset_in_title = m_cap_no_colon.start(2)
                else:
                    target = title_text
                    target_offset_in_title = 0

                # Where does title_text begin inside flat_text?
                base_offset = flat_text.find(title_text)
                if base_offset < 0:
                    base_offset = 0  # defensive fallback

                # Same ALWAYS_LOWER list as in title_needs_capitalization_fix
                ALWAYS_LOWER = {
                    "a", "an", "the",
                    "and", "but", "or", "nor", "for", "so", "yet",
                    "of", "in", "on", "at", "by", "to", "up",
                    "from", "into", "onto", "over", "under",
                    "with", "within", "without", "about",
                    "between", "before", "after", "around",
                    "through", "during", "inside", "outside",
                    "as", "than", "via", "per",
                    "is", "are", "was", "were", "be", "been", "being", "am",
                    "do", "does", "did",
                    "has", "have", "had",
                    "can", "could", "shall", "should",
                    "will", "would", "may", "might", "must",
                }

                # Scan the target segment for alphabetic words and mark the ones
                # that should have been capitalized but weren't.
                for idx, match in enumerate(
                    re.finditer(r"[A-Za-z]+(?:'[A-Za-z]+)?", target)
                ):
                    w = match.group(0)
                    lower = w.lower()

                    # First word in the segment: must be capitalized
                    if idx == 0:
                        needs_cap = not w[0].isupper()
                    else:
                        # Interior words: only enforce caps on content words
                        if lower in ALWAYS_LOWER:
                            needs_cap = False
                        else:
                            needs_cap = not w[0].isupper()

                    if not needs_cap:
                        continue

                    # Compute character offsets in flat_text for this word
                    start_in_title = target_offset_in_title + match.start()
                    end_in_title = target_offset_in_title + match.end()
                    start = base_offset + start_in_title
                    end = base_offset + end_in_title

                    # Gray-highlight the mis-capitalized word
                    title_marks.append({
                        "start": start,
                        "end": end,
                        "color": WD_COLOR_INDEX.GRAY_25,
                    })

                # Attach the usual yellow label at the end of the title paragraph
                anchor_pos = len(flat_text) + 1
                title_marks.append({
                    "start": anchor_pos,
                    "end": anchor_pos,
                    "note": cap_note,
                    "color": None,
                    "label": True,
                })
                if cap_note not in labels_used:
                    labels_used.append(cap_note)

            # ------------------------------------
            # Rule C: Text title formatting in the student essay title
            # ------------------------------------
            # Also enforce proper formatting of the teacher-supplied text title
            spans = compute_quote_spans(flat_text)
            title_format_marks = collect_text_title_format_marks(
                paragraph=p,
                flat_text=flat_text,
                segments=seg,
                spans=spans,
                config=config,
                labels_used=labels_used,
            )
            if title_format_marks:
                title_marks.extend(title_format_marks)

            # Apply title-related marks and always rebuild the title paragraph,
            # even when there are no new title issues, so stale labels disappear.
            title_sentences = [(0, len(flat_text))] if flat_text else None
            apply_marks(p, flat_text, seg, title_marks, sentences=title_sentences, paragraph_index=new_idx)

            # Skip further analysis of the title paragraph
            continue

        # Determine paragraph role for praise logic
        paragraph_role = get_paragraph_role(new_idx, intro_idx, total_real_paras, config=config)
        if paragraph_role != "body":
            prev_body_last_sentence_content_words = None
        
        # Foundation 4: Track if we see a body paragraph
        if config.mode == "foundation_4" and paragraph_role == "body":
            saw_body_para = True

        # =====================================================================
        # FOUNDATION ASSIGNMENT 1 — EXTRA PARAGRAPHS (AFTER INTRO)
        # =====================================================================
        # Foundation Assignment 1 requires ONLY the first sentence of the essay.
        #
        # Any paragraphs written after the intro paragraph (body paragraphs,
        # conclusion, etc.) are "extra content" and should be visually penalized:
        #   1. Red highlight + strikethrough across the ENTIRE paragraph
        #   2. The position of the last extra paragraph is tracked globally so we
        #      can attach a single yellow label at the very end (see label
        #      insertion code after this paragraph loop)
        #   3. No other analytic rules are applied (we skip normal analysis)
        #
        # This handles cases like Foundation_HW1_Alex where the student writes
        # extra paragraphs after the intro paragraph.
        #
        # Logic: paragraph_role != "intro" catches ALL non-intro paragraphs
        # (body, conclusion, and any other content after the intro).
        # =====================================================================
        if config.mode == "foundation_1" and paragraph_role != "intro":
            prev_body_last_sentence_content_words = None
            # Get the flat text for this extra paragraph
            flat_text, seg = flatten_paragraph(p, skip_vysti=False)
            
            # Only process if there's actual content (not just whitespace)
            if flat_text.strip():
                # Apply red highlight + strikethrough to the ENTIRE paragraph
                # This matches the "deleted word" style (see lines ~2411-2418)
                marks = [
                    {
                        "start": 0,
                        "end": len(flat_text),
                        "note": None,
                        "color": WD_COLOR_INDEX.RED,  # Red highlight
                        "strike": True,  # Strikethrough
                    }
                ]
                
                # Track this extra paragraph's end position for label placement
                # Each extra paragraph overwrites the previous target, so we end up
                # with the position of the LAST extra paragraph in the document
                FOUNDATION1_LABEL_TARGET = (new_idx, len(flat_text))
                
                # Apply the marks and skip normal analysis entirely
                # (no weak verbs, no quotation rules, no off-topic checks, etc.)
                apply_marks(p, flat_text, seg, marks, sentences=None, paragraph_index=new_idx)
                continue
            else:
                # Empty paragraph (only whitespace) - skip it entirely
                continue

        # =====================================================================
        # FOUNDATION ASSIGNMENT 2 — ONLY INTRO PARAGRAPH IS ANALYZED
        # =====================================================================
        # For Foundation 2, we only want to mark the introduction (first sentence + thesis).
        # Any later body / conclusion paragraphs should be left completely unmarked by the
        # automated rules so the focus stays on intro skills.
        if config.mode == "foundation_2" and paragraph_role != "intro":
            prev_body_last_sentence_content_words = None
            continue

        # All structural quotation rules now live inside analyze_text,
        # using paragraph_index, total_paragraphs, and intro_idx.
        marks, flat_text, seg, sentences, last_sentence_content_words = analyze_text(
            p,
            paragraph_index=new_idx,
            total_paragraphs=total_real_paras,
            labels_used=labels_used,
            intro_idx=intro_idx,
            config=config,
            prev_body_last_sentence_content_words=(
                prev_body_last_sentence_content_words if paragraph_role == "body" else None
            ),
        )
        if paragraph_role == "body":
            prev_body_last_sentence_content_words = last_sentence_content_words

        seen = set()
        for mm in marks:
            note = mm.get("note")
            if not note:
                continue
            # Only count issues that actually belong in the summary
            if note in rules or note in labels_used:
                key = (note, mm.get("start"), mm.get("end"))
                if key not in seen:
                    issue_counts[note] += 1
                    seen.add(key)


        # Count rule-breaks in this paragraph (labels and plain highlights),
        # but ignore TURQUOISE "soft" suggestions and pure device-highlighting
        # when deciding whether to recommend a full rewrite.
        rule_break_count = 0
        for m in marks:
            # Skip minor TURQUOISE nudges (e.g. "this", extra "and", pronoun clarifications)
            if m.get("color") == WD_COLOR_INDEX.TURQUOISE:
                continue
            # Skip GREEN thesis-device highlighting; it's diagnostic, not an error
            if m.get("device_highlight"):
                continue

            rule_break_count += 1

        needs_rewrite_practice = rule_break_count >= 5

        if marks:
            apply_marks(p, flat_text, seg, marks, sentences=sentences, paragraph_index=new_idx)

        # If this paragraph has many rule-breaks, add a red "rewrite" label
        # at the very end of the paragraph, after other yellow labels.
        if needs_rewrite_practice and flat_text and not getattr(config, "student_mode", False):
            # Main label text only: white font, red highlight, underlined, no bold.
            # The leading space keeps it visually separated from the paragraph text.
            lbl = p.add_run(" * Rewrite this paragraph for practice  *")
            enforce_font(lbl)
            lbl.font.highlight_color = WD_COLOR_INDEX.RED
            lbl.font.bold = False
            lbl.font.underline = True
            lbl.font.color.rgb = RGBColor(255, 255, 255)
            lbl._element.set("data-vysti", "yes")

    # =====================================================================
    # FOUNDATION ASSIGNMENT 4 — MISSING FIRST BODY TOPIC SENTENCE
    # =====================================================================
    # For Foundation Assignment 4, if the student wrote a valid thesis but
    # NO first body paragraph topic sentence, attach a yellow label at the
    # END of the thesis sentence.
    # =====================================================================
    if config.mode == "foundation_4":
        if THESIS_PARAGRAPH_INDEX is not None and not saw_body_para:
            # Get the paragraph that contains the thesis
            _, thesis_paragraph = real_paragraphs[THESIS_PARAGRAPH_INDEX]
            
            # Flatten the paragraph to get its text
            flat_text, _ = flatten_paragraph_without_labels(thesis_paragraph)
            
            # Define the assignment note
            assignment_note = "The assignment was to write the introduction and the first topic sentence"
            
            # Append the label as a run at the end of the thesis sentence
            # Format: " → The assignment was to write the introduction and the first topic sentence"
            # Style: bold, yellow highlight, black font, no underline
            lbl_run = thesis_paragraph.add_run(f" → {assignment_note}")
            enforce_font(lbl_run)
            lbl_run.font.bold = True
            lbl_run.font.underline = False
            lbl_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Keep font color black (default) - don't set lbl_run.font.color
            lbl_run._element.set("data-vysti", "yes")
            
            # Register this note in labels_used so it appears in the Issues/Explanation table
            if assignment_note not in labels_used:
                labels_used.append(assignment_note)

    # =====================================================================
    # FOUNDATION ASSIGNMENT 1 — INSERT SINGLE LABEL AT END
    # =====================================================================
    # For Foundation Assignment 1, after processing all paragraphs, insert a
    # single yellow label at the end of the last piece of extra content
    # (whether that's an extra sentence in the intro or an extra paragraph).
    # =====================================================================
    if config.mode == "foundation_1" and FOUNDATION1_LABEL_TARGET is not None:
        target_para_idx, target_char_pos = FOUNDATION1_LABEL_TARGET
        
        # Get the actual docx paragraph for that target
        _, target_paragraph = real_paragraphs[target_para_idx]
        
        # Define the standardized assignment note (shorter version)
        assignment_note = "The assignment is to write the first sentence"
        
        # Append the label as a run at the end of that paragraph
        # Format: " → The assignment is to write the first sentence"
        # Style: bold, yellow highlight, black font, no underline
        lbl_run = target_paragraph.add_run(f" → {assignment_note}")
        enforce_font(lbl_run)
        lbl_run.font.bold = True
        lbl_run.font.underline = False
        lbl_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        # Keep font color black (default) - don't set lbl_run.font.color
        lbl_run._element.set("data-vysti", "yes")
        
        # Register this note in labels_used so it appears in the Issues/Explanation table
        if assignment_note not in labels_used:
            labels_used.append(assignment_note)
    print(issue_counts.most_common(10))
    # After processing all real paragraphs and applying marks
    add_summary_table(doc, labels_used, rules, issue_counts=issue_counts)


    output_path = os.path.splitext(essay_path)[0] + "_marked.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Run the Vysti marker on one or more .docx essays.",
    )

    parser.add_argument(
        "essays",
        nargs="+",
        help="Path(s) to .docx essay file(s) to mark",
    )

    parser.add_argument(
        "--mode",
        default="textual_analysis",
        choices=[
            "textual_analysis",
            "intertextual_analysis",
            "reader_response",
            "no_title",
            "image_analysis",
            "argumentation",
            "analytic_frame",
            "foundation_1",
            "foundation_2",
            "foundation_3",
            "foundation_4",
            "foundation_5",
            "peel_paragraph",
        ],
        help="Assignment mode (default: textual_analysis)",
    )

    parser.add_argument(
        "--author",
        help="Primary author name (e.g. 'M. F. K. Fisher')",
    )

    parser.add_argument(
        "--title",
        help="Primary text title (e.g. 'Young Hunger')",
    )

    parser.add_argument(
        "--author2",
        help="Second author name (for intertextual/analytic_frame modes)",
    )

    parser.add_argument(
        "--title2",
        help="Second text title (for intertextual/analytic_frame modes)",
    )

    parser.add_argument(
        "--author3",
        help="Third author name (for intertextual/analytic_frame modes)",
    )

    parser.add_argument(
        "--title3",
        help="Third text title (for intertextual/analytic_frame modes)",
    )

    parser.add_argument(
        "--rules",
        default="Vysti Rules for Writing.xlsx",
        help="Path to the rules Excel file (default: Vysti Rules for Writing.xlsx)",
    )

    args = parser.parse_args()

    # Build config from mode + metadata
    config = get_preset_config(args.mode)
    config.author_name = args.author
    config.text_title = args.title
    config.author_name_2 = args.author2
    config.text_title_2 = args.title2
    config.author_name_3 = args.author3
    config.text_title_3 = args.title3

    for essay_path in args.essays:
        try:
            output_path = run_marker(
                essay_path,
                rules_path=args.rules,
                config=config,
            )
            print(f"Marked: {essay_path} -> {output_path}")
        except Exception as e:
            print(f"Error processing {essay_path}: {e}", file=sys.stderr)
