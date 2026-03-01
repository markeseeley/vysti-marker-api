# vysti_api.py

import os
import io
from io import BytesIO

import httpx
from collections import Counter
from fastapi import (
    FastAPI,
    File,
    UploadFile,
    Form,
    Depends,
    HTTPException,
    status,
)
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re
import difflib


app = FastAPI(title="Vysti Marker API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # loosened for dev; you can tighten this later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===== Lazy engine loader =====
_ENGINE = None

def get_engine():
    """
    Lazy-load the marker engine so the FastAPI app can boot even if spaCy/model
    import is slow or fails. Import errors become a 503 at request-time instead
    of killing the whole ASGI startup.
    """
    global _ENGINE
    if _ENGINE is not None:
        return _ENGINE
    try:
        from marker import mark_docx_bytes, extract_summary_metadata
        _ENGINE = (mark_docx_bytes, extract_summary_metadata)
        return _ENGINE
    except Exception as e:
        print("Failed to import marker engine:", repr(e))
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Marker engine failed to load on server. Check Render logs for the real import error.",
        )

# ===== Supabase config (from environment variables) =====
# ===== Supabase config (from environment variables) =====
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

auth_scheme = HTTPBearer(auto_error=False)


# ===== Pydantic models =====
class RevisionCheckRequest(BaseModel):
    label: str
    label_trimmed: str | None = None
    rewrite: str
    mode: str | None = None
    context_text: str | None = None
    original_sentence: str | None = None
    paragraph_index: int | None = None
    titles: list["TitleInfo"] | None = None


class TitleInfo(BaseModel):
    author: str
    title: str
    is_minor: bool = True


class MarkTextRequest(BaseModel):
    file_name: str
    text: str
    mode: str = "student"
    titles: list[TitleInfo] | None = None
    student_mode: bool = True
    include_summary_table: bool | None = False


class ExportDocxRequest(BaseModel):
    file_name: str
    text: str



async def get_current_user(
    cred: HTTPAuthorizationCredentials = Depends(auth_scheme),
):
    """
    Validate the Supabase JWT by calling Supabase's Auth API.
    Returns the user dict if valid; otherwise raises 401.
    """
    if cred is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing Authorization header",
        )

    token = cred.credentials

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise HTTPException(
            status_code=500,
            detail="Supabase config missing on server",
        )

    auth_url = f"{SUPABASE_URL}/auth/v1/user"

    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.get(
            auth_url,
            headers={
                "apikey": SUPABASE_ANON_KEY,
                "Authorization": f"Bearer {token}",
            },
        )

    if resp.status_code != 200:
        # Invalid or expired token
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired token",
        )

    return resp.json()


@app.get("/")
def read_root():
    return {"status": "ok", "message": "Vysti marker API is running"}


@app.post("/mark")
async def mark_essay(
    file: UploadFile = File(...),
    mode: str = Form("textual_analysis"),
    user: dict = Depends(get_current_user),  # <-- require Supabase auth
    include_summary_table: bool = Form(True),
    student_mode: bool | None = Form(None),

    # Primary work
    author: str | None = Form(None),
    title: str | None = Form(None),
    text_is_minor_work: bool | None = Form(None),

    # New metadata (optional)
    student_name: str | None = Form(None),
    assignment_name: str | None = Form(None),
    class_id: str | None = Form(None),

    # Second work (optional)
    author2: str | None = Form(None),
    title2: str | None = Form(None),
    text_is_minor_work_2: bool | None = Form(None),

    # Third work (optional)
    author3: str | None = Form(None),
    title3: str | None = Form(None),
    text_is_minor_work_3: bool | None = Form(None),

    # Rule toggles (optional)
    forbid_personal_pronouns: bool | None = Form(None),
    forbid_audience_reference: bool | None = Form(None),
    enforce_closed_thesis: bool | None = Form(None),
    require_body_evidence: bool | None = Form(None),
    allow_intro_summary_quotes: bool | None = Form(None),
    enforce_intro_quote_rule: bool | None = Form(None),
    enforce_long_quote_rule: bool | None = Form(None),
    enforce_contractions_rule: bool | None = Form(None),
    enforce_which_rule: bool | None = Form(None),
    enforce_weak_verbs_rule: bool | None = Form(None),
    enforce_fact_proof_rule: bool | None = Form(None),
    enforce_human_people_rule: bool | None = Form(None),
    enforce_vague_terms_rule: bool | None = Form(None),
    enforce_sva_rule: bool | None = Form(None),
    enforce_present_tense_rule: bool | None = Form(None),
    highlight_thesis_devices: bool | None = Form(None),
):
    """
    Mark a .docx essay using the Vysti engine.

    Accepts up to three works (author/title) plus:
      - text_is_minor_work, text_is_minor_work_2, text_is_minor_work_3
      - forbid_personal_pronouns
      - enforce_closed_thesis

    These map directly onto MarkerConfig in marker.py.
    """
    # 1. Basic validation
    if not file.filename.lower().endswith(".docx"):
        return JSONResponse(
            status_code=400,
            content={"error": "Please upload a .docx file"},
        )

    # 2. Validate class_id if provided
    class_id_validated = None
    if class_id:
        user_id = user.get("id") if isinstance(user, dict) else None
        if not user_id:
            return JSONResponse(
                status_code=400,
                content={"error": "Invalid user"},
            )
        
        # Verify class exists and belongs to user
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            db_url = f"{SUPABASE_URL}/rest/v1/classes"
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.get(
                    f"{db_url}?id=eq.{class_id}&user_id=eq.{user_id}&archived=eq.false",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if not data or len(data) == 0:
                        return JSONResponse(
                            status_code=400,
                            content={"error": "Class not found or does not belong to user"},
                        )
                    class_id_validated = class_id
                else:
                    return JSONResponse(
                        status_code=400,
                        content={"error": "Failed to validate class"},
                    )

    # 3. Read file bytes
    docx_bytes = await file.read()

    # 4. Build teacher_config from form fields (matches MarkerConfig)
    teacher_config: dict = {}

    # --- Works / titles ---
    if author:
        teacher_config["author_name"] = author
    if title:
        teacher_config["text_title"] = title
    if text_is_minor_work is not None:
        teacher_config["text_is_minor_work"] = text_is_minor_work

    if author2:
        teacher_config["author_name_2"] = author2
    if title2:
        teacher_config["text_title_2"] = title2
    if text_is_minor_work_2 is not None:
        teacher_config["text_is_minor_work_2"] = text_is_minor_work_2

    if author3:
        teacher_config["author_name_3"] = author3
    if title3:
        teacher_config["text_title_3"] = title3
    if text_is_minor_work_3 is not None:
        teacher_config["text_is_minor_work_3"] = text_is_minor_work_3

    # --- Rule overrides ---
    # These override the defaults chosen by get_preset_config(mode)
    if forbid_personal_pronouns is not None:
        teacher_config["forbid_personal_pronouns"] = forbid_personal_pronouns
    if forbid_audience_reference is not None:
        teacher_config["forbid_audience_reference"] = forbid_audience_reference
    if enforce_closed_thesis is not None:
        teacher_config["enforce_closed_thesis"] = enforce_closed_thesis
    if require_body_evidence is not None:
        teacher_config["require_body_evidence"] = require_body_evidence
    if allow_intro_summary_quotes is not None:
        teacher_config["allow_intro_summary_quotes"] = allow_intro_summary_quotes
    if enforce_intro_quote_rule is not None:
        teacher_config["enforce_intro_quote_rule"] = enforce_intro_quote_rule
    if enforce_long_quote_rule is not None:
        teacher_config["enforce_long_quote_rule"] = enforce_long_quote_rule
    if enforce_contractions_rule is not None:
        teacher_config["enforce_contractions_rule"] = enforce_contractions_rule
    if enforce_which_rule is not None:
        teacher_config["enforce_which_rule"] = enforce_which_rule
    if enforce_weak_verbs_rule is not None:
        teacher_config["enforce_weak_verbs_rule"] = enforce_weak_verbs_rule
    if enforce_fact_proof_rule is not None:
        teacher_config["enforce_fact_proof_rule"] = enforce_fact_proof_rule
    if enforce_human_people_rule is not None:
        teacher_config["enforce_human_people_rule"] = enforce_human_people_rule
    if enforce_vague_terms_rule is not None:
        teacher_config["enforce_vague_terms_rule"] = enforce_vague_terms_rule
    if enforce_sva_rule is not None:
        teacher_config["enforce_sva_rule"] = enforce_sva_rule
    if enforce_present_tense_rule is not None:
        teacher_config["enforce_present_tense_rule"] = enforce_present_tense_rule
    if highlight_thesis_devices is not None:
        teacher_config["highlight_thesis_devices"] = highlight_thesis_devices
    if student_mode is True:
        teacher_config["student_mode"] = True

    # 5. Call your engine
    mark_docx_bytes, _ = get_engine()
    marked_bytes, metadata = mark_docx_bytes(
        docx_bytes,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
        include_summary_table=include_summary_table,
    )

    # You can watch this in Render logs to conafirm the flags:
    print("Vysti metadata:", metadata)
    print("Teacher config used:", teacher_config)

    # ----- Extract examples from metadata -----
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []

    # ----- Count yellow labels from metadata -----
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []

    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl:
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        label_counter[lbl] += (cnt_i if cnt_i > 0 else 1)

    total_labels = sum(label_counter.values())



    # Log usage in Supabase mark_events (best-effort; do not break marking if this fails)
    mark_event_id = None
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            user_id = user.get("id") if isinstance(user, dict) else None
            db_url = f"{SUPABASE_URL}/rest/v1/mark_events?select=id"
            payload = {
                "user_id": user_id,
                "file_name": file.filename,
                "mode": mode,
                "bytes": len(docx_bytes),
                "student_name": student_name,
                "assignment_name": assignment_name,
                "class_id": class_id_validated,
                "total_labels": total_labels,
                "label_counts": dict(label_counter),
                "issues": issues,
            }

            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.post(
                    db_url,
                    json=payload,
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=representation",
                    },
                )
                # Capture mark_event_id from response
                if resp.status_code >= 200 and resp.status_code < 300:
                    resp_data = resp.json()
                    if resp_data and isinstance(resp_data, list) and len(resp_data) > 0:
                        mark_event_id = resp_data[0].get("id")
    except Exception as e:
        print("Failed to log mark_event:", repr(e))

    # Log examples to Supabase issue_examples (best-effort; do not break marking if this fails)
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and examples:
            user_id = user.get("id") if isinstance(user, dict) else None
            if user_id:
                example_rows = []
                for ex in examples:
                    if not isinstance(ex, dict):
                        continue
                    label = ex.get("label")
                    sentence = ex.get("sentence")
                    paragraph_index = ex.get("paragraph_index")
                    if not label or not sentence:
                        continue
                    example_row = {
                        "user_id": user_id,
                        "class_id": class_id_validated,
                        "assignment_name": assignment_name,
                        "student_name": student_name,
                        "mode": mode,
                        "file_name": file.filename,
                        "label": label,
                        "sentence": sentence,
                        "paragraph_index": paragraph_index,
                    }
                    # Include mark_event_id if we captured it
                    if mark_event_id:
                        example_row["mark_event_id"] = mark_event_id
                    example_rows.append(example_row)
                
                if example_rows:
                    db_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                    async with httpx.AsyncClient(timeout=5) as client:
                        await client.post(
                            db_url,
                            json=example_rows,
                            headers={
                                "apikey": SUPABASE_SERVICE_KEY,
                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                "Content-Type": "application/json",
                                "Prefer": "return=minimal",
                            },
                        )
    except Exception as e:
        print("Failed to log issue_examples:", repr(e))

    # 5. Stream the marked .docx back to the client
    base_name = file.filename.rsplit(".", 1)[0] if file.filename else "essay"
    output_filename = f"{base_name}_marked.docx"


    return StreamingResponse(
        io.BytesIO(marked_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
    )


@app.post("/export_docx")
async def export_docx(
    request: ExportDocxRequest,
    user: dict = Depends(get_current_user),
):
    """
    Export a clean .docx from plain text.
    This is for Student mode "Download revised essay" — no Vysti marks, no summary table.
    """
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="Missing text")

    # Build a clean document (NO marking)
    docx_bytes = build_doc_from_text(request.text)

    safe_name = request.file_name.strip() if request.file_name else "essay_revised.docx"
    if not safe_name.lower().endswith(".docx"):
        safe_name += ".docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@app.post("/ingest_marked")
async def ingest_marked_essay(
    file: UploadFile = File(...),
    student_name: str | None = Form(None),
    assignment_name: str | None = Form(None),
    mode: str = Form("imported_marked"),
    class_id: str | None = Form(None),
    user: dict = Depends(get_current_user),  # <-- require Supabase auth
):
    """
    Ingest an already-marked .docx file by extracting label counts from the summary table
    and logging it to mark_events.
    """
    # 1. Basic validation
    if not file.filename.lower().endswith(".docx"):
        return JSONResponse(
            status_code=400,
            content={"error": "Please upload a .docx file"},
        )

    # 2. Read file bytes
    docx_bytes = await file.read()

    # 3. Parse the document to extract summary metadata
    try:
        doc = Document(BytesIO(docx_bytes))
        _, extract_summary_metadata = get_engine()
        metadata = extract_summary_metadata(doc)
    except Exception as e:
        return JSONResponse(
            status_code=400,
            content={"error": f"Failed to parse document: {str(e)}"},
        )

    # 4. Validate that this is a marked document
    issues = metadata.get("issues", [])
    if not issues:
        return JSONResponse(
            status_code=400,
            content={"error": "This doc doesn't appear to be a Vysti-marked file"},
        )

    # 5. Build label_counts from issues
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl:
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        # Ensure count is at least 1
        label_counter[lbl] += max(cnt_i, 1)

    total_labels = sum(label_counter.values())

    # 6. Validate class_id if provided
    class_id_validated = None
    if class_id:
        user_id = user.get("id") if isinstance(user, dict) else None
        if not user_id:
            raise HTTPException(
                status_code=400,
                detail="Invalid user",
            )
        
        # Verify class exists and belongs to user
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            db_url = f"{SUPABASE_URL}/rest/v1/classes"
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.get(
                    f"{db_url}?id=eq.{class_id}&user_id=eq.{user_id}&archived=eq.false",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if not data or len(data) == 0:
                        raise HTTPException(
                            status_code=400,
                            detail="Class not found or does not belong to user",
                        )
                    class_id_validated = class_id
                else:
                    raise HTTPException(
                        status_code=400,
                        detail="Failed to validate class",
                    )

    # 7. Log to Supabase mark_events
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        raise HTTPException(
            status_code=500,
            detail="Supabase configuration missing on server",
        )

    user_id = user.get("id") if isinstance(user, dict) else None
    db_url = f"{SUPABASE_URL}/rest/v1/mark_events"
    payload = {
        "user_id": user_id,
        "file_name": file.filename,
        "mode": mode,
        "bytes": len(docx_bytes),
        "student_name": student_name,
        "assignment_name": assignment_name,
        "class_id": class_id_validated,
        "total_labels": total_labels,
        "label_counts": dict(label_counter),
        "issues": issues,
    }

    try:
        async with httpx.AsyncClient(timeout=5) as client:
            resp = await client.post(
                db_url,
                json=payload,
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json",
                    "Prefer": "return=minimal",
                },
            )
            # Check if the insert was successful
            if resp.status_code < 200 or resp.status_code >= 300:
                error_msg = resp.text or f"Supabase insert failed with status {resp.status_code}"
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to log mark event to database: {error_msg}",
                )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to log mark event: {str(e)}",
        )

    # 8. Return success response
    return JSONResponse(
        content={
            "ok": True,
            "total_labels": total_labels,
            "label_counts": dict(label_counter),
            "issues": issues,
        }
    )


def normalize_label(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip().lower()


def normalize_text(value: str | None) -> str:
    if not value:
        return ""
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    normalized = (
        normalized.replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201C", "\"")
        .replace("\u201D", "\"")
    )
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def split_into_paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    rewrite_pattern = r"\s*\*\s*Rewrite this paragraph for practice\s*\*\s*"
    normalized = re.sub(rewrite_pattern, "", normalized, flags=re.IGNORECASE)
    para_chunks = re.split(r"\n{2,}", normalized)
    paragraphs: list[str] = []
    for chunk in para_chunks:
        para_text = re.sub(r"\n+", " ", chunk).strip()
        if para_text:
            paragraphs.append(para_text)
    return paragraphs


def apply_rewrite_to_paragraph(
    paragraphs: list[str],
    paragraph_index: int | None,
    original_sentence: str,
    rewrite: str,
) -> tuple[list[str], bool, int | None]:
    if not paragraphs or not original_sentence:
        return paragraphs, False, paragraph_index

    if paragraph_index is not None and 0 <= paragraph_index < len(paragraphs):
        para_text = paragraphs[paragraph_index]
        if original_sentence in para_text:
            paragraphs[paragraph_index] = para_text.replace(original_sentence, rewrite, 1)
            return paragraphs, True, paragraph_index

    for idx, para_text in enumerate(paragraphs):
        if original_sentence in para_text:
            paragraphs[idx] = para_text.replace(original_sentence, rewrite, 1)
            return paragraphs, True, idx

    return paragraphs, False, paragraph_index


def build_doc_from_text(text: str) -> bytes:
    # Normalize newlines to \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Safety net: Remove rewrite-practice tag if it appears in the text
    rewrite_pattern = r"\s*\*\s*Rewrite this paragraph for practice\s*\*\s*"
    text = re.sub(rewrite_pattern, "", text, flags=re.IGNORECASE)

    # Split paragraphs on 2+ newlines
    para_chunks = re.split(r"\n{2,}", text)

    # Create document
    doc = Document()

    # Set default style to Times New Roman 12pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    # Set eastAsia font too
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Helper function to detect header-like lines
    def is_header_line(line_text: str) -> bool:
        """Detect if a line is likely a header (teacher name, date, course, etc.)"""
        text_lower = line_text.lower().strip()
        # Check for teacher titles
        if re.match(r"^(mr|ms|mrs|dr|prof)\.?\s+", text_lower):
            return True
        # Check for date patterns (e.g., "January 1, 2024" or "1/1/2024")
        if re.search(r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d+", text_lower):
            return True
        if re.search(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}", line_text):
            return True
        # Check for course/class keywords
        if "course" in text_lower or "class" in text_lower:
            return True
        # Short name-like lines (2-3 words, no sentence-ending punctuation)
        words = line_text.split()
        if len(words) <= 3 and not re.search(r"[.!?]$", line_text):
            return True
        return False

    # Helper function to check if text is a sentence (ends with .?!)
    def is_sentence(line_text: str) -> bool:
        """Check if text appears to be a sentence"""
        return bool(re.search(r"[.!?]$", line_text.strip()))

    # Track if we've found the essay title (first non-header short non-sentence line)
    title_found = False

    # Add paragraphs (collapsing single newlines within paragraphs to spaces)
    for para_chunk in para_chunks:
        # Collapse single newlines within paragraph to spaces
        para_text = re.sub(r"\n+", " ", para_chunk).strip()
        if not para_text:  # Skip empty paragraphs
            continue

        para = doc.add_paragraph(para_text)

        # Check if this is a header line
        is_header = is_header_line(para_text)
        is_sent = is_sentence(para_text)

        # Apply formatting based on paragraph type
        if is_header:
            # Header lines: no indentation, left-aligned
            para.paragraph_format.first_line_indent = Inches(0)
        elif not title_found and not is_header and not is_sent and len(para_text.split()) <= 10:
            # First non-header short non-sentence line: likely essay title - center it
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Inches(0)
            title_found = True
        else:
            # Prose paragraphs: apply MLA first-line indent (0.5")
            para.paragraph_format.first_line_indent = Inches(0.5)

    # If no paragraphs were created, add at least one
    if len(doc.paragraphs) == 0:
        para = doc.add_paragraph(text.strip() or "Empty document")
        para.paragraph_format.first_line_indent = Inches(0.5)

    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    docx_bytes = docx_buffer.getvalue()
    docx_buffer.close()
    return docx_bytes


def build_teacher_config_from_titles(titles: list[TitleInfo] | None) -> dict | None:
    if not titles:
        return {"highlight_thesis_devices": False}

    teacher_config: dict = {"highlight_thesis_devices": False}
    trimmed = titles[:3]
    if len(trimmed) > 0:
        t1 = trimmed[0]
        teacher_config["author_name"] = t1.author
        teacher_config["text_title"] = t1.title
        teacher_config["text_is_minor_work"] = t1.is_minor
    if len(trimmed) > 1:
        t2 = trimmed[1]
        teacher_config["author_name_2"] = t2.author
        teacher_config["text_title_2"] = t2.title
        teacher_config["text_is_minor_work_2"] = t2.is_minor
    if len(trimmed) > 2:
        t3 = trimmed[2]
        teacher_config["author_name_3"] = t3.author
        teacher_config["text_title_3"] = t3.title
        teacher_config["text_is_minor_work_3"] = t3.is_minor
    return teacher_config


def find_matching_examples(
    examples: list,
    label_value: str,
    paragraph_index: int | None,
    original_sentence: str | None,
) -> tuple[list[dict], int | None]:
    normalized_label = normalize_label(label_value)
    label_matches = [
        ex
        for ex in examples
        if isinstance(ex, dict)
        and normalize_label(ex.get("label")) == normalized_label
    ]

    if paragraph_index is not None:
        matches = [ex for ex in label_matches if ex.get("paragraph_index") == paragraph_index]
        return matches, paragraph_index

    if not original_sentence:
        return [], None

    normalized_original = normalize_text(original_sentence)
    best_ratio = 0.0
    best_example = None
    for ex in label_matches:
        sentence = ex.get("sentence") or ""
        ratio = difflib.SequenceMatcher(
            None,
            normalize_text(sentence),
            normalized_original,
        ).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_example = ex

    if best_example and best_ratio >= 0.7:
        target_index = best_example.get("paragraph_index")
        if target_index is None:
            return [best_example], None
        matches = [ex for ex in label_matches if ex.get("paragraph_index") == target_index]
        return matches, target_index

    return [], None


@app.post("/revision/check")
async def check_revision(
    request: RevisionCheckRequest,
    user: dict = Depends(get_current_user),
):
    """
    Check if a rewritten sentence still triggers a specific issue label.
    """
    # Validate rewrite: reject empty/whitespace
    if not request.rewrite or not request.rewrite.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite cannot be empty or whitespace only",
        )
    
    # Cap rewrite length (2000 chars)
    if len(request.rewrite) > 2000:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Rewrite exceeds maximum length of 2000 characters",
        )

    if not request.context_text or not request.original_sentence:
        return JSONResponse(
            content={
                "approved": False,
                "message": "Unable to verify rewrite without Preview context. Load the Preview and try again.",
            }
        )

    if normalize_text(request.rewrite) == normalize_text(request.original_sentence):
        return JSONResponse(
            content={
                "approved": False,
                "message": "No changes detected — edit the example before checking.",
            }
        )

    label_value = request.label_trimmed or request.label
    mode = request.mode or "textual_analysis"
    teacher_config = build_teacher_config_from_titles(request.titles)

    mark_docx_bytes, _ = get_engine()
    doc_before = build_doc_from_text(request.context_text)
    _, metadata_before = mark_docx_bytes(
        doc_before,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
    )

    examples_before = metadata_before.get("examples", []) if isinstance(metadata_before, dict) else []
    matches_before, target_paragraph_index = find_matching_examples(
        examples_before,
        label_value,
        request.paragraph_index,
        request.original_sentence,
    )

    if not matches_before:
        return JSONResponse(
            content={
                "approved": False,
                "message": "Could not locate this issue in the current Preview context. Click 'Recheck my essay' and try again.",
            }
        )

    before_local_count = len(matches_before)

    paragraphs = split_into_paragraphs(request.context_text)
    updated_paragraphs, replaced, used_paragraph_index = apply_rewrite_to_paragraph(
        paragraphs,
        request.paragraph_index,
        request.original_sentence,
        request.rewrite.strip(),
    )

    if not replaced:
        return JSONResponse(
            content={
                "approved": False,
                "message": "Could not apply rewrite to the current Preview text. Try 'Find in preview' and verify the sentence exists.",
            }
        )

    updated_text = "\n\n".join(updated_paragraphs)
    doc_after = build_doc_from_text(updated_text)
    _, metadata_after = mark_docx_bytes(
        doc_after,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
    )

    examples_after = metadata_after.get("examples", []) if isinstance(metadata_after, dict) else []
    effective_paragraph_index = (
        used_paragraph_index
        if used_paragraph_index is not None
        else target_paragraph_index
    )

    matches_after, _ = find_matching_examples(
        examples_after,
        label_value,
        effective_paragraph_index,
        request.original_sentence,
    )
    after_local_count = len(matches_after)

    if after_local_count < before_local_count:
        return JSONResponse(
            content={
                "approved": True,
                "message": "Looks good! Revision approved.",
                "before_local_count": before_local_count,
                "after_local_count": after_local_count,
            }
        )

    return JSONResponse(
        content={
            "approved": False,
            "message": "Still needs revision — the issue is still triggering here.",
            "before_local_count": before_local_count,
            "after_local_count": after_local_count,
        }
    )


@app.post("/mark_text")
async def mark_text(
    request: MarkTextRequest,
    user: dict = Depends(get_current_user),
):
    """
    Mark text content by creating a .docx in memory and running the marking pipeline.
    
    Request JSON:
    {
      "file_name": "OriginalFileName.docx",
      "text": "Full essay text with paragraphs",
      "mode": "student"
    }
    
    Returns the marked .docx bytes (same as /mark).
    """
    # 1. Create .docx from text
    docx_bytes = build_doc_from_text(request.text)
    
    # 2. Build teacher_config from request.titles
    teacher_config = build_teacher_config_from_titles(request.titles) or {}
    teacher_config["student_mode"] = request.student_mode
    
    # 3. Call mark_docx_bytes (same pipeline as /mark)
    mark_docx_bytes, _ = get_engine()
    mode = request.mode or "textual_analysis"
    marked_bytes, metadata = mark_docx_bytes(
        docx_bytes,
        mode=mode,
        teacher_config=teacher_config if teacher_config else None,
        include_summary_table=bool(request.include_summary_table),
    )
    
    # 4. Extract examples and issues from metadata
    examples = metadata.get("examples", []) if isinstance(metadata, dict) else []
    issues = metadata.get("issues", []) if isinstance(metadata, dict) else []
    
    # 4. Count labels
    label_counter = Counter()
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        lbl = issue.get("label")
        if not lbl:
            continue
        cnt = issue.get("count")
        try:
            cnt_i = int(cnt) if cnt is not None else 1
        except Exception:
            cnt_i = 1
        label_counter[lbl] += (cnt_i if cnt_i > 0 else 1)
    
    total_labels = sum(label_counter.values())
    
    # 5. Log to Supabase mark_events (best-effort)
    mark_event_id = None
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY:
            user_id = user.get("id") if isinstance(user, dict) else None
            db_url = f"{SUPABASE_URL}/rest/v1/mark_events?select=id"
            payload = {
                "user_id": user_id,
                "file_name": request.file_name,
                "mode": mode,
                "bytes": len(docx_bytes),
                "total_labels": total_labels,
                "label_counts": dict(label_counter),
                "issues": issues,
            }
            
            async with httpx.AsyncClient(timeout=5) as client:
                resp = await client.post(
                    db_url,
                    json=payload,
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                        "Content-Type": "application/json",
                        "Prefer": "return=representation",
                    },
                )
                if resp.status_code >= 200 and resp.status_code < 300:
                    resp_data = resp.json()
                    if resp_data and isinstance(resp_data, list) and len(resp_data) > 0:
                        mark_event_id = resp_data[0].get("id")
    except Exception as e:
        print("Failed to log mark_event:", repr(e))
    
    # 6. Log examples to Supabase issue_examples (best-effort)
    try:
        if SUPABASE_URL and SUPABASE_SERVICE_KEY and examples:
            user_id = user.get("id") if isinstance(user, dict) else None
            if user_id:
                example_rows = []
                for ex in examples:
                    if not isinstance(ex, dict):
                        continue
                    label = ex.get("label")
                    sentence = ex.get("sentence")
                    paragraph_index = ex.get("paragraph_index")
                    if not label or not sentence:
                        continue
                    example_row = {
                        "user_id": user_id,
                        "file_name": request.file_name,
                        "mode": mode,
                        "label": label,
                        "sentence": sentence,
                        "paragraph_index": paragraph_index,
                    }
                    if mark_event_id:
                        example_row["mark_event_id"] = mark_event_id
                    example_rows.append(example_row)
                
                if example_rows:
                    db_url = f"{SUPABASE_URL}/rest/v1/issue_examples"
                    async with httpx.AsyncClient(timeout=5) as client:
                        await client.post(
                            db_url,
                            json=example_rows,
                            headers={
                                "apikey": SUPABASE_SERVICE_KEY,
                                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                                "Content-Type": "application/json",
                                "Prefer": "return=minimal",
                            },
                        )
    except Exception as e:
        print("Failed to log issue_examples:", repr(e))
    
    # 7. Return marked .docx bytes
    base_name = request.file_name.rsplit(".", 1)[0] if request.file_name else "essay"
    output_filename = f"{base_name}_marked.docx"
    
    return StreamingResponse(
        io.BytesIO(marked_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
    )
